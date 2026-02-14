"""
BEA Lab - Document Upload API
Uploads are automatically pushed to GitHub: papers/evaluated/integrated/
"""
import os, uuid, json, base64, logging, hashlib, time, hmac, random
from datetime import datetime
from pathlib import Path
from typing import Optional, List

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel
from sqlalchemy import create_engine, Column, String, Text, DateTime, Integer, JSON, Boolean, text
from sqlalchemy.orm import declarative_base, sessionmaker

DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./beatrix.db")
UPLOAD_DIR = Path(os.getenv("UPLOAD_DIR", "./uploads"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
MAX_FILE_SIZE = 50 * 1024 * 1024
PORT = int(os.getenv("PORT", "8000"))

# GitHub config
GH_TOKEN = os.getenv("GH_TOKEN", "")
GH_REPO = os.getenv("GH_REPO", "FehrAdvice-Partners-AG/complementarity-context-framework")
GH_UPLOAD_PATH = os.getenv("GH_UPLOAD_PATH", "papers/evaluated/integrated")

# Auth config
JWT_SECRET = os.getenv("JWT_SECRET", "bea-lab-secret-key-change-me-2026")
JWT_EXPIRY = int(os.getenv("JWT_EXPIRY", "86400"))

# Registration domain restriction (comma-separated, empty = all allowed)
# Example: "fehradvice.com,bea-lab.io" → only these domains can register
ALLOWED_EMAIL_DOMAINS = [d.strip().lower() for d in os.getenv("ALLOWED_EMAIL_DOMAINS", "").split(",") if d.strip()]

# Auto-admin emails (comma-separated)
# Example: "gerhard.fehr@fehradvice.com" → these users get admin on registration
ADMIN_EMAILS = [e.strip().lower() for e in os.getenv("ADMIN_EMAILS", "").split(",") if e.strip()]

# Email verification via Resend API (resend.com)
REQUIRE_EMAIL_VERIFICATION = os.getenv("REQUIRE_EMAIL_VERIFICATION", "true").lower() in ("true", "1", "yes")
RESEND_API_KEY = os.getenv("RESEND_API_KEY", "")
EMAIL_FROM = os.getenv("EMAIL_FROM", "BEATRIX Lab <noreply@bea-lab.io>")
APP_URL = os.getenv("APP_URL", "https://www.bea-lab.io")
LINKEDIN_CLIENT_ID = os.getenv("LINKEDIN_CLIENT_ID", "")
LINKEDIN_CLIENT_SECRET = os.getenv("LINKEDIN_CLIENT_SECRET", "")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
# 3-Tier Model Strategy for cost optimization
ANTHROPIC_MODEL_HEAVY = os.getenv("ANTHROPIC_MODEL_HEAVY", "claude-opus-4-6")      # APE Review, Deep Research, Complex Analysis
ANTHROPIC_MODEL_STANDARD = os.getenv("ANTHROPIC_MODEL_STANDARD", "claude-sonnet-4-5")  # RAG Chat, Fact-Check, Belief Analysis
ANTHROPIC_MODEL_LIGHT = os.getenv("ANTHROPIC_MODEL_LIGHT", "claude-haiku-4-5")     # Intent, Triage, Metadata Extraction
# Legacy fallback
ANTHROPIC_MODEL = ANTHROPIC_MODEL_STANDARD
GH_CONTEXT_REPO = os.getenv("GH_CONTEXT_REPO", "FehrAdvice-Partners-AG/complementarity-context-framework")
VOYAGE_API_KEY = os.getenv("VOYAGE_API_KEY", "")
VOYAGE_MODEL = "voyage-3-lite"  # 512 dimensions, fast, free tier 200M tokens/month

# NotebookLM Enterprise API config
GCP_PROJECT_NUMBER = os.getenv("GCP_PROJECT_NUMBER", "368877792942")
GCP_NOTEBOOKLM_LOCATION = os.getenv("GCP_NOTEBOOKLM_LOCATION", "eu")
GCP_OAUTH_CLIENT_ID = os.getenv("GCP_OAUTH_CLIENT_ID", "")
GCP_OAUTH_CLIENT_SECRET = os.getenv("GCP_OAUTH_CLIENT_SECRET", "")
GCP_REFRESH_TOKEN = os.getenv("GCP_REFRESH_TOKEN", "")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("bea-lab")

def hash_password(password, salt=None):
    if salt is None:
        salt = base64.b64encode(os.urandom(16)).decode()
    hashed = hashlib.pbkdf2_hmac('sha256', password.encode(), salt.encode(), 100000)
    return base64.b64encode(hashed).decode(), salt

def verify_password(password, stored_hash, salt):
    computed, _ = hash_password(password, salt)
    return hmac.compare_digest(computed, stored_hash)

def _b64url_encode(data):
    return base64.urlsafe_b64encode(data).rstrip(b'=').decode()

def _b64url_decode(s):
    s += '=' * (4 - len(s) % 4)
    return base64.urlsafe_b64decode(s)

def create_jwt(payload):
    header = _b64url_encode(json.dumps({"alg": "HS256", "typ": "JWT"}).encode())
    payload_enc = _b64url_encode(json.dumps(payload).encode())
    sig_input = f"{header}.{payload_enc}".encode()
    sig = _b64url_encode(hmac.new(JWT_SECRET.encode(), sig_input, hashlib.sha256).digest())
    return f"{header}.{payload_enc}.{sig}"

def verify_jwt(token):
    try:
        parts = token.split('.')
        if len(parts) != 3: return None
        sig_input = f"{parts[0]}.{parts[1]}".encode()
        expected_sig = _b64url_encode(hmac.new(JWT_SECRET.encode(), sig_input, hashlib.sha256).digest())
        if not hmac.compare_digest(parts[2], expected_sig): return None
        payload = json.loads(_b64url_decode(parts[1]))
        if payload.get("exp", 0) < time.time(): return None
        return payload
    except: return None

async def require_auth(request: Request):
    auth = request.headers.get("Authorization", "")
    token = auth.replace("Bearer ", "") if auth.startswith("Bearer ") else ""
    if not token: token = request.cookies.get("bea_token", "")
    payload = verify_jwt(token)
    if not payload: raise HTTPException(401, "Nicht autorisiert")
    return payload

# ═══════════════════════════════════════════════════════════════════════════
# PERMISSION ENGINE v1.0 — RBAC + ReBAC + ABAC Hybrid
# ═══════════════════════════════════════════════════════════════════════════
# Architecture: Google Zanzibar-inspired with Policy-as-Code on GitHub
# - Role Hierarchy: owner > partner > senior_consultant > consultant > researcher > guest
# - Resource Roles: owner, contributor, viewer per entity
# - Attribute Rules: email domain, verification status
# - Audit Trail: logged for critical actions
# ═══════════════════════════════════════════════════════════════════════════

# Role hierarchy — higher index = more power
ROLE_HIERARCHY = {
    "guest": 0,
    "researcher": 1,
    "consultant": 2,
    "senior_consultant": 3,
    "partner": 4,
    "owner": 5,
}

# Permission → minimum required role
# Matches data/config/permissions.yaml on GitHub (Single Source of Truth)
PERMISSION_MAP = {
    # Platform
    "platform.admin_dashboard":     "owner",
    "platform.manage_users":        "owner",
    "platform.manage_roles":        "owner",
    "platform.manage_settings":     "owner",
    "platform.view_analytics":      "partner",
    "platform.manage_permissions":  "owner",
    # Projects
    "project.create":               "consultant",
    "project.read":                 "researcher",
    "project.read_own":             "researcher",
    "project.update":               "partner",
    "project.update_own":           "consultant",
    "project.delete":               "owner",
    "project.manage_team":          "partner",
    "project.assign_code":          "consultant",
    # Leads
    "lead.create":                  "consultant",
    "lead.read":                    "consultant",
    "lead.read_own":                "consultant",
    "lead.update":                  "partner",
    "lead.update_own":              "consultant",
    "lead.delete":                  "owner",
    "lead.manage_pipeline":         "consultant",
    "lead.view_financials":         "senior_consultant",
    # CRM
    "crm.read":                     "consultant",
    "crm.write":                    "senior_consultant",
    "crm.delete":                   "owner",
    "crm.export":                   "partner",
    # Documents
    "document.upload":              "researcher",
    "document.read":                "researcher",
    "document.classify":            "consultant",
    "document.classify_bulk":       "senior_consultant",
    "document.delete":              "partner",
    "document.manage_embeddings":   "owner",
    # Models
    "model.create":                 "consultant",
    "model.read":                   "researcher",
    "model.update":                 "consultant",
    "model.delete":                 "partner",
    # Chat
    "chat.use":                     "researcher",
    "chat.intent":                  "consultant",
    "chat.history":                 "researcher",
    "chat.clear":                   "researcher",
    # User
    "user.read_own_profile":        "guest",
    "user.read_profiles":           "consultant",
    "user.manage_others":           "owner",
}

# Permissions that require @fehradvice.com
FA_ONLY_PERMISSIONS = {
    "project.create", "project.assign_code",
    "lead.create", "lead.read", "lead.update", "lead.update_own",
    "lead.manage_pipeline", "lead.view_financials",
    "crm.read", "crm.write", "crm.delete", "crm.export",
    "chat.intent",
}

# Permissions that MUST be audit-logged
AUDIT_PERMISSIONS = {
    "project.create", "project.delete", "project.manage_team",
    "lead.create", "lead.delete", "lead.manage_pipeline",
    "crm.write", "crm.delete",
    "platform.manage_users", "platform.manage_roles",
    "document.delete", "document.manage_embeddings",
}

def resolve_user_role(user_payload: dict) -> str:
    """Determine effective system role from JWT payload.
    Auto-assignment rules:
      1. is_admin + ADMIN_EMAILS → owner
      2. @fehradvice.com + role=partner → partner
      3. @fehradvice.com + role=senior_management → senior_consultant
      4. @fehradvice.com + email_verified → consultant
      5. email_verified → researcher
      6. else → guest
    """
    email = (user_payload.get("sub") or "").lower()
    is_admin = user_payload.get("admin", False)
    stored_role = (user_payload.get("role") or "researcher").lower()
    is_fa = email.endswith("@fehradvice.com")

    # Owner: admins from ADMIN_EMAILS
    if is_admin or email in [e.lower() for e in ADMIN_EMAILS]:
        return "owner"

    # Partner: explicitly set
    if stored_role == "partner" and is_fa:
        return "partner"

    # Senior Consultant
    if stored_role in ("senior_management", "senior_consultant") and is_fa:
        return "senior_consultant"

    # Consultant: any FehrAdvice employee
    if is_fa:
        return "consultant"

    # Researcher: verified external
    return "researcher"

def has_permission(user_payload: dict, permission: str, resource_role: str = None) -> bool:
    """Check if user has a specific permission.
    
    Args:
        user_payload: JWT payload dict
        permission: Permission string (e.g. 'project.create')
        resource_role: Optional resource-level role (e.g. 'project.owner')
    
    Returns:
        True if permitted, False otherwise
    """
    if permission not in PERMISSION_MAP:
        return False  # Unknown permission = denied

    email = (user_payload.get("sub") or "").lower()
    user_role = resolve_user_role(user_payload)
    user_level = ROLE_HIERARCHY.get(user_role, 0)
    required_role = PERMISSION_MAP[permission]
    required_level = ROLE_HIERARCHY.get(required_role, 999)

    # Check @fehradvice.com requirement
    if permission in FA_ONLY_PERMISSIONS and not email.endswith("@fehradvice.com"):
        return False

    # Check role hierarchy
    if user_level >= required_level:
        return True

    # Check resource-level override (e.g. project.owner can always update their project)
    if resource_role:
        domain = permission.split(".")[0]
        # Resource owners get elevated permissions on their resources
        if resource_role == f"{domain}.owner":
            return True
        # Contributors get update_own permissions
        if resource_role == f"{domain}.contributor" and permission.endswith("_own"):
            return True

    return False

def audit_log(user_payload: dict, permission: str, resource_type: str = None,
              resource_id: str = None, action_detail: str = None, success: bool = True):
    """Log permission-critical actions for audit trail."""
    if permission not in AUDIT_PERMISSIONS:
        return
    try:
        email = user_payload.get("sub", "unknown")
        role = resolve_user_role(user_payload)
        log_entry = {
            "timestamp": datetime.utcnow().isoformat(),
            "user": email,
            "role": role,
            "permission": permission,
            "resource_type": resource_type,
            "resource_id": resource_id,
            "detail": action_detail,
            "success": success
        }
        logger.info(f"AUDIT: {json.dumps(log_entry)}")
        # TODO: Persist to audit_log table in PostgreSQL
    except Exception as e:
        logger.error(f"Audit log failed: {e}")

def require_permission(permission: str):
    """FastAPI dependency factory that checks a specific permission.
    
    Usage:
        @app.post("/api/projects")
        async def create_project(user=Depends(require_permission("project.create"))):
            ...
    """
    async def _check(request: Request):
        # First authenticate
        auth = request.headers.get("Authorization", "")
        token = auth.replace("Bearer ", "") if auth.startswith("Bearer ") else ""
        if not token:
            token = request.cookies.get("bea_token", "")
        payload = verify_jwt(token)
        if not payload:
            raise HTTPException(401, "Nicht autorisiert")

        if not has_permission(payload, permission):
            user_role = resolve_user_role(payload)
            required_role = PERMISSION_MAP.get(permission, "?")
            email = payload.get("sub", "?")
            is_fa_issue = permission in FA_ONLY_PERMISSIONS and not email.endswith("@fehradvice.com")

            if is_fa_issue:
                msg = f"Keine Berechtigung: '{permission}' ist nur für FehrAdvice-Mitarbeiter verfügbar"
            else:
                msg = f"Keine Berechtigung: '{permission}' erfordert mindestens Rolle '{required_role}' (aktuelle Rolle: '{user_role}')"

            audit_log(payload, permission, success=False, action_detail=msg)
            raise HTTPException(403, msg)

        return payload
    return _check

def get_user_permissions(user_payload: dict) -> dict:
    """Return all permissions for a user — used for frontend UI gating."""
    user_role = resolve_user_role(user_payload)
    user_level = ROLE_HIERARCHY.get(user_role, 0)
    email = (user_payload.get("sub") or "").lower()
    is_fa = email.endswith("@fehradvice.com")

    permissions = {}
    for perm, required_role in PERMISSION_MAP.items():
        required_level = ROLE_HIERARCHY.get(required_role, 999)
        if perm in FA_ONLY_PERMISSIONS and not is_fa:
            permissions[perm] = False
        else:
            permissions[perm] = user_level >= required_level
    return {
        "role": user_role,
        "role_level": user_level,
        "email": email,
        "is_fehradvice": is_fa,
        "permissions": permissions
    }

def _user_crm_payload(user):
    """Compute CRM access for JWT. Rules:
    - Only @fehradvice.com emails eligible
    - Senior Management / Partner / Admin: CRM auto-enabled
    - lead_management: see own leads (viewer + owner_code)
    - Others need explicit crm_access=True in DB
    """
    email = (user.email or "").lower()
    is_fa = email.endswith("@fehradvice.com")
    role = (user.role or "researcher").lower()
    senior_roles = ("senior_management", "partner")
    auto_crm = role in senior_roles or user.is_admin
    effective_crm = user.crm_access or (auto_crm and is_fa)
    effective_lead_mgmt = getattr(user, 'lead_management', False) or auto_crm
    return {
        "crm_access": effective_crm and is_fa,
        "crm_role": user.crm_role or ("admin" if user.is_admin else ("manager" if role in senior_roles else "none")),
        "crm_owner_code": user.crm_owner_code or "",
        "lead_management": bool(effective_lead_mgmt and effective_crm and is_fa),
    }

Base = declarative_base()
_engine = None
_SessionLocal = None

class User(Base):
    __tablename__ = "users"
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    email = Column(String(320), unique=True, nullable=False, index=True)
    name = Column(String(200), nullable=True)
    password_hash = Column(String(500), nullable=False)
    password_salt = Column(String(100), nullable=False)
    is_active = Column(Boolean, default=True)
    is_admin = Column(Boolean, default=False)
    email_verified = Column(Boolean, default=False)
    verification_token = Column(String(200), nullable=True)
    verification_sent_at = Column(DateTime, nullable=True)
    reset_token = Column(String(200), nullable=True)
    reset_sent_at = Column(DateTime, nullable=True)
    # Profile fields
    position = Column(String(200), nullable=True)
    company = Column(String(200), nullable=True)
    bio = Column(Text, nullable=True)
    phone = Column(String(50), nullable=True)
    linkedin_url = Column(String(500), nullable=True)
    linkedin_id = Column(String(100), nullable=True)
    profile_photo_url = Column(String(1000), nullable=True)
    expertise = Column(JSON, nullable=True)  # ["Behavioral Economics", "Strategy", ...]
    role = Column(String(50), default="researcher")  # researcher, sales, operations
    crm_access = Column(Boolean, default=False)  # explicitly enabled for CRM
    crm_role = Column(String(30), default="none")  # none, viewer, manager, admin
    crm_owner_code = Column(String(20), nullable=True)  # OWN-GF, OWN-EB, etc.
    lead_management = Column(Boolean, default=False)  # access to lead management suite
    created_at = Column(DateTime, default=datetime.utcnow)
    last_login = Column(DateTime, nullable=True)

class UserInsight(Base):
    __tablename__ = "user_insights"
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_email = Column(String(320), nullable=False, index=True)
    # Question & Answer
    question_text = Column(Text, nullable=False)
    question_category = Column(String(50), nullable=True)  # onboarding, login, voluntary
    answer_text = Column(Text, nullable=True)
    choice_index = Column(Integer, nullable=True)  # 1,2,3,4(custom),5(beatrix)
    # Behavioral Metadata
    latency_ms = Column(Integer, nullable=True)  # Time from question shown to answer
    session_number = Column(Integer, default=1)  # nth login/session
    question_number = Column(Integer, default=1)  # nth question in this session (1=mandatory, 2-3=nudged, 4+=voluntary)
    was_mandatory = Column(Boolean, default=True)
    was_nudged = Column(Boolean, default=False)
    skipped = Column(Boolean, default=False)
    # Extracted Insights
    domain_signal = Column(String(20), nullable=True)  # REL/FIN/HLT/ENV/POL/ORG/EDU/OTH
    thinking_style = Column(String(30), nullable=True)  # theoretical/practical/analytical/creative
    abstraction_level = Column(String(20), nullable=True)  # abstract/concrete/mixed
    autonomy_signal = Column(String(20), nullable=True)  # high(option4)/low(option5)/moderate(1-3)
    # Timestamps
    created_at = Column(DateTime, default=datetime.utcnow)
    context = Column(String(20), default="login")  # onboarding, login, model_building

class Document(Base):
    __tablename__ = "documents"
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    title = Column(String(500), nullable=False)
    content = Column(Text, nullable=True)
    source_type = Column(String(20), nullable=False)
    file_type = Column(String(10), nullable=True)
    file_path = Column(String(1000), nullable=True)
    file_size = Column(Integer, nullable=True)
    database_target = Column(String(50), nullable=False, default="knowledge_base")
    category = Column(String(50), nullable=True)
    language = Column(String(10), nullable=True)
    tags = Column(JSON, nullable=True)
    doc_metadata = Column("metadata", JSON, nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    status = Column(String(20), default="indexed")
    github_url = Column(String(1000), nullable=True)
    uploaded_by = Column(String(320), nullable=True)
    content_hash = Column(String(64), nullable=True, index=True)

def get_db():
    global _engine, _SessionLocal
    if _engine is None:
        _engine = create_engine(DATABASE_URL, pool_pre_ping=True)
        _SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=_engine)
        try:
            Base.metadata.create_all(bind=_engine)
            try:
                from sqlalchemy import text
                with _engine.connect() as conn:
                    conn.execute(text("ALTER TABLE documents ADD COLUMN IF NOT EXISTS github_url VARCHAR(1000)"))
                    conn.execute(text("ALTER TABLE documents ADD COLUMN IF NOT EXISTS uploaded_by VARCHAR(320)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS email_verified BOOLEAN DEFAULT FALSE"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS verification_token VARCHAR(200)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS verification_sent_at TIMESTAMP"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS reset_token VARCHAR(200)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS reset_sent_at TIMESTAMP"))
                    conn.execute(text("ALTER TABLE documents ADD COLUMN IF NOT EXISTS content_hash VARCHAR(64)"))
                    # Profile fields
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS position VARCHAR(200)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS company VARCHAR(200)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS bio TEXT"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS phone VARCHAR(50)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS linkedin_url VARCHAR(500)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS linkedin_id VARCHAR(100)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS profile_photo_url VARCHAR(1000)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS expertise JSON"))
                    # Chat messages table
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS chat_messages (
                        id VARCHAR PRIMARY KEY, user_email VARCHAR(320) NOT NULL,
                        role VARCHAR(20) NOT NULL, content TEXT NOT NULL,
                        sources JSON, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    conn.execute(text("ALTER TABLE chat_messages ADD COLUMN IF NOT EXISTS session_id VARCHAR(50)"))
                    try:
                        conn.execute(text("CREATE INDEX IF NOT EXISTS idx_chat_session ON chat_messages (session_id)"))
                    except Exception:
                        pass
                    # Chat sessions metadata table
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS chat_sessions (
                        id VARCHAR(50) PRIMARY KEY, user_email VARCHAR(320) NOT NULL,
                        session_type VARCHAR(30) DEFAULT 'general',
                        entities JSON DEFAULT '{}',
                        context JSON DEFAULT '{}',
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    # Session history columns
                    for col_def in [
                        "ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS title VARCHAR(120)",
                        "ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS summary TEXT",
                        "ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS msg_count INTEGER DEFAULT 0",
                        "ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS duration_min FLOAT DEFAULT 0",
                        "ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS closed_at TIMESTAMP",
                    ]:
                        try:
                            conn.execute(text(col_def))
                        except Exception:
                            pass
                    conn.commit()
                    # User insights / behavioral profiling table
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS user_insights (
                        id VARCHAR PRIMARY KEY, user_email VARCHAR(320) NOT NULL,
                        question_text TEXT NOT NULL, question_category VARCHAR(50),
                        answer_text TEXT, choice_index INTEGER,
                        latency_ms INTEGER, session_number INTEGER DEFAULT 1,
                        question_number INTEGER DEFAULT 1, was_mandatory BOOLEAN DEFAULT TRUE,
                        was_nudged BOOLEAN DEFAULT FALSE, skipped BOOLEAN DEFAULT FALSE,
                        domain_signal VARCHAR(20), thinking_style VARCHAR(30),
                        abstraction_level VARCHAR(20), autonomy_signal VARCHAR(20),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        context VARCHAR(20) DEFAULT 'login')"""))
                    # pgvector for semantic search (AlphaGo-style)
                    try:
                        conn.execute(text("CREATE EXTENSION IF NOT EXISTS vector"))
                        conn.execute(text("ALTER TABLE documents ADD COLUMN IF NOT EXISTS embedding vector(512)"))
                        conn.commit()
                        logger.info("pgvector enabled with embedding column")
                    except Exception as ve:
                        logger.warning(f"pgvector not available: {ve}")
                        conn.rollback()
                        # Fallback: use TEXT column to store embeddings as JSON
                        try:
                            conn.execute(text("ALTER TABLE documents ADD COLUMN IF NOT EXISTS embedding_json TEXT"))
                            conn.commit()
                            logger.info("Using JSON embedding fallback (no pgvector)")
                        except:
                            conn.rollback()
                    # Auto-verify existing admin users
                    conn.execute(text("UPDATE users SET email_verified = TRUE WHERE is_admin = TRUE AND email_verified = FALSE"))
                    # Role field for tab visibility
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS role VARCHAR(50) DEFAULT 'researcher'"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS crm_access BOOLEAN DEFAULT FALSE"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS crm_role VARCHAR(30) DEFAULT 'none'"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS crm_owner_code VARCHAR(20)"))
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS lead_management BOOLEAN DEFAULT FALSE"))
                    # Knowledge calibration table (fact-check results per user)
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS knowledge_checks (
                        id VARCHAR PRIMARY KEY, user_email VARCHAR(320) NOT NULL,
                        chat_doc_id VARCHAR REFERENCES documents(id),
                        claim_text TEXT NOT NULL,
                        claim_topic VARCHAR(100),
                        customer_code VARCHAR(50),
                        verification_status VARCHAR(20) NOT NULL DEFAULT 'unverified',
                        confidence_score FLOAT DEFAULT 0,
                        evidence TEXT,
                        evidence_source VARCHAR(500),
                        is_quantitative BOOLEAN DEFAULT FALSE,
                        user_certainty VARCHAR(20) DEFAULT 'stated',
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    # Belief & Bias Analysis table
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS belief_analyses (
                        id VARCHAR PRIMARY KEY,
                        user_email VARCHAR(320) NOT NULL,
                        chat_doc_id VARCHAR,
                        belief_text TEXT NOT NULL,
                        belief_category VARCHAR(50),
                        bias_type VARCHAR(50),
                        bias_source VARCHAR(20) DEFAULT 'user',
                        informed_score FLOAT DEFAULT 0.5,
                        reasoning_type VARCHAR(30) DEFAULT 'mixed',
                        evidence_basis VARCHAR(30) DEFAULT 'none',
                        customer_code VARCHAR(50),
                        context TEXT,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    # Auto-enable CRM for FehrAdvice admins (Senior Management) – runs every startup
                    conn.execute(text("""UPDATE users SET crm_access = TRUE, crm_role = 'admin' 
                        WHERE is_admin = TRUE AND email LIKE '%%@fehradvice.com'"""))
                    # Also enable for known senior management
                    conn.execute(text("""UPDATE users SET crm_access = TRUE, crm_role = 'admin', crm_owner_code = 'OWN-GF'
                        WHERE email = 'gerhard.fehr@fehradvice.com'"""))
                    # Ensure known admins have is_admin flag
                    conn.execute(text("""UPDATE users SET is_admin = TRUE WHERE email IN ('gerhard.fehr@fehradvice.com', 'nora.gavazajsusuri@fehradvice.com') AND is_admin = FALSE"""))
                    # Set initial roles for sales users
                    conn.execute(text("UPDATE users SET role = 'sales' WHERE email IN ('nora.gavazajsusuri@fehradvice.com', 'maria.neumann@fehradvice.com') AND (role IS NULL OR role = 'researcher')"))
                    # Leads table
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS leads (
                        id VARCHAR PRIMARY KEY, company VARCHAR(500),
                        contact VARCHAR(200), email VARCHAR(320),
                        stage VARCHAR(50) DEFAULT 'kontakt',
                        value REAL DEFAULT 0, source VARCHAR(200),
                        notes TEXT, created_by VARCHAR(320),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    # ── CRM Tables ──
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS crm_companies (
                        id VARCHAR PRIMARY KEY,
                        name VARCHAR(500) NOT NULL,
                        domain VARCHAR(200),
                        industry VARCHAR(200),
                        size VARCHAR(50),
                        website VARCHAR(500),
                        address TEXT,
                        notes TEXT,
                        created_by VARCHAR(320),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS crm_contacts (
                        id VARCHAR PRIMARY KEY,
                        company_id VARCHAR REFERENCES crm_companies(id),
                        name VARCHAR(300) NOT NULL,
                        email VARCHAR(320),
                        phone VARCHAR(50),
                        position VARCHAR(200),
                        role_type VARCHAR(50) DEFAULT 'kontakt',
                        psi_profile JSON,
                        notes TEXT,
                        created_by VARCHAR(320),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS crm_deals (
                        id VARCHAR PRIMARY KEY,
                        company_id VARCHAR REFERENCES crm_companies(id),
                        contact_id VARCHAR REFERENCES crm_contacts(id),
                        title VARCHAR(500) NOT NULL,
                        stage VARCHAR(50) DEFAULT 'erstkontakt',
                        value REAL DEFAULT 0,
                        probability INTEGER DEFAULT 10,
                        source VARCHAR(200),
                        next_action VARCHAR(500),
                        next_action_date DATE,
                        owner VARCHAR(320),
                        context_id VARCHAR,
                        notes TEXT,
                        closed_at TIMESTAMP,
                        lost_reason TEXT,
                        created_by VARCHAR(320),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS crm_activities (
                        id VARCHAR PRIMARY KEY,
                        deal_id VARCHAR REFERENCES crm_deals(id),
                        company_id VARCHAR REFERENCES crm_companies(id),
                        contact_id VARCHAR REFERENCES crm_contacts(id),
                        type VARCHAR(30) NOT NULL,
                        subject VARCHAR(500),
                        description TEXT,
                        due_date TIMESTAMP,
                        done BOOLEAN DEFAULT FALSE,
                        created_by VARCHAR(320),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    # Contexts / Ausgangslage table
                    conn.execute(text("ALTER TABLE crm_deals ADD COLUMN IF NOT EXISTS lead_code VARCHAR(30)"))
                    # Zefix enrichment columns
                    conn.execute(text("ALTER TABLE crm_companies ADD COLUMN IF NOT EXISTS uid VARCHAR(20)"))
                    conn.execute(text("ALTER TABLE crm_companies ADD COLUMN IF NOT EXISTS legal_form VARCHAR(100)"))
                    conn.execute(text("ALTER TABLE crm_companies ADD COLUMN IF NOT EXISTS zefix_id VARCHAR(20)"))
                    conn.execute(text("ALTER TABLE crm_companies ADD COLUMN IF NOT EXISTS postal_code VARCHAR(10)"))
                    conn.execute(text("ALTER TABLE crm_companies ADD COLUMN IF NOT EXISTS locality VARCHAR(200)"))
                    conn.execute(text("ALTER TABLE crm_companies ADD COLUMN IF NOT EXISTS canton VARCHAR(50)"))
                    conn.execute(text("ALTER TABLE crm_companies ADD COLUMN IF NOT EXISTS zefix_description TEXT"))
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS contexts (
                        id VARCHAR PRIMARY KEY, client VARCHAR(500),
                        project VARCHAR(500), domain VARCHAR(20),
                        status VARCHAR(50) DEFAULT 'aktiv',
                        situation TEXT, goal TEXT, constraints TEXT,
                        created_by VARCHAR(320),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    # Ψ-Analyses: versioned context analyses per user
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS psi_analyses (
                        id VARCHAR(50) PRIMARY KEY,
                        lineage_id VARCHAR(30) NOT NULL,
                        version INTEGER NOT NULL DEFAULT 1,
                        question TEXT NOT NULL,
                        mode VARCHAR(20) DEFAULT 'schnell',
                        macro_context JSONB,
                        meso_context JSONB,
                        micro_context JSONB,
                        psi_profile JSONB,
                        parameters JSONB,
                        synthesis TEXT,
                        implications JSONB,
                        confidence VARCHAR(200),
                        customer_code VARCHAR(20),
                        project_slug VARCHAR(200),
                        parent_version_id VARCHAR(20),
                        created_by VARCHAR(320) NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    conn.execute(text("CREATE INDEX IF NOT EXISTS idx_psi_lineage ON psi_analyses(lineage_id)"))
                    conn.execute(text("CREATE INDEX IF NOT EXISTS idx_psi_user ON psi_analyses(created_by)"))
                    # Feedback table
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS feedback (
                        id VARCHAR PRIMARY KEY,
                        type VARCHAR(20) DEFAULT 'bug',
                        comment TEXT,
                        screenshot TEXT,
                        page VARCHAR(100),
                        url VARCHAR(1000),
                        viewport VARCHAR(50),
                        user_agent VARCHAR(500),
                        status VARCHAR(20) DEFAULT 'neu',
                        created_by VARCHAR(320),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS project_edits (
                        id SERIAL PRIMARY KEY,
                        slug VARCHAR(200) NOT NULL,
                        section VARCHAR(100) NOT NULL,
                        data JSONB NOT NULL,
                        edited_by VARCHAR(320),
                        synced_to_github BOOLEAN DEFAULT FALSE,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        UNIQUE(slug, section))"""))
                    # ── Task Management ──
                    conn.execute(text("""CREATE TABLE IF NOT EXISTS tasks (
                        id VARCHAR(50) PRIMARY KEY,
                        title VARCHAR(500) NOT NULL,
                        description TEXT,
                        assignee_type VARCHAR(20) NOT NULL DEFAULT 'consultant',
                        assignee VARCHAR(200),
                        customer_code VARCHAR(50),
                        project_slug VARCHAR(200),
                        lead_id VARCHAR(50),
                        deal_id VARCHAR(50),
                        priority VARCHAR(20) DEFAULT 'normal',
                        status VARCHAR(30) DEFAULT 'open',
                        due_date DATE,
                        due_time TIME,
                        category VARCHAR(50),
                        source VARCHAR(30) DEFAULT 'manual',
                        trigger_event VARCHAR(100),
                        action_type VARCHAR(100),
                        action_config JSONB,
                        escalation_after_days INTEGER,
                        waiting_since TIMESTAMP,
                        completed_at TIMESTAMP,
                        completed_by VARCHAR(320),
                        parent_task_id VARCHAR(50),
                        sort_order INTEGER DEFAULT 0,
                        created_by VARCHAR(320),
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)"""))
                    conn.execute(text("CREATE INDEX IF NOT EXISTS idx_tasks_assignee ON tasks(assignee_type, assignee)"))
                    conn.execute(text("CREATE INDEX IF NOT EXISTS idx_tasks_status ON tasks(status)"))
                    conn.execute(text("CREATE INDEX IF NOT EXISTS idx_tasks_due ON tasks(due_date)"))
                    conn.execute(text("CREATE INDEX IF NOT EXISTS idx_tasks_customer ON tasks(customer_code)"))
                    conn.execute(text("CREATE INDEX IF NOT EXISTS idx_tasks_project ON tasks(project_slug)"))
                    conn.execute(text("CREATE INDEX IF NOT EXISTS idx_tasks_lead ON tasks(lead_id)"))
                    conn.commit()
            except: pass
            logger.info(f"DB connected: {DATABASE_URL[:50]}...")
        except Exception as e:
            logger.error(f"DB init error: {e}")
            _engine = None; _SessionLocal = None
            raise HTTPException(503, "Datenbank nicht bereit")
    return _SessionLocal()

def push_to_github(filename, content_bytes, folder=None):
    """Push uploaded file to GitHub in the context repo.
    folder: custom path like 'books/' or 'studies/' – defaults to papers/evaluated/integrated/"""
    if not GH_TOKEN:
        logger.warning("GH_TOKEN not set, skipping GitHub push")
        return {"error": "GH_TOKEN not configured"}
    import urllib.request, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    base_path = folder if folder else GH_UPLOAD_PATH
    path = f"{base_path}/{filename}"
    url = f"https://api.github.com/repos/{GH_REPO}/contents/{path}"
    headers = {"Authorization": f"token {GH_TOKEN}", "Content-Type": "application/json",
               "Accept": "application/vnd.github.v3+json", "User-Agent": "BEATRIXLab"}
    sha = None
    try:
        req = urllib.request.Request(url, headers=headers)
        resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
        sha = resp.get("sha")
    except: pass
    payload = {"message": f"Upload via BEA Lab: {filename}",
               "content": base64.b64encode(content_bytes).decode(), "branch": "main"}
    if sha: payload["sha"] = sha
    try:
        req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), method="PUT", headers=headers)
        resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=30).read())
        logger.info(f"GitHub push OK: {path}")
        return {"url": resp.get("content", {}).get("html_url", ""),
                "sha": resp.get("content", {}).get("sha", ""), "path": path}
    except Exception as e:
        logger.error(f"GitHub push failed: {e}")
        return {"error": str(e)}

# ========== VECTOR EMBEDDING (AlphaGo Neural Network) ==========

def embed_texts(texts: list, input_type: str = "document") -> list:
    """Generate embeddings via Voyage AI. Returns list of 512-dim vectors."""
    if not VOYAGE_API_KEY or not texts:
        return []
    import urllib.request, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    # Truncate long texts (Voyage limit ~32k tokens)
    truncated = [t[:16000] for t in texts]
    payload = json.dumps({
        "input": truncated,
        "model": VOYAGE_MODEL,
        "input_type": input_type  # "document" for storage, "query" for search
    }).encode()
    req = urllib.request.Request("https://api.voyageai.com/v1/embeddings", data=payload, method="POST",
        headers={"Authorization": f"Bearer {VOYAGE_API_KEY}", "Content-Type": "application/json"})
    try:
        resp = json.loads(urllib.request.urlopen(req, context=ctx).read())
        return [d["embedding"] for d in resp.get("data", [])]
    except Exception as e:
        logger.error(f"Voyage AI embedding error: {e}")
        return []

def embed_single(text: str, input_type: str = "document") -> list:
    """Embed a single text. Returns 512-dim vector or empty list."""
    results = embed_texts([text], input_type)
    return results[0] if results else []

def embed_document(db, doc_id: str):
    """Generate and store embedding for a single document."""
    if not VOYAGE_API_KEY: return False
    from sqlalchemy import text as sql_text
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc or not doc.content: return False
    embed_input = f"{doc.title or ''}\n{doc.content[:15000]}"
    vec = embed_single(embed_input, "document")
    if not vec: return False
    try:
        # Try pgvector first
        vec_str = f"[{','.join(str(v) for v in vec)}]"
        db.execute(text("UPDATE documents SET embedding = :vec WHERE id = :id"), {"vec": vec_str, "id": doc_id})
        db.commit()
        logger.info(f"Embedded doc (pgvector): {doc.title[:50]}")
        return True
    except Exception:
        db.rollback()
        try:
            # Fallback: store as JSON text
            vec_json = json.dumps(vec)
            db.execute(text("UPDATE documents SET embedding_json = :vec WHERE id = :id"), {"vec": vec_json, "id": doc_id})
            db.commit()
            logger.info(f"Embedded doc (JSON): {doc.title[:50]}")
            return True
        except Exception as e:
            logger.error(f"Store embedding error: {e}")
            db.rollback()
            return False

def embed_all_documents(db):
    """Embed all documents that don't have embeddings yet."""
    if not VOYAGE_API_KEY:
        logger.info("No VOYAGE_API_KEY, skipping embeddings")
        return 0
    from sqlalchemy import text as sql_text
    # Find docs without embeddings - try pgvector column first, then JSON
    try:
        result = db.execute(text("SELECT id, title, content FROM documents WHERE embedding IS NULL AND content IS NOT NULL"))
        rows = result.fetchall()
        use_pgvector = True
    except Exception:
        db.rollback()
        try:
            result = db.execute(text("SELECT id, title, content FROM documents WHERE (embedding_json IS NULL OR embedding_json = '') AND content IS NOT NULL"))
            rows = result.fetchall()
            use_pgvector = False
        except Exception as e:
            db.rollback()
            logger.warning(f"Cannot query embeddings: {e}")
            return 0
    if not rows:
        logger.info("All documents already embedded")
        return 0
    logger.info(f"Embedding {len(rows)} documents (pgvector={use_pgvector})...")
    count = 0
    batch_size = 8
    for i in range(0, len(rows), batch_size):
        batch = rows[i:i+batch_size]
        texts = [f"{r[1] or ''}\n{(r[2] or '')[:15000]}" for r in batch]
        vectors = embed_texts(texts, "document")
        for j, vec in enumerate(vectors):
            if vec:
                try:
                    if use_pgvector:
                        vec_str = f"[{','.join(str(v) for v in vec)}]"
                        db.execute(text("UPDATE documents SET embedding = :vec WHERE id = :id"), {"vec": vec_str, "id": batch[j][0]})
                    else:
                        db.execute(text("UPDATE documents SET embedding_json = :vec WHERE id = :id"), {"vec": json.dumps(vec), "id": batch[j][0]})
                    count += 1
                except: pass
        db.commit()
    logger.info(f"Embedded {count}/{len(rows)} documents")
    return count

# ── AUTO-SAVE CHAT TO KNOWLEDGE BASE ──

def gh_put_file(repo: str, path: str, content: str, message: str) -> dict:
    """Generic GitHub file PUT – creates or updates a file in any repo."""
    if not GH_TOKEN:
        return {"error": "no token"}
    import urllib.request, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    url = f"https://api.github.com/repos/{repo}/contents/{path}"
    headers = {"Authorization": f"token {GH_TOKEN}", "Content-Type": "application/json",
               "Accept": "application/vnd.github.v3+json", "User-Agent": "BEATRIXLab"}
    # Get existing SHA if file exists
    sha = None
    try:
        req = urllib.request.Request(url, headers=headers)
        resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
        sha = resp.get("sha")
    except: pass
    payload = {"message": message, "content": base64.b64encode(content.encode('utf-8')).decode(), "branch": "main"}
    if sha: payload["sha"] = sha
    try:
        req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"), method="PUT", headers=headers)
        resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=30).read())
        return {"ok": True, "url": resp.get("content", {}).get("html_url", ""), "sha": resp.get("content", {}).get("sha", "")}
    except Exception as e:
        logger.warning(f"GitHub PUT failed {path}: {e}")
        return {"error": str(e)}


def auto_save_chat_to_kb(question: str, answer: str, user_email: str, intent: str = "chat",
                          metadata: dict = None):
    """Save a chat Q&A pair as a searchable document in the Knowledge Base AND push to GitHub.
    Every conversation becomes organizational knowledge – in PostgreSQL for fast search,
    and on GitHub for permanence and team access."""
    import hashlib, yaml

    # Skip trivially short or empty answers
    if not answer or len(answer.strip()) < 50:
        return None
    if not question or len(question.strip()) < 5:
        return None

    # Build content: Question + Answer combined for best searchability
    content = f"FRAGE: {question.strip()}\n\nANTWORT: {answer.strip()}"

    # Deduplicate: hash the content
    content_hash = hashlib.sha256(content.encode()).hexdigest()[:16]

    # Build title from question (first 100 chars)
    title = f"Chat: {question.strip()[:100]}"

    # Tags for filtering
    user_short = user_email.split('@')[0] if user_email else 'unknown'
    tags = ["chat-insight", f"intent:{intent}", f"user:{user_short}"]
    customer_code = ""
    project_slug = ""
    if metadata:
        if metadata.get("customer_code"):
            customer_code = metadata["customer_code"]
            tags.append(f"customer:{customer_code}")
        if metadata.get("project_slug"):
            project_slug = metadata["project_slug"]
            tags.append(f"project:{project_slug}")

    now = datetime.utcnow()
    doc_meta = {
        "source": "chat",
        "intent": intent,
        "user_email": user_email,
        "question": question.strip()[:500],
        "timestamp": now.isoformat(),
    }
    if metadata:
        doc_meta.update(metadata)

    db = get_db()
    try:
        # Check for duplicate (same hash already exists)
        existing = db.query(Document).filter(Document.content_hash == content_hash).first()
        if existing:
            logger.debug(f"Chat insight already in KB: {content_hash}")
            return existing.id

        # ── 1) Save to PostgreSQL (fast search) ──
        doc = Document(
            title=title,
            content=content,
            source_type="chat_insight",
            file_type="chat",
            database_target="knowledge_base",
            category=intent,
            language="de",
            tags=tags,
            doc_metadata=doc_meta,
            uploaded_by=user_email,
            content_hash=content_hash,
            status="indexed"
        )
        db.add(doc)
        db.commit()
        doc_id = doc.id
        logger.info(f"💬→📚 Chat saved to KB: '{title[:60]}' (intent={intent}, user={user_short})")

        # Embed for vector search (non-blocking)
        try:
            embed_document(db, doc_id)
        except Exception as e:
            logger.warning(f"Embedding deferred: {e}")

        # ── 2) Push to GitHub (permanent knowledge) ──
        try:
            year_month = now.strftime("%Y-%m")
            # Build YAML for GitHub
            gh_data = {
                "id": doc_id,
                "title": title,
                "intent": intent,
                "user": user_email,
                "timestamp": now.isoformat(),
                "customer_code": customer_code or None,
                "project_slug": project_slug or None,
                "tags": tags,
                "question": question.strip()[:2000],
                "answer": answer.strip()[:5000],
            }
            yaml_content = yaml.dump(gh_data, allow_unicode=True, default_flow_style=False, sort_keys=False)
            gh_path = f"data/knowledge/chat-insights/{year_month}/{content_hash}.yaml"
            gh_result = gh_put_file(
                GH_CONTEXT_REPO, gh_path, yaml_content,
                f"💬 Chat insight: {intent} – {question.strip()[:60]}"
            )
            if gh_result.get("ok"):
                # Update DB with GitHub URL
                try:
                    doc.github_url = gh_result.get("url", "")
                    doc.status = "indexed+github"
                    db.commit()
                    logger.info(f"💬→🐙 Chat insight pushed to GitHub: {gh_path}")
                except: pass
            else:
                logger.warning(f"GitHub push skipped: {gh_result.get('error','unknown')}")

            # ── 3) Also push to CUSTOMER folder if customer_code is known ──
            if customer_code:
                try:
                    # Slugify question for readable filename
                    import re as _re
                    q_slug = _re.sub(r'[^a-z0-9]+', '-', question.strip()[:60].lower()).strip('-')[:40]
                    customer_path = f"data/customers/{customer_code}/insights/{year_month}_{q_slug}.yaml"
                    gh_cust = gh_put_file(
                        GH_CONTEXT_REPO, customer_path, yaml_content,
                        f"💬→👤 {customer_code.upper()} insight: {question.strip()[:50]}"
                    )
                    if gh_cust.get("ok"):
                        logger.info(f"💬→👤 Insight assigned to customer: {customer_path}")
                    else:
                        logger.warning(f"Customer push skipped: {gh_cust.get('error','')[:80]}")
                except Exception as e:
                    logger.warning(f"Customer folder push failed: {e}")

            # ── 4) Also push to PROJECT folder if project_slug is known ──
            if project_slug:
                try:
                    prj_q_slug = _re.sub(r'[^a-z0-9]+', '-', question.strip()[:60].lower()).strip('-')[:40]
                    project_path = f"data/projects/{project_slug}/insights/{year_month}_{prj_q_slug}.yaml"
                    gh_prj = gh_put_file(
                        GH_CONTEXT_REPO, project_path, yaml_content,
                        f"💬→📋 {project_slug} insight: {question.strip()[:50]}"
                    )
                    if gh_prj.get("ok"):
                        logger.info(f"💬→📋 Insight assigned to project: {project_path}")
                except Exception as e:
                    logger.warning(f"Project folder push failed: {e}")
        except Exception as e:
            logger.warning(f"GitHub push for chat insight failed: {e}")

        return doc_id
    except Exception as e:
        logger.error(f"Auto-save chat to KB failed: {e}")
        db.rollback()
        return None
    finally:
        db.close()

# ── FACT-CHECK ENGINE: Knowledge Calibration ──

FACT_CHECK_SYSTEM = """Du bist ein präziser Fact-Check-Analyst. Analysiere die USER-NACHRICHT und extrahiere alle überprüfbaren Behauptungen (Claims).

REGELN:
1. Extrahiere NUR konkrete, überprüfbare Aussagen – keine Meinungen oder Fragen
2. Kategorisiere jeden Claim:
   - Quantitative Claims: Zahlen, Prozente, Beträge, Daten (z.B. "72% Zufriedenheit", "Budget 50k")
   - Faktische Claims: Strukturen, Zusammenhänge, Zustände (z.B. "UBS hat 3 Divisionen", "Kröll ist Head of CX")
   - Kausale Claims: Ursache-Wirkungs-Behauptungen (z.B. "Standardisierung führt zu Abwanderung")
3. Bewerte die User-Sicherheit:
   - "certain": User präsentiert als Fakt ("es sind 72%")
   - "hedged": User ist unsicher ("ich glaube es sind etwa 72%")
   - "questioning": User fragt ("sind es 72%?")

Antworte NUR mit JSON (keine Backticks):
{
  "claims": [
    {
      "text": "Die Kundenzufriedenheit liegt bei 72%",
      "topic": "kundenzufriedenheit",
      "type": "quantitative",
      "customer": "zkb",
      "user_certainty": "certain",
      "verifiable": true
    }
  ]
}

Wenn keine überprüfbaren Claims in der Nachricht sind, antworte: {"claims": []}
Maximal 5 Claims pro Nachricht. Fokus auf die wichtigsten."""

def fact_check_user_claims(user_message: str, user_email: str, chat_doc_id: str = None,
                            customer_code: str = None):
    """Extract claims from user message, cross-reference with KB, store calibration data.
    Runs as background enrichment – does not block chat response."""
    import urllib.request, ssl, re

    if not ANTHROPIC_API_KEY:
        return []
    if not user_message or len(user_message.strip()) < 20:
        return []

    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    # ── Step 1: Extract claims using Claude ──
    try:
        payload = json.dumps({
            "model": ANTHROPIC_MODEL_STANDARD,
            "max_tokens": 1500,
            "system": FACT_CHECK_SYSTEM,
            "messages": [{"role": "user", "content": f"USER-NACHRICHT:\n{user_message[:3000]}"}]
        }).encode()
        req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
            headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
                "User-Agent": "BEATRIXLab/3.20-factcheck"
            })
        resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=30).read())
        answer = resp["content"][0]["text"]

        # Parse claims
        try:
            # Try raw JSON first
            parsed = json.loads(answer.strip())
        except:
            json_match = re.search(r'\{.*\}', answer, re.DOTALL)
            if json_match:
                parsed = json.loads(json_match.group())
            else:
                return []

        claims = parsed.get("claims", [])
        if not claims:
            return []

        logger.info(f"🔍 Fact-check: {len(claims)} claims extracted from {user_email.split('@')[0]}")

    except Exception as e:
        logger.warning(f"Fact-check extraction failed: {e}")
        return []

    # ── Step 2: Cross-reference each claim against KB ──
    from sqlalchemy import text as sql_text
    db = get_db()
    results = []
    try:
        for claim in claims[:5]:
            claim_text = claim.get("text", "")
            if not claim_text or not claim.get("verifiable", True):
                continue

            # Search KB for evidence
            kb_results = search_knowledge_base(db, claim_text)
            top_score = kb_results[0][0] if kb_results else 0
            top_content = kb_results[0][1].content[:500] if kb_results else ""
            top_source = kb_results[0][1].title if kb_results else ""

            # Determine verification status based on KB match
            if top_score >= 20:
                verification = "verified"  # Strong KB match
            elif top_score >= 8:
                verification = "partially_verified"  # Weak match, needs review
            elif top_score > 0:
                verification = "unverified"  # Some results but not matching
            else:
                verification = "novel"  # Nothing in KB – new knowledge

            confidence = min(top_score / 30.0, 1.0)  # Normalize to 0-1

            # Store in DB
            check_id = str(uuid.uuid4())
            try:
                db.execute(text("""
                    INSERT INTO knowledge_checks (id, user_email, chat_doc_id, claim_text,
                        claim_topic, customer_code, verification_status, confidence_score,
                        evidence, evidence_source, is_quantitative, user_certainty)
                    VALUES (:id, :email, :doc_id, :claim, :topic, :customer, :status,
                        :confidence, :evidence, :source, :is_quant, :certainty)
                """), {
                    "id": check_id,
                    "email": user_email,
                    "doc_id": chat_doc_id,
                    "claim": claim_text[:1000],
                    "topic": claim.get("topic", "")[:100],
                    "customer": customer_code or claim.get("customer", ""),
                    "status": verification,
                    "confidence": confidence,
                    "evidence": top_content[:2000] if top_score > 5 else None,
                    "source": top_source[:500] if top_score > 5 else None,
                    "is_quant": claim.get("type") == "quantitative",
                    "certainty": claim.get("user_certainty", "stated")
                })
                results.append({
                    "id": check_id,
                    "claim": claim_text,
                    "status": verification,
                    "confidence": round(confidence, 2),
                    "topic": claim.get("topic", ""),
                    "type": claim.get("type", "factual"),
                    "user_certainty": claim.get("user_certainty", "stated")
                })
            except Exception as e:
                logger.warning(f"Store knowledge check failed: {e}")
                db.rollback()
                continue

        db.commit()
        status_counts = {}
        for r in results:
            status_counts[r["status"]] = status_counts.get(r["status"], 0) + 1
        logger.info(f"🔍 Fact-check complete: {status_counts} for {user_email.split('@')[0]}")

        # ── Sync user profile to GitHub after fact-check ──
        try:
            sync_user_profile_to_github(user_email)
        except Exception as e:
            logger.warning(f"GitHub profile sync failed: {e}")

    except Exception as e:
        logger.error(f"Fact-check KB search failed: {e}")
        db.rollback()
    finally:
        db.close()

    return results


def sync_user_profile_to_github(user_email: str):
    """Build comprehensive user profile from all DB sources and push to GitHub.
    Creates data/team/{user-slug}/profile.yaml with:
    - Identity (name, role, company, expertise)
    - Knowledge Calibration (accuracy, overconfidence, strengths, gaps)
    - Behavioral Profile (Ψ-insights from onboarding)
    - Activity (chat stats, customers worked on)
    """
    import yaml, re as _re
    from sqlalchemy import text as sql_text

    if not GH_TOKEN:
        return

    db = get_db()
    try:
        # ── 1) User identity from users table ──
        user = db.query(User).filter(User.email == user_email).first()
        if not user:
            return

        user_slug = _re.sub(r'[^a-z0-9]+', '-', user_email.split('@')[0].lower()).strip('-')
        now = datetime.utcnow()

        identity = {
            "email": user.email,
            "name": user.name or user_email.split('@')[0].replace('.', ' ').title(),
            "role": user.role or "researcher",
            "company": user.company or "FehrAdvice & Partners AG",
            "position": user.position or None,
            "expertise": user.expertise or [],
            "linkedin_url": user.linkedin_url or None,
            "crm_role": user.crm_role or "none",
            "member_since": user.created_at.isoformat() if user.created_at else None,
            "last_login": user.last_login.isoformat() if user.last_login else None,
        }

        # ── 2) Knowledge Calibration from knowledge_checks ──
        calibration = {"total_claims": 0}
        try:
            stats = db.execute(text("""
                SELECT COUNT(*),
                       SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END),
                       SUM(CASE WHEN verification_status='partially_verified' THEN 1 ELSE 0 END),
                       SUM(CASE WHEN verification_status='unverified' THEN 1 ELSE 0 END),
                       SUM(CASE WHEN verification_status='novel' THEN 1 ELSE 0 END),
                       SUM(CASE WHEN is_quantitative=TRUE THEN 1 ELSE 0 END),
                       AVG(confidence_score)
                FROM knowledge_checks WHERE user_email=:e
            """), {"e": user_email}).fetchone()

            total_claims = stats[0] or 0
            verified = stats[1] or 0
            partial = stats[2] or 0
            unverified = stats[3] or 0
            novel = stats[4] or 0
            quant = stats[5] or 0

            checkable = verified + partial + unverified
            accuracy = round(verified / checkable * 100, 1) if checkable > 0 else None

            # Overconfidence
            overconfident_pct = 0.0
            certain_total = db.execute(text(
                "SELECT COUNT(*) FROM knowledge_checks WHERE user_email=:e AND user_certainty='certain'"
            ), {"e": user_email}).scalar() or 0
            if certain_total > 0:
                oc = db.execute(text("""
                    SELECT COUNT(*) FROM knowledge_checks
                    WHERE user_email=:e AND user_certainty='certain'
                      AND verification_status IN ('unverified', 'novel')
                """), {"e": user_email}).scalar() or 0
                overconfident_pct = round(oc / certain_total * 100, 1)

            # Label
            if total_claims < 5:
                label = "Zu wenig Daten"
            elif accuracy and accuracy >= 70 and overconfident_pct < 20:
                label = "Gut kalibriert"
            elif overconfident_pct >= 30:
                label = "Ueberschaetzt Wissen"
            elif accuracy and accuracy >= 50:
                label = "Vorsichtig"
            elif novel > verified:
                label = "Explorativ"
            else:
                label = "Wissensluecken"

            # Topic strengths & gaps
            strengths = db.execute(text("""
                SELECT claim_topic, COUNT(*) as cnt,
                       SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END) as v
                FROM knowledge_checks
                WHERE user_email=:e AND claim_topic IS NOT NULL AND claim_topic != ''
                GROUP BY claim_topic HAVING COUNT(*) >= 2
                ORDER BY (SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END)::float / COUNT(*)) DESC
                LIMIT 5
            """), {"e": user_email}).fetchall()

            gaps = db.execute(text("""
                SELECT claim_topic, COUNT(*) as cnt,
                       SUM(CASE WHEN verification_status IN ('unverified','novel') THEN 1 ELSE 0 END) as g
                FROM knowledge_checks
                WHERE user_email=:e AND claim_topic IS NOT NULL AND claim_topic != ''
                GROUP BY claim_topic HAVING COUNT(*) >= 2
                ORDER BY (SUM(CASE WHEN verification_status IN ('unverified','novel') THEN 1 ELSE 0 END)::float / COUNT(*)) DESC
                LIMIT 5
            """), {"e": user_email}).fetchall()

            # Customer knowledge
            cust_knowledge = db.execute(text("""
                SELECT customer_code, COUNT(*) as cnt,
                       SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END) as v,
                       SUM(CASE WHEN verification_status='novel' THEN 1 ELSE 0 END) as n
                FROM knowledge_checks
                WHERE user_email=:e AND customer_code IS NOT NULL AND customer_code != ''
                GROUP BY customer_code ORDER BY cnt DESC LIMIT 10
            """), {"e": user_email}).fetchall()

            calibration = {
                "total_claims": total_claims,
                "verified": verified,
                "partially_verified": partial,
                "unverified": unverified,
                "novel_knowledge": novel,
                "quantitative_claims": quant,
                "accuracy_pct": accuracy,
                "overconfidence_pct": overconfident_pct,
                "calibration_label": label,
                "strengths": [{"topic": s[0], "claims": s[1], "verified": s[2]} for s in strengths],
                "gaps": [{"topic": g[0], "claims": g[1], "unverified": g[2]} for g in gaps],
                "customer_knowledge": [
                    {"customer": c[0], "claims": c[1], "verified": c[2], "novel": c[3]}
                    for c in cust_knowledge
                ],
            }
        except Exception as e:
            logger.warning(f"Calibration data query failed: {e}")
            db.rollback()

        # ── 3) Behavioral Profile from user_insights (Ψ-dimensions) ──
        behavioral = {}
        try:
            insights = db.execute(text("""
                SELECT question_category, answer_text, domain_signal, thinking_style,
                       abstraction_level, autonomy_signal, latency_ms
                FROM user_insights WHERE user_email=:e
                ORDER BY created_at DESC LIMIT 20
            """), {"e": user_email}).fetchall()

            if insights:
                domains = {}
                styles = {}
                for ins in insights:
                    if ins[2]:  # domain_signal
                        domains[ins[2]] = domains.get(ins[2], 0) + 1
                    if ins[3]:  # thinking_style
                        styles[ins[3]] = styles.get(ins[3], 0) + 1

                avg_latency = None
                latencies = [ins[6] for ins in insights if ins[6]]
                if latencies:
                    avg_latency = round(sum(latencies) / len(latencies))

                behavioral = {
                    "total_insights": len(insights),
                    "domain_interests": dict(sorted(domains.items(), key=lambda x: -x[1])),
                    "thinking_styles": dict(sorted(styles.items(), key=lambda x: -x[1])),
                    "avg_response_latency_ms": avg_latency,
                }
        except Exception as e:
            logger.warning(f"Behavioral profile query failed: {e}")
            db.rollback()

        # ── 4) Activity stats from chat ──
        activity = {}
        try:
            chat_count = db.execute(text(
                "SELECT COUNT(*) FROM documents WHERE source_type='chat_insight' AND uploaded_by=:e"
            ), {"e": user_email}).scalar() or 0

            intent_dist = db.execute(text("""
                SELECT category, COUNT(*) FROM documents
                WHERE source_type='chat_insight' AND uploaded_by=:e
                GROUP BY category ORDER BY COUNT(*) DESC
            """), {"e": user_email}).fetchall()

            first_chat = db.execute(text(
                "SELECT MIN(created_at) FROM documents WHERE source_type='chat_insight' AND uploaded_by=:e"
            ), {"e": user_email}).scalar()

            activity = {
                "total_chat_insights": chat_count,
                "intent_distribution": {r[0]: r[1] for r in intent_dist},
                "first_interaction": first_chat.isoformat() if first_chat else None,
            }
        except Exception as e:
            logger.warning(f"Activity stats query failed: {e}")
            db.rollback()

        # ── 5) Belief & Bias Profile ──
        belief_profile = {}
        try:
            b_total = db.execute(text(
                "SELECT COUNT(*) FROM belief_analyses WHERE user_email=:e"
            ), {"e": user_email}).scalar() or 0
            if b_total > 0:
                avg_inf = db.execute(text(
                    "SELECT AVG(informed_score) FROM belief_analyses WHERE user_email=:e AND bias_source='user'"
                ), {"e": user_email}).scalar()
                avg_inf = float(avg_inf or 0.5)

                user_bias_count = db.execute(text(
                    "SELECT COUNT(*) FROM belief_analyses WHERE user_email=:e AND bias_source='user' AND belief_category='bias'"
                ), {"e": user_email}).scalar() or 0
                beatrix_bias_count = db.execute(text(
                    "SELECT COUNT(*) FROM belief_analyses WHERE user_email=:e AND bias_source='beatrix'"
                ), {"e": user_email}).scalar() or 0
                belief_count = b_total - user_bias_count - beatrix_bias_count

                top_biases = db.execute(text("""
                    SELECT bias_type, COUNT(*) FROM belief_analyses
                    WHERE user_email=:e AND bias_type IS NOT NULL
                    GROUP BY bias_type ORDER BY COUNT(*) DESC LIMIT 8
                """), {"e": user_email}).fetchall()

                reasoning_dist = db.execute(text("""
                    SELECT reasoning_type, COUNT(*) FROM belief_analyses
                    WHERE user_email=:e AND bias_source='user' GROUP BY reasoning_type
                """), {"e": user_email}).fetchall()

                evidence_dist = db.execute(text("""
                    SELECT evidence_basis, COUNT(*) FROM belief_analyses
                    WHERE user_email=:e AND bias_source='user' AND belief_category != 'bias'
                    GROUP BY evidence_basis
                """), {"e": user_email}).fetchall()

                belief_profile = {
                    "total_items": b_total,
                    "beliefs_extracted": belief_count,
                    "user_biases_detected": user_bias_count,
                    "beatrix_biases_detected": beatrix_bias_count,
                    "avg_informed_score": round(avg_inf, 2),
                    "reasoning_label": (
                        "Evidenzbasiert" if avg_inf >= 0.7
                        else "Gemischt" if avg_inf >= 0.4
                        else "Motiviert"
                    ),
                    "top_biases": [{"type": b[0], "count": b[1]} for b in top_biases],
                    "reasoning_distribution": {r[0]: r[1] for r in reasoning_dist},
                    "evidence_basis": {e[0]: e[1] for e in evidence_dist},
                }
        except Exception as e:
            logger.warning(f"Belief profile query failed: {e}")
            db.rollback()

        # ── 6) Build complete profile YAML ──
        profile = {
            "# BEATRIX User Profile": "Auto-generated from chat interactions, fact-checks, and bias analysis",
            "last_updated": now.isoformat(),
            "identity": identity,
            "knowledge_calibration": calibration,
            "belief_analysis": belief_profile,
            "behavioral_profile": behavioral,
            "activity": activity,
        }

        yaml_content = yaml.dump(profile, allow_unicode=True, default_flow_style=False, sort_keys=False)

        # ── 6) Push to GitHub ──
        gh_path = f"data/team/{user_slug}/profile.yaml"
        result = gh_put_file(
            GH_CONTEXT_REPO, gh_path, yaml_content,
            f"👤 Profile update: {identity['name']} – {calibration.get('calibration_label', 'initial')}"
        )
        if result.get("ok"):
            logger.info(f"👤→🐙 User profile synced to GitHub: {gh_path}")
        else:
            logger.warning(f"GitHub profile push failed: {result.get('error','')[:80]}")

        # ── 7) Push calibration details separately (claim-level data) ──
        try:
            recent_checks = db.execute(text("""
                SELECT claim_text, claim_topic, customer_code, verification_status,
                       confidence_score, is_quantitative, user_certainty, created_at
                FROM knowledge_checks WHERE user_email=:e
                ORDER BY created_at DESC LIMIT 50
            """), {"e": user_email}).fetchall()

            if recent_checks:
                checks_data = {
                    "last_updated": now.isoformat(),
                    "user": user_email,
                    "checks": [
                        {
                            "claim": r[0][:200],
                            "topic": r[1],
                            "customer": r[2],
                            "status": r[3],
                            "confidence": round(float(r[4] or 0), 2),
                            "quantitative": bool(r[5]),
                            "user_certainty": r[6],
                            "date": r[7].isoformat() if r[7] else None,
                        }
                        for r in recent_checks
                    ]
                }
                cal_yaml = yaml.dump(checks_data, allow_unicode=True, default_flow_style=False, sort_keys=False)
                cal_path = f"data/team/{user_slug}/calibration.yaml"
                gh_put_file(GH_CONTEXT_REPO, cal_path, cal_yaml,
                    f"🔍 Calibration update: {identity['name']} – {len(recent_checks)} checks")
        except Exception as e:
            logger.warning(f"Calibration detail push failed: {e}")

        # ── 9) Push beliefs & biases details to GitHub ──
        try:
            recent_beliefs = db.execute(text("""
                SELECT belief_text, belief_category, bias_type, bias_source,
                       informed_score, reasoning_type, evidence_basis, customer_code, created_at
                FROM belief_analyses WHERE user_email=:e
                ORDER BY created_at DESC LIMIT 50
            """), {"e": user_email}).fetchall()

            if recent_beliefs:
                beliefs_data = {
                    "last_updated": now.isoformat(),
                    "user": user_email,
                    "beliefs_and_biases": [
                        {
                            "text": r[0][:300],
                            "category": r[1],
                            "bias_type": r[2],
                            "source": r[3],
                            "informed_score": round(float(r[4] or 0.5), 2),
                            "reasoning": r[5],
                            "evidence_basis": r[6],
                            "customer": r[7],
                            "date": r[8].isoformat() if r[8] else None,
                        }
                        for r in recent_beliefs
                    ]
                }
                bel_yaml = yaml.dump(beliefs_data, allow_unicode=True, default_flow_style=False, sort_keys=False)
                bel_path = f"data/team/{user_slug}/beliefs.yaml"
                gh_put_file(GH_CONTEXT_REPO, bel_path, bel_yaml,
                    f"🧠 Beliefs update: {identity['name']} – {len(recent_beliefs)} items")
        except Exception as e:
            logger.warning(f"Beliefs detail push failed: {e}")

        # ── 9) Push belief/bias details to GitHub ──
        try:
            recent_beliefs = db.execute(text("""
                SELECT belief_text, belief_category, bias_type, bias_source,
                       informed_score, reasoning_type, evidence_basis, customer_code, context, created_at
                FROM belief_analyses WHERE user_email=:e
                ORDER BY created_at DESC LIMIT 50
            """), {"e": user_email}).fetchall()

            if recent_beliefs:
                beliefs_data = {
                    "last_updated": now.isoformat(),
                    "user": user_email,
                    "summary": belief_profile if belief_profile else {},
                    "items": [
                        {
                            "text": r[0][:300],
                            "category": r[1],
                            "bias_type": r[2],
                            "source": r[3],
                            "informed_score": round(float(r[4] or 0.5), 2),
                            "reasoning": r[5],
                            "evidence_basis": r[6],
                            "customer": r[7],
                            "context": (r[8] or "")[:200],
                            "date": r[9].isoformat() if r[9] else None,
                        }
                        for r in recent_beliefs
                    ]
                }
                bel_yaml = yaml.dump(beliefs_data, allow_unicode=True, default_flow_style=False, sort_keys=False)
                bel_path = f"data/team/{user_slug}/beliefs.yaml"
                gh_put_file(GH_CONTEXT_REPO, bel_path, bel_yaml,
                    f"🧠 Belief update: {identity['name']} – {len(recent_beliefs)} items, informed={belief_profile.get('avg_informed_score','?')}")
        except Exception as e:
            logger.warning(f"Belief detail push failed: {e}")

    except Exception as e:
        logger.error(f"sync_user_profile_to_github failed: {e}")
    finally:
        db.close()

# ── BIAS DETECTION ENGINE: Beliefs & Motivated Reasoning ──

BIAS_DETECTION_SYSTEM = """Du bist ein Experte fuer Behavioral Economics und kognitive Verzerrungen.
Analysiere den DIALOG (User-Nachricht + BEATRIX-Antwort) und extrahiere:

1. BELIEFS des Users – was glaubt/nimmt die Person an (explizit und implizit)?
2. BIASES des Users – welche kognitiven Verzerrungen zeigen sich?
3. BIASES von BEATRIX – wo koennte die KI-Antwort selbst verzerrt sein?
4. INFORMED vs. MOTIVATED Score (1.0 = rein evidenzbasiert/rational, 0.0 = rein motiviert/wunschgetrieben)

BIAS-TYPEN die du erkennen sollst:
- confirmation_bias: Sucht nur bestaetigenede Evidenz
- anchoring: Fixiert auf erste Zahl/Information
- availability_heuristic: Uebergewichtet leicht verfuegbare Beispiele
- overconfidence: Ueberschaetzt eigenes Wissen/Praezision
- status_quo_bias: Bevorzugt bestehende Loesung ohne Evidenz
- framing_effect: Schlussfolgerung haengt von Formulierung ab
- sunk_cost: Bewertet vergangene Investitionen statt Zukunftsnutzen
- halo_effect: Uebertraegt positiven Eindruck auf unverbundene Bereiche
- dunning_kruger: Kompetenzueberschaetzung bei geringem Wissen
- groupthink: Unkritische Uebernahme von Gruppenmeinung
- narrative_fallacy: Konstruiert kausale Geschichte aus Korrelation
- base_rate_neglect: Ignoriert Grundwahrscheinlichkeiten
- survivorship_bias: Beruecksichtigt nur erfolgreiche Faelle
- self_serving: Attribution von Erfolg auf sich, Misserfolg auf andere
- optimism_bias: Systematische Ueberschaetzung positiver Outcomes
- beatrix_anchoring: BEATRIX verankert auf KB-Daten ohne Kontext
- beatrix_confirmation: BEATRIX bestaetigt User unkritisch
- beatrix_knowledge_gap: BEATRIX antwortet trotz fehlender Evidenz

REASONING-TYPEN:
- informed: Basiert auf Daten, Evidenz, strukturierter Analyse
- motivated: Getrieben von Wuenschen, Identitaet, Selbstbild
- mixed: Teils evidenzbasiert, teils motiviert
- exploratory: Offene Suche ohne vorgefasste Meinung

EVIDENCE-BASIS:
- empirical: Beruft sich auf Daten/Studien
- experiential: Basiert auf eigener Erfahrung
- theoretical: Basiert auf Modell/Theorie
- anecdotal: Einzelbeispiel als Beweis
- none: Keine Evidenzbasis erkennbar

Antworte NUR mit JSON (keine Backticks, kein Markdown):
{
  "beliefs": [
    {
      "text": "Die Kundenzufriedenheit ist der wichtigste KPI",
      "category": "business_assumption",
      "implicit": false,
      "informed_score": 0.6,
      "reasoning_type": "mixed",
      "evidence_basis": "experiential",
      "bias_type": null
    }
  ],
  "user_biases": [
    {
      "bias_type": "anchoring",
      "description": "Fixiert auf 72% als Baseline ohne zu hinterfragen",
      "severity": 0.4,
      "belief_ref": "Die Zufriedenheit liegt bei 72%"
    }
  ],
  "beatrix_biases": [
    {
      "bias_type": "beatrix_confirmation",
      "description": "BEATRIX bestaetigt die 72% ohne eigene Quelle zu nennen",
      "severity": 0.3
    }
  ],
  "overall_informed_score": 0.65,
  "reasoning_pattern": "Der User kombiniert Daten mit Annahmen. Staerke bei quantitativen Aussagen, Schwaeche bei kausalen Schlussfolgerungen."
}

Wenn der Dialog trivial ist (Begruessung, einfache Frage ohne Beliefs): {"beliefs":[], "user_biases":[], "beatrix_biases":[], "overall_informed_score": null, "reasoning_pattern": null}
Maximal 5 Beliefs und 3 Biases pro Seite."""

def analyze_beliefs_and_biases(user_message: str, beatrix_response: str, user_email: str,
                                chat_doc_id: str = None, customer_code: str = None):
    """Analyze a chat exchange for beliefs, biases (user + BEATRIX), and reasoning quality.
    Stores results in belief_analyses table and syncs to user profile."""
    import urllib.request, ssl, re

    if not ANTHROPIC_API_KEY:
        return []
    # Skip short/trivial messages
    if not user_message or len(user_message.strip()) < 30:
        return []
    if not beatrix_response or len(beatrix_response.strip()) < 30:
        return []

    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    # ── Step 1: Analyze dialog with Claude ──
    dialog = f"USER-NACHRICHT:\n{user_message.strip()[:3000]}\n\nBEATRIX-ANTWORT:\n{beatrix_response.strip()[:3000]}"
    try:
        payload = json.dumps({
            "model": ANTHROPIC_MODEL_STANDARD,
            "max_tokens": 2000,
            "system": BIAS_DETECTION_SYSTEM,
            "messages": [{"role": "user", "content": dialog}]
        }).encode()
        req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
            headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
                "User-Agent": "BEATRIXLab/3.20-bias"
            })
        resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=45).read())
        answer = resp["content"][0]["text"]

        try:
            parsed = json.loads(answer.strip())
        except:
            json_match = re.search(r'\{.*\}', answer, re.DOTALL)
            if json_match:
                parsed = json.loads(json_match.group())
            else:
                return []

        beliefs = parsed.get("beliefs", [])
        user_biases = parsed.get("user_biases", [])
        beatrix_biases = parsed.get("beatrix_biases", [])
        overall_score = parsed.get("overall_informed_score")
        reasoning_pattern = parsed.get("reasoning_pattern")

        if not beliefs and not user_biases and not beatrix_biases:
            return []

        total_items = len(beliefs) + len(user_biases) + len(beatrix_biases)
        logger.info(f"🧠 Bias analysis: {len(beliefs)} beliefs, {len(user_biases)} user biases, "
                     f"{len(beatrix_biases)} BEATRIX biases, informed={overall_score}")

    except Exception as e:
        logger.warning(f"Bias analysis extraction failed: {e}")
        return []

    # ── Step 2: Store in DB ──
    from sqlalchemy import text as sql_text
    db = get_db()
    results = []
    try:
        # Store user beliefs
        for belief in beliefs[:5]:
            try:
                bid = str(uuid.uuid4())
                db.execute(text("""
                    INSERT INTO belief_analyses (id, user_email, chat_doc_id, belief_text,
                        belief_category, bias_type, bias_source, informed_score,
                        reasoning_type, evidence_basis, customer_code, context)
                    VALUES (:id, :email, :doc_id, :text, :cat, :bias, 'user',
                        :score, :reasoning, :evidence, :customer, :context)
                """), {
                    "id": bid, "email": user_email, "doc_id": chat_doc_id,
                    "text": belief.get("text", "")[:1000],
                    "cat": belief.get("category", "")[:50],
                    "bias": belief.get("bias_type") or None,
                    "score": belief.get("informed_score", 0.5),
                    "reasoning": belief.get("reasoning_type", "mixed")[:30],
                    "evidence": belief.get("evidence_basis", "none")[:30],
                    "customer": customer_code or "",
                    "context": reasoning_pattern or ""
                })
                results.append({"type": "belief", "text": belief.get("text", ""),
                               "informed_score": belief.get("informed_score", 0.5)})
            except Exception as e:
                logger.warning(f"Store belief failed: {e}")
                db.rollback()

        # Store user biases
        for bias in user_biases[:3]:
            try:
                bid = str(uuid.uuid4())
                db.execute(text("""
                    INSERT INTO belief_analyses (id, user_email, chat_doc_id, belief_text,
                        belief_category, bias_type, bias_source, informed_score,
                        reasoning_type, evidence_basis, customer_code, context)
                    VALUES (:id, :email, :doc_id, :text, 'bias', :bias, 'user',
                        :score, 'motivated', 'none', :customer, :desc)
                """), {
                    "id": bid, "email": user_email, "doc_id": chat_doc_id,
                    "text": bias.get("belief_ref", bias.get("description", ""))[:1000],
                    "bias": bias.get("bias_type", "unknown")[:50],
                    "score": 1.0 - bias.get("severity", 0.5),  # Invert: high severity = low informed
                    "customer": customer_code or "",
                    "desc": bias.get("description", "")[:500]
                })
                results.append({"type": "user_bias", "bias": bias.get("bias_type"),
                               "severity": bias.get("severity", 0.5)})
            except Exception as e:
                logger.warning(f"Store user bias failed: {e}")
                db.rollback()

        # Store BEATRIX biases (self-critique)
        for bias in beatrix_biases[:3]:
            try:
                bid = str(uuid.uuid4())
                db.execute(text("""
                    INSERT INTO belief_analyses (id, user_email, chat_doc_id, belief_text,
                        belief_category, bias_type, bias_source, informed_score,
                        reasoning_type, evidence_basis, customer_code, context)
                    VALUES (:id, :email, :doc_id, :text, 'bias', :bias, 'beatrix',
                        :score, 'motivated', 'none', :customer, :desc)
                """), {
                    "id": bid, "email": user_email, "doc_id": chat_doc_id,
                    "text": bias.get("description", "")[:1000],
                    "bias": bias.get("bias_type", "unknown")[:50],
                    "score": 1.0 - bias.get("severity", 0.5),
                    "customer": customer_code or "",
                    "desc": bias.get("description", "")[:500]
                })
                results.append({"type": "beatrix_bias", "bias": bias.get("bias_type"),
                               "severity": bias.get("severity", 0.5)})
            except Exception as e:
                logger.warning(f"Store beatrix bias failed: {e}")
                db.rollback()

        db.commit()
        logger.info(f"🧠 Stored {len(results)} belief/bias items for {user_email.split('@')[0]}")

    except Exception as e:
        logger.error(f"Bias analysis storage failed: {e}")
        db.rollback()
    finally:
        db.close()

    return results

def vector_search(db, query: str, limit: int = 8) -> list:
    """Semantic search. Tries pgvector, falls back to JSON + Python cosine similarity."""
    if not VOYAGE_API_KEY: return []
    from sqlalchemy import text as sql_text
    query_vec = embed_single(query, "query")
    if not query_vec: return []
    # Try pgvector first
    try:
        vec_str = f"[{','.join(str(v) for v in query_vec)}]"
        result = db.execute(text("""
            SELECT id, 1 - (embedding <=> :vec::vector) as similarity
            FROM documents WHERE embedding IS NOT NULL AND content IS NOT NULL
            ORDER BY embedding <=> :vec::vector LIMIT :lim
        """), {"vec": vec_str, "lim": limit})
        rows = result.fetchall()
        scored = []
        for row in rows:
            doc = db.query(Document).filter(Document.id == row[0]).first()
            if doc: scored.append((float(row[1]) * 100, doc))
        if scored:
            logger.info(f"Vector search (pgvector): {len(scored)} results")
            return scored
    except Exception:
        db.rollback()
    # Fallback: JSON embeddings + Python cosine similarity
    try:
        result = db.execute(text("SELECT id, embedding_json FROM documents WHERE embedding_json IS NOT NULL AND embedding_json != '' AND content IS NOT NULL"))
        rows = result.fetchall()
        if not rows: return []
        import math
        def cosine_sim(a, b):
            dot = sum(x*y for x,y in zip(a,b))
            na = math.sqrt(sum(x*x for x in a))
            nb = math.sqrt(sum(x*x for x in b))
            return dot / (na * nb) if na and nb else 0
        scored = []
        for row in rows:
            try:
                doc_vec = json.loads(row[1])
                sim = cosine_sim(query_vec, doc_vec)
                doc = db.query(Document).filter(Document.id == row[0]).first()
                if doc and sim > 0.3:
                    scored.append((sim * 100, doc))
            except: continue
        scored.sort(key=lambda x: -x[0])
        logger.info(f"Vector search (JSON cosine): {len(scored)} results")
        return scored[:limit]
    except Exception as e:
        logger.warning(f"Vector search error: {e}")
        return []

def fulltext_search(db, query: str, limit: int = 8) -> list:
    """PostgreSQL full-text search with ts_rank. Much better than keyword matching."""
    from sqlalchemy import text as sql_text
    try:
        # Use plainto_tsquery for natural language queries (handles German + English)
        result = db.execute(text("""
            SELECT id, 
                   ts_rank_cd(
                       setweight(to_tsvector('simple', COALESCE(title, '')), 'A') ||
                       setweight(to_tsvector('simple', COALESCE(category, '')), 'B') ||
                       setweight(to_tsvector('simple', COALESCE(content, '')), 'C'),
                       plainto_tsquery('simple', :query)
                   ) as rank
            FROM documents
            WHERE content IS NOT NULL
              AND (
                  to_tsvector('simple', COALESCE(title, '')) ||
                  to_tsvector('simple', COALESCE(content, ''))
              ) @@ plainto_tsquery('simple', :query)
            ORDER BY rank DESC
            LIMIT :lim
        """), {"query": query, "lim": limit})
        rows = result.fetchall()
        scored = []
        for row in rows:
            doc = db.query(Document).filter(Document.id == row[0]).first()
            if doc:
                # Scale rank to 0-100 range, boost EBF answers
                score = float(row[1]) * 1000
                if doc.source_type == "ebf_answer":
                    score *= 3
                scored.append((score, doc))
        logger.info(f"Fulltext search for '{query[:30]}': {len(scored)} results")
        return scored
    except Exception as e:
        logger.warning(f"Fulltext search error: {e}")
        return []

# ========== END VECTOR EMBEDDING ==========

app = FastAPI(title="BEA Lab Upload API", version="3.7.0")

def send_verification_email(email, name, token):
    """Send verification email via Resend API"""
    if not RESEND_API_KEY:
        logger.warning("RESEND_API_KEY not set, skipping verification email")
        return False
    import urllib.request, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    verify_url = f"{APP_URL}/api/verify/{token}"
    html = f"""<div style="font-family:system-ui,sans-serif;max-width:480px;margin:0 auto;padding:40px 24px">
    <div style="text-align:center;margin-bottom:32px">
        <h1 style="font-size:24px;font-weight:800;color:#0a1628;margin:0">BEATRIX <span style="color:#5b8af5">Lab</span></h1>
        <p style="color:#666;font-size:14px;margin:8px 0 0">Strategic Intelligence Suite</p>
    </div>
    <p style="font-size:15px;color:#333">Hallo {name or 'dort'},</p>
    <p style="font-size:15px;color:#333;line-height:1.6">Bitte bestätige deine E-Mail-Adresse, um dein BEATRIX Lab Konto zu aktivieren:</p>
    <div style="text-align:center;margin:32px 0">
        <a href="{verify_url}" style="display:inline-block;padding:14px 36px;background:#5b8af5;color:white;text-decoration:none;border-radius:10px;font-weight:700;font-size:15px">E-Mail bestätigen</a>
    </div>
    <p style="font-size:12px;color:#999;line-height:1.5">Falls der Button nicht funktioniert, kopiere diesen Link:<br>
    <a href="{verify_url}" style="color:#5b8af5;word-break:break-all">{verify_url}</a></p>
    <p style="font-size:12px;color:#999;margin-top:24px">Dieser Link ist 24 Stunden gültig.</p>
    <hr style="border:none;border-top:1px solid #eee;margin:32px 0">
    <p style="font-size:11px;color:#aaa;text-align:center">FehrAdvice &amp; Partners AG · Zürich</p>
</div>"""
    payload = json.dumps({
        "from": EMAIL_FROM,
        "to": [email],
        "subject": "BEATRIX Lab – E-Mail bestätigen",
        "html": html,
        "text": f"Hallo {name},\n\nBitte bestätige deine E-Mail: {verify_url}\n\nDieser Link ist 24 Stunden gültig."
    }).encode()
    try:
        req = urllib.request.Request("https://api.resend.com/emails", data=payload, method="POST",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}", "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.4"})
        resp = json.loads(urllib.request.urlopen(req, context=ctx).read())
        logger.info(f"Verification email sent to {email} via Resend: {resp.get('id','?')}")
        return True
    except Exception as e:
        logger.error(f"Failed to send verification email to {email}: {e}")
        return False
app.add_middleware(CORSMiddleware, allow_origins=["https://bea-lab.io", "https://www.bea-lab.io", "https://bea-lab-frontend.vercel.app", "http://localhost:3000"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
FRONTEND_DIR = Path(__file__).parent / "frontend"

class RegisterRequest(BaseModel):
    email: str
    password: str
    name: Optional[str] = None

class LoginRequest(BaseModel):
    email: str
    password: str

class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str

class ResetPasswordRequest(BaseModel):
    email: str

class ResetPasswordConfirm(BaseModel):
    token: str
    new_password: str

class ProfileUpdate(BaseModel):
    name: Optional[str] = None
    position: Optional[str] = None
    company: Optional[str] = None
    bio: Optional[str] = None
    phone: Optional[str] = None
    linkedin_url: Optional[str] = None
    expertise: Optional[List[str]] = None

class TextUploadRequest(BaseModel):
    title: str = "Untitled"
    content: str
    category: Optional[str] = "general"
    language: Optional[str] = "de"
    tags: Optional[List[str]] = []
    database: Optional[str] = "knowledge_base"

class DocumentResponse(BaseModel):
    id: str; title: str; source_type: str; file_type: Optional[str] = None
    database_target: str; status: str; created_at: str; github_url: Optional[str] = None; metadata: Optional[dict] = None
    metadata: Optional[dict] = None; uploaded_by: Optional[str] = None

def extract_text(file_path, file_type):
    try:
        if file_type == "pdf":
            import fitz; doc = fitz.open(file_path); text = "\n".join(page.get_text() for page in doc); doc.close(); return text.replace("\x00", "").strip()
        elif file_type == "docx":
            from docx import Document as DocxDoc; doc = DocxDoc(file_path); return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).replace("\x00", "")
        elif file_type in ("txt", "md", "csv", "json", "ics"):
            with open(file_path, "r", encoding="utf-8") as f: return f.read().replace("\x00", "")
    except Exception as e: return f"[Extraction error: {e}]"
    return ""

@app.get("/")
async def root():
    index_path = FRONTEND_DIR / "index.html"
    if index_path.exists():
        r = FileResponse(str(index_path))
        r.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        r.headers["Pragma"] = "no-cache"
        r.headers["Expires"] = "0"
        return r
    return {"message": "BEA Lab Upload API", "docs": "/docs"}

@app.get("/static/{filepath:path}")
async def static_files(filepath):
    file_path = FRONTEND_DIR / filepath
    if file_path.exists() and file_path.is_file(): return FileResponse(str(file_path))
    raise HTTPException(404, "Not found")

@app.post("/api/register")
async def register(request: RegisterRequest):
    email = request.email.strip().lower()
    if not email or '@' not in email: raise HTTPException(400, "Ungültige E-Mail-Adresse")
    if len(request.password) < 6: raise HTTPException(400, "Passwort muss mindestens 6 Zeichen haben")
    if ALLOWED_EMAIL_DOMAINS:
        domain = email.split('@')[-1]
        if domain not in ALLOWED_EMAIL_DOMAINS:
            allowed = ", ".join(f"@{d}" for d in ALLOWED_EMAIL_DOMAINS)
            raise HTTPException(403, f"Registrierung nur mit folgenden E-Mail-Domains erlaubt: {allowed}")
    db = get_db()
    try:
        if db.query(User).filter(User.email == email).first():
            raise HTTPException(409, "E-Mail bereits registriert")
        pw_hash, pw_salt = hash_password(request.password)
        is_admin = email in ADMIN_EMAILS
        verification_token = base64.urlsafe_b64encode(os.urandom(32)).decode().rstrip('=')
        skip_verification = is_admin or not REQUIRE_EMAIL_VERIFICATION or not RESEND_API_KEY
        user = User(
            email=email, name=request.name or email.split('@')[0],
            password_hash=pw_hash, password_salt=pw_salt, is_admin=is_admin,
            email_verified=skip_verification,
            verification_token=None if skip_verification else verification_token,
            verification_sent_at=None if skip_verification else datetime.utcnow()
        )
        db.add(user); db.commit(); db.refresh(user)
        if not skip_verification:
            sent = send_verification_email(email, user.name, verification_token)
            if not sent:
                # Fallback: auto-verify if SMTP fails
                user.email_verified = True; user.verification_token = None; db.commit()
                logger.warning(f"SMTP failed, auto-verified {email}")
            else:
                logger.info(f"Verification email sent to {email}")
                return JSONResponse({"status": "verification_required", "message": "Registrierung erfolgreich! Bitte prüfe dein E-Mail-Postfach und bestätige deine E-Mail-Adresse."})
        # Admin or no verification required → direct login
        token = create_jwt({**{"sub": user.email, "name": user.name, "uid": user.id, "admin": user.is_admin, "role": user.role or "researcher", "iat": int(time.time()), "exp": int(time.time()) + JWT_EXPIRY}, **_user_crm_payload(user)})
        logger.info(f"New user: {email} (verified={user.email_verified})")
        resp = JSONResponse({"token": token, "expires_in": JWT_EXPIRY, "user": {"email": user.email, "name": user.name}})
        resp.set_cookie(key="bea_token", value=token, max_age=JWT_EXPIRY, httponly=True, samesite="lax")
        return resp
    except HTTPException: raise
    except Exception as e: db.rollback(); raise HTTPException(500, f"Fehler: {e}")
    finally: db.close()

@app.post("/api/login")
async def login(request: LoginRequest):
    email = request.email.strip().lower()
    db = get_db()
    try:
        user = db.query(User).filter(User.email == email).first()
        if not user or not verify_password(request.password, user.password_hash, user.password_salt):
            raise HTTPException(401, "E-Mail oder Passwort falsch")
        if not user.is_active: raise HTTPException(403, "Konto deaktiviert")
        if REQUIRE_EMAIL_VERIFICATION and not user.email_verified:
            raise HTTPException(403, "E-Mail noch nicht bestätigt. Bitte prüfe dein Postfach.")
        user.last_login = datetime.utcnow(); db.commit()
        token = create_jwt({**{"sub": user.email, "name": user.name, "uid": user.id, "admin": user.is_admin, "role": user.role or "researcher", "iat": int(time.time()), "exp": int(time.time()) + JWT_EXPIRY}, **_user_crm_payload(user)})
        logger.info(f"Login: {email}")
        resp = JSONResponse({"token": token, "expires_in": JWT_EXPIRY, "user": {"email": user.email, "name": user.name}})
        resp.set_cookie(key="bea_token", value=token, max_age=JWT_EXPIRY, httponly=True, samesite="lax")
        return resp
    except HTTPException: raise
    except Exception as e: raise HTTPException(500, f"Login-Fehler: {e}")
    finally: db.close()

@app.get("/api/verify/{token}")
async def verify_email(token: str):
    db = get_db()
    try:
        user = db.query(User).filter(User.verification_token == token).first()
        if not user:
            html = _verify_page("Ungültiger Link", "Dieser Bestätigungslink ist ungültig oder wurde bereits verwendet.", False)
            return HTMLResponse(html)
        # Check 24h expiry
        if user.verification_sent_at and (datetime.utcnow() - user.verification_sent_at).total_seconds() > 86400:
            html = _verify_page("Link abgelaufen", "Dieser Bestätigungslink ist abgelaufen. Bitte melde dich an und fordere einen neuen Link an.", False)
            return HTMLResponse(html)
        user.email_verified = True
        user.verification_token = None
        db.commit()
        logger.info(f"Email verified: {user.email}")
        html = _verify_page("E-Mail bestätigt!", f"Deine E-Mail-Adresse ({user.email}) wurde erfolgreich bestätigt. Du kannst dich jetzt anmelden.", True)
        return HTMLResponse(html)
    except Exception as e:
        logger.error(f"Verify error: {e}")
        html = _verify_page("Fehler", "Ein Fehler ist aufgetreten. Bitte versuche es erneut.", False)
        return HTMLResponse(html)
    finally: db.close()

@app.post("/api/resend-verification")
async def resend_verification(request: LoginRequest):
    email = request.email.strip().lower()
    db = get_db()
    try:
        user = db.query(User).filter(User.email == email).first()
        if not user: raise HTTPException(404, "E-Mail nicht gefunden")
        if not verify_password(request.password, user.password_hash, user.password_salt):
            raise HTTPException(401, "E-Mail oder Passwort falsch")
        if user.email_verified: raise HTTPException(400, "E-Mail bereits bestätigt")
        # Rate limit: max once every 2 minutes
        if user.verification_sent_at and (datetime.utcnow() - user.verification_sent_at).total_seconds() < 120:
            raise HTTPException(429, "Bitte warte 2 Minuten, bevor du einen neuen Link anforderst.")
        new_token = base64.urlsafe_b64encode(os.urandom(32)).decode().rstrip('=')
        user.verification_token = new_token
        user.verification_sent_at = datetime.utcnow()
        db.commit()
        sent = send_verification_email(email, user.name, new_token)
        if not sent: raise HTTPException(500, "E-Mail konnte nicht gesendet werden")
        return {"message": "Neuer Bestätigungslink wurde gesendet."}
    except HTTPException: raise
    except Exception as e: raise HTTPException(500, f"Fehler: {e}")
    finally: db.close()

@app.put("/api/admin/users/{user_id}/verify")
async def admin_verify_user(user_id: str, user=Depends(require_permission("platform.manage_users"))):
    db = get_db()
    try:
        target = db.query(User).filter(User.id == user_id).first()
        if not target: raise HTTPException(404, "Benutzer nicht gefunden")
        target.email_verified = True; target.verification_token = None; db.commit()
        return {"email": target.email, "email_verified": True}
    finally: db.close()

@app.post("/api/change-password")
async def change_password(request: ChangePasswordRequest, user=Depends(require_auth)):
    db = get_db()
    try:
        db_user = db.query(User).filter(User.email == user["sub"]).first()
        if not db_user: raise HTTPException(404, "Benutzer nicht gefunden")
        if not verify_password(request.current_password, db_user.password_hash, db_user.password_salt):
            raise HTTPException(401, "Aktuelles Passwort ist falsch")
        if len(request.new_password) < 6:
            raise HTTPException(400, "Neues Passwort muss mindestens 6 Zeichen haben")
        pw_hash, pw_salt = hash_password(request.new_password)
        db_user.password_hash = pw_hash; db_user.password_salt = pw_salt; db.commit()
        logger.info(f"Password changed: {user['sub']}")
        return {"message": "Passwort erfolgreich geändert"}
    except HTTPException: raise
    except Exception as e: raise HTTPException(500, f"Fehler: {e}")
    finally: db.close()

# ── Profile ──────────────────────────────────────────────────────────────
@app.get("/api/profile")
async def get_profile(user=Depends(require_auth)):
    db = get_db()
    try:
        u = db.query(User).filter(User.email == user["sub"]).first()
        if not u: raise HTTPException(404, "Benutzer nicht gefunden")
        return {
            "email": u.email, "name": u.name, "position": u.position, "company": u.company,
            "bio": u.bio, "phone": u.phone, "linkedin_url": u.linkedin_url,
            "linkedin_connected": bool(u.linkedin_id), "profile_photo_url": u.profile_photo_url,
            "expertise": u.expertise or [], "is_admin": u.is_admin, "role": u.role or "researcher",
            "created_at": u.created_at.isoformat() if u.created_at else None,
            "last_login": u.last_login.isoformat() if u.last_login else None
        }
    finally: db.close()

@app.put("/api/profile")
async def update_profile(request: ProfileUpdate, user=Depends(require_auth)):
    db = get_db()
    try:
        u = db.query(User).filter(User.email == user["sub"]).first()
        if not u: raise HTTPException(404, "Benutzer nicht gefunden")
        if request.name is not None: u.name = request.name.strip()[:200]
        if request.position is not None: u.position = request.position.strip()[:200]
        if request.company is not None: u.company = request.company.strip()[:200]
        if request.bio is not None: u.bio = request.bio.strip()[:2000]
        if request.phone is not None: u.phone = request.phone.strip()[:50]
        if request.linkedin_url is not None:
            url = request.linkedin_url.strip()
            if url and not url.startswith("http"): url = "https://" + url
            u.linkedin_url = url[:500] if url else None
        if request.expertise is not None: u.expertise = request.expertise[:20]
        db.commit()
        logger.info(f"Profile updated: {user['sub']}")
        return {"message": "Profil aktualisiert"}
    except HTTPException: raise
    except Exception as e: db.rollback(); raise HTTPException(500, f"Fehler: {e}")
    finally: db.close()

@app.post("/api/profile/photo")
async def upload_profile_photo(file: UploadFile = File(...), user=Depends(require_auth)):
    ext = file.filename.split(".")[-1].lower() if file.filename else ""
    if ext not in {"jpg", "jpeg", "png", "webp"}: raise HTTPException(400, "Nur JPG, PNG oder WebP erlaubt")
    content = await file.read()
    if len(content) > 5_000_000: raise HTTPException(400, "Max 5 MB")
    photo_id = str(uuid.uuid4())
    photo_path = UPLOAD_DIR / f"photos/{photo_id}.{ext}"
    photo_path.parent.mkdir(parents=True, exist_ok=True)
    with open(photo_path, "wb") as f: f.write(content)
    db = get_db()
    try:
        u = db.query(User).filter(User.email == user["sub"]).first()
        if not u: raise HTTPException(404)
        u.profile_photo_url = f"/api/photos/{photo_id}.{ext}"
        db.commit()
        return {"photo_url": u.profile_photo_url}
    finally: db.close()

@app.get("/api/photos/{filename}")
async def serve_photo(filename: str):
    photo_path = UPLOAD_DIR / f"photos/{filename}"
    if not photo_path.exists(): raise HTTPException(404)
    return FileResponse(str(photo_path))

# ── LinkedIn OAuth ───────────────────────────────────────────────────────
@app.get("/api/auth/linkedin")
async def linkedin_auth(user=Depends(require_auth)):
    if not LINKEDIN_CLIENT_ID: raise HTTPException(501, "LinkedIn nicht konfiguriert")
    state = base64.urlsafe_b64encode(os.urandom(16)).decode().rstrip("=")
    redirect_uri = f"{APP_URL}/api/auth/linkedin/callback"
    url = (f"https://www.linkedin.com/oauth/v2/authorization?response_type=code"
           f"&client_id={LINKEDIN_CLIENT_ID}&redirect_uri={redirect_uri}"
           f"&state={state}&scope=openid%20profile%20email")
    return {"auth_url": url, "state": state}

@app.get("/api/auth/linkedin/callback")
async def linkedin_callback(code: str = "", state: str = "", error: str = ""):
    err_msg = error or "cancelled"
    if error or not code:
        return HTMLResponse(f"<script>window.opener.postMessage({{type:'linkedin_error',error:'{err_msg}'}},'*');window.close();</script>")
    redirect_uri = f"{APP_URL}/api/auth/linkedin/callback"
    import urllib.request, urllib.parse, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    # Exchange code for token
    try:
        token_data = urllib.parse.urlencode({"grant_type": "authorization_code", "code": code, "redirect_uri": redirect_uri, "client_id": LINKEDIN_CLIENT_ID, "client_secret": LINKEDIN_CLIENT_SECRET}).encode()
        req = urllib.request.Request("https://www.linkedin.com/oauth/v2/accessToken", data=token_data, method="POST", headers={"Content-Type": "application/x-www-form-urlencoded", "User-Agent": "BEATRIXLab/3.5"})
        token_resp = json.loads(urllib.request.urlopen(req, context=ctx).read())
        access_token = token_resp["access_token"]
    except Exception as e:
        logger.error(f"LinkedIn token exchange failed: {e}")
        return HTMLResponse(f"<script>window.opener.postMessage({{type:'linkedin_error',error:'token_exchange_failed'}},'*');window.close();</script>")
    # Fetch profile using OpenID userinfo
    try:
        req = urllib.request.Request("https://api.linkedin.com/v2/userinfo", headers={"Authorization": f"Bearer {access_token}", "User-Agent": "BEATRIXLab/3.5"})
        profile = json.loads(urllib.request.urlopen(req, context=ctx).read())
    except Exception as e:
        logger.error(f"LinkedIn profile fetch failed: {e}")
        return HTMLResponse(f"<script>window.opener.postMessage({{type:'linkedin_error',error:'profile_fetch_failed'}},'*');window.close();</script>")
    linkedin_data = {
        "sub": profile.get("sub", ""),
        "name": profile.get("name", ""),
        "given_name": profile.get("given_name", ""),
        "family_name": profile.get("family_name", ""),
        "email": profile.get("email", ""),
        "picture": profile.get("picture", ""),
    }
    return HTMLResponse(f"<script>window.opener.postMessage({{type:'linkedin_success',data:{json.dumps(linkedin_data)}}},'*');window.close();</script>")

@app.post("/api/profile/linkedin")
async def save_linkedin_data(data: dict, user=Depends(require_auth)):
    db = get_db()
    try:
        u = db.query(User).filter(User.email == user["sub"]).first()
        if not u: raise HTTPException(404)
        if data.get("sub"): u.linkedin_id = str(data["sub"])
        if data.get("name") and not u.name: u.name = data["name"]
        if data.get("picture"): u.profile_photo_url = data["picture"]
        db.commit()
        return {"message": "LinkedIn-Daten gespeichert", "linkedin_connected": True}
    finally: db.close()

@app.post("/api/profile/linkedin/disconnect")
async def disconnect_linkedin(user=Depends(require_auth)):
    db = get_db()
    try:
        u = db.query(User).filter(User.email == user["sub"]).first()
        if not u: raise HTTPException(404)
        u.linkedin_id = None
        db.commit()
        return {"message": "LinkedIn getrennt"}
    finally: db.close()

# ── Behavioral Insights / Ψ-Profiling ──────────────────────────────────

INSIGHT_QUESTION_POOL = [
    {"id": "iq-001", "text": "Wenn du eine Entscheidung unter Unsicherheit triffst — worauf verlässt du dich?", "category": "decision_style", "domain": "ORG",
     "choices": [
        {"icon": "📊", "label": "Daten & Evidenz", "signal": {"style": "analytical", "domain": "FIN"}},
        {"icon": "🧭", "label": "Intuition & Erfahrung", "signal": {"style": "intuitive", "domain": "ORG"}},
        {"icon": "⚖️", "label": "Beides — je nach Kontext", "signal": {"style": "adaptive", "domain": "ORG"}},
        {"icon": "🗣️", "label": "Ich frage mein Netzwerk", "signal": {"style": "social", "domain": "ORG"}},
        {"icon": "🎲", "label": "Ich entscheide schnell und korrigiere", "signal": {"style": "experimental", "domain": "ORG"}},
     ]},
    {"id": "iq-002", "text": "Welches Verhaltensphänomen fasziniert dich am meisten?", "category": "interest_probe", "domain": "FIN",
     "choices": [
        {"icon": "💰", "label": "Warum Menschen nicht sparen", "signal": {"style": "practical", "domain": "FIN"}},
        {"icon": "🏥", "label": "Warum wir ungesund leben", "signal": {"style": "practical", "domain": "HLT"}},
        {"icon": "🌍", "label": "Warum Gesellschaften nicht handeln beim Klima", "signal": {"style": "systemic", "domain": "ENV"}},
        {"icon": "🏢", "label": "Warum Organisationen sich nicht verändern", "signal": {"style": "structural", "domain": "ORG"}},
        {"icon": "🕊️", "label": "Wie Glaube und Kultur Entscheidungen formen", "signal": {"style": "cultural", "domain": "REL"}},
     ]},
    {"id": "iq-003", "text": "Du darfst ein Nudge-Experiment weltweit durchführen — welchen Bereich wählst du?", "category": "domain_preference", "domain": "HLT",
     "choices": [
        {"icon": "🏥", "label": "Gesundheit & Prävention", "signal": {"style": "practical", "domain": "HLT"}},
        {"icon": "💰", "label": "Finanzen & Altersvorsorge", "signal": {"style": "practical", "domain": "FIN"}},
        {"icon": "🌱", "label": "Umwelt & Nachhaltigkeit", "signal": {"style": "visionary", "domain": "ENV"}},
        {"icon": "🎓", "label": "Bildung & Lernen", "signal": {"style": "developmental", "domain": "EDU"}},
        {"icon": "🏛️", "label": "Demokratie & Partizipation", "signal": {"style": "systemic", "domain": "POL"}},
     ]},
    {"id": "iq-004", "text": "Was ist der wichtigste Hebel für Verhaltensänderung?", "category": "theory_preference", "domain": "ORG",
     "choices": [
        {"icon": "💡", "label": "Die richtige Information", "signal": {"style": "rational", "domain": "EDU"}},
        {"icon": "🏗️", "label": "Die richtige Architektur", "signal": {"style": "structural", "domain": "ORG"}},
        {"icon": "⏰", "label": "Der richtige Zeitpunkt", "signal": {"style": "contextual", "domain": "ORG"}},
        {"icon": "👥", "label": "Die richtigen sozialen Normen", "signal": {"style": "social", "domain": "ORG"}},
        {"icon": "❤️", "label": "Die richtige Motivation", "signal": {"style": "motivational", "domain": "ORG"}},
     ]},
    {"id": "iq-005", "text": "Zwei Strategien stehen zur Wahl — welche nimmst du?", "category": "risk_preference", "domain": "FIN",
     "choices": [
        {"icon": "🎯", "label": "70% sicher, moderater Impact", "signal": {"style": "risk_averse", "domain": "FIN"}},
        {"icon": "🚀", "label": "30% aber 10x Impact", "signal": {"style": "risk_seeking", "domain": "FIN"}},
        {"icon": "🔄", "label": "Sequenziell testen, dann skalieren", "signal": {"style": "experimental", "domain": "ORG"}},
        {"icon": "🤝", "label": "Beide kombinieren als Portfolio", "signal": {"style": "portfolio", "domain": "FIN"}},
        {"icon": "📋", "label": "Mehr Daten sammeln bevor ich entscheide", "signal": {"style": "analytical", "domain": "FIN"}},
     ]},
    {"id": "iq-006", "text": "Prägen kulturelle Werte das wirtschaftliche Verhalten stärker als Anreize?", "category": "worldview", "domain": "REL",
     "choices": [
        {"icon": "🕊️", "label": "Ja — Kultur ist der tiefste Treiber", "signal": {"style": "cultural", "domain": "REL"}},
        {"icon": "📈", "label": "Nein — Anreize dominieren immer", "signal": {"style": "economic", "domain": "FIN"}},
        {"icon": "🔀", "label": "Es ist komplementär", "signal": {"style": "integrative", "domain": "REL"}},
        {"icon": "🔬", "label": "Kommt auf den Kontext an", "signal": {"style": "contextual", "domain": "ORG"}},
        {"icon": "🧬", "label": "Biologie und Evolution prägen am stärksten", "signal": {"style": "evolutionary", "domain": "HLT"}},
     ]},
    {"id": "iq-007", "text": "Welcher Auftrag wäre am spannendsten für dich?", "category": "scope_preference", "domain": "POL",
     "choices": [
        {"icon": "🏛️", "label": "Eine Regierung beraten", "signal": {"style": "macro", "domain": "POL"}},
        {"icon": "🚀", "label": "Ein Startup transformieren", "signal": {"style": "micro", "domain": "ORG"}},
        {"icon": "🌐", "label": "Multilaterale Organisation (UN, WHO)", "signal": {"style": "global", "domain": "POL"}},
        {"icon": "🏦", "label": "Eine Grossbank neu denken", "signal": {"style": "structural", "domain": "FIN"}},
        {"icon": "🎓", "label": "Ein Bildungssystem redesignen", "signal": {"style": "developmental", "domain": "EDU"}},
     ]},
    {"id": "iq-008", "text": "Willingness, Ability, Capacity — wo scheitern die meisten Veränderungsprojekte?", "category": "bcm_understanding", "domain": "ORG",
     "choices": [
        {"icon": "❤️", "label": "Willingness — die Bereitschaft fehlt", "signal": {"style": "motivational", "domain": "ORG"}},
        {"icon": "🧠", "label": "Ability — die Fähigkeit fehlt", "signal": {"style": "capability", "domain": "EDU"}},
        {"icon": "🏗️", "label": "Capacity — die Struktur verhindert es", "signal": {"style": "structural", "domain": "ORG"}},
        {"icon": "🔗", "label": "Am Zusammenspiel aller drei", "signal": {"style": "systemic", "domain": "ORG"}},
        {"icon": "📏", "label": "Am falschen Messen — man weiss nicht wo", "signal": {"style": "analytical", "domain": "ORG"}},
     ]},
    {"id": "iq-009", "text": "Wie sollte ein gutes Verhaltensmodell sein?", "category": "abstraction_preference", "domain": "EDU",
     "choices": [
        {"icon": "🔢", "label": "Mathematisch präzise", "signal": {"style": "formal", "domain": "FIN"}},
        {"icon": "🎨", "label": "Intuitiv verständlich", "signal": {"style": "intuitive", "domain": "EDU"}},
        {"icon": "🔬", "label": "Empirisch validiert", "signal": {"style": "empirical", "domain": "ORG"}},
        {"icon": "🛠️", "label": "Direkt anwendbar in der Praxis", "signal": {"style": "practical", "domain": "ORG"}},
        {"icon": "🌊", "label": "Flexibel und kontextabhängig", "signal": {"style": "adaptive", "domain": "ORG"}},
     ]},
    {"id": "iq-010", "text": "Welches Bias ist am gefährlichsten in strategischen Entscheidungen?", "category": "bias_awareness", "domain": "ORG",
     "choices": [
        {"icon": "🦚", "label": "Overconfidence", "signal": {"style": "metacognitive", "domain": "ORG"}},
        {"icon": "🪨", "label": "Status Quo Bias", "signal": {"style": "structural", "domain": "ORG"}},
        {"icon": "💸", "label": "Sunk Cost Fallacy", "signal": {"style": "economic", "domain": "FIN"}},
        {"icon": "👥", "label": "Groupthink", "signal": {"style": "social", "domain": "ORG"}},
        {"icon": "🔍", "label": "Confirmation Bias", "signal": {"style": "analytical", "domain": "ORG"}},
     ]},
    {"id": "iq-011", "text": "Wenn du ein Behavioral Design Projekt startest — womit beginnst du?", "category": "design_approach", "domain": "ORG",
     "choices": [
        {"icon": "👤", "label": "Mit der Zielgruppe verstehen", "signal": {"style": "empathic", "domain": "ORG"}},
        {"icon": "🎯", "label": "Mit dem gewünschten Verhalten", "signal": {"style": "goal_oriented", "domain": "ORG"}},
        {"icon": "📊", "label": "Mit der Datenlage", "signal": {"style": "analytical", "domain": "FIN"}},
        {"icon": "🗺️", "label": "Mit der Entscheidungsumgebung", "signal": {"style": "structural", "domain": "ORG"}},
        {"icon": "📚", "label": "Mit der Literatur und Evidenz", "signal": {"style": "theoretical", "domain": "EDU"}},
     ]},
    {"id": "iq-012", "text": "Was treibt Unternehmen stärker?", "category": "motivation_theory", "domain": "FIN",
     "choices": [
        {"icon": "📈", "label": "Der Wunsch nach Wachstum", "signal": {"style": "growth", "domain": "FIN"}},
        {"icon": "🛡️", "label": "Die Angst vor Verlust", "signal": {"style": "loss_averse", "domain": "FIN"}},
        {"icon": "🏆", "label": "Der Wettbewerb mit anderen", "signal": {"style": "competitive", "domain": "ORG"}},
        {"icon": "🔄", "label": "Der Druck sich anzupassen", "signal": {"style": "adaptive", "domain": "ORG"}},
        {"icon": "💡", "label": "Die Vision einzelner Führungspersonen", "signal": {"style": "visionary", "domain": "ORG"}},
     ]},
]

@app.get("/api/insight/question")
async def get_insight_question(context: str = "login", user=Depends(require_auth)):
    """Get next insight question for user. Returns question + metadata about what's expected."""
    db = get_db()
    try:
        email = user["sub"]
        # Count how many insights this user already has
        total_insights = db.execute(text("SELECT COUNT(*) FROM user_insights WHERE user_email = :e"), {"e": email}).scalar() or 0
        # Count insights in current session (last 10 minutes)
        session_insights = db.execute(text(
            "SELECT COUNT(*) FROM user_insights WHERE user_email = :e AND created_at > NOW() - INTERVAL '10 minutes'"
        ), {"e": email}).scalar() or 0

        # Determine question number and mandatory/nudge status
        q_number = session_insights + 1
        is_mandatory = q_number <= 1
        is_nudged = q_number in [2, 3]
        is_voluntary = q_number >= 4

        # Get IDs of already-answered questions
        answered = db.execute(text(
            "SELECT question_text FROM user_insights WHERE user_email = :e AND skipped = FALSE"
        ), {"e": email}).fetchall()
        answered_texts = {r[0] for r in answered}

        # Pick a question not yet answered, personalized by profile
        profile = db.query(User).filter(User.email == email).first()
        expertise = (profile.expertise or []) if profile else []

        # Prioritize questions matching user expertise domains
        domain_map = {
            "Behavioral Economics": ["FIN", "ORG"],
            "Strategy": ["ORG", "POL"],
            "Decision Architecture": ["ORG", "FIN"],
            "Healthcare": ["HLT"],
            "Finance": ["FIN"],
            "Education": ["EDU"],
            "Religion": ["REL"],
            "Environment": ["ENV"],
            "Politics": ["POL"],
        }
        preferred_domains = set()
        for exp in expertise:
            for key, domains in domain_map.items():
                if key.lower() in exp.lower():
                    preferred_domains.update(domains)
        if not preferred_domains:
            preferred_domains = {"ORG", "FIN"}

        # Sort: preferred domain first, then others
        available = [q for q in INSIGHT_QUESTION_POOL if q["text"] not in answered_texts]
        if not available:
            available = INSIGHT_QUESTION_POOL  # Reset if all answered

        preferred = [q for q in available if q["domain"] in preferred_domains]
        other = [q for q in available if q["domain"] not in preferred_domains]

        random.shuffle(preferred)
        random.shuffle(other)
        ordered = preferred + other
        question = ordered[0] if ordered else INSIGHT_QUESTION_POOL[0]

        # Nudge messages
        nudge_msg = None
        if is_nudged:
            nudges = [
                "Noch eine Frage? Hilft BEATRIX, dich besser zu verstehen.",
                "Eine weitere Frage stärkt dein Ψ-Profil.",
                "Je mehr BEATRIX über deinen Denkstil weiss, desto besser die Modelle.",
            ]
            nudge_msg = random.choice(nudges)

        return {
            "question": question["text"],
            "question_id": question["id"],
            "category": question["category"],
            "choices": [{"icon": c["icon"], "label": c["label"]} for c in question.get("choices", [])],
            "question_number": q_number,
            "is_mandatory": is_mandatory,
            "is_nudged": is_nudged,
            "is_voluntary": is_voluntary,
            "nudge_message": nudge_msg,
            "total_answered": total_insights,
            "can_skip": not is_mandatory,
            "session_number": total_insights + 1,
        }
    except Exception as e:
        logger.error(f"Insight question error: {e}")
        # Fallback
        fallback_q = INSIGHT_QUESTION_POOL[0]
        return {
            "question": fallback_q["text"],
            "question_id": fallback_q["id"],
            "category": fallback_q["category"],
            "choices": [{"icon": c["icon"], "label": c["label"]} for c in fallback_q.get("choices", [])],
            "question_number": 1, "is_mandatory": True, "is_nudged": False,
            "is_voluntary": False, "nudge_message": None, "total_answered": 0,
            "can_skip": False, "session_number": 1,
        }
    finally: db.close()

class InsightAnswer(BaseModel):
    question_text: str
    question_id: Optional[str] = None
    answer_text: str
    choice_index: Optional[int] = None
    latency_ms: Optional[int] = None
    question_number: int = 1
    was_mandatory: bool = True
    was_nudged: bool = False
    skipped: bool = False
    context: str = "login"

@app.post("/api/insight/answer")
async def submit_insight_answer(request: InsightAnswer, user=Depends(require_auth)):
    """Submit answer to an insight question. Extracts behavioral metadata."""
    db = get_db()
    try:
        email = user["sub"]

        # Extract behavioral signals from the selected choice
        domain_signal = None
        thinking_style = None
        abstraction_level = "concise"  # Choice clicks are always concise
        autonomy_signal = "moderate"

        # Look up signal from question pool based on question_id and choice_index
        question_data = None
        for q in INSIGHT_QUESTION_POOL:
            if q["id"] == request.question_id:
                question_data = q; break

        if question_data and request.choice_index and question_data.get("choices"):
            idx = request.choice_index - 1  # 1-based to 0-based
            if 0 <= idx < len(question_data["choices"]):
                chosen = question_data["choices"][idx]
                signal = chosen.get("signal", {})
                domain_signal = signal.get("domain")
                thinking_style = signal.get("style")

        # Latency adds decision speed signal
        speed_suffix = ""
        if request.latency_ms and request.latency_ms < 3000:
            speed_suffix = "_fast"
        elif request.latency_ms and request.latency_ms > 20000:
            speed_suffix = "_deliberate"
        if thinking_style and speed_suffix:
            thinking_style = thinking_style + speed_suffix

        insight = UserInsight(
            user_email=email,
            question_text=request.question_text,
            question_category=request.question_id,
            answer_text=request.answer_text,
            choice_index=request.choice_index,
            latency_ms=request.latency_ms,
            session_number=1,  # TODO: calculate from login count
            question_number=request.question_number,
            was_mandatory=request.was_mandatory,
            was_nudged=request.was_nudged,
            skipped=request.skipped,
            domain_signal=domain_signal,
            thinking_style=thinking_style,
            abstraction_level=abstraction_level,
            autonomy_signal=autonomy_signal,
            context=request.context,
        )
        db.add(insight); db.commit()
        logger.info(f"Insight recorded: {email} q#{request.question_number} domain={domain_signal} style={thinking_style}")
        return {"status": "recorded", "domain_signal": domain_signal, "thinking_style": thinking_style}
    except Exception as e:
        db.rollback()
        logger.error(f"Insight answer error: {e}")
        raise HTTPException(500, f"Fehler: {e}")
    finally: db.close()

@app.get("/api/insight/profile")
async def get_insight_profile(user=Depends(require_auth)):
    """Get aggregated behavioral Ψ-profile for user."""
    db = get_db()
    try:
        email = user["sub"]
        insights = db.execute(text(
            "SELECT * FROM user_insights WHERE user_email = :e ORDER BY created_at DESC"
        ), {"e": email}).fetchall()

        if not insights:
            return {"has_profile": False, "total_insights": 0, "message": "Noch kein Ψ-Profil. Beantworte Fragen um dein Profil aufzubauen."}

        # Aggregate signals
        domains = {}
        styles = {}
        abstractions = {}
        autonomy_signals = {}
        total_latency = []
        voluntary_count = 0

        for row in insights:
            r = dict(row._mapping)
            if r.get("domain_signal"):
                domains[r["domain_signal"]] = domains.get(r["domain_signal"], 0) + 1
            if r.get("thinking_style"):
                base_style = r["thinking_style"].split("_")[0]
                styles[base_style] = styles.get(base_style, 0) + 1
            if r.get("abstraction_level"):
                abstractions[r["abstraction_level"]] = abstractions.get(r["abstraction_level"], 0) + 1
            if r.get("autonomy_signal"):
                autonomy_signals[r["autonomy_signal"]] = autonomy_signals.get(r["autonomy_signal"], 0) + 1
            if r.get("latency_ms"):
                total_latency.append(r["latency_ms"])
            if not r.get("was_mandatory") and not r.get("was_nudged") and not r.get("skipped"):
                voluntary_count += 1

        # Primary signals
        primary_domain = max(domains, key=domains.get) if domains else None
        primary_style = max(styles, key=styles.get) if styles else None
        avg_latency = sum(total_latency) / len(total_latency) if total_latency else None

        # Engagement score (0-100)
        engagement = min(100, len(insights) * 10 + voluntary_count * 15)

        return {
            "has_profile": True,
            "total_insights": len(insights),
            "primary_domain": primary_domain,
            "domain_distribution": domains,
            "primary_thinking_style": primary_style,
            "thinking_styles": styles,
            "abstraction_distribution": abstractions,
            "autonomy_distribution": autonomy_signals,
            "avg_decision_latency_ms": round(avg_latency) if avg_latency else None,
            "voluntary_answers": voluntary_count,
            "engagement_score": engagement,
            "decision_speed": "fast" if avg_latency and avg_latency < 5000 else "deliberate" if avg_latency and avg_latency > 20000 else "moderate",
            "answers": [
                {
                    "question": dict(row._mapping).get("question_text", ""),
                    "answer": dict(row._mapping).get("answer_text", ""),
                    "choice_index": dict(row._mapping).get("choice_index"),
                    "domain": dict(row._mapping).get("domain_signal"),
                    "style": dict(row._mapping).get("thinking_style"),
                    "latency_ms": dict(row._mapping).get("latency_ms"),
                    "created_at": str(dict(row._mapping).get("created_at", "")),
                }
                for row in insights if not dict(row._mapping).get("skipped")
            ],
        }
    except Exception as e:
        logger.error(f"Insight profile error: {e}")
        return {"has_profile": False, "total_insights": 0, "error": str(e)}
    finally: db.close()

@app.post("/api/insight/skip")
async def skip_insight_question(data: dict, user=Depends(require_auth)):
    """Record that user skipped a question."""
    db = get_db()
    try:
        insight = UserInsight(
            user_email=user["sub"],
            question_text=data.get("question_text", ""),
            question_category=data.get("question_id"),
            skipped=True,
            question_number=data.get("question_number", 1),
            was_mandatory=False,
            context=data.get("context", "login"),
        )
        db.add(insight); db.commit()
        return {"status": "skipped"}
    except Exception as e:
        db.rollback()
        return {"status": "error", "detail": str(e)}
    finally: db.close()

@app.post("/api/forgot-password")
async def forgot_password(request: ResetPasswordRequest):
    email = request.email.strip().lower()
    db = get_db()
    try:
        user = db.query(User).filter(User.email == email).first()
        # Always return success to prevent email enumeration
        if not user:
            return {"message": "Falls ein Konto mit dieser E-Mail existiert, wurde ein Reset-Link gesendet."}
        if not user.is_active:
            return {"message": "Falls ein Konto mit dieser E-Mail existiert, wurde ein Reset-Link gesendet."}
        # Rate limit: max once every 2 minutes
        if user.reset_sent_at and (datetime.utcnow() - user.reset_sent_at).total_seconds() < 120:
            raise HTTPException(429, "Bitte warte 2 Minuten, bevor du einen neuen Link anforderst.")
        reset_token = base64.urlsafe_b64encode(os.urandom(32)).decode().rstrip('=')
        user.reset_token = reset_token
        user.reset_sent_at = datetime.utcnow()
        db.commit()
        sent = send_reset_email(email, user.name, reset_token)
        if not sent:
            logger.warning(f"Reset email failed for {email}")
        return {"message": "Falls ein Konto mit dieser E-Mail existiert, wurde ein Reset-Link gesendet."}
    except HTTPException: raise
    except Exception as e: raise HTTPException(500, f"Fehler: {e}")
    finally: db.close()

@app.get("/api/reset/{token}")
async def reset_page(token: str):
    db = get_db()
    try:
        user = db.query(User).filter(User.reset_token == token).first()
        if not user:
            html = _verify_page("Ungültiger Link", "Dieser Reset-Link ist ungültig oder wurde bereits verwendet.", False)
            return HTMLResponse(html)
        if user.reset_sent_at and (datetime.utcnow() - user.reset_sent_at).total_seconds() > 3600:
            html = _verify_page("Link abgelaufen", "Dieser Reset-Link ist abgelaufen. Bitte fordere einen neuen an.", False)
            return HTMLResponse(html)
        return HTMLResponse(_reset_page(token))
    finally: db.close()

@app.post("/api/reset-password")
async def reset_password(request: ResetPasswordConfirm):
    db = get_db()
    try:
        user = db.query(User).filter(User.reset_token == request.token).first()
        if not user: raise HTTPException(400, "Ungültiger oder abgelaufener Link")
        if user.reset_sent_at and (datetime.utcnow() - user.reset_sent_at).total_seconds() > 3600:
            raise HTTPException(400, "Link abgelaufen. Bitte fordere einen neuen an.")
        if len(request.new_password) < 6:
            raise HTTPException(400, "Passwort muss mindestens 6 Zeichen haben")
        pw_hash, pw_salt = hash_password(request.new_password)
        user.password_hash = pw_hash; user.password_salt = pw_salt
        user.reset_token = None; user.reset_sent_at = None; db.commit()
        logger.info(f"Password reset: {user.email}")
        return {"message": "Passwort erfolgreich zurückgesetzt"}
    except HTTPException: raise
    except Exception as e: raise HTTPException(500, f"Fehler: {e}")
    finally: db.close()

def _verify_page(title, message, success):
    color = "#34d399" if success else "#f87171"
    icon = "✓" if success else "✗"
    return f"""<!DOCTYPE html><html lang="de"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>{title} – BEATRIX Lab</title>
<link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;600;800&display=swap" rel="stylesheet">
<style>body{{font-family:'Plus Jakarta Sans',sans-serif;background:#0a1628;color:#e4e9f2;display:flex;align-items:center;justify-content:center;min-height:100vh;margin:0}}
.card{{max-width:440px;padding:48px 40px;text-align:center}}
.icon{{font-size:56px;margin-bottom:20px;color:{color}}}
h1{{font-size:24px;font-weight:800;margin-bottom:12px}}h1 span{{color:#5b8af5}}
p{{font-size:15px;color:#8899b8;line-height:1.6;margin-bottom:28px}}
a{{display:inline-block;padding:14px 36px;background:#5b8af5;color:white;text-decoration:none;border-radius:10px;font-weight:700;font-size:15px}}
a:hover{{background:#7ba3ff}}</style></head>
<body><div class="card"><div class="icon">{icon}</div><h1>BEATRIX <span>Lab</span></h1><h2>{title}</h2><p>{message}</p>
<a href="{APP_URL}">Zum Login →</a></div></body></html>"""

def _reset_page(token):
    return f"""<!DOCTYPE html><html lang="de"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Passwort zurücksetzen – BEATRIX Lab</title>
<link href="https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@400;600;800&display=swap" rel="stylesheet">
<style>body{{font-family:'Plus Jakarta Sans',sans-serif;background:#0a1628;color:#e4e9f2;display:flex;align-items:center;justify-content:center;min-height:100vh;margin:0}}
.card{{max-width:440px;padding:48px 40px;text-align:center}}
h1{{font-size:24px;font-weight:800;margin-bottom:8px}}h1 span{{color:#5b8af5}}
p{{font-size:14px;color:#8899b8;margin-bottom:24px}}
input{{width:100%;padding:14px;background:#111d33;border:1px solid #1e2d4a;border-radius:10px;color:#e4e9f2;font-family:inherit;font-size:15px;outline:none;margin-bottom:12px;box-sizing:border-box}}
input:focus{{border-color:#5b8af5;box-shadow:0 0 0 3px rgba(91,138,245,0.15)}}
button{{width:100%;padding:14px;background:#5b8af5;color:white;border:none;border-radius:10px;font-weight:700;font-size:15px;cursor:pointer;font-family:inherit}}
button:hover{{background:#7ba3ff}}
.msg{{font-size:13px;padding:10px;border-radius:8px;margin-bottom:12px;display:none}}
.msg.error{{display:block;color:#f87171;background:rgba(248,113,113,0.08)}}
.msg.success{{display:block;color:#34d399;background:rgba(52,211,153,0.08)}}
</style></head>
<body><div class="card">
<h1>BEATRIX <span>Lab</span></h1>
<p>Neues Passwort festlegen</p>
<div class="msg" id="msg"></div>
<input type="password" id="pw1" placeholder="Neues Passwort (mind. 6 Zeichen)">
<input type="password" id="pw2" placeholder="Passwort bestätigen">
<button onclick="resetPw()">Passwort speichern</button>
<script>
async function resetPw() {{
    const pw1=document.getElementById('pw1').value, pw2=document.getElementById('pw2').value, msg=document.getElementById('msg');
    msg.className='msg';
    if(pw1.length<6){{msg.className='msg error';msg.textContent='Mind. 6 Zeichen';return}}
    if(pw1!==pw2){{msg.className='msg error';msg.textContent='Passwörter stimmen nicht überein';return}}
    try{{
        const r=await fetch('/api/reset-password',{{method:'POST',headers:{{'Content-Type':'application/json'}},body:JSON.stringify({{token:'{token}',new_password:pw1}})}});
        const d=await r.json();
        if(r.ok){{msg.className='msg success';msg.textContent='✓ Passwort geändert! Du wirst weitergeleitet...';setTimeout(()=>window.location.href='{APP_URL}',2000)}}
        else{{msg.className='msg error';msg.textContent=d.detail||'Fehler'}}
    }}catch(e){{msg.className='msg error';msg.textContent='Verbindungsfehler'}}
}}
</script>
</div></body></html>"""

def send_reset_email(email, name, token):
    """Send password reset email via Resend API"""
    if not RESEND_API_KEY:
        logger.warning("RESEND_API_KEY not set, skipping reset email")
        return False
    import urllib.request, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    reset_url = f"{APP_URL}/api/reset/{token}"
    html = f"""<div style="font-family:system-ui,sans-serif;max-width:480px;margin:0 auto;padding:40px 24px">
    <div style="text-align:center;margin-bottom:32px">
        <h1 style="font-size:24px;font-weight:800;color:#0a1628;margin:0">BEATRIX <span style="color:#5b8af5">Lab</span></h1>
        <p style="color:#666;font-size:14px;margin:8px 0 0">Strategic Intelligence Suite</p>
    </div>
    <p style="font-size:15px;color:#333">Hallo {name or 'dort'},</p>
    <p style="font-size:15px;color:#333;line-height:1.6">Du hast ein neues Passwort für dein BEATRIX Lab Konto angefordert:</p>
    <div style="text-align:center;margin:32px 0">
        <a href="{reset_url}" style="display:inline-block;padding:14px 36px;background:#5b8af5;color:white;text-decoration:none;border-radius:10px;font-weight:700;font-size:15px">Passwort zurücksetzen</a>
    </div>
    <p style="font-size:12px;color:#999;line-height:1.5">Falls du kein neues Passwort angefordert hast, ignoriere diese E-Mail.<br>
    <a href="{reset_url}" style="color:#5b8af5;word-break:break-all">{reset_url}</a></p>
    <p style="font-size:12px;color:#999;margin-top:24px">Dieser Link ist 1 Stunde gültig.</p>
    <hr style="border:none;border-top:1px solid #eee;margin:32px 0">
    <p style="font-size:11px;color:#aaa;text-align:center">FehrAdvice &amp; Partners AG · Zürich</p>
</div>"""
    payload = json.dumps({
        "from": EMAIL_FROM,
        "to": [email],
        "subject": "BEATRIX Lab – Passwort zurücksetzen",
        "html": html,
        "text": f"Hallo {name},\n\nSetze dein Passwort zurück: {reset_url}\n\nDieser Link ist 1 Stunde gültig."
    }).encode()
    try:
        req = urllib.request.Request("https://api.resend.com/emails", data=payload, method="POST",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}", "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.4"})
        resp = json.loads(urllib.request.urlopen(req, context=ctx).read())
        logger.info(f"Reset email sent to {email} via Resend: {resp.get('id','?')}")
        return True
    except Exception as e:
        logger.error(f"Failed to send reset email to {email}: {e}")
        return False

from fastapi.responses import HTMLResponse

@app.get("/api/auth/check")
async def check_auth(user=Depends(require_auth)):
    perms = get_user_permissions(user)
    return {
        "authenticated": True,
        "email": user.get("sub"),
        "name": user.get("name"),
        "role": perms["role"],
        "role_level": perms["role_level"],
        "is_fehradvice": perms["is_fehradvice"],
        "permissions": perms["permissions"]
    }

@app.get("/api/auth/permissions")
async def auth_permissions(user=Depends(require_auth)):
    """Return full permission map for frontend UI gating."""
    return get_user_permissions(user)

@app.get("/api/health")
async def health():
    db_ok = False
    try: db = get_db(); db.close(); db_ok = True
    except: pass
    return {"status": "ok" if db_ok else "degraded", "database": "connected" if db_ok else "unavailable", "github": "configured" if GH_TOKEN else "not configured", "github_repo": GH_REPO, "timestamp": datetime.utcnow().isoformat()}

@app.get("/api/admin/settings")
async def admin_settings(user=Depends(require_permission("platform.admin_dashboard"))):
    return {"allowed_email_domains": ALLOWED_EMAIL_DOMAINS or ["*"], "registration": "restricted" if ALLOWED_EMAIL_DOMAINS else "open", "jwt_expiry_hours": JWT_EXPIRY // 3600, "email_verification": REQUIRE_EMAIL_VERIFICATION, "smtp_configured": bool(RESEND_API_KEY)}

@app.get("/api/admin/kb-stats")
async def admin_kb_stats(user=Depends(require_permission("platform.view_analytics"))):
    """Knowledge Base statistics – shows how many chat insights have been captured."""
    from sqlalchemy import text as sql_text
    db = get_db()
    try:
        total = db.execute(text("SELECT COUNT(*) FROM documents")).scalar()
        chat_insights = db.execute(text("SELECT COUNT(*) FROM documents WHERE source_type='chat_insight'")).scalar()
        by_intent = db.execute(text("SELECT category, COUNT(*) as cnt FROM documents WHERE source_type='chat_insight' GROUP BY category ORDER BY cnt DESC")).fetchall()
        by_user = db.execute(text("SELECT uploaded_by, COUNT(*) as cnt FROM documents WHERE source_type='chat_insight' GROUP BY uploaded_by ORDER BY cnt DESC")).fetchall()
        recent = db.execute(text("SELECT title, category, uploaded_by, created_at FROM documents WHERE source_type='chat_insight' ORDER BY created_at DESC LIMIT 10")).fetchall()
        return {
            "total_documents": total,
            "chat_insights": chat_insights,
            "other_documents": total - chat_insights,
            "by_intent": [{"intent": r[0], "count": r[1]} for r in by_intent],
            "by_user": [{"user": r[0], "count": r[1]} for r in by_user],
            "recent": [{"title": r[0], "intent": r[1], "user": r[2], "created_at": r[3].isoformat() if r[3] else None} for r in recent]
        }
    finally: db.close()

@app.get("/api/admin/knowledge-calibration")
async def admin_knowledge_calibration(user=Depends(require_permission("platform.view_analytics"))):
    """Knowledge Calibration Dashboard – shows team knowledge profiles based on fact-checked claims."""
    from sqlalchemy import text as sql_text
    db = get_db()
    try:
        # Total checks
        total = db.execute(text("SELECT COUNT(*) FROM knowledge_checks")).scalar() or 0
        if total == 0:
            return {"total_checks": 0, "users": [], "summary": "Noch keine Fact-Checks. Daten werden automatisch bei Chat-Interaktionen gesammelt."}

        # Per-user calibration profile
        user_stats = db.execute(text("""
            SELECT user_email,
                   COUNT(*) as total_claims,
                   SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END) as verified,
                   SUM(CASE WHEN verification_status='partially_verified' THEN 1 ELSE 0 END) as partial,
                   SUM(CASE WHEN verification_status='unverified' THEN 1 ELSE 0 END) as unverified,
                   SUM(CASE WHEN verification_status='novel' THEN 1 ELSE 0 END) as novel,
                   SUM(CASE WHEN is_quantitative=TRUE THEN 1 ELSE 0 END) as quantitative,
                   SUM(CASE WHEN user_certainty='certain' THEN 1 ELSE 0 END) as stated_certain,
                   AVG(confidence_score) as avg_confidence
            FROM knowledge_checks
            GROUP BY user_email
            ORDER BY total_claims DESC
        """)).fetchall()

        users = []
        for r in user_stats:
            total_claims = r[1]
            verified = r[2] or 0
            partial = r[3] or 0
            novel = r[4] or 0
            unverified = r[5] or 0
            certain = r[7] or 0

            # Knowledge Accuracy Score: verified / (verified + unverified)
            checkable = verified + partial + unverified
            accuracy = round(verified / checkable * 100, 1) if checkable > 0 else None

            # Overconfidence: certain claims that are unverified
            overconfident = 0
            if certain > 0:
                oc_result = db.execute(text("""
                    SELECT COUNT(*) FROM knowledge_checks
                    WHERE user_email = :email
                      AND user_certainty = 'certain'
                      AND verification_status IN ('unverified', 'novel')
                """), {"email": r[0]}).scalar() or 0
                overconfident = round(oc_result / certain * 100, 1)

            # Top topics
            topics = db.execute(text("""
                SELECT claim_topic, COUNT(*) as cnt,
                       SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END) as v
                FROM knowledge_checks
                WHERE user_email = :email AND claim_topic IS NOT NULL AND claim_topic != ''
                GROUP BY claim_topic ORDER BY cnt DESC LIMIT 5
            """), {"email": r[0]}).fetchall()

            # Top customers
            customers = db.execute(text("""
                SELECT customer_code, COUNT(*) as cnt,
                       SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END) as v
                FROM knowledge_checks
                WHERE user_email = :email AND customer_code IS NOT NULL AND customer_code != ''
                GROUP BY customer_code ORDER BY cnt DESC LIMIT 5
            """), {"email": r[0]}).fetchall()

            users.append({
                "email": r[0],
                "name": r[0].split("@")[0].replace(".", " ").title(),
                "total_claims": total_claims,
                "verified": verified,
                "partially_verified": partial,
                "unverified": unverified,
                "novel_knowledge": novel,
                "quantitative_claims": r[6] or 0,
                "accuracy_pct": accuracy,
                "overconfidence_pct": overconfident,
                "avg_kb_confidence": round(float(r[8] or 0), 2),
                "calibration_label": (
                    "Gut kalibriert" if accuracy and accuracy >= 70 and overconfident < 20
                    else "Überschätzt Wissen" if overconfident >= 30
                    else "Vorsichtig" if accuracy and accuracy >= 50
                    else "Explorativ" if novel > verified
                    else "Zu wenig Daten" if total_claims < 5
                    else "Wissenslücken"
                ),
                "top_topics": [{"topic": t[0], "claims": t[1], "verified": t[2]} for t in topics],
                "top_customers": [{"customer": c[0], "claims": c[1], "verified": c[2]} for c in customers]
            })

        # Global summary
        all_verified = sum(u["verified"] for u in users)
        all_novel = sum(u["novel_knowledge"] for u in users)
        return {
            "total_checks": total,
            "total_verified": all_verified,
            "total_novel": all_novel,
            "knowledge_growth_rate": round(all_novel / total * 100, 1) if total > 0 else 0,
            "users": users
        }
    finally: db.close()

@app.get("/api/admin/knowledge-checks")
async def admin_knowledge_checks(user=Depends(require_permission("platform.view_analytics")), email: str = None, status: str = None, limit: int = 50):
    """Recent fact-check results, optionally filtered by user or status."""
    from sqlalchemy import text as sql_text
    db = get_db()
    try:
        query = "SELECT id, user_email, claim_text, claim_topic, customer_code, verification_status, confidence_score, evidence_source, is_quantitative, user_certainty, created_at FROM knowledge_checks WHERE 1=1"
        params = {"lim": limit}
        if email:
            query += " AND user_email = :email"
            params["email"] = email
        if status:
            query += " AND verification_status = :status"
            params["status"] = status
        query += " ORDER BY created_at DESC LIMIT :lim"
        rows = db.execute(text(query), params).fetchall()
        return [
            {
                "id": r[0], "user": r[1], "claim": r[2], "topic": r[3],
                "customer": r[4], "status": r[5], "confidence": round(float(r[6] or 0), 2),
                "source": r[7], "quantitative": r[8], "user_certainty": r[9],
                "created_at": r[10].isoformat() if r[10] else None
            } for r in rows
        ]
    finally: db.close()

@app.get("/api/admin/belief-analysis")
async def admin_belief_analysis(user=Depends(require_permission("platform.view_analytics"))):
    """Belief & Bias Dashboard – shows team reasoning patterns and cognitive biases."""
    from sqlalchemy import text as sql_text
    db = get_db()
    try:
        total = db.execute(text("SELECT COUNT(*) FROM belief_analyses")).scalar() or 0
        if total == 0:
            return {"total": 0, "users": [], "summary": "Noch keine Belief-Analysen. Werden automatisch bei Chat-Interaktionen erfasst."}

        # Per-user belief profile
        user_stats = db.execute(text("""
            SELECT user_email,
                   COUNT(*) as total,
                   AVG(informed_score) as avg_informed,
                   SUM(CASE WHEN bias_source='user' AND belief_category='bias' THEN 1 ELSE 0 END) as user_biases,
                   SUM(CASE WHEN bias_source='beatrix' THEN 1 ELSE 0 END) as beatrix_biases,
                   SUM(CASE WHEN bias_source='user' AND belief_category != 'bias' THEN 1 ELSE 0 END) as beliefs
            FROM belief_analyses GROUP BY user_email ORDER BY total DESC
        """)).fetchall()

        users = []
        for r in user_stats:
            # Top bias types for this user
            bias_types = db.execute(text("""
                SELECT bias_type, COUNT(*) as cnt FROM belief_analyses
                WHERE user_email=:e AND bias_type IS NOT NULL
                GROUP BY bias_type ORDER BY cnt DESC LIMIT 5
            """), {"e": r[0]}).fetchall()

            # Reasoning distribution
            reasoning = db.execute(text("""
                SELECT reasoning_type, COUNT(*) FROM belief_analyses
                WHERE user_email=:e AND bias_source='user' GROUP BY reasoning_type
            """), {"e": r[0]}).fetchall()

            # Evidence basis distribution
            evidence = db.execute(text("""
                SELECT evidence_basis, COUNT(*) FROM belief_analyses
                WHERE user_email=:e AND bias_source='user' AND belief_category != 'bias'
                GROUP BY evidence_basis
            """), {"e": r[0]}).fetchall()

            avg_informed = float(r[2] or 0.5)
            users.append({
                "email": r[0],
                "name": r[0].split("@")[0].replace(".", " ").title(),
                "total_items": r[1],
                "beliefs": r[5] or 0,
                "user_biases": r[3] or 0,
                "beatrix_biases": r[4] or 0,
                "avg_informed_score": round(avg_informed, 2),
                "reasoning_label": (
                    "Evidenzbasiert" if avg_informed >= 0.7
                    else "Gemischt" if avg_informed >= 0.4
                    else "Motiviert"
                ),
                "top_biases": [{"type": b[0], "count": b[1]} for b in bias_types],
                "reasoning_distribution": {r2[0]: r2[1] for r2 in reasoning},
                "evidence_basis": {e[0]: e[1] for e in evidence},
            })

        # Global bias frequency
        global_biases = db.execute(text("""
            SELECT bias_type, bias_source, COUNT(*) FROM belief_analyses
            WHERE bias_type IS NOT NULL GROUP BY bias_type, bias_source ORDER BY COUNT(*) DESC LIMIT 15
        """)).fetchall()

        return {
            "total": total,
            "users": users,
            "global_biases": [{"type": b[0], "source": b[1], "count": b[2]} for b in global_biases]
        }
    finally: db.close()

@app.get("/api/admin/beliefs")
async def admin_beliefs(user=Depends(require_permission("platform.view_analytics")), email: str = None, source: str = None, limit: int = 50):
    """Recent belief/bias items, optionally filtered."""
    from sqlalchemy import text as sql_text
    db = get_db()
    try:
        q = "SELECT id, user_email, belief_text, belief_category, bias_type, bias_source, informed_score, reasoning_type, evidence_basis, customer_code, created_at FROM belief_analyses WHERE 1=1"
        params = {"lim": limit}
        if email: q += " AND user_email=:email"; params["email"] = email
        if source: q += " AND bias_source=:source"; params["source"] = source
        q += " ORDER BY created_at DESC LIMIT :lim"
        rows = db.execute(text(q), params).fetchall()
        return [{
            "id": r[0], "user": r[1], "text": r[2][:200], "category": r[3],
            "bias_type": r[4], "source": r[5], "informed_score": round(float(r[6] or 0.5), 2),
            "reasoning": r[7], "evidence": r[8], "customer": r[9],
            "created_at": r[10].isoformat() if r[10] else None
        } for r in rows]
    finally: db.close()

@app.get("/api/user/knowledge-profile")
async def user_knowledge_profile(user=Depends(require_auth)):
    """Own knowledge profile – shows the user their calibration data (self-awareness tool)."""
    from sqlalchemy import text as sql_text
    db = get_db()
    email = user["sub"]
    try:
        total = db.execute(text("SELECT COUNT(*) FROM knowledge_checks WHERE user_email=:e"), {"e": email}).scalar() or 0
        if total == 0:
            return {"total_claims": 0, "message": "Noch keine Daten. Chatte mit BEATRIX und dein Wissensprofil baut sich automatisch auf."}

        stats = db.execute(text("""
            SELECT
                SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END),
                SUM(CASE WHEN verification_status='partially_verified' THEN 1 ELSE 0 END),
                SUM(CASE WHEN verification_status='unverified' THEN 1 ELSE 0 END),
                SUM(CASE WHEN verification_status='novel' THEN 1 ELSE 0 END),
                SUM(CASE WHEN is_quantitative=TRUE THEN 1 ELSE 0 END),
                AVG(confidence_score)
            FROM knowledge_checks WHERE user_email=:e
        """), {"e": email}).fetchone()

        verified, partial, unverified, novel, quant, avg_conf = stats
        checkable = (verified or 0) + (partial or 0) + (unverified or 0)
        accuracy = round((verified or 0) / checkable * 100, 1) if checkable > 0 else None

        # Recent checks
        recent = db.execute(text("""
            SELECT claim_text, verification_status, claim_topic, customer_code, confidence_score, created_at
            FROM knowledge_checks WHERE user_email=:e ORDER BY created_at DESC LIMIT 10
        """), {"e": email}).fetchall()

        # Strengths = topics with high verified rate
        strengths = db.execute(text("""
            SELECT claim_topic, COUNT(*) as cnt,
                   SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END) as v
            FROM knowledge_checks
            WHERE user_email=:e AND claim_topic IS NOT NULL AND claim_topic != ''
            GROUP BY claim_topic HAVING COUNT(*) >= 2
            ORDER BY (SUM(CASE WHEN verification_status='verified' THEN 1 ELSE 0 END)::float / COUNT(*)) DESC
            LIMIT 5
        """), {"e": email}).fetchall()

        # Gaps = topics with low verified rate
        gaps = db.execute(text("""
            SELECT claim_topic, COUNT(*) as cnt,
                   SUM(CASE WHEN verification_status IN ('unverified','novel') THEN 1 ELSE 0 END) as g
            FROM knowledge_checks
            WHERE user_email=:e AND claim_topic IS NOT NULL AND claim_topic != ''
            GROUP BY claim_topic HAVING COUNT(*) >= 2
            ORDER BY (SUM(CASE WHEN verification_status IN ('unverified','novel') THEN 1 ELSE 0 END)::float / COUNT(*)) DESC
            LIMIT 5
        """), {"e": email}).fetchall()

        # ── Belief & Bias data ──
        belief_stats = {"total": 0, "avg_informed_score": None, "reasoning_label": None,
                        "top_biases": [], "recent_beliefs": []}
        try:
            b_total = db.execute(text(
                "SELECT COUNT(*) FROM belief_analyses WHERE user_email=:e"
            ), {"e": email}).scalar() or 0
            if b_total > 0:
                avg_inf = db.execute(text(
                    "SELECT AVG(informed_score) FROM belief_analyses WHERE user_email=:e AND bias_source='user'"
                ), {"e": email}).scalar()
                avg_inf = float(avg_inf or 0.5)

                top_biases = db.execute(text("""
                    SELECT bias_type, COUNT(*) FROM belief_analyses
                    WHERE user_email=:e AND bias_type IS NOT NULL
                    GROUP BY bias_type ORDER BY COUNT(*) DESC LIMIT 5
                """), {"e": email}).fetchall()

                recent_b = db.execute(text("""
                    SELECT belief_text, belief_category, bias_type, bias_source,
                           informed_score, reasoning_type, evidence_basis, created_at
                    FROM belief_analyses WHERE user_email=:e ORDER BY created_at DESC LIMIT 10
                """), {"e": email}).fetchall()

                belief_stats = {
                    "total": b_total,
                    "avg_informed_score": round(avg_inf, 2),
                    "reasoning_label": (
                        "Evidenzbasiert" if avg_inf >= 0.7
                        else "Gemischt" if avg_inf >= 0.4
                        else "Motiviert"
                    ),
                    "top_biases": [{"type": b[0], "count": b[1]} for b in top_biases],
                    "recent_beliefs": [{
                        "text": r[0][:150], "category": r[1], "bias": r[2],
                        "source": r[3], "informed_score": round(float(r[4] or 0.5), 2),
                        "reasoning": r[5], "evidence": r[6],
                        "created_at": r[7].isoformat() if r[7] else None
                    } for r in recent_b]
                }
        except Exception:
            db.rollback()

        return {
            "total_claims": total,
            "verified": verified or 0,
            "partially_verified": partial or 0,
            "unverified": unverified or 0,
            "novel_knowledge": novel or 0,
            "quantitative_claims": quant or 0,
            "accuracy_pct": accuracy,
            "avg_kb_confidence": round(float(avg_conf or 0), 2),
            "strengths": [{"topic": s[0], "claims": s[1], "verified": s[2]} for s in strengths],
            "gaps": [{"topic": g[0], "claims": g[1], "unverified": g[2]} for g in gaps],
            "recent_checks": [
                {"claim": r[0][:150], "status": r[1], "topic": r[2], "customer": r[3],
                 "confidence": round(float(r[4] or 0), 2), "created_at": r[5].isoformat() if r[5] else None}
                for r in recent
            ],
            "belief_analysis": belief_stats
        }
    finally: db.close()

@app.get("/api/admin/users")
async def admin_users(user=Depends(require_permission("platform.manage_users"))):
    db = get_db()
    try:
        users = db.query(User).order_by(User.created_at.desc()).all()
        return [{"id": u.id, "email": u.email, "name": u.name, "is_active": u.is_active, "is_admin": u.is_admin, "role": u.role or "researcher", "email_verified": u.email_verified, "crm_access": u.crm_access or False, "crm_role": u.crm_role or "none", "crm_owner_code": u.crm_owner_code or "", "lead_management": getattr(u, 'lead_management', False) or False, "created_at": u.created_at.isoformat() if u.created_at else None, "last_login": u.last_login.isoformat() if u.last_login else None} for u in users]
    finally: db.close()

@app.put("/api/admin/users/{user_id}/toggle-active")
async def toggle_user_active(user_id: str, user=Depends(require_permission("platform.manage_users"))):
    db = get_db()
    try:
        target = db.query(User).filter(User.id == user_id).first()
        if not target: raise HTTPException(404, "Benutzer nicht gefunden")
        target.is_active = not target.is_active; db.commit()
        return {"email": target.email, "is_active": target.is_active}
    finally: db.close()

@app.put("/api/admin/users/{user_id}/role")
async def set_user_role(user_id: str, request: Request, user=Depends(require_permission("platform.manage_roles"))):
    db = get_db()
    try:
        data = await request.json()
        role = data.get("role", "researcher")
        if role not in ("researcher", "sales", "operations", "senior_management", "partner"):
            raise HTTPException(400, "Ungültige Rolle")
        target = db.query(User).filter(User.id == user_id).first()
        if not target: raise HTTPException(404, "Benutzer nicht gefunden")
        target.role = role
        # Auto-enable CRM for senior roles (if @fehradvice.com)
        if role in ("senior_management", "partner") and target.email.lower().endswith("@fehradvice.com"):
            if not target.crm_access:
                target.crm_access = True
                target.lead_management = True
                if not target.crm_role or target.crm_role == "none":
                    target.crm_role = "manager"
                logger.info(f"Auto-enabled CRM for senior role: {target.email}")
        db.commit()
        logger.info(f"Role changed: {target.email} → {role}")
        return {"email": target.email, "role": role, "crm_access": target.crm_access or False}
    finally: db.close()

@app.put("/api/admin/users/{user_id}/crm-access")
async def set_user_crm_access(user_id: str, request: Request, user=Depends(require_permission("platform.manage_users"))):
    """Toggle CRM access for a user. Only admins. Only @fehradvice.com emails eligible."""
    db = get_db()
    try:
        data = await request.json()
        target = db.query(User).filter(User.id == user_id).first()
        if not target: raise HTTPException(404, "Benutzer nicht gefunden")
        if not target.email.lower().endswith("@fehradvice.com"):
            raise HTTPException(400, "CRM nur für FehrAdvice-Mitarbeiter verfügbar")
        if "crm_access" in data:
            target.crm_access = bool(data["crm_access"])
        if "crm_role" in data:
            if data["crm_role"] not in ("none", "viewer", "manager", "admin"):
                raise HTTPException(400, "Ungültige CRM-Rolle")
            target.crm_role = data["crm_role"]
        if "crm_owner_code" in data:
            target.crm_owner_code = data["crm_owner_code"] or None
        if "lead_management" in data:
            target.lead_management = bool(data["lead_management"])
        db.commit()
        logger.info(f"CRM access changed: {target.email} → access={target.crm_access}, role={target.crm_role}, owner={target.crm_owner_code}, leads={target.lead_management}")
        return {"email": target.email, "crm_access": target.crm_access, "crm_role": target.crm_role, "crm_owner_code": target.crm_owner_code, "lead_management": target.lead_management or False}
    finally: db.close()

@app.put("/api/admin/users/{user_id}/reset-password")
async def admin_reset_password(user_id: str, request: Request, user=Depends(require_permission("platform.manage_users"))):
    """Admin resets a user's password."""
    db = get_db()
    try:
        data = await request.json()
        new_pw = data.get("password", "")
        if len(new_pw) < 8: raise HTTPException(400, "Passwort muss mindestens 8 Zeichen haben")
        target = db.query(User).filter(User.id == user_id).first()
        if not target: raise HTTPException(404, "Benutzer nicht gefunden")
        pw_hash, pw_salt = hash_password(new_pw)
        target.password_hash = pw_hash
        target.password_salt = pw_salt
        db.commit()
        logger.info(f"Admin password reset: {target.email}")

        # Auto-send notification email if requested
        data_notify = data.get("notify", True)
        email_sent = False
        if data_notify and RESEND_API_KEY:
            email_sent = send_admin_notification(
                to_email=target.email,
                to_name=target.name or target.email.split("@")[0],
                subject="🔐 BEATRIX – Dein Passwort wurde zurückgesetzt",
                template="password_reset",
                context={"new_password": new_pw}
            )

        return {"email": target.email, "status": "password_reset", "email_sent": email_sent}
    finally: db.close()

# ═══════════════════════════════════════════════════════════════
# ADMIN: Send notification email
# ═══════════════════════════════════════════════════════════════
def send_admin_notification(to_email: str, to_name: str, subject: str, template: str = "generic", context: dict = None):
    """Send admin notification email via Resend. Supports templates: password_reset, welcome, generic."""
    if not RESEND_API_KEY:
        logger.warning("RESEND_API_KEY not set, cannot send notification")
        return False
    import urllib.request, ssl
    ctx_ssl = ssl.create_default_context(); ctx_ssl.check_hostname = False; ctx_ssl.verify_mode = ssl.CERT_NONE
    context = context or {}

    # ── Templates ──
    if template == "password_reset":
        body_html = f"""
        <p style="margin:0 0 12px;color:#e2e8f0">Hallo {to_name},</p>
        <p style="margin:0 0 16px;color:#94a3b8">Dein BEATRIX-Passwort wurde soeben zurückgesetzt.</p>
        <div style="background:rgba(0,0,0,0.3);border-radius:8px;padding:16px;margin:16px 0">
            <div style="font-size:11px;color:#64748b;text-transform:uppercase;letter-spacing:0.5px;margin-bottom:8px">Neue Login-Daten</div>
            <div style="color:#e2e8f0"><strong>E-Mail:</strong> {to_email}</div>
            <div style="color:#e2e8f0;margin-top:4px"><strong>Passwort:</strong> {context.get('new_password', '(nicht angezeigt)')}</div>
            <div style="color:#e2e8f0;margin-top:4px"><strong>URL:</strong> <a href="{APP_URL}" style="color:#2a7f8e">{APP_URL.replace('https://','')}</a></div>
        </div>
        <p style="margin:16px 0 0;font-size:12px;color:#64748b">Bitte ändere dein Passwort nach dem ersten Login.</p>"""
    elif template == "welcome":
        body_html = f"""
        <p style="margin:0 0 12px;color:#e2e8f0">Hallo {to_name},</p>
        <p style="margin:0 0 16px;color:#94a3b8">Willkommen bei BEATRIX! Dein Account wurde eingerichtet.</p>
        <div style="background:rgba(0,0,0,0.3);border-radius:8px;padding:16px;margin:16px 0">
            <div style="color:#e2e8f0"><strong>URL:</strong> <a href="{APP_URL}" style="color:#2a7f8e">{APP_URL.replace('https://','')}</a></div>
        </div>
        <p style="margin:16px 0 0;font-size:12px;color:#64748b">{context.get('message', '')}</p>"""
    else:  # generic
        body_html = f"""
        <p style="margin:0 0 12px;color:#e2e8f0">Hallo {to_name},</p>
        <p style="margin:0 0 16px;color:#94a3b8">{context.get('message', 'Dies ist eine Benachrichtigung von BEATRIX.')}</p>"""

    html = f"""<div style="font-family:system-ui,-apple-system,sans-serif;max-width:520px;margin:0 auto;padding:32px;background:#0a1628;color:#e2e8f0;border-radius:16px">
    <div style="text-align:center;margin-bottom:24px">
        <div style="font-size:28px;font-weight:700;color:#2a7f8e">BEATRIX</div>
        <div style="font-size:12px;color:#64748b;letter-spacing:1px">BEHAVIORAL ECONOMICS LAB</div>
    </div>
    <div style="background:rgba(42,127,142,0.08);border:1px solid rgba(42,127,142,0.2);border-radius:12px;padding:24px;margin:20px 0">
        {body_html}
    </div>
    <div style="text-align:center;margin-top:24px">
        <a href="{APP_URL}" style="display:inline-block;background:#2a7f8e;color:white;padding:12px 32px;border-radius:8px;text-decoration:none;font-weight:600">Jetzt einloggen →</a>
    </div>
    <div style="text-align:center;margin-top:24px;font-size:11px;color:#475569">BEATRIX · FehrAdvice &amp; Partners AG · Zürich</div>
    </div>"""

    payload = json.dumps({
        "from": EMAIL_FROM,
        "to": [to_email],
        "subject": subject,
        "html": html,
    }).encode()
    try:
        req = urllib.request.Request("https://api.resend.com/emails", data=payload, method="POST",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}", "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7"})
        resp = json.loads(urllib.request.urlopen(req, context=ctx_ssl).read())
        logger.info(f"Admin notification sent to {to_email}: {resp.get('id','?')}")
        return True
    except Exception as e:
        logger.error(f"Failed to send notification to {to_email}: {e}")
        return False

@app.post("/api/admin/notify")
async def admin_send_notification(request: Request, user=Depends(require_permission("platform.manage_users"))):
    """Admin sends a notification email to any user. Templates: password_reset, welcome, generic."""
    data = await request.json()
    to_email = data.get("email", "").strip()
    if not to_email: raise HTTPException(400, "E-Mail-Adresse erforderlich")
    to_name = data.get("name", to_email.split("@")[0])
    subject = data.get("subject", "BEATRIX – Benachrichtigung")
    template = data.get("template", "generic")
    context = data.get("context", {})
    sent = send_admin_notification(to_email, to_name, subject, template, context)
    if not sent: raise HTTPException(500, "E-Mail konnte nicht gesendet werden")
    return {"status": "sent", "email": to_email, "template": template}

@app.put("/api/admin/users/{user_id}/toggle-admin")
async def admin_toggle_admin(user_id: str, request: Request, user=Depends(require_permission("platform.manage_roles"))):
    """Promote/demote a user to/from admin. Only existing admins can do this."""
    db = get_db()
    try:
        data = await request.json()
        target = db.query(User).filter(User.id == user_id).first()
        if not target: raise HTTPException(404, "Benutzer nicht gefunden")
        # Don't allow removing own admin
        if target.email.lower() == user.get("sub", "").lower() and not data.get("is_admin", True):
            raise HTTPException(400, "Eigenen Admin-Status kann man nicht entfernen")
        target.is_admin = bool(data.get("is_admin", False))
        # Auto-enable CRM for new FehrAdvice admins
        if target.is_admin and target.email.lower().endswith("@fehradvice.com"):
            target.crm_access = True
            if not target.crm_role or target.crm_role == "none":
                target.crm_role = "admin"
        db.commit()
        logger.info(f"Admin toggle: {target.email} → admin={target.is_admin}")
        return {"email": target.email, "is_admin": target.is_admin}
    finally: db.close()

# ── Session Context Management ──────────────────────────────────────────
SESSION_TYPES = {
    "lead": {"label": "Lead Management", "icon": "🎯", "color": "251,191,36"},
    "project": {"label": "Projekt", "icon": "📋", "color": "52,211,153"},
    "model": {"label": "BCM Modell", "icon": "🔬", "color": "139,170,255"},
    "context": {"label": "Kontext-Analyse", "icon": "🧠", "color": "232,163,61"},
    "research": {"label": "Research", "icon": "📚", "color": "168,162,255"},
    "general": {"label": "Allgemein", "icon": "💬", "color": "42,127,142"},
}

SESSION_RULES = {
    "lead": """SESSION-REGELN (Lead Management):
- Du arbeitest in einer LEAD-SESSION. Alle Nachrichten beziehen sich auf Lead/Sales.
- Behalte den Kundenkontext: Wenn ein Kunde erkannt wurde, gilt er für die ganze Session.
- Wenn der User ein Bild schickt, lies es SOFORT aus und extrahiere Kontaktdaten, Firmeninfos etc.
- Frage NICHT nach Infos die bereits in der Konversation stehen!
- Nutze bekannte Kundendaten um Felder automatisch zu befüllen.
- Ansprechpartner, Werte und Details aus früheren Nachrichten MERKEN und wiederverwenden.""",

    "project": """SESSION-REGELN (Projekt):
- Du arbeitest in einer PROJEKT-SESSION. Alle Nachrichten beziehen sich auf Projektverwaltung.
- Behalte den Projekt- und Kundenkontext über Nachrichten hinweg.
- Wenn Bilder/Dokumente geteilt werden, extrahiere relevante Projektinfos.
- Schlage das nächste Projektkürzel vor wenn nötig.
- DUPLIKAT-CHECK: Prüfe ob ein ähnliches Projekt bereits existiert.""",

    "model": """SESSION-REGELN (BCM Modell):
- Du arbeitest in einer MODELL-SESSION für Behavioral Competence Modelling.
- Behalte die Ψ-Dimensionen und Kontext-Vektoren über Nachrichten hinweg.
- Baue iterativ auf vorherigen Antworten auf.""",

    "context": """SESSION-REGELN (Kontext-Analyse):
- Du arbeitest in einer KONTEXT-SESSION zur Analyse von Ausgangslage/Kundensituation.
- Sammle systematisch Informationen über mehrere Nachrichten hinweg.
- Fasse periodisch zusammen was du bereits weisst.""",

    "research": """SESSION-REGELN (Research):
- Du arbeitest in einer RESEARCH-SESSION.
- Nutze die Knowledge Base für fundierte, wissenschaftliche Antworten.
- Verknüpfe verschiedene Quellen und Konzepte.""",

    "general": """SESSION-REGELN (Allgemein):
- Beantworte Fragen basierend auf dem Kontext der Konversation.
- Wenn sich ein spezifischer Modus ergibt (Lead, Projekt etc.), wechsle dorthin.""",
}

def get_or_create_session(db, session_id: str, user_email: str) -> dict:
    """Load or create a chat session with its metadata."""
    if not session_id:
        return {"id": None, "type": "general", "entities": {}, "context": {}}
    try:
        row = db.execute(text(
            "SELECT id, session_type, entities, context FROM chat_sessions WHERE id = :sid"
        ), {"sid": session_id}).fetchone()
        if row:
            entities = row[2] if isinstance(row[2], dict) else json.loads(row[2] or '{}')
            context = row[3] if isinstance(row[3], dict) else json.loads(row[3] or '{}')
            return {"id": row[0], "type": row[1] or "general", "entities": entities, "context": context}
        # Create new session
        db.execute(text(
            "INSERT INTO chat_sessions (id, user_email, session_type) VALUES (:sid, :email, 'general')"
        ), {"sid": session_id, "email": user_email})
        db.commit()
        return {"id": session_id, "type": "general", "entities": {}, "context": {}}
    except Exception as e:
        logger.warning(f"Session load/create failed: {e}")
        return {"id": session_id, "type": "general", "entities": {}, "context": {}}

def update_session(db, session_id: str, session_type: str = None, entities: dict = None, context: dict = None):
    """Update session metadata."""
    if not session_id:
        return
    try:
        updates = ["updated_at = CURRENT_TIMESTAMP"]
        params = {"sid": session_id}
        if session_type:
            updates.append("session_type = :stype")
            params["stype"] = session_type
        if entities is not None:
            updates.append("entities = :ent")
            params["ent"] = json.dumps(entities, ensure_ascii=False)
        if context is not None:
            updates.append("context = :ctx")
            params["ctx"] = json.dumps(context, ensure_ascii=False)
        db.execute(text(f"UPDATE chat_sessions SET {', '.join(updates)} WHERE id = :sid"), params)
        db.commit()
    except Exception as e:
        logger.warning(f"Session update failed: {e}")

def get_session_history(db, session_id: str, limit: int = 10) -> list:
    """Load recent messages from this session for conversation context.
    Truncates long messages to avoid token overflow (e.g. base64 images in content)."""
    if not session_id:
        return []
    try:
        rows = db.execute(text(
            "SELECT role, content FROM chat_messages WHERE session_id = :sid ORDER BY created_at DESC LIMIT :lim"
        ), {"sid": session_id, "lim": limit}).fetchall()
        history = []
        total_chars = 0
        max_total = 24000  # ~6k tokens budget for history
        for r in reversed(rows):
            content = r[1] or ""
            # Truncate individual messages
            if len(content) > 3000:
                content = content[:3000] + "\n[... gekürzt ...]"
            total_chars += len(content)
            if total_chars > max_total:
                break
            history.append({"role": r[0], "content": content})
        return history
    except Exception:
        return []

# ── BEATRIX Chat (RAG) ──────────────────────────────────────────────────
class ChatMessage(Base):
    __tablename__ = "chat_messages"
    id = Column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    user_email = Column(String(320), nullable=False, index=True)
    role = Column(String(20), nullable=False)  # "user" or "assistant"
    content = Column(Text, nullable=False)
    sources = Column(JSON, nullable=True)
    session_id = Column(String(50), nullable=True, index=True)
    created_at = Column(DateTime, default=datetime.utcnow)

class ChatRequest(BaseModel):
    question: str
    database: Optional[str] = None
    project_slug: Optional[str] = None
    session_id: Optional[str] = None
    attachments: Optional[list] = None  # [{type, name, content_text, image_base64, media_type}]
    analysis_mode: Optional[str] = "fast"  # 'fast' (default, auto-decide) or 'deep' (force deep path)


# ── Chat File Upload: Extract text / prepare image for Claude Vision ──
# If PDF is a scientific paper → runs full paper pipeline (KB + GitHub + Embedding)

def detect_book(text: str, filename: str = "") -> dict:
    """Detect if text is from a book (not a paper).
    Books have: ISBN, chapters, table of contents, publisher, preface/foreword, very long text.
    Returns {is_book: bool, confidence: float, signals: [str], score: int, isbn: str|None, publisher: str|None}"""
    if not text or len(text) < 1000:
        return {"is_book": False, "confidence": 0, "signals": [], "score": 0, "isbn": None, "publisher": None}

    import re
    text_lower = text[:20000].lower()
    text_end = text[-5000:].lower() if len(text) > 5000 else ""
    full_lower = text_lower + " " + text_end
    signals = []
    book_score = 0

    # ── ISBN extraction (STRONG: 4 points) ──
    isbn_13 = re.findall(r'(?:isbn[-\s:]?\s*(?:13)?[-\s:]?\s*)(97[89][-\s]?\d[-\s]?\d{2}[-\s]?\d{4,6}[-\s]?\d{1,3}[-\s]?\d)', text[:15000], re.IGNORECASE)
    isbn_10 = re.findall(r'(?:isbn[-\s:]?\s*(?:10)?[-\s:]?\s*)(\d[-\s]?\d{2}[-\s]?\d{4,6}[-\s]?\d{1,3}[-\s]?[\dxX])', text[:15000], re.IGNORECASE)
    # Also catch bare ISBN patterns
    isbn_bare_13 = re.findall(r'\b(97[89]\d{10})\b', text[:15000])
    isbn_bare_dash = re.findall(r'\b(97[89]-\d-\d{2,5}-\d{4,5}-\d)\b', text[:15000])
    isbn = None
    if isbn_13:
        isbn = re.sub(r'[-\s]', '', isbn_13[0])
        signals.append("isbn_13")
        book_score += 4
    elif isbn_10:
        isbn = re.sub(r'[-\s]', '', isbn_10[0])
        signals.append("isbn_10")
        book_score += 4
    elif isbn_bare_13:
        isbn = isbn_bare_13[0]
        signals.append("isbn_bare")
        book_score += 4
    elif isbn_bare_dash:
        isbn = re.sub(r'[-\s]', '', isbn_bare_dash[0])
        signals.append("isbn_dash")
        book_score += 4

    # ── Chapter structure (STRONG: 3 points) ──
    chapter_patterns = [r'\bchapter\s+\d', r'\bkapitel\s+\d', r'\bchapter\s+[ivxl]+\b',
                        r'\bteil\s+[ivxl]+\b', r'\bpart\s+[ivxl]+\b', r'\bpart\s+\d']
    chapter_count = sum(1 for p in chapter_patterns if re.search(p, full_lower))
    # Also count numbered chapters in text
    numbered_chapters = len(re.findall(r'\bchapter\s+\d+', full_lower))
    if numbered_chapters >= 3 or chapter_count >= 2:
        signals.append(f"chapter_structure({numbered_chapters}ch)")
        book_score += 3
    elif chapter_count >= 1:
        signals.append("chapter_ref")
        book_score += 1

    # ── Table of Contents (STRONG: 3 points) ──
    toc_patterns = ["table of contents", "contents\n", "inhaltsverzeichnis", "inhalt\n"]
    if any(p in text_lower for p in toc_patterns):
        signals.append("table_of_contents")
        book_score += 3

    # ── Publisher (MEDIUM: 2 points) ──
    publishers = ["oxford university press", "cambridge university press", "princeton university press",
                  "harvard university press", "mit press", "yale university press", "stanford university press",
                  "university of chicago press", "columbia university press", "springer", "wiley",
                  "elsevier", "routledge", "sage publications", "palgrave macmillan", "academic press",
                  "mcgraw-hill", "pearson", "cengage", "norton", "penguin", "random house",
                  "verlag", "de gruyter", "nomos", "c.h. beck", "mohr siebeck", "campus verlag",
                  "schäffer-poeschel", "gabler verlag", "haupt verlag"]
    publisher_found = None
    for pub in publishers:
        if pub in full_lower:
            publisher_found = pub.title()
            signals.append(f"publisher:{pub}")
            book_score += 2
            break

    # ── Book-specific sections (MEDIUM: 1 point each) ──
    book_sections = {
        "preface": ["preface", "vorwort", "geleitwort"],
        "foreword": ["foreword", "foreword by"],
        "index": ["\nindex\n", "sachregister", "stichwortverzeichnis", "personenregister", "name index", "subject index"],
        "about_author": ["about the author", "über den autor", "about the authors", "biographical note"],
        "acknowledgments_book": ["acknowledgments\n", "acknowledgements\n", "danksagung\n"],
        "copyright_page": ["all rights reserved", "printed in", "library of congress", "cataloging-in-publication",
                           "© 20", "© 19", "first published", "first edition", "second edition", "reprinted"],
        "dedication": ["dedicated to", "für meine", "to my wife", "to my family", "this book is dedicated"],
    }
    for name, patterns in book_sections.items():
        if any(p in full_lower for p in patterns):
            signals.append(name)
            book_score += 1

    # ── Very long text (books are typically >100k chars) ──
    if len(text) > 300000:
        signals.append(f"very_long({len(text)//1000}k)")
        book_score += 2
    elif len(text) > 150000:
        signals.append(f"long({len(text)//1000}k)")
        book_score += 1

    # ── Anti-paper signals that suggest book ──
    # Papers rarely have >10 chapters or >200k chars
    if len(text) > 200000 and not any(s.startswith("isbn") for s in signals):
        # Very long doc without ISBN – check for chapter structure
        all_chapters = len(re.findall(r'\bchapter\b', full_lower))
        if all_chapters >= 5:
            signals.append(f"long_with_chapters({all_chapters})")
            book_score += 2

    # Decision: score >= 5 = book
    is_book = book_score >= 5
    confidence = min(book_score / 12.0, 1.0)

    return {
        "is_book": is_book, "confidence": round(confidence, 2),
        "signals": signals, "score": book_score,
        "isbn": isbn, "publisher": publisher_found
    }


def lookup_isbn(isbn: str) -> dict:
    """Lookup book metadata via Open Library API using ISBN."""
    if not isbn:
        return {}
    try:
        import urllib.request, ssl, json
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE
        clean = isbn.replace("-", "").replace(" ", "")
        url = f"https://openlibrary.org/api/books?bibkeys=ISBN:{clean}&jscmd=data&format=json"
        req = urllib.request.Request(url, headers={"User-Agent": "BEATRIX/3.7"})
        resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=10).read())
        key = f"ISBN:{clean}"
        if key not in resp:
            return {}
        data = resp[key]
        return {
            "title": data.get("title", ""),
            "authors": [a.get("name", "") for a in data.get("authors", [])],
            "publishers": [p.get("name", "") for p in data.get("publishers", [])],
            "publish_date": data.get("publish_date", ""),
            "number_of_pages": data.get("number_of_pages"),
            "subjects": [s.get("name", "") for s in data.get("subjects", [])][:8],
            "cover": data.get("cover", {}).get("medium", ""),
            "url": data.get("url", ""),
            "isbn": clean
        }
    except Exception as e:
        logger.warning(f"ISBN lookup failed for {isbn}: {e}")
        return {}


def detect_study_report(text: str, filename: str = "") -> dict:
    """Detect if text is a study/report (not a paper, not a book).
    Studies/Reports: Consulting reports, government studies, think tank publications,
    industry analyses, white papers, policy briefs. No peer review, no ISBN,
    but empirical data, executive summaries, institutional branding.
    Returns {is_study: bool, confidence: float, signals: [str], score: int, org: str|None, study_type: str|None}"""
    if not text or len(text) < 500:
        return {"is_study": False, "confidence": 0, "signals": [], "score": 0, "org": None, "study_type": None}

    import re
    text_lower = text[:15000].lower()
    text_end = text[-5000:].lower() if len(text) > 5000 else ""
    full_lower = text_lower + " " + text_end
    fn_lower = (filename or "").lower()
    signals = []
    study_score = 0
    org_found = None
    study_type = None

    # ── STRONG: Executive Summary (3 points) ──
    exec_patterns = ["executive summary", "management summary", "zusammenfassung für entscheidungsträger",
                     "key findings", "key takeaways", "kernaussagen", "zentrale ergebnisse",
                     "highlights auf einen blick", "die wichtigsten ergebnisse"]
    if any(p in full_lower for p in exec_patterns):
        signals.append("executive_summary")
        study_score += 3

    # ── STRONG: Known organizations (3 points) ──
    consulting = ["mckinsey", "boston consulting", "bcg", "bain & company", "bain &", "deloitte",
                  "pwc", "pricewaterhousecoopers", "kpmg", "ernst & young", "ernst and young",
                  "accenture", "roland berger", "oliver wyman", "a.t. kearney",
                  "capgemini", "booz allen", "strategy&", "fehradvice", "fehr advice"]
    think_tanks = ["brookings", "rand corporation", "world economic forum",
                   "peterson institute", "chatham house", "carnegie endowment",
                   "heritage foundation", "cato institute", "bertelsmann stiftung",
                   "ifo institut", "diw berlin", "zew mannheim", "zew zentrum"]
    intl_orgs = ["world bank", "weltbank", "oecd", "european commission", "eu commission",
                 "europäische kommission", "united nations",
                 "world health organization", "international monetary fund",
                 "european central bank", "ezb ", "ecb ",
                 "bank for international settlements"]
    gov_agencies = ["bundesamt", "bundesministerium", "seco", "staatssekretariat",
                    "european parliament", "europäisches parlament", "congressional",
                    "government accountability", "bundesanstalt",
                    "bundesnetzagentur", "statistisches bundesamt", "destatis",
                    "bundesamt für statistik"]

    for org_list, org_type in [(consulting, "consulting"), (think_tanks, "think_tank"),
                                (intl_orgs, "international_org"), (gov_agencies, "government")]:
        for org in org_list:
            if org in full_lower:
                org_found = org.strip().title()
                study_type = org_type
                signals.append(f"org:{org.strip()}")
                study_score += 3
                break
        if org_found:
            break

    # ── STRONG: Study/Report title patterns in filename (2 points) ──
    report_fn_patterns = ["studie", "study", "report", "analyse", "analysis", "survey",
                          "whitepaper", "white_paper", "briefing", "outlook", "barometer",
                          "monitor", "index_report", "benchmark"]
    if any(p in fn_lower for p in report_fn_patterns):
        signals.append("filename_report")
        study_score += 2

    # ── MEDIUM: Report structure signals (1-2 points each) ──
    report_sections = {
        "methodology_section": (["methodology", "methodik", "our approach", "unser ansatz",
                                  "research methodology", "forschungsmethodik", "how we conducted",
                                  "survey methodology", "befragungsmethodik"], 2),
        "sample_description": (["sample size", "stichprobe", "n = ", "n=", "respondents",
                                 "befragte", "teilnehmer", "survey of ", "befragung von",
                                 "we surveyed", "wir haben befragt"], 2),
        "recommendations": (["recommendations", "empfehlungen", "handlungsempfehlungen",
                              "implications for", "implikationen für", "we recommend",
                              "wir empfehlen", "action points", "massnahmen"], 2),
        "disclaimer": (["disclaimer", "haftungsausschluss", "this report does not",
                         "the views expressed", "die in diesem bericht",
                         "for informational purposes only", "nicht als anlageberatung"], 1),
        "about_org": (["about us", "über uns", "who we are", "wer wir sind",
                        "our firm", "unser unternehmen", "about the authors",
                        "about the institute", "über das institut"], 1),
        "commissioned": (["commissioned by", "im auftrag von", "beauftragt von",
                           "prepared for", "erstellt für", "in cooperation with",
                           "in zusammenarbeit mit", "funded by", "gefördert von"], 2),
        "figures_charts": (["figure 1", "abbildung 1", "exhibit 1", "chart 1",
                            "figure 2", "abbildung 2", "see figure", "siehe abbildung",
                            "source:", "quelle:"], 1),
        "copyright_report": (["© 20", "all rights reserved", "proprietary and confidential",
                               "vertraulich", "nicht zur weitergabe", "for internal use"], 1),
    }

    for name, (patterns, points) in report_sections.items():
        if any(p in full_lower for p in patterns):
            signals.append(name)
            study_score += points

    # ── MEDIUM: Typical report phrases (1 point) ──
    report_phrases = {
        "survey_language": ["survey results", "umfrageergebnisse", "befragungsergebnisse",
                            "our survey", "unsere umfrage", "our research shows",
                            "unsere forschung zeigt", "the data shows", "die daten zeigen"],
        "industry_focus": ["industry report", "branchenreport", "market analysis", "marktanalyse",
                           "industry trends", "branchentrends", "sector analysis", "sektoranalyse",
                           "market outlook", "marktausblick"],
        "policy_language": ["policy brief", "policy implications", "politische implikationen",
                            "regulatory", "regulatorisch", "legislation", "gesetzgebung",
                            "compliance", "governance framework"],
        "benchmarking": ["benchmark", "best practice", "best practices", "peer comparison",
                         "vergleichsanalyse", "benchmarking", "ranking", "scorecard"],
        "year_in_title": ["2024", "2025", "2026", "q1 ", "q2 ", "q3 ", "q4 ",
                          "annual report", "jahresbericht", "annual review"],
    }

    for name, patterns in report_phrases.items():
        if any(p in full_lower for p in patterns):
            signals.append(name)
            study_score += 1

    # ── Length heuristic: reports are typically 10k-200k chars ──
    text_len = len(text)
    if 10000 < text_len < 200000:
        signals.append(f"report_length({text_len//1000}k)")
        study_score += 1

    # ── Anti-signals: if it has ISBN → book, if it has DOI/peer review → paper ──
    if any(p in full_lower for p in ["isbn", "isbn-13", "isbn-10"]):
        study_score -= 3
    if any(p in full_lower for p in ["peer-review", "submitted to", "accepted for publication"]):
        study_score -= 2

    # Decision: score >= 5 = study/report
    is_study = study_score >= 5
    confidence = min(max(study_score, 0) / 12.0, 1.0)

    # Classify study type if not already set
    if is_study and not study_type:
        if any(s in signals for s in ["policy_language"]):
            study_type = "policy_brief"
        elif any(s in signals for s in ["industry_focus", "benchmarking"]):
            study_type = "industry_report"
        elif any(s in signals for s in ["survey_language", "sample_description"]):
            study_type = "survey_study"
        else:
            study_type = "report"

    return {
        "is_study": is_study, "confidence": round(confidence, 2),
        "signals": signals, "score": study_score,
        "org": org_found, "study_type": study_type
    }


def detect_scientific_paper(text: str, filename: str = "") -> dict:
    """Content-based detection: Is this an academic/scientific document?
    
    Logic: TWO-PASS approach
    1. Check for NON-paper signals (invoice, report, slides, letter) → if strong → NOT paper
    2. Check for ANY academic signals → if found → paper
    3. If ambiguous: PDF with 3+ pages of prose → lean toward paper
    
    Returns {is_paper: bool, confidence: float, signals: [str], score: int}"""
    if not text or len(text) < 300:
        return {"is_paper": False, "confidence": 0, "signals": ["too_short"], "score": 0}

    text_start = text[:10000].lower()
    text_end = text[-4000:].lower() if len(text) > 4000 else ""
    text_lower = text_start + " " + text_end
    signals = []
    paper_score = 0
    not_paper_score = 0

    # ══════════════════════════════════════════════
    # PASS 1: Is this clearly NOT a paper?
    # ══════════════════════════════════════════════
    not_paper_signals = {
        "invoice": ["rechnung", "invoice", "rechnungsnummer", "zahlungsziel", "mwst",
                     "ust-id", "bankverbindung", "iban"],
        "slides": ["slide ", "folie ", "powerpoint", "präsentation erstellt"],
        "email": ["betreff:", "subject:", "von:", "an:", "gesendet:", "weitergeleitet"],
        "contract": ["vertrag", "vertragspartner", "kündigungsfrist", "laufzeit",
                      "vertragsgegenstand", "allgemeine geschäftsbedingungen"],
        "marketing": ["call to action", "landing page", "conversion rate",
                       "jetzt kaufen", "unverbindliches angebot"],
        "internal_report": ["protokoll der", "tagesordnung", "nächste schritte",
                            "action items", "meeting notes", "besprechungsprotokoll"],
    }

    for name, patterns in not_paper_signals.items():
        if sum(1 for p in patterns if p in text_lower) >= 2:
            signals.append(f"NOT_PAPER:{name}")
            not_paper_score += 3

    # If clearly not a paper → stop early
    if not_paper_score >= 3:
        return {"is_paper": False, "confidence": 0.9, "signals": signals, "score": 0,
                "reason": "non-academic document detected"}

    # ══════════════════════════════════════════════
    # PASS 2: Academic content signals
    # ══════════════════════════════════════════════

    # ── STRONG (2 points) ──
    strong = {
        "abstract": ["abstract\n", "abstract:", "abstract ", "\nabstract",
                      "zusammenfassung\n", "zusammenfassung:", "\nzusammenfassung"],
        "references_section": ["references\n", "\nreferences", "bibliography\n",
                               "\nbibliography", "literaturverzeichnis", "quellenverzeichnis",
                               "works cited", "literatur\n", "\nliteratur"],
        "doi": ["doi:", "doi.org/", "https://doi", "10.1016/", "10.1038/",
                "10.1111/", "10.1257/", "10.2139/", "10.1007/"],
        "journal_venue": ["journal of ", "proceedings of ", "conference on ",
                          "zeitschrift für", "quarterly journal", "economic review",
                          "annual review", "review of ", "american economic",
                          "econometrica", "published in", "erschienen in"],
        "peer_review": ["peer-review", "submitted to", "accepted for publication",
                        "forthcoming in", "under review", "revised version",
                        "working paper", "discussion paper", "arbeitspapier"],
        "preprint_repo": ["ssrn.com", "arxiv.org", "nber.org", "repec.org",
                          "social science research network",
                          "electronic copy available", "available at ssrn"],
    }

    # ── MEDIUM (1 point) ──
    medium = {
        "et_al": ["et al.", "et al,", "et al "],
        "methodology": ["methodology", "research method", "empirical study",
                        "empirische studie", "randomized controlled",
                        "experimental design", "field experiment",
                        "natural experiment", "quasi-experiment", "research design",
                        "methodik", "forschungsdesign", "untersuchungsmethode"],
        "hypothesis": ["hypothesis", "hypothes", "h1:", "h2:", "hypothese",
                        "forschungsfrage", "research question"],
        "findings": ["findings", "results show", "ergebnisse zeigen", "we find that",
                     "our results", "the evidence suggests", "our findings",
                     "befunde", "resultate"],
        "literature": ["literature review", "related work", "prior research",
                       "literaturüberblick", "previous studies", "existing literature",
                       "forschungsstand", "stand der forschung"],
        "academic_citations": ["(20", "(19", "(18", "pp.", "vol.", "no.", "p. ", "ibid",
                                "vgl.", "ebd.", "a.a.o."],
        "keywords": ["keywords:", "key words:", "schlüsselwörter", "jel classification",
                     "jel codes", "schlagworte:"],
        "acknowledgments": ["acknowledgment", "acknowledgement", "danksagung",
                            "we thank", "we are grateful", "financial support from",
                            "förderung durch"],
        "author_affiliation": ["university", "universität", "institute", "institut",
                               "department of", "school of", "faculty of",
                               "business school", "hochschule", "akademie",
                               "forschungsinstitut", "max-planck", "fraunhofer"],
        "statistical": ["p < ", "p<0.", "p=0.", "n = ", "sd = ",
                        "statistically significant", "regression", "correlation",
                        "standard error", "confidence interval", "treatment effect",
                        "coefficient", "t-stat", "robust", "signifikant"],
        "academic_sections": ["introduction\n", "\n1. introduction",
                              "\n1 introduction", "conclusion\n", "\nconclusion",
                              "discussion\n", "\ndiscussion", "appendix\n",
                              "\nappendix", "supplementary material",
                              "einleitung\n", "\neinleitung", "schluss\n",
                              "schlussbemerkung", "fazit\n", "\nfazit",
                              "kapitel ", "abschnitt "],
        "data_sources": ["dataset", "data set", "survey data", "panel data",
                         "cross-section", "time series", "sample size",
                         "observations", "stichprobe", "datensatz", "erhebung"],
        "contribution": ["contribution", "we contribute", "this paper",
                         "in this paper", "this study", "we show that",
                         "we argue", "we propose", "in dieser arbeit",
                         "diese studie", "vorliegende arbeit", "diese arbeit",
                         "der vorliegende beitrag", "dieser beitrag",
                         "diese untersuchung", "gegenstand dieser"],
        "academic_language": ["furthermore", "moreover", "nevertheless",
                              "notwithstanding", "henceforth", "thereby",
                              "thereof", "insofar", "darüber hinaus",
                              "nichtsdestotrotz", "indes", "mithin",
                              "gleichwohl", "insbesondere", "hingegen"],
        "theoretical": ["theory", "theorem", "proposition", "framework",
                        "paradigm", "theorie", "paradigma", "modell",
                        "konzept", "ansatz", "these", "antithese"],
        "footnotes_endnotes": ["footnote", "endnote", "fußnote", "anmerkung",
                                "fn.", "fn "],
    }

    for name, patterns in strong.items():
        if any(p in text_lower for p in patterns):
            signals.append(name)
            paper_score += 2
    for name, patterns in medium.items():
        if any(p in text_lower for p in patterns):
            signals.append(name)
            paper_score += 1

    # ══════════════════════════════════════════════
    # PASS 3: Structural heuristic for ambiguous cases
    # ══════════════════════════════════════════════
    # Long prose document (>5000 chars) with at least SOME academic flavor
    if paper_score >= 2 and len(text) > 5000:
        signals.append("long_prose_with_signals")
        paper_score += 1

    # Decision: score >= 3 = paper (aggressive: better to index than to miss)
    is_paper = paper_score >= 3
    confidence = min(paper_score / 10.0, 1.0)

    return {"is_paper": is_paper, "confidence": round(confidence, 2),
            "signals": signals, "score": paper_score}


def extract_paper_metadata(text: str, filename: str = "") -> dict:
    """Extract author, year, title from paper text content.
    Returns {author: str, year: str, title: str, doc_type: str}"""
    import re
    lines = [l.strip() for l in text[:5000].split("\n") if l.strip()]
    author = ""
    year = ""
    title = ""
    doc_type = "PAP"  # default

    # ── Title: usually first substantial line (>10 chars, not all-caps header like "NBER") ──
    for line in lines[:15]:
        if len(line) < 8: continue
        if re.match(r'^(NBER|SSRN|WORKING PAPER|DISCUSSION PAPER|WP\s)', line, re.I): continue
        if re.match(r'^(Electronic copy|Available at|Downloaded from)', line, re.I): continue
        if re.match(r'^https?://', line): continue
        if re.match(r'^\d+$', line): continue  # page numbers
        if '@' in line: continue  # email lines
        # Skip lines that are clearly author affiliations
        if re.search(r'(university|universität|institute|institut|department|school of|faculty)', line, re.I) and len(line) < 80:
            continue
        title = line
        break

    # ── Author: look for name patterns near the top ──
    # Strategy: find lines between title and abstract/venue that look like author names
    author_blacklist = {'vol', 'no', 'pp', 'ed', 'eds', 'jan', 'feb', 'mar', 'apr', 'may', 'jun',
                        'jul', 'aug', 'sep', 'oct', 'nov', 'dec', 'the', 'and', 'for', 'doi',
                        'abstract', 'introduction', 'econometrica', 'quarterly', 'journal',
                        'review', 'american', 'economic', 'political', 'science', 'annual',
                        'university', 'institute', 'department', 'working', 'paper', 'discussion',
                        'national', 'bureau', 'research', 'electronic', 'copy', 'available'}

    def is_valid_author(name):
        parts = name.strip().split()
        if not parts: return False
        for p in parts:
            if p.lower().rstrip('.,') in author_blacklist: return False
            if len(p) < 2: return False
        return True

    # First: try "FirstName LastName and FirstName LastName" pattern (most reliable)
    text_top = text[:3000]
    # Pattern 1: "FirstName [M.] LastName and FirstName [M.] LastName" (most reliable)
    m = re.search(r'(?:^|\n)\s*([A-Z][a-zäöüé]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-zäöüéA-ZÄÖÜß]+)\s+(?:and|und|&)\s+([A-Z][a-zäöüé]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-zäöüéA-ZÄÖÜß]+)', text_top)
    if m and is_valid_author(m.group(1)):
        author = m.group(1)
    else:
        # Search only AFTER the title line for single-author patterns
        title_pos = text_top.find(title) + len(title) if title else 0
        search_text = text_top[title_pos:]
        author_patterns = [
            # "FirstName M. LastName" or "FirstName LastName" on own line
            r'(?:^|\n)\s*((?:[A-Z][a-zäöüé]+\.?\s+)(?:[A-Z]\.?\s+)?[A-Z][a-zäöüéA-ZÄÖÜß]+)\s*[\*†\n]',
            # "LastName, FirstName" pattern
            r'(?:^|\n)\s*([A-Z][a-zäöüé]+,\s+[A-Z][a-zäöüé]+)',
        ]
        for pat in author_patterns:
            m = re.search(pat, search_text)
            if m:
                candidate = m.group(1).strip().rstrip('*†,0123456789 ')
                if is_valid_author(candidate):
                    author = candidate
                    break

    # Fallback: extract from filename if pattern like "Author_Year_Title" or "Author (Year)"
    if not author and filename:
        fn_clean = re.sub(r'\.\w+$', '', filename)  # remove extension
        # Pattern: Author_Year or Author-Year or Author (Year)
        m = re.match(r'^([A-Za-zäöüÄÖÜß]+)[\s_-]+(\d{4})', fn_clean)
        if m:
            author = m.group(1)

    # ── Year: find 4-digit year (1900-2029) ──
    year_candidates = re.findall(r'\b(19\d{2}|20[0-2]\d)\b', text[:5000])
    if year_candidates:
        # Prefer years in range 1990-2029, take the first one
        for yc in year_candidates:
            if 1990 <= int(yc) <= 2029:
                year = yc
                break
        if not year:
            year = year_candidates[0]

    # Fallback: year from filename
    if not year and filename:
        m = re.search(r'(19\d{2}|20[0-2]\d)', filename)
        if m: year = m.group(1)

    # ── Doc type detection ──
    text_lower = text[:5000].lower()
    if any(s in text_lower for s in ['ssrn', 'social science research network']):
        doc_type = "SSRN"
    elif any(s in text_lower for s in ['working paper', 'discussion paper', 'arbeitspapier']):
        doc_type = "WP"
    elif any(s in text_lower for s in ['nber', 'national bureau']):
        doc_type = "NBER"
    elif any(s in filename.lower() for s in ['[ape]', 'ape_', 'ape-']):
        doc_type = "APE"
    elif any(s in text_lower for s in ['book review', 'isbn', 'publisher', 'verlag', 'chapter ']):
        doc_type = "BOOK"

    # ── Clean up author: take last name only ──
    if author:
        parts = author.replace(',', ' ').split()
        # If "FirstName LastName" → take LastName
        if len(parts) >= 2:
            author = parts[-1]  # last word = surname
        else:
            author = parts[0]

    return {"author": author, "year": year, "title": title, "doc_type": doc_type}


def generate_canonical_key(author: str, year: str, title: str, doc_type: str = "PAP", ext: str = "pdf") -> str:
    """Generate a canonical filename for papers/books/studies.
    Format: [TYPE]_Author_Year_Short-Title.ext
    Example: PAP_Fehr_2002_Reciprocity-Social-Norms.pdf
    
    This is the SINGLE SOURCE OF TRUTH for paper naming.
    Frontend replicates this exact algorithm."""
    import re

    # Clean author: capitalize, ascii-safe
    a = (author or "Unknown").strip()
    a = a.replace("ä", "ae").replace("ö", "oe").replace("ü", "ue")
    a = a.replace("Ä", "Ae").replace("Ö", "Oe").replace("Ü", "Ue")
    a = a.replace("ß", "ss")
    a = re.sub(r'[^A-Za-z]', '', a)
    if a: a = a[0].upper() + a[1:]
    if not a: a = "Unknown"

    # Clean year
    y = (year or "0000").strip()[:4]
    if not re.match(r'^\d{4}$', y): y = "0000"

    # Clean title: take first 5 meaningful words, hyphenate
    t = (title or "Untitled").strip()
    # Remove non-alphanumeric except spaces
    t = re.sub(r'[^\w\s-]', '', t)
    words = t.split()
    # Filter out short stop words
    stop = {'a','an','the','of','in','on','at','to','for','and','or','is','are','was','by','with','from','der','die','das','und','oder','ein','eine','von','zu','im','am','des','den','dem','für','als','auf','nach','über','bei','mit'}
    meaningful = [w for w in words if w.lower() not in stop]
    if not meaningful: meaningful = words[:5]
    short_title = "-".join(meaningful[:5])
    # Capitalize each word
    short_title = "-".join(w.capitalize() for w in short_title.split("-"))
    # Limit length
    if len(short_title) > 60: short_title = short_title[:60].rsplit("-", 1)[0]
    if not short_title: short_title = "Untitled"

    # Clean doc_type
    dtype = (doc_type or "PAP").upper()[:4]

    # Clean extension
    e = (ext or "pdf").lower().lstrip(".")

    return f"{dtype}_{a}_{y}_{short_title}.{e}"


def run_paper_pipeline(filename: str, content_bytes: bytes, text_content: str, user_email: str) -> dict:
    """Full paper pipeline: Hash check → DB → GitHub papers/ → Embedding.
    Returns {ok, doc_id, github_url, duplicate, title, canonical_key}"""
    content_hash = hashlib.sha256(content_bytes).hexdigest()
    db = get_db()
    try:
        # 1. Duplicate check
        existing = db.query(Document).filter(Document.content_hash == content_hash).first()
        if existing:
            return {
                "ok": True, "duplicate": True,
                "doc_id": existing.id, "title": existing.title,
                "github_url": existing.github_url,
                "message": f"Paper bereits in KB: \"{existing.title}\" ({existing.created_at.strftime('%d.%m.%Y')})"
            }

        # 2. Extract metadata and generate canonical filename
        ext = filename.split(".")[-1].lower() if filename else "pdf"
        meta = extract_paper_metadata(text_content, filename)
        canonical = generate_canonical_key(meta["author"], meta["year"], meta["title"], meta["doc_type"], ext)
        logger.info(f"Paper canonical key: '{filename}' → '{canonical}' (author={meta['author']}, year={meta['year']}, type={meta['doc_type']})")

        # 3. Save file locally
        file_id = str(uuid.uuid4())
        file_path = UPLOAD_DIR / f"{file_id}.{ext}"
        with open(file_path, "wb") as f:
            f.write(content_bytes)

        # 4. Push to GitHub papers/evaluated/integrated/ with canonical name
        gh_result = push_to_github(canonical, content_bytes)
        github_url = gh_result.get("url", None)
        gh_status = "indexed+github" if github_url else "indexed"

        # 5. Store in Knowledge Base (PostgreSQL) – title = canonical key
        doc = Document(
            id=file_id,
            title=canonical,
            content=text_content,
            source_type="file",
            file_type=ext,
            file_path=str(file_path),
            file_size=len(content_bytes),
            database_target="knowledge_base",
            status=gh_status,
            github_url=github_url,
            uploaded_by=user_email,
            content_hash=content_hash,
            doc_metadata={
                "original_filename": filename,
                "canonical_key": canonical,
                "extracted_author": meta["author"],
                "extracted_year": meta["year"],
                "extracted_title": meta["title"],
                "doc_type": meta["doc_type"],
                "content_length": len(text_content),
                "github": gh_result,
                "upload_source": "chat",
                "auto_detected": True
            }
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        # 5. Vector Embedding
        try:
            embed_document(db, doc.id)
        except Exception as e:
            logger.warning(f"Paper embedding failed: {e}")

        logger.info(f"Paper pipeline complete: {filename} → {canonical} → KB:{doc.id} GH:{github_url}")
        return {
            "ok": True, "duplicate": False,
            "doc_id": doc.id, "title": canonical,
            "canonical_key": canonical,
            "github_url": github_url, "gh_status": gh_status,
            "extracted_metadata": meta,
            "message": f"Paper \"{canonical}\" in KB + GitHub gespeichert"
        }

    except Exception as e:
        db.rollback()
        logger.error(f"Paper pipeline error: {e}")
        return {"ok": False, "error": str(e), "message": f"Paper-Pipeline Fehler: {e}"}
    finally:
        db.close()


@app.post("/api/chat/upload")
async def chat_upload(file: UploadFile = File(...), user=Depends(require_auth)):
    """Upload a file for chat context. PDFs are auto-detected as papers → full pipeline."""
    ext = (file.filename.split(".")[-1].lower() if file.filename else "").strip()
    content_bytes = await file.read()

    if len(content_bytes) > 20 * 1024 * 1024:
        raise HTTPException(400, "Datei zu gross (max 20 MB für Chat)")

    image_types = {"png", "jpg", "jpeg", "gif", "webp"}
    doc_types = {"pdf", "txt", "md", "csv", "json", "docx", "ics"}

    if ext in image_types:
        import base64 as b64
        media_map = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "gif": "image/gif", "webp": "image/webp"}
        return {
            "ok": True,
            "type": "image",
            "name": file.filename,
            "media_type": media_map.get(ext, "image/png"),
            "image_base64": b64.b64encode(content_bytes).decode(),
            "size": len(content_bytes)
        }
    elif ext in doc_types:
        # Save temp, extract text
        temp_path = UPLOAD_DIR / f"chat_temp_{uuid.uuid4().hex[:8]}.{ext}"
        with open(temp_path, "wb") as f:
            f.write(content_bytes)
        text = extract_text(str(temp_path), ext)
        try:
            os.remove(temp_path)
        except:
            pass

        # ── PDF Paper Detection ──
        paper_result = None
        if ext == "pdf":
            detection = detect_scientific_paper(text, file.filename)
            logger.info(f"Paper detection for '{file.filename}': score={detection['score']}, signals={detection['signals']}")

            if detection["is_paper"]:
                # Run full paper pipeline!
                paper_result = run_paper_pipeline(
                    file.filename, content_bytes, text,
                    user.get("sub", user.get("email", ""))
                )
                logger.info(f"Paper pipeline result: {paper_result}")

        # Truncate for chat context
        text_for_chat = text
        if len(text_for_chat) > 50000:
            text_for_chat = text_for_chat[:50000] + f"\n\n[... gekürzt, {len(text)} Zeichen total]"

        result = {
            "ok": True,
            "type": "document",
            "name": file.filename,
            "content_text": text_for_chat,
            "size": len(content_bytes),
            "chars": len(text)
        }

        # Add paper info if detected
        if paper_result:
            result["paper_detected"] = True
            result["paper"] = paper_result
        elif ext == "pdf":
            result["paper_detected"] = False

        return result
    else:
        raise HTTPException(400, f"Dateityp .{ext} nicht unterstützt. Erlaubt: {', '.join(sorted(image_types | doc_types))}")

def search_knowledge_base(db, query: str, database: Optional[str] = None, limit: int = 8):
    """Three-tier hybrid search: Vector (if available) + PostgreSQL Fulltext + Keyword fallback."""
    
    # === TIER 1: Vector search via Voyage AI (best, if available) ===
    vector_results = vector_search(db, query, limit=limit) if VOYAGE_API_KEY else []

    # === TIER 2: PostgreSQL full-text search (great, always available) ===
    ft_results = fulltext_search(db, query, limit=limit)

    # === TIER 3: Keyword fallback (basic, always works) ===
    words = [w.lower() for w in query.split() if len(w) > 2]
    keyword_results = []
    if words:
        docs = db.query(Document).filter(Document.content.isnot(None))
        if database: docs = docs.filter(Document.database_target == database)
        docs = docs.all()
        for doc in docs:
            if not doc.content: continue
            text_lower = doc.content.lower()
            score = sum(text_lower.count(w) for w in words)
            if doc.title:
                score += sum(10 for w in words if w in doc.title.lower())
            if doc.tags:
                tags_lower = " ".join(doc.tags).lower()
                score += sum(5 for w in words if w in tags_lower)
            if doc.source_type == "ebf_answer":
                score = int(score * 3)
            if score > 0:
                keyword_results.append((score, doc))
        keyword_results.sort(key=lambda x: -x[0])

    # === MERGE: Combine all tiers ===
    merged = {}
    # Vector gets highest weight
    for score, doc in vector_results:
        merged[doc.id] = {"doc": doc, "score": score * 2.0}
    # Fulltext second
    for score, doc in ft_results:
        if doc.id in merged:
            merged[doc.id]["score"] += score
        else:
            merged[doc.id] = {"doc": doc, "score": score}
    # Keyword third
    max_kw = max((s for s, _ in keyword_results), default=1) or 1
    for score, doc in keyword_results:
        normalized = (score / max_kw) * 50  # Scale keyword to max 50
        if doc.id in merged:
            merged[doc.id]["score"] += normalized * 0.3
        else:
            merged[doc.id] = {"doc": doc, "score": normalized * 0.3}

    combined = [(m["score"], m["doc"]) for m in merged.values()]
    combined.sort(key=lambda x: -x[0])
    
    tiers_used = []
    if vector_results: tiers_used.append(f"vector:{len(vector_results)}")
    if ft_results: tiers_used.append(f"fulltext:{len(ft_results)}")
    if keyword_results: tiers_used.append(f"keyword:{len(keyword_results)}")
    logger.info(f"Hybrid search: {' + '.join(tiers_used)} → {len(combined)} combined results")
    
    return combined[:limit]

def build_context(results, max_chars=12000):
    """Build context string from search results."""
    context_parts = []
    total = 0
    for score, doc in results:
        text = doc.content or ""
        available = max_chars - total
        if available <= 0: break
        chunk = text[:available]
        source_info = f"[Quelle: {doc.title}"
        if doc.category: source_info += f", Kategorie: {doc.category}"
        if doc.tags: source_info += f", Tags: {', '.join(doc.tags)}"
        source_info += f", Typ: {doc.source_type}]"
        context_parts.append(f"{source_info}\n{chunk}")
        total += len(chunk) + len(source_info)
    return "\n\n---\n\n".join(context_parts)

def create_github_issue(question: str, user_email: str, kb_context: str = "") -> dict:
    """Create a GitHub Issue on the context repo to trigger Claude Code.
    If kb_context is provided, includes existing KB answer as warm start."""
    import urllib.request, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    body_parts = [f"**Frage von:** {user_email}", f"@claude {question}"]
    if kb_context:
        body_parts.append(f"\n---\n\n<details><summary>📚 KONTEXT: Existierende KB-Antwort (als Ausgangspunkt nutzen und verbessern)</summary>\n\n{kb_context[:3000]}\n\n</details>")
    payload = json.dumps({
        "title": f"BEATRIX: {question[:100]}",
        "body": "\n\n".join(body_parts),
        "labels": ["beatrix-question"]
    }).encode()
    req = urllib.request.Request(
        f"https://api.github.com/repos/{GH_CONTEXT_REPO}/issues",
        data=payload, method="POST",
        headers={"Authorization": f"token {GH_TOKEN}", "Content-Type": "application/json", "Accept": "application/vnd.github.v3+json"})
    resp = json.loads(urllib.request.urlopen(req, context=ctx).read())
    return {"issue_number": resp["number"], "html_url": resp["html_url"]}

def poll_github_answer(issue_number: int) -> dict:
    """Poll a GitHub Issue for Claude Code's answer."""
    import urllib.request, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    req = urllib.request.Request(
        f"https://api.github.com/repos/{GH_CONTEXT_REPO}/issues/{issue_number}/comments",
        headers={"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3+json"})
    comments = json.loads(urllib.request.urlopen(req, context=ctx).read())
    for comment in reversed(comments):
        login = comment.get("user", {}).get("login", "")
        body = comment.get("body", "")
        if login != "claude[bot]": continue
        if "Claude finished" in body:
            parts = body.split("---", 1)
            answer = parts[1].strip() if len(parts) > 1 else body
            return {"status": "done", "answer": answer, "comment_url": comment.get("html_url")}
        if "encountered an error" in body:
            return {"status": "error", "answer": "BEATRIX konnte die Analyse nicht abschliessen. Bitte versuche es erneut."}
        if "working" in body.lower() or "Working" in body:
            return {"status": "processing", "progress": body[:300]}
    return {"status": "waiting"}

def store_ebf_answer(db, question: str, answer: str, issue_url: str = ""):
    """Store a Claude Code EBF answer back into the Knowledge Base for fast retrieval."""
    # Extract keywords from question for tags
    stop_words = {"was", "ist", "das", "die", "der", "den", "dem", "ein", "eine", "wie", "und", "oder",
                  "von", "mit", "für", "auf", "aus", "bei", "nach", "über", "unter", "sind", "wird",
                  "kann", "hat", "haben", "sein", "werden", "nicht", "auch", "aber", "als", "nur",
                  "the", "what", "how", "are", "is", "in", "of", "and", "for", "to", "a", "an"}
    words = [w.strip("?!.,;:()[]") for w in question.split()]
    tags = list(set(w for w in words if len(w) > 2 and w.lower() not in stop_words))[:10]
    tags.append("ebf-answer")
    tags.append("claude-code")

    doc = Document(
        title=f"EBF: {question[:200]}",
        content=f"FRAGE: {question}\n\nANTWORT:\n{answer}",
        source_type="ebf_answer",
        database_target="knowledge_base",
        category="EBF Framework",
        tags=tags,
        status="indexed",
        github_url=issue_url,
        doc_metadata={"type": "ebf_cached_answer", "question": question, "generated_at": datetime.utcnow().isoformat()}
    )
    db.add(doc); db.commit()
    # Embed for vector search
    embed_document(db, doc.id)
    logger.info(f"EBF answer stored + embedded in KB: {question[:60]}")
    return doc.id

def fast_path_answer(question: str, context: str, sources: list) -> str:
    """Fast path: Use Claude API with existing KB context for instant answer."""
    import urllib.request, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    system_prompt = """Du bist BEATRIX, die Strategic Intelligence Suite von FehrAdvice & Partners AG.
Du bist spezialisiert auf das Evidence-Based Framework (EBF), Behavioral Economics, das Behavioral Competence Model (BCM) und Decision Architecture.

Dir steht vorhandenes Wissen aus der BEATRIX Knowledge Base zur Verfügung. Dieses Wissen wurde zuvor durch tiefgehende Analyse des EBF-Frameworks erarbeitet.

Deine Aufgabe:
- Beantworte Fragen basierend auf dem bereitgestellten Kontext
- Antworte präzise, wissenschaftlich fundiert und praxisorientiert
- Nenne Quellen wenn du aus dem Kontext zitierst
- Wenn der Kontext nicht ausreicht, sage das ehrlich
- Antworte auf Deutsch, es sei denn die Frage ist auf Englisch

Stil: Professionell, klar, auf den Punkt. Wie ein Senior Berater bei FehrAdvice."""

    user_message = f"""Hier ist relevanter Kontext aus der BEATRIX Knowledge Base (vorberechnetes EBF-Wissen):

{context}

---

Frage: {question}"""

    payload = json.dumps({
        "model": ANTHROPIC_MODEL_STANDARD,
        "max_tokens": 3000,
        "system": system_prompt,
        "messages": [{"role": "user", "content": user_message}]
    }).encode()
    req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
        headers={
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
            "Content-Type": "application/json",
            "User-Agent": "BEATRIXLab/3.8"
        })
    resp = json.loads(urllib.request.urlopen(req, context=ctx).read())
    return resp["content"][0]["text"]

# Relevance threshold: minimum score to use fast path
FAST_PATH_THRESHOLD = 15

# Freshness: EBF answers older than this are considered stale → force deep refresh
KB_FRESHNESS_DAYS = 90

# ── BEATRIX INTENT ROUTER: Universal natural language interface ──

CUSTOMERS_LIST = """ubs, erste-bank, lukb, alpla, a1-telekom, raiffeisen, zkb, gkb, valiant, julius-baer,
vontobel, migros-bank, postfinance, css, helsana, swica, sanitas, kpt, concordia,
groupe-mutuel, assura, atupri, visana, bmw, orf, srg, ringier-medien-schweiz,
lindt-copacking, porr, neon, revolut, peek-cloppenburg, philoro, sob, spo, bfe,
economiesuisse, prio-swiss, bekb, localsearch, zindel-united, plusminus, awe-sg"""

# ═══════════════════════════════════════════════════════════
# RICH CUSTOMER CONTEXT (loaded from GitHub at startup)
# ═══════════════════════════════════════════════════════════
_CUSTOMER_CONTEXT_CACHE = {"loaded": False, "context": "", "customers": {}, "contacts": {}, "leads_summary": ""}

def load_customer_context_from_github():
    """Load customer-registry, contacts, and lead-database from GitHub for rich context."""
    global _CUSTOMER_CONTEXT_CACHE
    if _CUSTOMER_CONTEXT_CACHE["loaded"]:
        return _CUSTOMER_CONTEXT_CACHE

    import urllib.request, ssl, base64, yaml
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    gh_headers = {"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"} if GH_TOKEN else {}

    def fetch_yaml(path):
        try:
            req = urllib.request.Request(
                f"https://api.github.com/repos/{GH_REPO}/contents/{path}",
                headers=gh_headers)
            resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
            return yaml.safe_load(base64.b64decode(resp["content"]).decode("utf-8"))
        except Exception as e:
            logger.warning(f"Failed to load {path}: {e}")
            return None

    try:
        # 1. Customer Registry
        reg = fetch_yaml("data/customer-registry.yaml")
        customers = {}
        if reg and reg.get("customers"):
            for c in reg["customers"]:
                code = c.get("code", "").lower()
                customers[code] = {
                    "name": c.get("name", ""),
                    "short": c.get("short_name", ""),
                    "industry": c.get("industry", ""),
                    "country": c.get("country", ""),
                    "type": c.get("type", ""),
                    "status": c.get("status", "")
                }

        # 2. Customer Contacts
        contacts_data = fetch_yaml("data/customer-contacts.yaml")
        contacts = {}
        if contacts_data and contacts_data.get("customers"):
            for cc in contacts_data["customers"]:
                cid = cc.get("customer_id", "").lower()
                contacts[cid] = {
                    "fa_owner": cc.get("fa_owner", ""),
                    "status": cc.get("relationship_status", ""),
                    "contacts": []
                }
                for con in cc.get("contacts", []):
                    contacts[cid]["contacts"].append({
                        "name": con.get("name", ""),
                        "role": con.get("role", ""),
                        "primary": con.get("is_primary", False),
                        "email": con.get("email", "")
                    })

        # 3. Lead Database (summary only - it's 160K)
        leads = fetch_yaml("data/sales/lead-database.yaml")
        lead_summaries = {}
        if leads and leads.get("leads"):
            for lead in leads.get("leads", []):
                # company can be a dict with name/short_name or a string
                company = lead.get("company", {})
                if isinstance(company, dict):
                    cid = (company.get("short_name") or company.get("name", "")).lower().replace(" ", "-")
                elif isinstance(company, str):
                    cid = company.lower().replace(" ", "-")
                else:
                    cid = str(lead.get("id", "")).lower()
                if not cid:
                    continue
                if cid not in lead_summaries:
                    lead_summaries[cid] = []
                lead_summaries[cid].append({
                    "id": lead.get("id", ""),
                    "opportunity": lead.get("opportunity") or lead.get("notes", "")[:60] if isinstance(lead.get("notes",""), str) else "",
                    "stage": lead.get("stage", ""),
                    "value": lead.get("value") or lead.get("estimated_value_chf", 0),
                    "owner": lead.get("owner") or lead.get("fa_owner", "")
                })

        # 4. Projects per customer (from data/projects/)
        project_history = {}
        try:
            proj_list_req = urllib.request.Request(
                f"https://api.github.com/repos/{GH_REPO}/contents/data/projects",
                headers=gh_headers)
            proj_dirs = json.loads(urllib.request.urlopen(proj_list_req, context=ctx, timeout=15).read())
            for pdir in proj_dirs:
                if pdir.get("type") != "dir":
                    continue
                try:
                    purl = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{pdir['name']}/project.yaml"
                    preq = urllib.request.Request(purl, headers={**gh_headers, "Accept": "application/vnd.github.v3.raw"})
                    pdata = yaml.safe_load(urllib.request.urlopen(preq, context=ctx, timeout=10).read().decode()) or {}
                    pmeta = pdata.get("metadata", {})
                    pclient = pdata.get("client", {})
                    pproj = pdata.get("project", {})
                    ptimeline = pdata.get("timeline", {})
                    pteam = pdata.get("team", {})
                    ccode = (pclient.get("customer_code") or "").lower()
                    if not ccode:
                        continue
                    if ccode not in project_history:
                        project_history[ccode] = []
                    project_history[ccode].append({
                        "code": pmeta.get("project_code", ""),
                        "slug": pmeta.get("slug", pdir["name"]),
                        "name": pproj.get("name", ""),
                        "type": pproj.get("type", ""),
                        "description": (pproj.get("description") or "")[:100],
                        "status": pmeta.get("status", ""),
                        "category": pmeta.get("project_category", ""),
                        "start": ptimeline.get("start_date", ""),
                        "end": ptimeline.get("end_date", ""),
                        "budget": ptimeline.get("budget_chf", 0),
                        "owner": pteam.get("fa_owner", ""),
                        "created": pmeta.get("created", "")
                    })
                except Exception:
                    pass
        except Exception as e:
            logger.warning(f"Failed to load projects: {e}")

        # 5. Project sequences (for next code prediction)
        project_sequences = {}
        try:
            seq_data = fetch_yaml("data/config/project-sequences.yaml")
            if seq_data:
                project_sequences = seq_data.get("sequences", {})
        except Exception:
            pass

        # 6. Build context string
        lines = ["BEKANNTE KUNDEN MIT DETAILS:"]
        all_codes = set(list(customers.keys()) + list(contacts.keys()))
        for code in sorted(all_codes):
            c = customers.get(code, {})
            cn = contacts.get(code, {})
            name = c.get("name") or c.get("short", code.upper())
            parts = [f"- {code.upper()}: {name}"]
            if c.get("industry"): parts.append(f"Branche={c['industry']}")
            if c.get("country"): parts.append(f"Land={c['country']}")
            if c.get("status"): parts.append(f"Status={c['status']}")
            if cn.get("fa_owner"): parts.append(f"Owner={cn['fa_owner']}")

            # Contacts
            if cn.get("contacts"):
                contact_strs = []
                for con in cn["contacts"]:
                    cs = con["name"]
                    if con.get("role"): cs += f" ({con['role']})"
                    if con.get("primary"): cs += " [PRIMARY]"
                    contact_strs.append(cs)
                parts.append(f"Kontakte: {', '.join(contact_strs)}")

            # Existing leads
            if code in lead_summaries:
                ls = lead_summaries[code]
                lead_strs = [f"{l['id']}:{l.get('stage','?')}" for l in ls[:3]]
                parts.append(f"Leads: {', '.join(lead_strs)}")

            # Existing projects
            if code in project_history:
                projs = sorted(project_history[code], key=lambda x: x.get("created",""), reverse=True)
                proj_strs = []
                for p in projs[:5]:
                    ps = f"{p['code']}:{p['name'][:30]}"
                    if p.get("status"): ps += f"({p['status']})"
                    if p.get("start"): ps += f"[{p['start']}]"
                    proj_strs.append(ps)
                parts.append(f"Projekte: {', '.join(proj_strs)}")

            lines.append(" | ".join(parts))

        # Add codes without details
        extra_codes = [c.strip() for c in CUSTOMERS_LIST.split(",") if c.strip().lower() not in all_codes]
        if extra_codes:
            lines.append(f"\nWeitere Kunden (ohne Details): {', '.join(extra_codes)}")

        context_str = "\n".join(lines)

        _CUSTOMER_CONTEXT_CACHE = {
            "loaded": True,
            "context": context_str,
            "customers": customers,
            "contacts": contacts,
            "leads_summary": lead_summaries,
            "projects": project_history,
            "sequences": project_sequences
        }
        logger.info(f"Customer context loaded: {len(customers)} customers, {len(contacts)} with contacts, {sum(len(v) for v in lead_summaries.values())} leads, {sum(len(v) for v in project_history.values())} projects")
        return _CUSTOMER_CONTEXT_CACHE

    except Exception as e:
        logger.error(f"Customer context loading failed: {e}")
        _CUSTOMER_CONTEXT_CACHE["loaded"] = True  # Don't retry on every request
        _CUSTOMER_CONTEXT_CACHE["context"] = f"Kundencodes: {CUSTOMERS_LIST}"
        return _CUSTOMER_CONTEXT_CACHE

def get_customer_context():
    """Get cached customer context string for prompts."""
    cache = load_customer_context_from_github()
    return cache["context"]

def get_customer_detail(customer_code):
    """Get detailed info for a specific customer (for enriching prompts)."""
    cache = load_customer_context_from_github()
    code = customer_code.lower()
    c = cache["customers"].get(code, {})
    cn = cache["contacts"].get(code, {})
    leads = cache.get("leads_summary", {}).get(code, [])
    projects = cache.get("projects", {}).get(code, [])
    sequences = cache.get("sequences", {})

    if not c and not cn and not projects:
        return None

    detail = f"Kunde: {c.get('name', code.upper())}"
    if c.get("industry"): detail += f"\nBranche: {c['industry']}"
    if c.get("country"):
        country = c['country']
        detail += f"\nLand: {country}"
        currency = COUNTRY_CURRENCY.get(country, "CHF")
        detail += f"\nWährung: {currency}"
        if currency != "CHF":
            fx = convert_to_chf(1.0, currency)
            detail += f" (1 {currency} = {fx['fx_rate']} CHF, Stand: {fx['fx_date']})"
    if cn.get("fa_owner"): detail += f"\nFehrAdvice Owner: {cn['fa_owner']}"
    if cn.get("contacts"):
        detail += "\nKontakte:"
        for con in cn["contacts"]:
            detail += f"\n  - {con['name']}"
            if con.get("role"): detail += f" ({con['role']})"
            if con.get("primary"): detail += " [Hauptkontakt]"
    if leads:
        detail += f"\nBestehende Leads ({len(leads)}):"
        for l in leads[:5]:
            opp = l.get('opportunity', '?')
            if isinstance(opp, str):
                opp = opp[:50]
            detail += f"\n  - {l['id']}: {opp} (Stage: {l.get('stage', '?')}, Wert: {l.get('value', 0)})"

    # Project history
    if projects:
        sorted_projs = sorted(projects, key=lambda x: x.get("created",""), reverse=True)
        detail += f"\n\nBESTEHENDE PROJEKTE ({len(projects)}):"
        for p in sorted_projs:
            detail += f"\n  - {p['code']}: {p['name']}"
            if p.get("type"): detail += f" ({p['type']})"
            if p.get("status"): detail += f" [{p['status']}]"
            if p.get("start"): detail += f" Start: {p['start']}"
            if p.get("end"): detail += f" Ende: {p['end']}"
            if p.get("budget"): detail += f" Budget: {p['budget']} CHF"
            if p.get("description"): detail += f"\n    Beschreibung: {p['description']}"
        last = sorted_projs[0]
        detail += f"\n  LETZTES PROJEKT: {last['code']} - {last['name']}"

    # Next project code
    prefix = code[:3].upper()
    next_seq = sequences.get(prefix, 1)
    # Also check from project codes
    for p in projects:
        pc = p.get("code", "")
        if pc.upper().startswith(prefix):
            try:
                num = int(pc[len(prefix):])
                if num >= next_seq:
                    next_seq = num + 1
            except ValueError:
                pass
    detail += f"\nNÄCHSTES PROJEKTKÜRZEL: {prefix}{next_seq:03d}"

    return detail


# ═══════════════════════════════════════════════════════════
# CURRENCY ENGINE – Country-based detection, live FX rates
# ═══════════════════════════════════════════════════════════

REPORTING_CURRENCY = "CHF"  # FehrAdvice default

COUNTRY_CURRENCY = {
    "AT":"EUR","DE":"EUR","FR":"EUR","IT":"EUR","ES":"EUR","NL":"EUR","BE":"EUR",
    "LU":"EUR","PT":"EUR","IE":"EUR","FI":"EUR","GR":"EUR","SK":"EUR","SI":"EUR",
    "EE":"EUR","LV":"EUR","LT":"EUR","CY":"EUR","MT":"EUR","HR":"EUR",
    "CH":"CHF","LI":"CHF",
    "GB":"GBP","US":"USD","CA":"CAD","AU":"AUD","NZ":"NZD",
    "SE":"SEK","NO":"NOK","DK":"DKK","PL":"PLN","CZ":"CZK","HU":"HUF",
    "RO":"RON","BG":"BGN","JP":"JPY","CN":"CNY","SG":"SGD","HK":"HKD",
    "AE":"AED","SA":"SAR","IN":"INR","BR":"BRL","MX":"MXN","ZA":"ZAR",
    "TR":"TRY","IL":"ILS","KR":"KRW","TW":"TWD","TH":"THB",
}

CURRENCY_SYMBOLS = {
    "CHF":"CHF","EUR":"\u20ac","USD":"$","GBP":"\u00a3","SEK":"kr","NOK":"kr","DKK":"kr",
    "PLN":"z\u0142","CZK":"K\u010d","HUF":"Ft","CAD":"CA$","AUD":"A$","JPY":"\u00a5","CNY":"\u00a5",
    "SGD":"S$","HKD":"HK$","AED":"AED","INR":"\u20b9","BRL":"R$",
    "NZD":"NZ$","MXN":"MX$","ZAR":"R","TRY":"\u20ba","ILS":"\u20aa","KRW":"\u20a9",
}

CURRENCY_DECIMALS = {"JPY": 0, "HUF": 0, "KRW": 0}  # Most are 2

# Cache: {"rates": {}, "base": "CHF", "date": "...", "fetched_at": datetime}
_FX_CACHE = {"rates": {}, "date": None, "fetched_at": None}

def fetch_exchange_rates():
    """Fetch daily exchange rates from open.er-api.com (free, ECB-sourced). Cached for 6h."""
    global _FX_CACHE
    from datetime import timedelta
    import urllib.request as _urlreq, ssl as _ssl

    # Return cache if fresh (< 6 hours)
    if _FX_CACHE["fetched_at"] and (datetime.utcnow() - _FX_CACHE["fetched_at"]) < timedelta(hours=6):
        return _FX_CACHE

    try:
        _ctx = _ssl.create_default_context(); _ctx.check_hostname = False; _ctx.verify_mode = _ssl.CERT_NONE
        req = _urlreq.Request("https://open.er-api.com/v6/latest/CHF",
            headers={"User-Agent": "BEATRIXLab/3.21"})
        resp = json.loads(_urlreq.urlopen(req, context=_ctx, timeout=10).read())

        if resp.get("result") == "success":
            _FX_CACHE = {
                "rates": resp.get("rates", {}),
                "date": resp.get("time_last_update_utc", "")[:16],
                "fetched_at": datetime.utcnow()
            }
            # Add inverse rates (X -> CHF) for convenience
            _FX_CACHE["to_chf"] = {}
            for cur, rate in _FX_CACHE["rates"].items():
                if rate > 0:
                    _FX_CACHE["to_chf"][cur] = round(1.0 / rate, 6)
            _FX_CACHE["to_chf"]["CHF"] = 1.0
            logger.info(f"FX rates updated: {resp.get('time_last_update_utc', '')[:16]} | {len(_FX_CACHE['rates'])} currencies")
            return _FX_CACHE
    except Exception as e:
        logger.warning(f"FX rate fetch failed: {e}")

    # Fallback: hardcoded rates if API fails
    if not _FX_CACHE["rates"]:
        _FX_CACHE = {
            "rates": {"CHF": 1.0, "EUR": 1.095, "USD": 1.303, "GBP": 0.954},
            "to_chf": {"CHF": 1.0, "EUR": 0.913, "USD": 0.768, "GBP": 1.048},
            "date": "fallback",
            "fetched_at": datetime.utcnow()
        }
    return _FX_CACHE


def get_customer_currency(customer_code: str) -> str:
    """Detect currency from customer country. Falls back to CHF."""
    cache = load_customer_context_from_github()
    c = cache["customers"].get(customer_code.lower(), {})
    country = c.get("country", "CH")
    return COUNTRY_CURRENCY.get(country, "CHF")


def convert_to_chf(amount: float, from_currency: str) -> dict:
    """Convert amount to CHF using daily rates. Returns conversion details."""
    if not amount or from_currency == "CHF":
        return {
            "amount_original": amount, "currency": "CHF",
            "amount_chf": amount, "fx_rate": 1.0,
            "fx_date": datetime.utcnow().strftime("%Y-%m-%d")
        }

    fx = fetch_exchange_rates()
    to_chf_rate = fx.get("to_chf", {}).get(from_currency.upper())

    if not to_chf_rate:
        logger.warning(f"No FX rate for {from_currency}, treating as CHF")
        return {
            "amount_original": amount, "currency": from_currency,
            "amount_chf": amount, "fx_rate": 1.0, "fx_date": "unknown"
        }

    amount_chf = round(amount * to_chf_rate, 2)
    return {
        "amount_original": amount, "currency": from_currency,
        "amount_chf": amount_chf, "fx_rate": round(to_chf_rate, 6),
        "fx_date": fx.get("date", "")
    }


def convert_amount(amount: float, from_currency: str, to_currency: str) -> dict:
    """Convert between any two currencies via CHF cross-rate."""
    if from_currency == to_currency:
        return {"amount": amount, "currency": to_currency, "fx_rate": 1.0}

    fx = fetch_exchange_rates()
    to_chf = fx.get("to_chf", {}).get(from_currency.upper(), 1.0)
    from_chf = fx.get("rates", {}).get(to_currency.upper(), 1.0)

    rate = to_chf * from_chf
    decimals = CURRENCY_DECIMALS.get(to_currency, 2)
    return {
        "amount": round(amount * rate, decimals),
        "currency": to_currency, "fx_rate": round(rate, 6),
        "fx_date": fx.get("date", "")
    }


def format_currency(amount, currency: str = "CHF") -> str:
    """Format amount with currency symbol. E.g. CHF 40'000 or EUR 36'520"""
    if not amount: return ""
    amount = float(amount)
    decimals = CURRENCY_DECIMALS.get(currency, 2)
    symbol = CURRENCY_SYMBOLS.get(currency, currency)
    if decimals == 0:
        formatted = f"{int(amount):,}".replace(",", "'")
    else:
        formatted = f"{amount:,.{decimals}f}".replace(",", "'")
    return f"{symbol} {formatted}"


INTENT_ROUTER_SYSTEM = """Du bist der BEATRIX Intent-Router. Analysiere die User-Nachricht und bestimme den Intent.

VERFÜGBARE INTENTS (Session-Typen):
- "lead" (Pipeline): Neuer Kontakt, Anfrage, Angebot in Arbeit, Akquise, Neugeschäft, Pitch, Offerte, Erstgespräch
- "project" (Mandat): Kunde hat beauftragt, Projekt läuft, Deliverables, Timeline, Workshop, Lieferung, Auftrag
- "company": Neue Firma/Unternehmen anlegen, Firmen-Stammdaten erfassen (INTERN - nicht als Session-Typ)
- "model" (Modell): BCM bauen, Axiome definieren, Intervention designen, Nudge, Wirkung, Behavior Change
- "context" (Kontext Ψ): Situation analysieren, Zielgruppe verstehen, 8 Ψ-Dimensionen, Stakeholder, Umfeld, Kultur
- "knowledge" (Literatur): Paper suchen, Studie zitieren, Theorie, Forschung, Evidenz, Quelle, Referenz (→ nutzt KB)
- "general" (Fragen): Alles andere, schnelle Antwort, Erklärung, Hilfe, Smalltalk

WICHTIGE UNTERSCHEIDUNGEN:
- Pipeline vs Mandat: "Noch kein Auftrag" → lead, "Auftrag erteilt, Arbeit läuft" → project
- Modell vs Kontext Ψ: "BCM/Intervention bauen" → model, "Situation verstehen" → context
- Lead vs Company: "Neue Firma anlegen" → company, "Neuer Lead mit Budget" → lead

UNTERSCHEIDUNG LEAD vs COMPANY:
- "Neue Firma Hofer Reisen" → company (Stammdaten anlegen)
- "Neuer Lead Hofer Reisen, Budget 50k" → lead (Opportunity mit Wert/Pipeline)
- "Lead für UBS" → lead (UBS existiert schon)
- "Lege Kaufland als Firma an" → company

Antworte NUR mit einem JSON-Objekt (keine Markdown-Backticks nötig):
{"intent": "...", "confidence": 0.0-1.0, "entities": {"customer": "...", "project": "..."}, "session_type": "..."}

session_type ist der Frontend-Key: lead, project, model, context, research (für knowledge), general
Bei "company" setze session_type auf "lead" (ist Teil der Pipeline).

Erkenne Kunden auch bei ungefährer Nennung ("UBS" → "ubs", "Luzerner KB" → "lukb", "Erste" → "erste-bank").
Extrahiere wenn möglich: customer (code), project (name/slug), und andere Schlüssel-Entitäten."""

# Intent → Session-Type Mapping (Backend-Intent → Frontend-SessionType)
INTENT_TO_SESSION_TYPE = {
    "lead": "lead",           # Pipeline
    "project": "project",     # Mandat
    "company": "lead",        # Teil der Pipeline
    "task": "project",        # Tasks gehören zu Mandaten
    "model": "model",         # Modell
    "context": "context",     # Kontext Ψ
    "knowledge": "research",  # Literatur (KB/RAG)
    "general": "general",     # Fragen
}

def get_session_type(intent: str) -> str:
    """Map backend intent to frontend session_type."""
    return INTENT_TO_SESSION_TYPE.get(intent, "general")

def get_domain_prompts():
    """Build domain prompts with rich customer context from GitHub."""
    ctx = get_customer_context()
    return {
    "project": f"""Du bist BEATRIX, die Projekt-Assistentin von FehrAdvice & Partners AG.
Der User moechte ein Projekt eroeffnen oder ein bestehendes Projekt ergaenzen.

Deine Aufgabe: Extrahiere strukturierte Projektdaten aus dem Gespraech.
NUTZE die bekannten Kundendaten (Kontakte, Owner, Branche) um Felder AUTOMATISCH zu befuellen!
Frage NUR nach was du nicht aus den bekannten Daten ableiten kannst.

FELDER:
- customer_code: Kundencode (z.B. "ubs", "lukb", "erste-bank")
- name: Projektname (kurz, praegnant)
- type: beratung | workshop | studie | intervention | keynote | training | retainer
- project_category: mandat (bezahlt) | lead (Akquise) | probono (unentgeltlich)
- description: Kurzbeschreibung / Ziele
- start_date: YYYY-MM-DD
- end_date: YYYY-MM-DD
- budget: Zahl in Kundenwaehrung (nur bei mandat)
- currency: Waehrungscode (CHF, EUR, USD, GBP, ...) - automatisch aus Kundenland
- budget_chf: Umgerechnet in CHF (FehrAdvice Reporting-Waehrung)
- billing_type: fixed | time_material | retainer
- fa_owner: Kuerzel (GF=Gerhard Fehr, AF=Andrea Fehr)
- fa_team: Liste von Team-Mitgliedern

WAEHRUNG: Die Waehrung richtet sich nach dem Land des Kunden:
- CH/LI = CHF, AT/DE/EU = EUR, GB = GBP, US = USD
- Wenn der User eine Zahl OHNE Waehrung nennt, nimm die Kundenwaehrung
- Bei Nicht-CHF-Kunden: Zeige Betrag in Kundenwaehrung UND CHF-Aequivalent

{{ctx}}

REGELN:
1. Antworte mit JSON in ```json ... ``` Tags
2. JSON-Struktur: {{"intent":"project", "action":"create"|"update", "status":"complete"|"need_info", "data":{{...felder...}}, "missing":[...], "message":"...", "confidence":0.0-1.0}}
3. Pflichtfelder: customer_code, name, project_category
4. Sinnvolle Defaults: type="beratung", billing="fixed", fa_owner aus Kundendaten, start_date=heute
5. AUTOMATISCH befuellen: fa_owner aus Kundendaten, Branche, Land, naechstes Projektkuerzel
6. Deutsch, kurz, professionell
7. Wenn der Kunde bekannt ist: zeige was du schon weisst und frage NUR nach dem Projektspezifischen
8. WICHTIGSTE REGEL - DUPLIKAT-CHECK: Wenn der Kunde bestehende Projekte hat, frage ZUERST ob das gewuenschte Projekt nicht bereits existiert! Z.B.: "Fuer Zindel gibt es bereits ZIN004 Organisationsentwicklung - meinst du dieses Projekt, oder brauchst du wirklich ein NEUES Projekt?" Erst wenn der User bestaetigt dass es ein neues Projekt sein soll, weiter mit Eroeffnung.
9. Bei bestehendem Projekt: Frage ob der User das bestehende Projekt oeffnen/bearbeiten moechte statt ein neues zu eroeffnen
10. Bei echtem neuem Projekt: Frage ob es mit einem bestehenden zusammenhaengt oder ein Folgeprojekt ist
11. PROJEKTKUERZEL: Schlage das naechste Kuerzel vor (z.B. wenn letztes ZIN004 war → ZIN005). Das Kuerzel steht in den Kundendaten unter NAECHSTES PROJEKTKUERZEL
12. Bei Folgeprojekten: Uebernimm relevante Daten (Typ, Team, Billing) vom Vorgaengerprojekt als Default""",

    "lead": f"""Du bist BEATRIX, die Sales-Assistentin von FehrAdvice & Partners AG.
Der User moechte einen Lead/eine Opportunity erfassen oder die Sales-Pipeline bearbeiten.

WICHTIG: Du KENNST bereits alle Kundendaten! Wenn der User einen bekannten Kunden nennt,
befuelle AUTOMATISCH: customer_code, customer_name, fa_owner, bestehende Kontakte.
Frage NUR nach was du nicht weisst (Opportunity, Wert, naechste Schritte).

WAEHRUNG: Die Waehrung richtet sich nach dem Land des Kunden:
- Schweiz (CH) = CHF, Oesterreich/Deutschland/EU (AT, DE, ...) = EUR, UK = GBP, USA = USD
- FehrAdvice Reporting-Waehrung ist immer CHF
- Wenn der Kunde NICHT aus der Schweiz ist, speichere den Betrag in Kundenwaehrung UND rechne in CHF um
- Wenn der User eine Zahl OHNE Waehrung nennt, nimm die Kundenwaehrung (oder CHF als Default)
- estimated_value: Betrag in Kundenwaehrung
- currency: ISO-Code (CHF, EUR, USD, GBP, ...)
- estimated_value_chf: Umgerechnet in CHF (bei CHF-Kunden = gleicher Wert)

FELDER:
- lead_code: AUTOMATISCH GENERIERT nach Schema L-{{KUNDE}}-{{B/N}}-{{JJ}}-{{NNN}} (z.B. L-UBS-B-26-003)
- customer_code: Kundencode
- customer_name: Voller Firmenname
- is_new_customer: true wenn unbekannter Kunde, false wenn in Datenbank
- contact_person: Ansprechpartner (nutze bekannte Kontakte als Default!)
- contact_email: E-Mail
- contact_role: Position/Rolle
- opportunity: Kurzbeschreibung der Opportunity
- type: beratung | workshop | studie | intervention | keynote | training | retainer
- estimated_value: Geschaetzter Wert in Kundenwaehrung
- currency: Waehrungscode (CHF, EUR, USD, GBP, ...)
- estimated_value_chf: Geschaetzter Wert umgerechnet in CHF
- probability: Gewinnwahrscheinlichkeit (0-100%)
- stage: initial_contact | qualification | proposal | negotiation | won | lost
- next_action: Naechster Schritt
- next_action_date: YYYY-MM-DD
- fa_owner: Kuerzel (aus Kundendaten uebernehmen!)
- source: referral | event | inbound | outbound | existing_client
- notes: Zusaetzliche Infos

{{ctx}}

REGELN:
1. Antworte mit JSON in ```json ... ``` Tags
2. JSON: {{"intent":"lead", "action":"create"|"update", "status":"complete"|"need_info", "data":{{...}}, "missing":[...], "message":"...", "confidence":0.0-1.0}}
3. Pflichtfelder: customer_name, opportunity, stage
4. AUTOMATISCH aus Datenbank: customer_code, fa_owner, is_new_customer=false, contact_person (Primary), Branche, Waehrung
5. LEAD-CODE AUTOMATISCH GENERIEREN nach Schema: L-{{KUNDE}}-{{B/N}}-{{JJ}}-{{NNN}}
   - L = Lead (immer)
   - KUNDE = Kundenkuerzel aus prefix_mapping (z.B. UBS, HEL, GEB, SBB)
   - B = Bestandskunde, N = Neukunde (aus is_new_customer ableiten)
   - JJ = aktuelles Jahr zweistellig (z.B. 26)
   - NNN = fortlaufende Nummer (starte mit 001)
   Beispiele: L-UBS-B-26-003, L-HEL-B-26-001, L-GEB-N-26-001
   GENERIERE den lead_code IMMER AUTOMATISCH! Frage NICHT danach!
6. Bei Bestandskunden: Zeige bekannte Kontakte und frage "Ist [Name] der richtige Ansprechpartner?"
7. Bei bestehenden Leads: Erwaehne sie kurz ("Es gibt bereits X Leads fuer diesen Kunden")
8. Frage gezielt nach: Opportunity, geschaetzter Wert, naechster Schritt
9. Bei Bestandskunden mit Projekten: Erwaehne bestehende Projekte und frage ob der Lead damit zusammenhaengt
10. Bei Nicht-CHF-Kunden: Zeige Betrag in Kundenwaehrung UND CHF-Aequivalent
11. Deutsch, professionell, effizient""",

    "company": f"""Du bist BEATRIX, die CRM-Assistentin von FehrAdvice & Partners AG.
Der User moechte eine neue Firma/ein neues Unternehmen als Kunden anlegen.

FELDER:
- name: Voller Firmenname (z.B. "Swisscom AG") - PFLICHT
- code: Kuerzel 3-5 Buchstaben (z.B. "SWC") - wird auto-generiert wenn nicht angegeben
- short_name: Kurzname (z.B. "Swisscom")
- industry: Branche (finance, insurance, retail, fmcg, construction, packaging, public_sector, telecom, media, automotive, pharma_health, manufacturing, technology, energy, transport, creative, ngo, consulting)
- country: Laendercode (CH, AT, DE, SE, UK, US) - Default CH
- website: URL
- fa_owner: Kuerzel (GF=Gerhard Fehr, AF=Andrea Fehr)
- notes: Zusaetzliche Infos

{ctx}

REGELN:
1. Antworte mit JSON in ```json ... ``` Tags
2. JSON: {{"intent":"company", "action":"create", "status":"complete"|"need_info", "data":{{...}}, "missing":[...], "message":"...", "confidence":0.0-1.0}}
3. Pflichtfeld: name
4. PRUEFE ob die Firma schon existiert! Wenn ja, sage dem User: "{{Name}} existiert bereits als {{code}}. Moechtest du die Firma oeffnen oder Daten ergaenzen?"
5. Code auto-generieren: Erste 3-5 Buchstaben des Firmennamens, nur Grossbuchstaben
6. AUTOMATISCH ableiten: Branche aus Firmenname wenn moeglich (z.B. "Kantonalbank" → finance, "Versicherung" → insurance)
7. Land: CH als Default, AT fuer oesterreichische Firmen, DE fuer deutsche
8. Deutsch, professionell, effizient
9. Nach Anlage: Frage ob gleich ein Lead fuer die neue Firma erstellt werden soll""",

    "task": f"""Du bist BEATRIX, die Task-Managerin von FehrAdvice & Partners AG.
Der User moechte eine Aufgabe erstellen, ein Todo setzen, einen Follow-up planen oder eine Erinnerung setzen.

ES GIBT 4 TYPEN VON AUFGABEN (assignee_type):
- "consultant": Berater-Aufgabe (inhaltliche Arbeit: Proposal schreiben, Workshop vorbereiten, Analyse)
- "bdm": BDM-Team/Administration (operative Arbeit: Vertrag, Rechnung, Reise, Termin, Dokumente)
- "beatrix": BEATRIX-Automation (System: Report generieren, Daten aufbereiten, Reminder senden)
- "external": Externe Partei (Kunde liefert Daten, Partner schickt Vertrag, wartet auf Freigabe)

FELDER:
- title: Aufgabentitel (kurz, klar, aktionsorientiert)
- description: Detailbeschreibung (optional)
- assignee_type: consultant | bdm | beatrix | external
- assignee: Kuerzel (GF, AF, ...) | "BDM" | "system" | Name der externen Person
- customer_code: Kundencode (optional)
- project_slug: Projektbezug (optional)
- lead_id: Lead-Bezug (optional)
- priority: low | normal | high | urgent
- due_date: YYYY-MM-DD
- category: sales | delivery | admin | vertrag | rechnung | reise | termin | dokument | zugang | data | allgemein
- escalation_after_days: Bei external – nach wie vielen Tagen eskalieren (optional)

{{ctx}}

REGELN:
1. Antworte mit JSON in ```json ... ``` Tags
2. JSON: {{"intent":"task", "action":"create", "status":"complete"|"need_info", "data":{{...}}, "missing":[...], "message":"...", "confidence":0.0-1.0}}
3. Pflichtfelder: title, assignee_type
4. AUTOMATISCH ableiten:
   - "Ruf X an" / "Proposal schreiben" / "Workshop vorbereiten" → consultant
   - "Vertrag aufsetzen" / "Rechnung stellen" / "Reise buchen" / "Termin koordinieren" → bdm
   - "Erstelle Report" / "Aktualisiere Daten" → beatrix
   - "Kunde soll liefern" / "Warten auf Freigabe" → external
5. Erkenne Kunden und Projekte automatisch aus dem Kontext
6. Setze sinnvolle Defaults: priority=normal, assignee aus Kundendaten
7. Bei "erinnere mich morgen" → due_date = morgen, assignee = aktueller User
8. Bei externen Tasks: Frage nach escalation_after_days
9. Deutsch, kurz, professionell""",

    "model": f"""Du bist BEATRIX, die Modell-Spezialistin von FehrAdvice & Partners AG.
Der User moechte ein Behavioral-Modell erstellen oder bearbeiten, basierend auf dem Evidence-Based Framework (EBF).

DU KANNST:
- BCM (Behavioral Competence Model) Analysen erstellen
- Psi-Dimensionen (8 psychologische Kontextdimensionen) definieren
- Kontextvektoren berechnen/beschreiben
- Behavioral Objectives formulieren
- Interventionsdesign vorschlagen
- EBF-Axiome anwenden

FELDER:
- customer_code: Fuer welchen Kunden
- project_slug: Fuer welches Projekt (optional)
- model_type: bcm | psi_profile | context_vector | intervention | behavioral_objectives
- target_behavior: Welches Verhalten soll veraendert werden?
- target_group: Zielgruppe
- psi_dimensions: Liste der relevanten Psi-Dimensionen mit Einschaetzung
  (autonomy, certainty, fairness, belonging, status, trust, competence, meaning)
- context_factors: Relevante Kontextfaktoren
- current_state: Ausgangslage / Ist-Zustand
- desired_state: Ziel-Zustand
- hypotheses: Verhaltenshypothesen
- interventions: Vorgeschlagene Interventionen
- evidence_base: Wissenschaftliche Grundlage

{{ctx}}

REGELN:
1. Antworte mit JSON in ```json ... ``` Tags
2. JSON: {{"intent":"model", "action":"create"|"analyze", "status":"complete"|"need_info", "data":{{...}}, "missing":[...], "message":"...", "confidence":0.0-1.0}}
3. Pflichtfelder: model_type, target_behavior ODER target_group
4. AUTOMATISCH aus Datenbank: Kundendetails, Branche, bekannte Kontakte
5. Nutze echtes EBF/BCM-Wissen - keine generischen Antworten
6. Deutsch, wissenschaftlich fundiert aber verstaendlich""",

    "context": f"""Du bist BEATRIX, die Kontext-Analystin von FehrAdvice & Partners AG.
Der User moechte die Ausgangslage eines Kunden oder Projekts erfassen, beschreiben oder analysieren.

DU ERFASST:
- Branchenkontext und Marktumfeld
- Organisationskultur und -struktur
- Aktuelle Herausforderungen und Pain Points
- Stakeholder-Landschaft
- Bestehende Initiativen und Massnahmen
- Regulatorisches Umfeld
- Bisherige Zusammenarbeit mit FehrAdvice

FELDER:
- customer_code: Kundencode
- project_slug: Projektbezug (optional)
- context_type: market | organization | stakeholder | regulatory | behavioral | competitive
- summary: Zusammenfassung der Ausgangslage
- challenges: Liste der Herausforderungen
- opportunities: Chancen
- stakeholders: Relevante Stakeholder
- market_factors: Marktfaktoren
- behavioral_patterns: Beobachtete Verhaltensmuster
- data_sources: Woher stammen die Infos
- assessment: Eigene Einschaetzung/Hypothese

{{ctx}}

REGELN:
1. Antworte mit JSON in ```json ... ``` Tags
2. JSON: {{"intent":"context", "action":"capture"|"analyze"|"update", "status":"complete"|"need_info", "data":{{...}}, "missing":[...], "message":"...", "confidence":0.0-1.0}}
3. Pflichtfelder: customer_code, summary
4. AUTOMATISCH aus Datenbank: Branche, Land, bekannte Kontakte/Stakeholder
5. Strukturiere die Antwort so, dass sie direkt in ein Projekt-YAML fliessen kann
6. Deutsch, analytisch, auf den Punkt"""
    }


def call_claude_json(system_prompt, messages, today, model=ANTHROPIC_MODEL_STANDARD, max_tokens=2500):
    """Call Claude API and extract JSON from response."""
    import urllib.request, ssl, re
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    payload = json.dumps({
        "model": model,
        "max_tokens": max_tokens,
        "system": system_prompt + f"\n\nHeute ist: {today}",
        "messages": messages
    }).encode()
    req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
        headers={
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
            "Content-Type": "application/json",
            "User-Agent": "BEATRIXLab/3.20"
        })
    resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=45).read())
    answer = resp["content"][0]["text"]

    # Try to extract JSON
    json_match = re.search(r'```json\s*(.*?)\s*```', answer, re.DOTALL)
    if json_match:
        try:
            return json.loads(json_match.group(1)), answer
        except json.JSONDecodeError:
            pass
    # Try raw JSON parse
    try:
        return json.loads(answer.strip()), answer
    except:
        pass
    return None, answer


@app.post("/api/chat/intent")
async def chat_intent(request: Request, user=Depends(require_permission("chat.intent"))):
    """Universal intent router – detects what the user wants and routes to domain specialist."""
    body = await request.json()
    message = body.get("message", "").strip()
    history = body.get("history", [])
    current_draft = body.get("current_draft", {})
    forced_intent = body.get("intent", "")  # Frontend can force an intent
    _session_id = body.get("session_id", "")

    if not message:
        return JSONResponse({"error": "message required"}, status_code=400)
    if not ANTHROPIC_API_KEY:
        return JSONResponse({"error": "Claude API nicht konfiguriert"}, status_code=501)

    # ── Session Context ──
    db_sess = get_db()
    try:
        session = get_or_create_session(db_sess, _session_id, user["sub"])
        # Save user message for session history
        _intent_user_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="user", content=message)
        db_sess.add(_intent_user_msg); db_sess.commit()
    finally:
        db_sess.close()
    session_type = session.get("type", "general")
    session_entities = session.get("entities", {})

    today = datetime.utcnow().strftime("%Y-%m-%d")

    # ── Step 1: Route intent ──
    intent = forced_intent
    entities = {}
    # If session already has a type, use it as default
    if not intent and session_type != "general":
        intent = session_type

    if not intent:
        # Quick classification call (Haiku = fast, 1-2s vs 3-5s Sonnet)
        try:
            router_messages = [{"role": "user", "content": message}]
            parsed, raw = call_claude_json(INTENT_ROUTER_SYSTEM, router_messages, today,
                                           model=ANTHROPIC_MODEL_LIGHT, max_tokens=500)
            if parsed:
                intent = parsed.get("intent", "general")
                entities = parsed.get("entities", {})
                logger.info(f"Intent routed: {intent} (confidence: {parsed.get('confidence',0)}) for: {message[:60]}")
            else:
                intent = "general"
        except Exception as e:
            logger.warning(f"Router failed: {e}, defaulting to general")
            intent = "general"

    # ── Step 2: Handle knowledge/general via existing KB path ──
    if intent in ("knowledge", "general"):
        # 🔍 Background fact-check
        import threading
        _fc_msg, _fc_user, _fc_cust = message, user["sub"], entities.get("customer", "")
        threading.Thread(target=lambda: fact_check_user_claims(_fc_msg, _fc_user, customer_code=_fc_cust), daemon=True).start()
        return {
            "ok": True,
            "intent": intent,
            "status": "redirect_kb",
            "message": "",
            "entities": entities,
            "session_id": _session_id
        }

    # ── Update session type + merge entities ──
    merged_entities = {**session_entities, **entities}
    db_upd = get_db()
    try:
        update_session(db_upd, _session_id, session_type=get_session_type(intent), entities=merged_entities)
    finally:
        db_upd.close()
    session_entities = merged_entities

    # ── Step 3: Domain specialist ──
    _prompts = get_domain_prompts()
    domain_prompt = _prompts.get(intent, _prompts.get("project"))

    # Inject session rules and entities
    _sess_rules = SESSION_RULES.get(intent, SESSION_RULES.get("general", ""))
    domain_prompt += f"\n\n{_sess_rules}"
    if session_entities:
        domain_prompt += f"\n\nSESSION-ENTITÄTEN (bereits erkannt): {json.dumps(session_entities, ensure_ascii=False)}"
        domain_prompt += "\nNutze diese bekannten Daten! Frage NICHT erneut danach."

    # ── Project context enrichment ──
    project_slug = body.get("project_slug", "")
    if project_slug:
        try:
            prj_ctx = load_project_context_from_github(project_slug)
            if prj_ctx.get("project"):
                p = prj_ctx["project"]
                client = p.get("client", {})
                project_info = p.get("project", {})
                objective = p.get("objective", {})
                scope = p.get("scope", {})
                fa_scope = p.get("fehradvice_scope", {})
                meta = p.get("metadata", {})
                ctx_lines = [f"\n\nPROJEKTKONTEXT (Projekt: {meta.get('project_code', project_slug)}):"]
                if client.get("name"): ctx_lines.append(f"Kunde: {client['name']}")
                if client.get("customer_code"): ctx_lines.append(f"Kundencode: {client['customer_code']}")
                if client.get("industry"): ctx_lines.append(f"Branche: {client['industry']}")
                if client.get("country"): ctx_lines.append(f"Land: {client['country']}")
                if project_info.get("name"): ctx_lines.append(f"Projektname: {project_info['name']}")
                if project_info.get("description"): ctx_lines.append(f"Beschreibung: {project_info['description']}")
                if project_info.get("type"): ctx_lines.append(f"Typ: {project_info['type']}")
                if objective.get("summary"): ctx_lines.append(f"Ziel: {objective['summary']}")
                if scope.get("in_scope"):
                    in_s = scope["in_scope"]
                    ctx_lines.append(f"In Scope: {', '.join(in_s) if isinstance(in_s, list) else in_s}")
                beh_obj = fa_scope.get("behavioral_objectives", [])
                if beh_obj: ctx_lines.append(f"Behavioral Objectives: {', '.join(beh_obj) if isinstance(beh_obj, list) else beh_obj}")
                ten_c = fa_scope.get("ten_c_focus", [])
                if ten_c: ctx_lines.append(f"10C Fokus: {', '.join(ten_c) if isinstance(ten_c, list) else ten_c}")
                # Previous insights
                for ins in prj_ctx.get("insights", [])[-3:]:
                    ctx_lines.append(f"Bisherige Erkenntnis: {ins.get('question', '')[:80]} → {ins.get('answer', '')[:150]}")
                domain_prompt += "\n".join(ctx_lines)
                domain_prompt += f"\n\nWICHTIG: Setze in deiner Antwort project_slug='{project_slug}' und customer_code='{client.get('customer_code', '')}' in die data-Felder."
                logger.info(f"Intent enriched with project context: {project_slug}")
        except Exception as e:
            logger.warning(f"Project context enrichment failed: {e}")

    # ── Customer-specific enrichment ──
    detected_customer = entities.get("customer", "")
    if detected_customer:
        try:
            detail = get_customer_detail(detected_customer)
            if detail:
                domain_prompt += f"\n\nSPEZIFISCHE KUNDENDATEN FÜR '{detected_customer.upper()}':\n{detail}\n\nNutze diese Daten um Felder AUTOMATISCH zu befüllen. Frage NICHT nach Infos die hier stehen!"
                logger.info(f"Intent enriched with customer detail: {detected_customer}")
        except Exception as e:
            logger.warning(f"Customer enrichment failed: {e}")

    # Build messages from DB session history (primary source)
    db_hist2 = get_db()
    try:
        _sh = get_session_history(db_hist2, _session_id, limit=10)
    finally:
        db_hist2.close()
    messages = []
    for h in _sh[-8:]:
        messages.append({"role": h["role"], "content": h["content"]})
    if history and not _sh:
        for h in history[-10:]:
            messages.append({"role": h["role"], "content": h["content"]})

    # Add draft context
    draft_note = ""
    if current_draft:
        draft_note = f"\n\nAKTUELLER ENTWURF:\n{json.dumps(current_draft, ensure_ascii=False, indent=2)}\nDer User möchte diesen Entwurf anpassen."

    messages.append({"role": "user", "content": message})

    try:
        parsed, raw = call_claude_json(domain_prompt + draft_note, messages, today)

        if parsed:
            resp_message = parsed.get("message", "")
            resp_data = parsed.get("data", parsed.get("project", {}))

            # 🚀 Move ALL post-response tasks to background thread
            # These were blocking the response by ~8-12 seconds!
            def _background_enrichment():
                try:
                    meta = {"intent": intent}
                    if resp_data.get("customer_code"): meta["customer_code"] = resp_data["customer_code"]
                    if resp_data.get("project_slug"): meta["project_slug"] = resp_data["project_slug"]
                    save_content = resp_message
                    if resp_data:
                        save_content = f"{resp_message}\n\nStrukturierte Daten ({intent}): {json.dumps(resp_data, ensure_ascii=False)}"
                    doc_id = auto_save_chat_to_kb(message, save_content, user["sub"], intent=intent, metadata=meta)
                    # Fact-check
                    try:
                        fact_check_user_claims(message, user["sub"], chat_doc_id=doc_id,
                                               customer_code=resp_data.get("customer_code", ""))
                    except Exception: pass
                    # Bias analysis
                    try:
                        analyze_beliefs_and_biases(message, resp_message, user["sub"],
                            chat_doc_id=doc_id, customer_code=resp_data.get("customer_code", ""))
                    except Exception: pass
                except Exception as e:
                    logger.warning(f"Background enrichment failed: {e}")

            import threading
            threading.Thread(target=_background_enrichment, daemon=True).start()

            return {
                "ok": True,
                "intent": parsed.get("intent", intent),
                "action": parsed.get("action", "create"),
                "status": parsed.get("status", "need_info"),
                "data": resp_data,
                "missing": parsed.get("missing", []),
                "message": resp_message,
                "confidence": parsed.get("confidence", 0),
                "entities": entities,
                "raw": raw
            }
        else:
            return {
                "ok": True,
                "intent": intent,
                "status": "need_info",
                "data": {},
                "missing": [],
                "message": raw,
                "confidence": 0,
                "entities": entities
            }
    except Exception as e:
        logger.error(f"Domain specialist error ({intent}): {e}")
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/api/chat/intent/stream")
async def chat_intent_stream(request: Request, user=Depends(require_permission("chat.intent"))):
    """SSE streaming version of intent chat – first token appears in ~2s instead of ~10s.

    SSE event types:
      {type: "intent", intent: "lead", entities: {...}}     — after routing
      {type: "token", text: "..."}                          — streamed text chunks
      {type: "data", data: {...}, status: "...", ...}       — final structured data
      {type: "done"}                                        — end
    """
    import http.client, ssl as _ssl, re as _re

    body = await request.json()
    message = body.get("message", "").strip()
    history = body.get("history", [])
    current_draft = body.get("current_draft", {})
    forced_intent = body.get("intent", "")
    attachments = body.get("attachments", [])
    _session_id = body.get("session_id", "")
    _analysis_mode = body.get("analysis_mode", "fast")  # 'fast' or 'deep'

    if not message and not attachments:
        return JSONResponse({"error": "message required"}, status_code=400)
    if not ANTHROPIC_API_KEY:
        return JSONResponse({"error": "Claude API nicht konfiguriert"}, status_code=501)

    # ── Session Context ──
    db_sess = get_db()
    try:
        session = get_or_create_session(db_sess, _session_id, user["sub"])
        # Save user message so session history works across messages
        user_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="user", content=message)
        db_sess.add(user_msg); db_sess.commit()
    finally:
        db_sess.close()
    session_type = session.get("type", "general")
    session_entities = session.get("entities", {})

    today = datetime.utcnow().strftime("%Y-%m-%d")

    # ── Step 1: Route intent (Haiku, fast) ──
    intent = forced_intent
    entities = {}
    # If session already has a type, use it as default
    if not intent and session_type != "general":
        intent = session_type
    if not intent:
        try:
            router_messages = [{"role": "user", "content": message}]
            parsed, raw = call_claude_json(INTENT_ROUTER_SYSTEM, router_messages, today,
                                           model=ANTHROPIC_MODEL_LIGHT, max_tokens=500)
            if parsed:
                intent = parsed.get("intent", "general")
                entities = parsed.get("entities", {})
            else:
                intent = "general"
        except Exception:
            intent = "general"

    # ── Update session type + merge entities ──
    if intent not in ("knowledge", "general"):
        merged_entities = {**session_entities, **entities}  # new entities override old
        db_upd = get_db()
        try:
            update_session(db_upd, _session_id, session_type=get_session_type(intent), entities=merged_entities)
        finally:
            db_upd.close()
        session_entities = merged_entities

    # ── Step 2: Redirect knowledge/general ──
    if intent in ("knowledge", "general"):
        import threading
        _fc_msg, _fc_user, _fc_cust = message, user["sub"], entities.get("customer", "")
        threading.Thread(target=lambda: fact_check_user_claims(_fc_msg, _fc_user, customer_code=_fc_cust), daemon=True).start()

        async def _redirect_gen():
            yield f"data: {json.dumps({'type': 'intent', 'intent': intent, 'entities': entities, 'status': 'redirect_kb', 'session_id': _session_id, 'analysis_mode': _analysis_mode})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        return StreamingResponse(_redirect_gen(), media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"})

    # ── Step 3: Build domain prompt (same enrichment as non-streaming) ──
    _prompts = get_domain_prompts()
    domain_prompt = _prompts.get(intent, _prompts.get("project"))

    # Inject session rules and entities
    session_rules = SESSION_RULES.get(intent, SESSION_RULES.get("general", ""))
    domain_prompt += f"\n\n{session_rules}"
    if session_entities:
        domain_prompt += f"\n\nSESSION-ENTITÄTEN (bereits erkannt): {json.dumps(session_entities, ensure_ascii=False)}"
        domain_prompt += "\nNutze diese bekannten Daten! Frage NICHT erneut danach."

    # Project context enrichment
    project_slug = body.get("project_slug", "")
    if project_slug:
        try:
            prj_ctx = load_project_context_from_github(project_slug)
            if prj_ctx.get("project"):
                p = prj_ctx["project"]
                client = p.get("client", {})
                project_info = p.get("project", {})
                meta = p.get("metadata", {})
                ctx_lines = [f"\n\nPROJEKTKONTEXT (Projekt: {meta.get('project_code', project_slug)}):"]
                if client.get("name"): ctx_lines.append(f"Kunde: {client['name']}")
                if project_info.get("name"): ctx_lines.append(f"Projektname: {project_info['name']}")
                domain_prompt += "\n".join(ctx_lines)
        except Exception: pass

    # Customer enrichment
    detected_customer = entities.get("customer", "")
    if detected_customer:
        try:
            detail = get_customer_detail(detected_customer)
            if detail:
                domain_prompt += f"\n\nSPEZIFISCHE KUNDENDATEN FÜR '{detected_customer.upper()}':\n{detail}\n\nNutze diese Daten um Felder AUTOMATISCH zu befüllen. Frage NICHT nach Infos die hier stehen!"
        except Exception: pass

    # ── Streaming prompt modification ──
    # Tell Claude to output message text FIRST, then JSON block
    stream_instruction = """

AUSGABEFORMAT FÜR STREAMING:
Schreibe ZUERST deine Nachricht als normalen Text (wird dem User live angezeigt).
Dann schreibe auf einer NEUEN Zeile exakt: <<<JSON>>>
Dann das JSON-Objekt mit den strukturierten Daten:
{"intent": "...", "action": "...", "status": "...", "data": {...}, "missing": [...], "confidence": ...}

Beispiel:
Ich habe den Lead für Helvetia angelegt. Folgende Daten fehlen noch:
- Geschätztes Projektvolumen
- Konkrete Opportunity

<<<JSON>>>
{"intent": "lead", "action": "create", "status": "need_info", "data": {"customer_code": "helvetia"}, "missing": ["value", "opportunity"], "confidence": 0.6}"""

    domain_prompt += stream_instruction

    # Build messages from DB session history (primary source of truth)
    db_hist = get_db()
    try:
        session_history = get_session_history(db_hist, _session_id, limit=10)
    finally:
        db_hist.close()
    
    messages = []
    # DB history = persistent across reloads
    for h in session_history[-8:]:
        messages.append({"role": h["role"], "content": h["content"]})
    # If frontend has newer messages not yet in DB, append them
    if history and not session_history:
        for h in history[-10:]:
            messages.append({"role": h["role"], "content": h["content"]})
    draft_note = ""
    if current_draft:
        draft_note = f"\n\nAKTUELLER ENTWURF:\n{json.dumps(current_draft, ensure_ascii=False, indent=2)}\nDer User möchte diesen Entwurf anpassen."

    # Build user message content with attachments (images + documents)
    user_content = []
    doc_context = ""
    for att in attachments:
        if att.get("type") == "image" and att.get("image_base64"):
            user_content.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": att.get("media_type", "image/png"),
                    "data": att["image_base64"]
                }
            })
        elif att.get("type") == "document" and att.get("content_text"):
            doc_context += f"\n\n--- DOKUMENT: {att.get('name', 'Datei')} ---\n{att['content_text'][:30000]}\n--- ENDE DOKUMENT ---\n"
    if doc_context:
        message += f"\n\nHochgeladene Dokumente:{doc_context}"
    user_content.append({"type": "text", "text": message})
    messages.append({"role": "user", "content": user_content})

    async def event_generator():
        ctx = _ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = _ssl.CERT_NONE

        # Send intent event immediately
        yield f"data: {json.dumps({'type': 'intent', 'intent': intent, 'entities': entities, 'session_id': _session_id, 'session_type': get_session_type(intent)})}\n\n"

        # Stream domain specialist
        payload = json.dumps({
            "model": ANTHROPIC_MODEL_STANDARD,
            "max_tokens": 2500,
            "stream": True,
            "system": domain_prompt + draft_note + f"\n\nHeute ist: {today}",
            "messages": messages
        }).encode()

        full_text = []
        json_started = False
        try:
            conn = http.client.HTTPSConnection("api.anthropic.com", timeout=90, context=ctx)
            conn.request("POST", "/v1/messages", body=payload, headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
                "User-Agent": "BEATRIXLab/3.22-stream"
            })
            resp = conn.getresponse()

            if resp.status != 200:
                error_body = resp.read().decode()
                logger.error(f"Intent stream error: {resp.status} {error_body[:200]}")
                yield f"data: {json.dumps({'type': 'error', 'text': 'Claude API Fehler'})}\n\n"
                return

            buffer = ""
            while True:
                chunk = resp.read(1024)
                if not chunk:
                    break
                buffer += chunk.decode("utf-8", errors="replace")

                while "\n\n" in buffer:
                    event_str, buffer = buffer.split("\n\n", 1)
                    for line in event_str.split("\n"):
                        if line.startswith("data: "):
                            data_str = line[6:]
                            if data_str.strip() == "[DONE]":
                                continue
                            try:
                                event = json.loads(data_str)
                                event_type = event.get("type", "")

                                if event_type == "content_block_delta":
                                    delta = event.get("delta", {})
                                    if delta.get("type") == "text_delta":
                                        text_chunk = delta.get("text", "")
                                        full_text.append(text_chunk)

                                        # Check if we've hit the JSON delimiter
                                        current = "".join(full_text)
                                        if "<<<JSON>>>" in current and not json_started:
                                            json_started = True
                                            # Don't stream JSON tokens to frontend
                                            # But stream any text before the delimiter that wasn't sent yet
                                            continue

                                        if not json_started:
                                            # Stream visible text to frontend
                                            yield f"data: {json.dumps({'type': 'token', 'text': text_chunk})}\n\n"

                            except json.JSONDecodeError:
                                pass

            conn.close()

            # ── Parse complete response ──
            complete_text = "".join(full_text)
            resp_message = complete_text
            parsed_data = {}

            if "<<<JSON>>>" in complete_text:
                parts = complete_text.split("<<<JSON>>>", 1)
                resp_message = parts[0].strip()
                json_part = parts[1].strip()
                try:
                    parsed_data = json.loads(json_part)
                except:
                    # Try extracting JSON from ```json blocks
                    json_match = _re.search(r'```json\s*(.*?)\s*```', json_part, _re.DOTALL)
                    if json_match:
                        try: parsed_data = json.loads(json_match.group(1))
                        except: pass
            else:
                # Fallback: try to parse entire response as JSON
                json_match = _re.search(r'```json\s*(.*?)\s*```', complete_text, _re.DOTALL)
                if json_match:
                    try:
                        parsed_data = json.loads(json_match.group(1))
                        resp_message = parsed_data.get("message", resp_message)
                    except: pass

            # Send structured data event
            resp_data = parsed_data.get("data", parsed_data) if isinstance(parsed_data, dict) else {}
            yield f"data: {json.dumps({'type': 'data', 'intent': parsed_data.get('intent', intent), 'action': parsed_data.get('action', 'create'), 'status': parsed_data.get('status', 'need_info'), 'data': resp_data, 'missing': parsed_data.get('missing', []), 'confidence': parsed_data.get('confidence', 0), 'message': resp_message, 'entities': entities})}\n\n"

            # Save assistant message to chat_messages for session history
            if resp_message:
                db_save = get_db()
                try:
                    asst_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="assistant", content=resp_message)
                    db_save.add(asst_msg); db_save.commit()
                except Exception:
                    pass
                finally:
                    db_save.close()

            # Background tasks
            def _bg():
                try:
                    meta = {"intent": intent}
                    if isinstance(resp_data, dict):
                        if resp_data.get("customer_code"): meta["customer_code"] = resp_data["customer_code"]
                        if resp_data.get("project_slug"): meta["project_slug"] = resp_data["project_slug"]
                    save_content = resp_message
                    if resp_data:
                        save_content = f"{resp_message}\n\nStrukturierte Daten ({intent}): {json.dumps(resp_data, ensure_ascii=False)}"
                    doc_id = auto_save_chat_to_kb(message, save_content, user["sub"], intent=intent, metadata=meta)
                    try: fact_check_user_claims(message, user["sub"], chat_doc_id=doc_id, customer_code=(resp_data.get("customer_code","") if isinstance(resp_data,dict) else ""))
                    except: pass
                    try: analyze_beliefs_and_biases(message, resp_message, user["sub"], chat_doc_id=doc_id)
                    except: pass
                except Exception as e:
                    logger.warning(f"Intent stream bg failed: {e}")
            import threading; threading.Thread(target=_bg, daemon=True).start()

            yield f"data: {json.dumps({'type': 'done'})}\n\n"

        except Exception as e:
            logger.error(f"Intent stream error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'text': str(e)})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"}
    )

# Keep legacy endpoint as alias
@app.post("/api/chat/project-intent")
async def chat_project_intent(request: Request, user=Depends(require_permission("chat.intent"))):
    """Legacy alias → routes through universal intent router with forced project intent."""
    from starlette.requests import Request as StarletteRequest
    body = await request.json()
    body["intent"] = "project"
    # Remap legacy fields
    if "current_draft" in body and "project" not in body.get("current_draft", {}):
        body["current_draft"] = body.get("current_draft", {})
    # Create a simple forwarding
    import urllib.request, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    today = datetime.utcnow().strftime("%Y-%m-%d")
    message = body.get("message", "")
    history = body.get("history", [])
    current_draft = body.get("current_draft", {})
    domain_prompt = get_domain_prompts()["project"]

    messages = []
    for h in history[-10:]:
        messages.append({"role": h["role"], "content": h["content"]})
    draft_note = ""
    if current_draft:
        draft_note = f"\n\nAKTUELLER ENTWURF:\n{json.dumps(current_draft, ensure_ascii=False, indent=2)}\nDer User möchte diesen Entwurf anpassen."
    messages.append({"role": "user", "content": message})

    try:
        parsed, raw = call_claude_json(domain_prompt + draft_note, messages, today)
        if parsed:
            return {
                "ok": True,
                "intent": "project",
                "status": parsed.get("status", "need_info"),
                "project": parsed.get("data", parsed.get("project", {})),
                "data": parsed.get("data", parsed.get("project", {})),
                "missing": parsed.get("missing", []),
                "message": parsed.get("message", ""),
                "confidence": parsed.get("confidence", 0),
                "action": parsed.get("action", "create")
            }
        return {"ok": True, "intent": "project", "status": "need_info", "project": {}, "data": {},
                "missing": [], "message": raw, "confidence": 0, "action": "create"}
    except Exception as e:
        return JSONResponse({"error": str(e)}, status_code=500)


@app.post("/api/chat/stream")
async def chat_stream(request: ChatRequest, user=Depends(require_auth)):
    """SSE streaming endpoint – streams Claude tokens to frontend in real-time."""
    import urllib.request, ssl, http.client

    question = request.question.strip()
    _session_id = request.session_id
    if not question:
        raise HTTPException(400, "Bitte stelle eine Frage")

    if not ANTHROPIC_API_KEY:
        raise HTTPException(501, "Claude API nicht konfiguriert")

    # ── Session Context ──
    db = get_db()
    try:
        session = get_or_create_session(db, _session_id, user["sub"])
        session_history = get_session_history(db, _session_id, limit=10)
        # Save user message BEFORE streaming so next request sees it
        user_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="user", content=question)
        db.add(user_msg); db.commit()
        results = search_knowledge_base(db, question)
        context = build_context(results) if results else ""
    finally:
        db.close()

    session_type = session.get("type", "general")
    session_entities = session.get("entities", {})
    session_rules = SESSION_RULES.get(session_type, SESSION_RULES["general"])

    system_prompt = f"""Du bist BEATRIX, die Strategic Intelligence Suite von FehrAdvice & Partners AG.
Du bist spezialisiert auf das Evidence-Based Framework (EBF), Behavioral Economics, das Behavioral Competence Model (BCM) und Decision Architecture.

Dir steht vorhandenes Wissen aus der BEATRIX Knowledge Base zur Verfügung.

{session_rules}

{f"SESSION-KONTEXT: Bekannte Entitäten: {json.dumps(session_entities, ensure_ascii=False)}" if session_entities else ""}

Deine Aufgabe:
- Beantworte Fragen basierend auf dem bereitgestellten Kontext UND der bisherigen Konversation
- Wenn der User Bilder schickt, lies sie IMMER aus und extrahiere alle relevanten Informationen
- Antworte präzise, wissenschaftlich fundiert und praxisorientiert
- Antworte auf Deutsch, es sei denn die Frage ist auf Englisch
- Halte dich an die Session-Regeln oben

Stil: Professionell, klar, auf den Punkt. Wie ein Senior Berater bei FehrAdvice."""

    user_message = f"""Kontext aus der BEATRIX Knowledge Base:

{context}

---

{question}"""

    # Build message list with session history
    messages = []
    for h in session_history[-8:]:  # Last 8 messages for context
        messages.append({"role": h["role"], "content": h["content"]})

    # Build current message content with attachments
    attachments = request.attachments or []
    user_content = []
    doc_context = ""

    for att in attachments:
        if att.get("type") == "image" and att.get("image_base64"):
            user_content.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": att.get("media_type", "image/png"),
                    "data": att["image_base64"]
                }
            })
        elif att.get("type") == "document" and att.get("content_text"):
            doc_context += f"\n\n--- DOKUMENT: {att.get('name', 'Datei')} ---\n{att['content_text'][:30000]}\n--- ENDE DOKUMENT ---\n"

    if doc_context:
        user_message += f"\n\nHochgeladene Dokumente:{doc_context}"

    user_content.append({"type": "text", "text": user_message})

    async def event_generator():
        """Generator that streams Claude API response as SSE events."""
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE

        payload = json.dumps({
            "model": ANTHROPIC_MODEL_STANDARD,
            "max_tokens": 4000,
            "stream": True,
            "system": system_prompt,
            "messages": messages + [{"role": "user", "content": user_content}]
        }).encode()

        full_text = []
        try:
            # Use http.client for streaming (urllib doesn't support chunked reading well)
            conn = http.client.HTTPSConnection("api.anthropic.com", timeout=90, context=ctx)
            conn.request("POST", "/v1/messages", body=payload, headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
                "User-Agent": "BEATRIXLab/3.11"
            })
            resp = conn.getresponse()

            if resp.status != 200:
                error_body = resp.read().decode()
                logger.error(f"Claude stream error: {resp.status} {error_body[:200]}")
                yield f"data: {json.dumps({'type': 'error', 'text': 'Claude API Fehler'})}\n\n"
                return

            buffer = ""
            while True:
                chunk = resp.read(1024)
                if not chunk:
                    break
                buffer += chunk.decode("utf-8", errors="replace")

                # Process complete SSE events from buffer
                while "\n\n" in buffer:
                    event_str, buffer = buffer.split("\n\n", 1)
                    for line in event_str.split("\n"):
                        if line.startswith("data: "):
                            data_str = line[6:]
                            if data_str.strip() == "[DONE]":
                                continue
                            try:
                                event = json.loads(data_str)
                                event_type = event.get("type", "")

                                if event_type == "content_block_delta":
                                    delta = event.get("delta", {})
                                    if delta.get("type") == "text_delta":
                                        text_chunk = delta.get("text", "")
                                        full_text.append(text_chunk)
                                        yield f"data: {json.dumps({'type': 'token', 'text': text_chunk})}\n\n"

                                elif event_type == "message_stop":
                                    yield f"data: {json.dumps({'type': 'done', 'text': ''})}\n\n"

                            except json.JSONDecodeError:
                                pass

            conn.close()

            # Store complete answer in DB + auto-save to KB
            complete_text = "".join(full_text)
            if complete_text:
                db2 = get_db()
                try:
                    assistant_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="assistant", content=complete_text, sources=[])
                    db2.add(assistant_msg)
                    db2.commit()
                except Exception:
                    pass
                finally:
                    db2.close()
                # 🚀 Background: save + analysis (don't delay 'done' event)
                def _bg_stream():
                    try:
                        doc_id = auto_save_chat_to_kb(question, complete_text, user["sub"], intent="knowledge")
                        try: fact_check_user_claims(question, user["sub"], chat_doc_id=doc_id)
                        except: pass
                        try: analyze_beliefs_and_biases(question, complete_text, user["sub"], chat_doc_id=doc_id)
                        except: pass
                    except Exception as e:
                        logger.warning(f"Chat KB auto-save failed (stream): {e}")
                import threading; threading.Thread(target=_bg_stream, daemon=True).start()

            # Final done event
            yield f"data: {json.dumps({'type': 'done', 'text': ''})}\n\n"

        except Exception as e:
            logger.error(f"Stream error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'text': str(e)})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        }
    )

@app.post("/api/chat")
async def chat(request: ChatRequest, user=Depends(require_auth)):
    question = request.question.strip()
    _session_id = request.session_id
    if not question:
        raise HTTPException(400, "Bitte stelle eine Frage")
    if len(question) > 5000:
        raise HTTPException(400, "Frage zu lang (max 5000 Zeichen)")

    db = get_db()
    try:
        # Store user message
        user_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="user", content=question)
        db.add(user_msg); db.commit()

        # === PATH DECISION: Search KB for existing EBF answers ===
        results = search_knowledge_base(db, question)
        top_score = results[0][0] if results else 0
        ebf_results = [(s, d) for s, d in results if d.source_type == "ebf_answer" or (d.tags and "ebf-answer" in d.tags)]
        ebf_score = ebf_results[0][0] if ebf_results else 0

        # Decision: Use fast path only if we have a strong EBF answer match
        # or if there's a very strong general match AND an EBF answer exists at all
        use_fast_path = (ebf_score >= FAST_PATH_THRESHOLD) or (top_score >= FAST_PATH_THRESHOLD * 3 and ebf_score > 0)

        # Freshness check: if best EBF answer is older than KB_FRESHNESS_DAYS → stale, force deep
        if use_fast_path and ebf_results:
            best_ebf_doc = ebf_results[0][1]
            doc_age = (datetime.utcnow() - best_ebf_doc.updated_at).days if best_ebf_doc.updated_at else 999
            if doc_age > KB_FRESHNESS_DAYS:
                logger.info(f"Chat: EBF answer stale ({doc_age}d > {KB_FRESHNESS_DAYS}d) → forcing deep refresh for: {question[:50]}")
                use_fast_path = False

        # User override: if analysis_mode='deep', force deep path regardless of score
        if getattr(request, 'analysis_mode', 'fast') == 'deep':
            logger.info(f"Chat: User forced DEEP mode for: {question[:50]}")
            use_fast_path = False

        # If fast path uses general docs (not EBF answers), prioritize EBF answers in context
        if use_fast_path and ebf_results:
            # Put EBF answers first in results
            non_ebf = [(s, d) for s, d in results if d not in [d2 for _, d2 in ebf_results]]
            results = ebf_results + non_ebf

        _ebf_age = (datetime.utcnow() - ebf_results[0][1].updated_at).days if ebf_results and ebf_results[0][1].updated_at else -1
        logger.info(f"Chat: q='{question[:50]}' | top={top_score} | ebf={ebf_score} | age={_ebf_age}d | fast={use_fast_path}")

        # === FAST PATH: Good KB match exists → Claude API with context (3 sec) ===
        if use_fast_path and ANTHROPIC_API_KEY:
            context = build_context(results)
            sources = [{"title": doc.title, "id": doc.id, "type": doc.source_type, "category": doc.category} for _, doc in results[:5]]
            try:
                answer = fast_path_answer(question, context, sources)
                assistant_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="assistant", content=answer, sources=sources)
                db.add(assistant_msg); db.commit()
                logger.info(f"FAST PATH answer for: {question[:50]}")
                # 🚀 Background: save + fact-check + bias analysis
                def _bg_fast():
                    try:
                        doc_id = auto_save_chat_to_kb(question, answer, user["sub"], intent="knowledge",
                                                         metadata={"project_slug": request.project_slug} if request.project_slug else None)
                        try: fact_check_user_claims(question, user["sub"], chat_doc_id=doc_id)
                        except: pass
                        try: analyze_beliefs_and_biases(question, answer, user["sub"], chat_doc_id=doc_id)
                        except: pass
                    except: pass
                import threading; threading.Thread(target=_bg_fast, daemon=True).start()
                return {
                    "status": "done",
                    "answer": answer,
                    "sources": sources,
                    "path": "fast",
                    "knowledge_score": top_score
                }
            except Exception as e:
                logger.warning(f"Fast path failed, falling through to deep path: {e}")

        # === DEEP PATH: No good match → GitHub Claude Code (4-5 min) ===
        if not GH_TOKEN:
            raise HTTPException(501, "Kein ausreichendes Wissen vorhanden und GitHub-Integration nicht konfiguriert")

        # Collect existing KB context as warm start for Claude Code
        _kb_context = ""
        if results:
            _ctx_parts = []
            for _sc, _doc in results[:3]:
                if _doc.content and len(_doc.content) > 50:
                    _ctx_parts.append(f"**{_doc.title}** (Score: {_sc:.0f}):\n{_doc.content[:800]}")
            if _ctx_parts:
                _kb_context = "\n\n---\n\n".join(_ctx_parts)

        gh = create_github_issue(question, user["sub"], kb_context=_kb_context)
        logger.info(f"DEEP PATH: GitHub Issue #{gh['issue_number']} for: {question[:50]} (kb_context: {len(_kb_context)} chars)")

        return {
            "status": "processing",
            "issue_number": gh["issue_number"],
            "issue_url": gh["html_url"],
            "path": "deep",
            "knowledge_score": top_score,
            "message": "Neue Frage! BEATRIX erarbeitet die Antwort mit dem Evidence-Based Framework..."
        }
    except HTTPException: raise
    except Exception as e:
        logger.error(f"Chat error: {e}")
        raise HTTPException(500, f"Konnte Frage nicht verarbeiten: {str(e)}")
    finally: db.close()


# ── PROJECT-SCOPED CHAT: Chat within project context, save to GitHub ──

def load_project_context_from_github(slug: str) -> dict:
    """Load project.yaml + any existing insights from GitHub."""
    import urllib.request, ssl, yaml as _yaml
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    gh = {"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab", "Accept": "application/vnd.github.v3.raw"}

    result = {"project": {}, "insights": [], "customer_context": ""}

    # 1. Load project.yaml
    try:
        url = f"https://api.github.com/repos/{GH_CONTEXT_REPO}/contents/data/projects/{slug}/project.yaml"
        req = urllib.request.Request(url, headers=gh)
        result["project"] = _yaml.safe_load(urllib.request.urlopen(req, context=ctx, timeout=15).read().decode()) or {}
    except Exception as e:
        logger.warning(f"Project context load failed for {slug}: {e}")

    # 2. Load existing insights
    try:
        url = f"https://api.github.com/repos/{GH_CONTEXT_REPO}/contents/data/projects/{slug}/insights"
        req = urllib.request.Request(url, headers={"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"})
        items = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
        for item in items[-5:]:
            try:
                req2 = urllib.request.Request(item["download_url"], headers=gh)
                content = urllib.request.urlopen(req2, context=ctx, timeout=10).read().decode()
                data = _yaml.safe_load(content) or {}
                result["insights"].append({"file": item["name"], "question": data.get("question", ""), "answer": data.get("answer", "")[:500]})
            except: pass
    except: pass

    # 3. Load customer context if available
    try:
        client = result["project"].get("client", {})
        cust_code = (client.get("short_name") or client.get("customer_code") or "").lower()
        if cust_code:
            url = f"https://api.github.com/repos/{GH_CONTEXT_REPO}/contents/data/customers/{cust_code}"
            req = urllib.request.Request(url, headers={"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"})
            items = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
            ctx_files = [i for i in items if i["name"].endswith(".yaml") and "context" in i["name"].lower()]
            for cf in ctx_files[:3]:
                try:
                    req2 = urllib.request.Request(cf["download_url"], headers=gh)
                    result["customer_context"] += urllib.request.urlopen(req2, context=ctx, timeout=10).read().decode()[:2000] + "\n"
                except: pass
    except: pass

    return result


def build_project_system_prompt(project_ctx: dict, slug: str) -> str:
    """Build a system prompt enriched with project context."""
    p = project_ctx.get("project", {})
    client = p.get("client", {})
    project = p.get("project", {})
    objective = p.get("objective", {})
    scope = p.get("scope", {})
    fehradvice_scope = p.get("fehradvice_scope", {})
    meta = p.get("metadata", {})

    ctx_parts = []
    ctx_parts.append(f"PROJEKT: {meta.get('project_code', slug)} - {project.get('name', '')}")
    if client.get("name"): ctx_parts.append(f"KUNDE: {client['name']}")
    if client.get("industry"): ctx_parts.append(f"BRANCHE: {client['industry']}")
    if client.get("country"): ctx_parts.append(f"LAND: {client['country']}")
    if objective.get("summary"): ctx_parts.append(f"ZIEL: {objective['summary']}")
    if project.get("description"): ctx_parts.append(f"BESCHREIBUNG: {project['description']}")
    if project.get("type"): ctx_parts.append(f"TYP: {project['type']}")
    if scope.get("in_scope"):
        in_s = scope["in_scope"]
        ctx_parts.append(f"IN SCOPE: {', '.join(in_s) if isinstance(in_s, list) else in_s}")
    beh_obj = fehradvice_scope.get("behavioral_objectives", [])
    if beh_obj: ctx_parts.append(f"BEHAVIORAL OBJECTIVES: {', '.join(beh_obj) if isinstance(beh_obj, list) else beh_obj}")
    ten_c = fehradvice_scope.get("ten_c_focus", [])
    if ten_c: ctx_parts.append(f"10C FOKUS: {', '.join(ten_c) if isinstance(ten_c, list) else ten_c}")

    insights = project_ctx.get("insights", [])
    insights_text = ""
    if insights:
        insights_text = "\n\nBISHERIGE ERKENNTNISSE ZU DIESEM PROJEKT:\n"
        for ins in insights:
            insights_text += f"- {ins.get('question', '')[:100]}: {ins.get('answer', '')[:200]}...\n"

    cust_ctx = project_ctx.get("customer_context", "")
    cust_text = ""
    if cust_ctx:
        cust_text = f"\n\nKUNDEN-KONTEXT:\n{cust_ctx[:3000]}"

    return f"""Du bist BEATRIX, die Strategic Intelligence Suite von FehrAdvice & Partners AG.

Du arbeitest gerade im Kontext eines spezifischen Kundenprojekts. Alle deine Antworten sollen sich auf dieses Projekt beziehen.

PROJEKTKONTEXT:
{chr(10).join(ctx_parts)}
{insights_text}{cust_text}

Deine Aufgabe:
- Beantworte Fragen im Kontext dieses Projekts
- Nutze dein Wissen ueber Behavioral Economics, EBF, BCM, Decision Architecture
- Wenn du Kontext-Analysen, Segment-Analysen oder Modelle erstellst, strukturiere sie klar
- Gib praxisorientierte, auf den Kunden zugeschnittene Antworten
- Antworte auf Deutsch, es sei denn die Frage ist auf Englisch

Stil: Professionell, klar, auf den Punkt. Wie ein Senior Berater bei FehrAdvice."""


@app.post("/api/chat/project/{slug}")
async def chat_project(slug: str, request: ChatRequest, user=Depends(require_auth)):
    """Project-scoped chat: enriched with project context, saves insights to GitHub."""
    import urllib.request, ssl, yaml as _yaml
    question = request.question.strip()
    _session_id = request.session_id
    if not question:
        raise HTTPException(400, "Bitte stelle eine Frage")

    project_ctx = load_project_context_from_github(slug)
    if not project_ctx.get("project"):
        raise HTTPException(404, f"Projekt '{slug}' nicht gefunden")

    db = get_db()
    try:
        results = search_knowledge_base(db, question)
        kb_context = build_context(results) if results else ""
        sources = [{"title": doc.title, "id": doc.id, "type": doc.source_type, "category": doc.category}
                   for _, doc in results[:5]] if results else []

        system_prompt = build_project_system_prompt(project_ctx, slug)
        if kb_context:
            system_prompt += f"\n\nZUSAETZLICHES WISSEN AUS DER BEATRIX KNOWLEDGE BASE:\n{kb_context[:4000]}"

        # Build message content with attachments
        attachments = request.attachments or []
        user_content = []
        doc_context = ""

        for att in attachments:
            if att.get("type") == "image" and att.get("image_base64"):
                user_content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": att.get("media_type", "image/png"),
                        "data": att["image_base64"]
                    }
                })
            elif att.get("type") == "document" and att.get("content_text"):
                doc_context += f"\n\n--- DOKUMENT: {att.get('name', 'Datei')} ---\n{att['content_text'][:30000]}\n--- ENDE DOKUMENT ---\n"

        full_question = question + (f"\n\nHochgeladene Dokumente:{doc_context}" if doc_context else "")
        user_content.append({"type": "text", "text": full_question})

        ctx_ssl = ssl.create_default_context(); ctx_ssl.check_hostname = False; ctx_ssl.verify_mode = ssl.CERT_NONE
        payload = json.dumps({
            "model": ANTHROPIC_MODEL_STANDARD,
            "max_tokens": 4000,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_content}]
        }).encode()
        req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
            headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
                "User-Agent": "BEATRIXLab/3.8"
            })
        resp = json.loads(urllib.request.urlopen(req, context=ctx_ssl, timeout=60).read())
        answer = resp["content"][0]["text"]

        user_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="user", content=f"[{slug}] {question}")
        db.add(user_msg); db.commit()
        assistant_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="assistant", content=answer, sources=sources)
        db.add(assistant_msg); db.commit()

        # Save insight to GitHub
        try:
            import re as _re
            now = datetime.utcnow()
            q_slug = _re.sub(r'[^a-z0-9]+', '-', question.strip()[:60].lower()).strip('-')[:40]
            insight_path = f"data/projects/{slug}/insights/{now.strftime('%Y-%m')}_{q_slug}.yaml"
            insight_yaml = _yaml.dump({
                "question": question,
                "answer": answer,
                "user": user.get("email", user.get("sub", "")),
                "timestamp": now.isoformat(),
                "project_slug": slug,
                "project_code": project_ctx["project"].get("metadata", {}).get("project_code", ""),
                "sources": [s.get("title", "") for s in sources[:3]],
                "type": "project_insight"
            }, default_flow_style=False, allow_unicode=True, sort_keys=False)

            gh_result = gh_put_file(
                GH_CONTEXT_REPO, insight_path, insight_yaml,
                f"prj-chat {slug}: {question.strip()[:60]}"
            )
            github_saved = gh_result.get("ok", False)
            logger.info(f"Project chat saved: {slug} -> {insight_path} (github={github_saved})")
        except Exception as e:
            github_saved = False
            logger.warning(f"Project insight GitHub push failed: {e}")

        # 🚀 Background save
        import threading
        _slug, _q, _a, _u, _cc = slug, question, answer, user.get("sub",""), project_ctx["project"].get("client",{}).get("customer_code","")
        threading.Thread(target=lambda: auto_save_chat_to_kb(_q, _a, _u, intent="project_chat", metadata={"project_slug":_slug,"customer_code":_cc}), daemon=True).start()

        return {
            "status": "done",
            "answer": answer,
            "sources": sources,
            "path": "project",
            "project_slug": slug,
            "github_saved": github_saved
        }
    except HTTPException: raise
    except Exception as e:
        logger.error(f"Project chat error: {e}")
        import traceback; traceback.print_exc()
        raise HTTPException(500, f"Fehler: {str(e)}")
    finally:
        db.close()


@app.post("/api/chat/project/{slug}/stream")
async def chat_project_stream(slug: str, request: ChatRequest, user=Depends(require_auth)):
    """Project-scoped streaming chat: enriched with project context, streams tokens via SSE."""
    import urllib.request, ssl, http.client, yaml as _yaml

    question = request.question.strip()
    _session_id = request.session_id
    if not question:
        raise HTTPException(400, "Bitte stelle eine Frage")
    if not ANTHROPIC_API_KEY:
        raise HTTPException(501, "Claude API nicht konfiguriert")

    # Load project context
    project_ctx = load_project_context_from_github(slug)
    if not project_ctx.get("project"):
        raise HTTPException(404, f"Projekt '{slug}' nicht gefunden")

    # Search KB
    db = get_db()
    try:
        results = search_knowledge_base(db, question)
        kb_context = build_context(results) if results else ""
        sources = [{"title": doc.title, "id": doc.id, "type": doc.source_type} for _, doc in results[:5]] if results else []
    finally:
        db.close()

    # Build enriched prompt
    system_prompt = build_project_system_prompt(project_ctx, slug)
    if kb_context:
        system_prompt += f"\n\nZUSAETZLICHES WISSEN AUS DER BEATRIX KNOWLEDGE BASE:\n{kb_context[:4000]}"

    async def event_generator():
        ctx_ssl = ssl.create_default_context()
        ctx_ssl.check_hostname = False
        ctx_ssl.verify_mode = ssl.CERT_NONE

        # Build message content with attachments
        attachments = request.attachments or []
        user_content = []
        doc_context = ""

        for att in attachments:
            if att.get("type") == "image" and att.get("image_base64"):
                user_content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": att.get("media_type", "image/png"),
                        "data": att["image_base64"]
                    }
                })
            elif att.get("type") == "document" and att.get("content_text"):
                doc_context += f"\n\n--- DOKUMENT: {att.get('name', 'Datei')} ---\n{att['content_text'][:30000]}\n--- ENDE DOKUMENT ---\n"

        # Add document text to question
        full_question = question
        if doc_context:
            full_question = f"{question}\n\nHochgeladene Dokumente:{doc_context}"

        user_content.append({"type": "text", "text": full_question})

        payload = json.dumps({
            "model": ANTHROPIC_MODEL_STANDARD,
            "max_tokens": 4000,
            "stream": True,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_content}]
        }).encode()

        full_text = []
        try:
            conn = http.client.HTTPSConnection("api.anthropic.com", timeout=90, context=ctx_ssl)
            conn.request("POST", "/v1/messages", body=payload, headers={
                "x-api-key": ANTHROPIC_API_KEY,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
                "User-Agent": "BEATRIXLab/3.20"
            })
            resp = conn.getresponse()

            if resp.status != 200:
                error_body = resp.read().decode()
                logger.error(f"Project stream error: {resp.status} {error_body[:200]}")
                yield f"data: {json.dumps({'type': 'error', 'text': 'Claude API Fehler'})}\n\n"
                return

            # Send sources first
            if sources:
                yield f"data: {json.dumps({'type': 'sources', 'sources': sources})}\n\n"

            buffer = ""
            while True:
                chunk = resp.read(1024)
                if not chunk:
                    break
                buffer += chunk.decode("utf-8", errors="replace")

                while "\n\n" in buffer:
                    event_str, buffer = buffer.split("\n\n", 1)
                    for line in event_str.split("\n"):
                        if line.startswith("data: "):
                            data_str = line[6:]
                            if data_str.strip() == "[DONE]":
                                continue
                            try:
                                event = json.loads(data_str)
                                event_type = event.get("type", "")

                                if event_type == "content_block_delta":
                                    delta = event.get("delta", {})
                                    if delta.get("type") == "text_delta":
                                        text_chunk = delta.get("text", "")
                                        full_text.append(text_chunk)
                                        yield f"data: {json.dumps({'type': 'token', 'text': text_chunk})}\n\n"

                                elif event_type == "message_stop":
                                    yield f"data: {json.dumps({'type': 'done', 'text': ''})}\n\n"

                            except json.JSONDecodeError:
                                pass

            conn.close()

            # Store complete answer
            complete_text = "".join(full_text)
            if complete_text:
                db2 = get_db()
                try:
                    user_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="user", content=f"[{slug}] {question}")
                    db2.add(user_msg)
                    assistant_msg = ChatMessage(user_email=user["sub"], session_id=_session_id, role="assistant", content=complete_text, sources=sources)
                    db2.add(assistant_msg)
                    db2.commit()
                except Exception:
                    pass
                finally:
                    db2.close()

                # Save insight to GitHub
                github_saved = False
                try:
                    import re as _re
                    now = datetime.utcnow()
                    q_slug = _re.sub(r'[^a-z0-9]+', '-', question.strip()[:60].lower()).strip('-')[:40]
                    insight_path = f"data/projects/{slug}/insights/{now.strftime('%Y-%m')}_{q_slug}.yaml"
                    insight_yaml = _yaml.dump({
                        "question": question,
                        "answer": complete_text,
                        "user": user.get("email", user.get("sub", "")),
                        "timestamp": now.isoformat(),
                        "project_slug": slug,
                        "project_code": project_ctx["project"].get("metadata", {}).get("project_code", ""),
                        "type": "project_insight"
                    }, default_flow_style=False, allow_unicode=True, sort_keys=False)
                    gh_result = gh_put_file(GH_CONTEXT_REPO, insight_path, insight_yaml,
                                            f"prj-chat {slug}: {question.strip()[:60]}")
                    github_saved = gh_result.get("ok", False)
                    logger.info(f"Project stream saved: {slug} -> {insight_path}")
                except Exception as e:
                    logger.warning(f"Project stream GitHub push failed: {e}")

                # 🚀 Background KB save
                import threading
                _q, _ct, _u, _sl = question, complete_text, user.get("sub",""), slug
                threading.Thread(target=lambda: auto_save_chat_to_kb(_q, _ct, _u, intent="project_chat", metadata={"project_slug":_sl}), daemon=True).start()

                # Send final meta
                yield f"data: {json.dumps({'type': 'meta', 'github_saved': github_saved, 'project_slug': slug})}\n\n"

            yield f"data: {json.dumps({'type': 'done', 'text': ''})}\n\n"

        except Exception as e:
            logger.error(f"Project stream error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'text': str(e)})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"}
    )


@app.get("/api/chat/poll/{issue_number}")
async def chat_poll(issue_number: int, user=Depends(require_auth)):
    """Poll for Claude Code's answer on a GitHub Issue. Stores answer in KB when done."""
    try:
        result = poll_github_answer(issue_number)
        if result["status"] == "done":
            db = get_db()
            try:
                # 1. Store in chat history
                assistant_msg = ChatMessage(
                    user_email=user["sub"], role="assistant",
                    content=result["answer"], session_id=None,
                    sources=[{"type": "github", "url": result.get("comment_url", "")}]
                )
                db.add(assistant_msg); db.commit()

                # 2. Store in Knowledge Base for future fast path!
                # Get the original question from the issue
                import urllib.request, ssl
                ctx2 = ssl.create_default_context(); ctx2.check_hostname = False; ctx2.verify_mode = ssl.CERT_NONE
                req = urllib.request.Request(
                    f"https://api.github.com/repos/{GH_CONTEXT_REPO}/issues/{issue_number}",
                    headers={"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3+json"})
                issue = json.loads(urllib.request.urlopen(req, context=ctx2).read())
                # Extract question from issue body (after @claude)
                body = issue.get("body", "")
                q_parts = body.split("@claude", 1)
                original_question = q_parts[1].strip() if len(q_parts) > 1 else issue.get("title", "").replace("BEATRIX: ", "")

                store_ebf_answer(db, original_question, result["answer"], result.get("comment_url", ""))
                logger.info(f"EBF answer cached from Issue #{issue_number}")
            finally: db.close()
        return result
    except Exception as e:
        logger.error(f"Poll error: {e}")
        return {"status": "error", "message": str(e)}

@app.get("/api/chat/history")
async def chat_history(limit: int = 50, session_id: str = None, user=Depends(require_auth)):
    db = get_db()
    try:
        q = db.query(ChatMessage).filter(ChatMessage.user_email == user["sub"])
        if session_id:
            q = q.filter(ChatMessage.session_id == session_id)
        msgs = q.order_by(ChatMessage.created_at.desc()).limit(limit).all()
        return [{"id": m.id, "role": m.role, "content": m.content, "sources": m.sources, "session_id": m.session_id, "created_at": m.created_at.isoformat()} for m in reversed(msgs)]
    finally: db.close()

@app.delete("/api/chat/history")
async def clear_chat_history(user=Depends(require_auth)):
    db = get_db()
    try:
        db.query(ChatMessage).filter(ChatMessage.user_email == user["sub"]).delete()
        db.commit()
        return {"message": "Chat-Verlauf gelöscht"}
    finally: db.close()

@app.get("/api/chat/session/{session_id}")
async def get_chat_session(session_id: str, user=Depends(require_auth)):
    """Get current session metadata."""
    db = get_db()
    try:
        session = get_or_create_session(db, session_id, user["sub"])
        history = get_session_history(db, session_id, limit=5)
        return {
            "session_id": session["id"],
            "type": session["type"],
            "entities": session["entities"],
            "context": session["context"],
            "recent_messages": len(history),
            "type_info": SESSION_TYPES.get(session["type"], SESSION_TYPES.get("general"))
        }
    finally:
        db.close()

@app.put("/api/chat/session/{session_id}/type")
async def set_chat_session_type(session_id: str, request: Request, user=Depends(require_auth)):
    """Manually set session type."""
    body = await request.json()
    new_type = body.get("type", "general")
    if new_type not in SESSION_TYPES:
        raise HTTPException(400, f"Unbekannter Session-Typ: {new_type}")
    db = get_db()
    try:
        get_or_create_session(db, session_id, user["sub"])
        update_session(db, session_id, session_type=new_type)
        return {"session_id": session_id, "type": new_type, "type_info": SESSION_TYPES[new_type]}
    finally:
        db.close()

@app.post("/api/chat/session/{session_id}/close")
async def close_chat_session(session_id: str, user=Depends(require_auth)):
    """Close/archive a chat session with auto-generated summary."""
    db = get_db()
    try:
        get_or_create_session(db, session_id, user["sub"])

        # Fetch messages for summary
        msgs = db.execute(text(
            "SELECT role, content, created_at FROM chat_messages WHERE session_id = :sid ORDER BY created_at ASC"
        ), {"sid": session_id}).fetchall()

        msg_count = len(msgs)
        user_msgs = [m for m in msgs if m[0] == "user"]
        assistant_msgs = [m for m in msgs if m[0] == "assistant"]

        # Auto-generate title from first user message
        title = ""
        if user_msgs:
            first_q = user_msgs[0][1][:120].strip()
            # Clean up: take first line or sentence
            for sep in ["\n", ". ", "? ", "! "]:
                if sep in first_q:
                    first_q = first_q.split(sep)[0] + sep.strip()[-1] if sep.strip() else first_q.split(sep)[0]
                    break
            title = first_q[:80]

        # Build summary: topics discussed
        topics = []
        for m in user_msgs[:5]:
            snippet = m[1][:60].strip().replace("\n", " ")
            if snippet and snippet not in topics:
                topics.append(snippet)
        summary = " | ".join(topics)[:300] if topics else "Keine Nachrichten"

        # Duration
        duration_min = 0
        if msgs:
            try:
                duration_min = round((msgs[-1][2] - msgs[0][2]).total_seconds() / 60, 1)
            except Exception:
                pass

        # Ensure columns exist, then update
        try:
            db.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS title VARCHAR(120)"))
            db.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS summary TEXT"))
            db.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS msg_count INTEGER DEFAULT 0"))
            db.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS duration_min FLOAT DEFAULT 0"))
            db.execute(text("ALTER TABLE chat_sessions ADD COLUMN IF NOT EXISTS closed_at TIMESTAMP"))
            db.commit()
        except Exception:
            db.rollback()

        # Mark session as closed with summary
        db.execute(text("""
            UPDATE chat_sessions
            SET session_type = 'closed', title = :title, summary = :summary,
                msg_count = :cnt, duration_min = :dur, closed_at = CURRENT_TIMESTAMP,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = :sid
        """), {"sid": session_id, "title": title, "summary": summary, "cnt": msg_count, "dur": duration_min})
        db.commit()
        logger.info(f"Session {session_id} closed by {user['sub']}: {msg_count} msgs, {duration_min}min, title='{title[:40]}'")
        return {
            "session_id": session_id, "status": "closed",
            "title": title, "summary": summary,
            "messages": msg_count, "duration_minutes": duration_min
        }
    except Exception as e:
        db.rollback()
        logger.warning(f"Session close error: {e}")
        return {"session_id": session_id, "status": "closed"}
    finally:
        db.close()

@app.delete("/api/chat/session/{session_id}")
async def delete_chat_session(session_id: str, user=Depends(require_auth)):
    """Permanently delete a chat session and all its messages."""
    db = get_db()
    try:
        # Verify ownership
        sess = db.execute(text(
            "SELECT id FROM chat_sessions WHERE id = :sid AND user_email = :email"
        ), {"sid": session_id, "email": user["sub"]}).fetchone()
        if not sess:
            raise HTTPException(404, "Session not found")

        # Delete messages first, then session
        db.execute(text("DELETE FROM chat_messages WHERE session_id = :sid"), {"sid": session_id})
        db.execute(text("DELETE FROM chat_sessions WHERE id = :sid"), {"sid": session_id})
        db.commit()
        logger.info(f"Session {session_id} deleted by {user['sub']}")
        return {"session_id": session_id, "status": "deleted"}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.warning(f"Session delete error: {e}")
        raise HTTPException(500, str(e))
    finally:
        db.close()


@app.get("/api/chat/session/{session_id}/messages")
async def chat_session_messages(session_id: str, user=Depends(require_auth)):
    """Get full conversation history of a session (own sessions only)."""
    db = get_db()
    try:
        # Verify ownership
        owns = db.execute(text(
            "SELECT COUNT(*) FROM chat_messages WHERE session_id = :sid AND user_email = :email"
        ), {"sid": session_id, "email": user["sub"]}).scalar()
        if not owns:
            raise HTTPException(403, "Kein Zugriff auf diese Session")
        msgs = db.execute(text(
            "SELECT role, content, sources, created_at FROM chat_messages WHERE session_id = :sid ORDER BY created_at ASC"
        ), {"sid": session_id}).fetchall()
        # Session metadata
        sess = db.execute(text(
            "SELECT session_type, title, summary, msg_count, duration_min, created_at, closed_at FROM chat_sessions WHERE id = :sid"
        ), {"sid": session_id}).fetchone()
        meta = {}
        if sess:
            meta = {
                "type": sess[0], "title": sess[1] or "", "summary": sess[2] or "",
                "msg_count": sess[3] or len(msgs), "duration_min": sess[4] or 0,
                "created_at": sess[5].isoformat() if sess[5] else "",
                "closed_at": sess[6].isoformat() if sess[6] else ""
            }
        return {
            "session_id": session_id, "meta": meta,
            "messages": [
                {"role": m[0], "content": m[1], "sources": m[2], "timestamp": m[3].isoformat() if m[3] else ""}
                for m in msgs
            ]
        }
    finally:
        db.close()

@app.get("/api/chat/sessions")
async def chat_sessions(user=Depends(require_auth)):
    """List all chat sessions for the current user. Auto-closes sessions inactive >24h."""
    db = get_db()
    try:
        from sqlalchemy import func as sql_func
        from datetime import timedelta

        # Auto-close sessions inactive for >24h
        try:
            cutoff = datetime.utcnow() - timedelta(hours=24)
            stale = db.execute(text("""
                SELECT cs.id FROM chat_sessions cs
                WHERE cs.user_email = :email
                  AND (cs.session_type IS NULL OR cs.session_type != 'closed')
                  AND cs.updated_at < :cutoff
            """), {"email": user["sub"], "cutoff": cutoff}).fetchall()

            for row in stale:
                sid = row[0]
                # Get message count and first user message for title
                msg_info = db.execute(text("""
                    SELECT COUNT(*), MIN(created_at), MAX(created_at)
                    FROM chat_messages WHERE session_id = :sid
                """), {"sid": sid}).fetchone()
                preview = db.execute(text(
                    "SELECT content FROM chat_messages WHERE session_id = :sid AND role = 'user' ORDER BY created_at ASC LIMIT 1"
                ), {"sid": sid}).fetchone()
                title = (preview[0][:80] if preview else "")
                dur = round((msg_info[2] - msg_info[1]).total_seconds() / 60, 1) if msg_info[1] and msg_info[2] else 0

                db.execute(text("""
                    UPDATE chat_sessions
                    SET session_type = 'closed', title = COALESCE(NULLIF(title,''), :title),
                        summary = COALESCE(NULLIF(summary,''), 'Auto-geschlossen nach 24h Inaktivität'),
                        msg_count = :cnt, duration_min = :dur, closed_at = CURRENT_TIMESTAMP,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE id = :sid
                """), {"sid": sid, "title": title, "cnt": msg_info[0] or 0, "dur": dur})
                logger.info(f"Auto-closed stale session {sid} (>24h inactive)")
            if stale:
                db.commit()
        except Exception as e:
            db.rollback()
            logger.warning(f"Auto-close error: {e}")

        sessions = db.query(
            ChatMessage.session_id,
            sql_func.count(ChatMessage.id).label("msg_count"),
            sql_func.min(ChatMessage.created_at).label("started_at"),
            sql_func.max(ChatMessage.created_at).label("last_msg_at")
        ).filter(
            ChatMessage.user_email == user["sub"],
            ChatMessage.session_id.isnot(None)
        ).group_by(ChatMessage.session_id).order_by(sql_func.max(ChatMessage.created_at).desc()).limit(50).all()
        result = []
        for s in sessions:
            st_row = db.execute(text(
                "SELECT session_type, title, summary, closed_at, updated_at FROM chat_sessions WHERE id = :sid"
            ), {"sid": s.session_id}).fetchone()
            st = st_row[0] if st_row else "general"
            title = (st_row[1] or "") if st_row else ""
            summary = (st_row[2] or "") if st_row else ""
            closed_at = st_row[3].isoformat() if st_row and st_row[3] else None
            updated_at = st_row[4].isoformat() if st_row and len(st_row) > 4 and st_row[4] else s.last_msg_at.isoformat()
            st_info = SESSION_TYPES.get(st, SESSION_TYPES.get("general", {}))

            # Auto-preview from first user message if no title
            if not title:
                preview_row = db.execute(text(
                    "SELECT content FROM chat_messages WHERE session_id = :sid AND role = 'user' ORDER BY created_at ASC LIMIT 1"
                ), {"sid": s.session_id}).fetchone()
                title = (preview_row[0][:80] + "...") if preview_row and len(preview_row[0]) > 80 else (preview_row[0] if preview_row else "")

            duration_sec = (s.last_msg_at - s.started_at).total_seconds() if s.last_msg_at and s.started_at else 0
            result.append({
                "session_id": s.session_id,
                "session_type": st, "session_icon": st_info.get("icon", "💬"),
                "session_label": st_info.get("label", "Allgemein"),
                "title": title, "summary": summary,
                "messages": s.msg_count,
                "started_at": s.started_at.isoformat(),
                "updated_at": updated_at,
                "last_message": s.last_msg_at.isoformat(),
                "closed_at": closed_at,
                "duration_minutes": round(duration_sec / 60, 1),
                "is_closed": st == "closed"
            })
        return result
    finally: db.close()

@app.get("/api/admin/chat/sessions")
async def admin_chat_sessions(limit: int = 100, email: str = None, user=Depends(require_permission("platform.view_analytics"))):
    """Admin: List all chat sessions across users with analytics."""
    db = get_db()
    try:
        from sqlalchemy import func as sql_func
        q = db.query(
            ChatMessage.session_id,
            ChatMessage.user_email,
            sql_func.count(ChatMessage.id).label("msg_count"),
            sql_func.min(ChatMessage.created_at).label("started_at"),
            sql_func.max(ChatMessage.created_at).label("last_msg_at")
        ).filter(ChatMessage.session_id.isnot(None))
        if email:
            q = q.filter(ChatMessage.user_email == email)
        sessions = q.group_by(ChatMessage.session_id, ChatMessage.user_email).order_by(sql_func.max(ChatMessage.created_at).desc()).limit(limit).all()

        # Get first user message per session for preview
        result = []
        for s in sessions:
            preview = db.query(ChatMessage.content).filter(
                ChatMessage.session_id == s.session_id,
                ChatMessage.role == "user"
            ).order_by(ChatMessage.created_at.asc()).first()
            # Get session type from chat_sessions table
            sess_type_row = db.execute(text(
                "SELECT session_type FROM chat_sessions WHERE id = :sid"
            ), {"sid": s.session_id}).fetchone()
            sess_type = sess_type_row[0] if sess_type_row else "general"
            duration_sec = (s.last_msg_at - s.started_at).total_seconds() if s.last_msg_at and s.started_at else 0
            st_info = SESSION_TYPES.get(sess_type, SESSION_TYPES.get("general", {}))
            result.append({
                "session_id": s.session_id,
                "user": s.user_email,
                "session_type": sess_type,
                "session_label": st_info.get("label", "Allgemein"),
                "session_icon": st_info.get("icon", "💬"),
                "messages": s.msg_count,
                "started_at": s.started_at.isoformat(),
                "last_message": s.last_msg_at.isoformat(),
                "duration_minutes": round(duration_sec / 60, 1),
                "preview": (preview.content[:100] + "...") if preview and len(preview.content) > 100 else (preview.content if preview else "")
            })
        return {"sessions": result, "total": len(result)}
    finally: db.close()

@app.get("/api/admin/chat/session/{session_id}")
async def admin_chat_session_detail(session_id: str, user=Depends(require_permission("platform.view_analytics"))):
    """Admin: Get full conversation of a specific session."""
    db = get_db()
    try:
        msgs = db.query(ChatMessage).filter(ChatMessage.session_id == session_id).order_by(ChatMessage.created_at.asc()).all()
        if not msgs:
            raise HTTPException(404, "Session nicht gefunden")
        return {
            "session_id": session_id,
            "user": msgs[0].user_email,
            "messages": [{"id": m.id, "role": m.role, "content": m.content, "sources": m.sources, "created_at": m.created_at.isoformat()} for m in msgs],
            "total_messages": len(msgs),
            "started_at": msgs[0].created_at.isoformat(),
            "ended_at": msgs[-1].created_at.isoformat()
        }
    finally: db.close()

@app.post("/api/upload", response_model=DocumentResponse)
async def upload_file(file: UploadFile = File(...), database: str = Form("knowledge_base"),
                      title: Optional[str] = Form(None), category: Optional[str] = Form(None),
                      language: Optional[str] = Form(None), tags: Optional[str] = Form(None),
                      user=Depends(require_auth)):
    ext = file.filename.split(".")[-1].lower() if file.filename else ""
    if ext not in {"pdf", "txt", "md", "docx", "csv", "json"}: raise HTTPException(400, f"Dateityp .{ext} nicht unterstuetzt")
    content_bytes = await file.read()
    if len(content_bytes) > MAX_FILE_SIZE: raise HTTPException(400, "Datei zu gross (max 50 MB)")
    # Parse tags from JSON string
    parsed_tags = []
    if tags:
        try: parsed_tags = json.loads(tags)
        except: parsed_tags = [t.strip() for t in tags.split(",") if t.strip()]
    # Duplicate check via SHA256 hash
    content_hash = hashlib.sha256(content_bytes).hexdigest()
    db = get_db()
    try:
        existing = db.query(Document).filter(Document.content_hash == content_hash).first()
        if existing:
            from fastapi.responses import JSONResponse
            return JSONResponse(status_code=409, content={
                "duplicate": True,
                "doc_id": str(existing.id),
                "title": existing.title or "Unbekannt",
                "date": existing.created_at.strftime('%d.%m.%Y %H:%M') if existing.created_at else "",
                "category": existing.category or "document",
                "tags": existing.tags if isinstance(existing.tags, list) else (existing.tags.split(",") if existing.tags else []),
                "message": f"Dieses Dokument wurde bereits am {existing.created_at.strftime('%d.%m.%Y')} erfolgreich in die BEATRIX-Wissensdatenbank aufgenommen."
            })
        file_id = str(uuid.uuid4()); file_path = UPLOAD_DIR / f"{file_id}.{ext}"
        with open(file_path, "wb") as f: f.write(content_bytes)
        text_content = extract_text(str(file_path), ext)

        # Book detection: books go to GitHub books/ folder
        book_det = detect_book(text_content)
        if book_det["is_book"]:
            logger.info(f"Book detected for '{file.filename}': score={book_det['score']}, isbn={book_det.get('isbn')}, signals={book_det['signals']}")
            book_meta = extract_paper_metadata(text_content, file.filename)
            book_canonical = generate_canonical_key(book_meta["author"], book_meta["year"], book_meta["title"], "BOOK", ext)
            logger.info(f"Book canonical: '{file.filename}' → '{book_canonical}'")
            gh_result = push_to_github(book_canonical, content_bytes, folder="books") or {}
            github_url = gh_result.get("url", None) if isinstance(gh_result, dict) else None
            gh_status = "indexed+github" if github_url else "indexed"
            doc_title = book_canonical
            doc = Document(
                id=file_id, title=doc_title, content=text_content,
                source_type="file", file_type=ext, file_path=str(file_path),
                file_size=len(content_bytes), database_target=database, status=gh_status,
                github_url=github_url, uploaded_by=user.get("sub"), content_hash=content_hash,
                category=category or "book", language=language or None, tags=parsed_tags or None,
                doc_metadata={
                    "original_filename": file.filename,
                    "content_length": len(text_content),
                    "github": gh_result or {},
                    "book_detected": True,
                    "book_score": book_det["score"],
                    "book_signals": book_det["signals"],
                    "isbn": book_det.get("isbn"),
                    "publisher": book_det.get("publisher"),
                }
            )
            db.add(doc); db.commit(); db.refresh(doc)
            try: embed_document(db, doc.id)
            except Exception as e: logger.warning(f"Embedding failed for {file.filename}: {e}")
            return DocumentResponse(id=doc.id, title=doc.title, source_type=doc.source_type, file_type=doc.file_type, database_target=doc.database_target, status=doc.status, created_at=doc.created_at.isoformat(), github_url=github_url)

        # Study/Report detection: studies go to GitHub studies/ folder
        study_det = detect_study_report(text_content, file.filename)
        if study_det["is_study"]:
            logger.info(f"Study/Report detected for '{file.filename}': score={study_det['score']}, type={study_det.get('study_type')}, org={study_det.get('org')}, signals={study_det['signals']}")
            study_meta = extract_paper_metadata(text_content, file.filename)
            study_type = study_det.get("study_type", "STU").upper()[:4] or "STU"
            study_canonical = generate_canonical_key(study_meta["author"], study_meta["year"], study_meta["title"], study_type, ext)
            logger.info(f"Study canonical: '{file.filename}' → '{study_canonical}'")
            gh_result = push_to_github(study_canonical, content_bytes, folder="studies") or {}
            github_url = gh_result.get("url", None) if isinstance(gh_result, dict) else None
            gh_status = "indexed+github" if github_url else "indexed"
            doc_title = study_canonical
            doc = Document(
                id=file_id, title=doc_title, content=text_content,
                source_type="file", file_type=ext, file_path=str(file_path),
                file_size=len(content_bytes), database_target=database, status=gh_status,
                github_url=github_url, uploaded_by=user.get("sub"), content_hash=content_hash,
                category=category or "study", language=language or None, tags=parsed_tags or None,
                doc_metadata={
                    "original_filename": file.filename,
                    "canonical_key": study_canonical,
                    "content_length": len(text_content),
                    "github": gh_result or {},
                    "study_detected": True,
                    "study_score": study_det["score"],
                    "study_signals": study_det["signals"],
                    "study_type": study_det.get("study_type"),
                    "organization": study_det.get("org"),
                }
            )
            db.add(doc); db.commit(); db.refresh(doc)
            try: embed_document(db, doc.id)
            except Exception as e: logger.warning(f"Embedding failed for {file.filename}: {e}")
            return DocumentResponse(id=doc.id, title=doc.title, source_type=doc.source_type, file_type=doc.file_type, database_target=doc.database_target, status=doc.status, created_at=doc.created_at.isoformat(), github_url=github_url)

        # Paper detection: only push papers to GitHub papers/evaluated/integrated/
        detection = detect_scientific_paper(text_content)
        logger.info(f"File upload paper detection for '{file.filename}': score={detection['score']}, is_paper={detection['is_paper']}, signals={detection['signals']}")

        gh_result = {}
        github_url = None
        canonical_key = None
        paper_meta = None
        if detection["is_paper"]:
            paper_meta = extract_paper_metadata(text_content, file.filename)
            canonical_key = generate_canonical_key(paper_meta["author"], paper_meta["year"], paper_meta["title"], paper_meta["doc_type"], ext)
            logger.info(f"Paper canonical: '{file.filename}' → '{canonical_key}'")
            gh_result = push_to_github(canonical_key, content_bytes) or {}
            github_url = gh_result.get("url", None) if isinstance(gh_result, dict) else None
            logger.info(f"Paper → GitHub push: {github_url}")

        gh_status = "indexed+github" if github_url else "indexed"
        doc_title = canonical_key or title or file.filename or "Unnamed"
        doc = Document(
            id=file_id, title=doc_title, content=text_content,
            source_type="file", file_type=ext, file_path=str(file_path),
            file_size=len(content_bytes), database_target=database, status=gh_status,
            github_url=github_url, uploaded_by=user.get("sub"), content_hash=content_hash,
            category=category or None, language=language or None, tags=parsed_tags or None,
            doc_metadata={
                "original_filename": file.filename,
                "canonical_key": canonical_key,
                "content_length": len(text_content),
                "github": gh_result or {},
                "paper_detected": detection["is_paper"],
                "paper_score": detection["score"],
                "paper_signals": detection["signals"],
                **({"extracted_author": paper_meta["author"], "extracted_year": paper_meta["year"], "extracted_title": paper_meta["title"], "doc_type": paper_meta["doc_type"]} if paper_meta else {})
            }
        )
        db.add(doc); db.commit(); db.refresh(doc)
        # Auto-embed for vector search
        try:
            embed_document(db, doc.id)
        except Exception as e:
            logger.warning(f"Embedding failed for {file.filename}: {e}")
        return DocumentResponse(id=doc.id, title=doc.title, source_type=doc.source_type, file_type=doc.file_type, database_target=doc.database_target, status=doc.status, created_at=doc.created_at.isoformat(), github_url=doc.github_url)
    except HTTPException: raise
    except Exception as e: db.rollback(); raise HTTPException(500, f"Datenbankfehler: {e}")
    finally: db.close()

@app.post("/api/text", response_model=DocumentResponse)
async def upload_text(request: TextUploadRequest, user=Depends(require_auth)):
    if not request.content.strip(): raise HTTPException(400, "Inhalt darf nicht leer sein")

    # Auto-extract title from content if not provided
    if not request.title or request.title in ["Untitled", ""]:
        lines = [l.strip() for l in request.content[:2000].split('\n') if l.strip() and len(l.strip()) > 3]
        # Skip common non-title headers
        skip_patterns = ["nber working paper", "working paper series", "discussion paper",
            "technical report", "research paper", "policy brief", "staff report",
            "occasional paper", "conference paper", "working paper no", "wp no",
            "issn", "isbn", "http", "www.", "doi:", "jel class", "keywords:",
            "abstract", "table of contents", "acknowledgment"]
        title_found = False
        for line in lines:
            lower = line.lower().strip()
            if any(skip in lower for skip in skip_patterns):
                continue
            if len(line) < 5:
                continue
            # Skip lines that are just numbers or dates
            if line.replace('.', '').replace('-', '').replace('/', '').strip().isdigit():
                continue
            request.title = line[:150]
            title_found = True
            break
        if not title_found:
            request.title = lines[0][:120] if lines else "Untitled"

    # Duplicate check via SHA256 hash of content
    content_hash = hashlib.sha256(request.content.encode("utf-8")).hexdigest()
    db = get_db()
    try:
        existing = db.query(Document).filter(Document.content_hash == content_hash).first()
        if existing:
            from fastapi.responses import JSONResponse
            return JSONResponse(status_code=409, content={
                "duplicate": True,
                "doc_id": str(existing.id),
                "title": existing.title or "Unbekannt",
                "date": existing.created_at.strftime('%d.%m.%Y %H:%M') if existing.created_at else "",
                "category": existing.category or "document",
                "tags": existing.tags if isinstance(existing.tags, list) else (existing.tags.split(",") if existing.tags else []),
                "message": f"Dieses Dokument wurde bereits am {existing.created_at.strftime('%d.%m.%Y')} erfolgreich in die BEATRIX-Wissensdatenbank aufgenommen."
            })

        # Paper detection: only push papers to GitHub papers/evaluated/integrated/
        detection = detect_scientific_paper(request.content, request.title)
        logger.info(f"Text upload paper detection for '{request.title}': score={detection['score']}, is_paper={detection['is_paper']}, signals={detection['signals']}")

        gh_result = {}
        github_url = None
        canonical_key = None
        paper_meta = None
        if detection["is_paper"]:
            paper_meta = extract_paper_metadata(request.content, request.title or "")
            canonical_key = generate_canonical_key(paper_meta["author"], paper_meta["year"], paper_meta["title"], paper_meta["doc_type"], "txt")
            logger.info(f"Text paper canonical: '{request.title}' → '{canonical_key}'")
            gh_result = push_to_github(canonical_key, request.content.encode("utf-8")) or {}
            github_url = gh_result.get("url", None) if isinstance(gh_result, dict) else None
            logger.info(f"Paper → GitHub push: {github_url}")

        status = "indexed+github" if github_url else "indexed"
        doc_title = canonical_key or request.title
        doc = Document(
            title=doc_title, content=request.content, source_type="text",
            database_target=request.database or "knowledge_base",
            category=request.category, language=request.language, tags=request.tags,
            status=status, github_url=github_url, uploaded_by=user.get("sub"),
            content_hash=content_hash,
            doc_metadata={
                "original_title": request.title,
                "canonical_key": canonical_key,
                "content_length": len(request.content),
                "github": gh_result or {},
                "paper_detected": detection["is_paper"],
                "paper_score": detection["score"],
                "paper_signals": detection["signals"],
                **({"extracted_author": paper_meta["author"], "extracted_year": paper_meta["year"], "extracted_title": paper_meta["title"], "doc_type": paper_meta["doc_type"]} if paper_meta else {})
            }
        )
        db.add(doc); db.commit(); db.refresh(doc)
        # Auto-embed for vector search
        embed_document(db, doc.id)
        return DocumentResponse(id=doc.id, title=doc.title, source_type=doc.source_type, file_type=None, database_target=doc.database_target, status=doc.status, created_at=doc.created_at.isoformat(), github_url=doc.github_url)
    except HTTPException: raise
    except Exception as e: db.rollback(); raise HTTPException(500, f"Datenbankfehler: {e}")
    finally: db.close()

@app.post("/api/canonical-key")
async def api_canonical_key(request: Request, user=Depends(require_auth)):
    """Generate canonical filename for a paper/book/study.
    Body: {author, year, title, doc_type?, ext?} or {text, filename?} for auto-extraction."""
    body = await request.json()

    # If text provided, auto-extract metadata
    if body.get("text"):
        meta = extract_paper_metadata(body["text"], body.get("filename", ""))
        author = body.get("author") or meta["author"]
        year = body.get("year") or meta["year"]
        title = body.get("title") or meta["title"]
        doc_type = body.get("doc_type") or meta["doc_type"]
    else:
        author = body.get("author", "")
        year = body.get("year", "")
        title = body.get("title", "")
        doc_type = body.get("doc_type", "PAP")

    ext = body.get("ext", "pdf")
    canonical = generate_canonical_key(author, year, title, doc_type, ext)

    return {
        "canonical_key": canonical,
        "author": author, "year": year, "title": title,
        "doc_type": doc_type, "ext": ext
    }


# ========== PAPER UPLOAD PRIORITY TOP 100 ==========

_paper_priority_cache = {"data": None, "ts": 0}

@app.get("/api/papers/upload-priority")
async def get_paper_upload_priority(user=Depends(require_auth)):
    """Serve the Top 100 paper upload priority list from GitHub (cached 5 min)."""
    import time
    now = time.time()
    if _paper_priority_cache["data"] and (now - _paper_priority_cache["ts"]) < 300:
        return _paper_priority_cache["data"]

    if not GH_TOKEN:
        raise HTTPException(500, "GitHub not configured")
    import urllib.request as ur, ssl
    ctx2 = ssl.create_default_context(); ctx2.check_hostname = False; ctx2.verify_mode = ssl.CERT_NONE
    gh = {"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3+json", "User-Agent": "BEATRIXLab"}
    try:
        url = f"https://api.github.com/repos/{GH_REPO}/contents/data/paper-upload-priority-top100.json"
        req = ur.Request(url, headers=gh)
        resp = json.loads(ur.urlopen(req, context=ctx2, timeout=15).read())
        data = json.loads(base64.b64decode(resp["content"]).decode())
        _paper_priority_cache["data"] = data
        _paper_priority_cache["ts"] = now
        return data
    except Exception as e:
        logger.error(f"Paper priority fetch failed: {e}")
        raise HTTPException(500, f"Konnte Priority-Liste nicht laden: {e}")


# ========== PAPER FINDER SOURCES (verified researcher PDFs) ==========

_paper_sources_cache = {"data": None, "ts": 0}

@app.get("/api/papers/sources")
async def get_paper_finder_sources(user=Depends(require_auth)):
    """Serve verified researcher PDF sources from GitHub (cached 5 min).
    Contains 352+ verified PDF URLs across 20+ researchers with discovery algorithm v3."""
    import time
    now = time.time()
    if _paper_sources_cache["data"] and (now - _paper_sources_cache["ts"]) < 300:
        return _paper_sources_cache["data"]

    if not GH_TOKEN:
        raise HTTPException(500, "GitHub not configured")
    import urllib.request as ur, ssl
    ctx2 = ssl.create_default_context(); ctx2.check_hostname = False; ctx2.verify_mode = ssl.CERT_NONE
    gh = {"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3+json", "User-Agent": "BEATRIXLab"}
    try:
        url = f"https://api.github.com/repos/{GH_REPO}/contents/data/paper-finder-sources.json"
        req = ur.Request(url, headers=gh)
        resp = json.loads(ur.urlopen(req, context=ctx2, timeout=15).read())
        data = json.loads(base64.b64decode(resp["content"]).decode())
        _paper_sources_cache["data"] = data
        _paper_sources_cache["ts"] = now
        return data
    except Exception as e:
        logger.error(f"Paper sources fetch failed: {e}")
        raise HTTPException(500, f"Konnte Sources nicht laden: {e}")

@app.get("/api/papers/sources/{researcher}")
async def get_researcher_sources(researcher: str, user=Depends(require_auth)):
    """Get verified PDF sources for a specific researcher."""
    all_sources = await get_paper_finder_sources(user)
    researchers = all_sources.get("researchers", {})
    # Try exact match first, then fuzzy
    if researcher in researchers:
        return {"researcher": researcher, **researchers[researcher]}
    # Fuzzy: try replacing hyphens/spaces
    normalized = researcher.lower().replace("-", "_").replace(" ", "_")
    for key, val in researchers.items():
        if key.lower() == normalized:
            return {"researcher": key, **val}
    available = sorted(researchers.keys())
    raise HTTPException(404, f"Researcher '{researcher}' not found. Available: {', '.join(available)}")


@app.post("/api/papers/find-pdf")
async def find_paper_pdf(request: Request, user=Depends(require_auth)):
    """Search for exact paper PDF. Strategy: exact title web search first, then DOI-based APIs."""
    import urllib.request as ur, ssl, urllib.parse, re
    body = await request.json()
    doi = body.get("doi", "")
    title = body.get("title", "")
    author = body.get("author", "")
    year = body.get("year", "")

    ctx2 = ssl.create_default_context(); ctx2.check_hostname = False; ctx2.verify_mode = ssl.CERT_NONE
    results = []
    seen = set()

    def add_result(source, url, rtype="pdf"):
        if url and url not in seen:
            seen.add(url)
            results.append({"source": source, "url": url, "type": rtype})

    def ddg_search(query, max_pdfs=5):
        """Search DuckDuckGo and return PDF URLs."""
        try:
            ddg_url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(query)}"
            req = ur.Request(ddg_url, headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"})
            resp_bytes = ur.urlopen(req, context=ctx2, timeout=10).read().decode(errors='replace')
            raw_links = re.findall(r'uddg=([^&"]+)', resp_bytes)
            found = []
            for raw in raw_links:
                link = urllib.parse.unquote(raw)
                if '.pdf' in link.lower() and link.startswith('http') and link not in seen:
                    domain = link.split('/')[2] if len(link.split('/')) > 2 else ''
                    short_domain = '.'.join(domain.split('.')[-2:]) if domain else 'Web'
                    found.append((f"Web ({short_domain})", link))
                    if len(found) >= max_pdfs:
                        break
            return found
        except Exception as e:
            logger.debug(f"DDG search error: {e}")
            return []

    # === STRATEGY: Exact title search first, then DOI-based APIs ===

    # 1. DuckDuckGo EXACT title search (finds university-hosted PDFs: Stanford, UCLA, UCSD etc.)
    if title:
        clean_title = title.split(":")[0].strip() if len(title.split(":")[0]) > 20 else title
        exact_q = '"' + clean_title + '" filetype:pdf'
        logger.info(f"PDF-Finder Step 1: DDG exact → {exact_q[:80]}")
        for source, url in ddg_search(exact_q, max_pdfs=4):
            add_result(source, url)

    # 2. DuckDuckGo author+year+title (broader, catches different title formats)
    if len([r for r in results if r["type"] == "pdf"]) < 3:
        import time as _time; _time.sleep(0.5)  # Small delay to avoid DDG rate limiting
        broad_q = f"{author} {year} {title} filetype:pdf".strip()[:150]
        logger.info(f"PDF-Finder Step 2: DDG broad → {broad_q[:80]}")
        for source, url in ddg_search(broad_q, max_pdfs=3):
            add_result(source, url)

    # 3. Unpaywall (DOI only – reliable official OA versions)
    if doi:
        try:
            url = f"https://api.unpaywall.org/v2/{urllib.parse.quote(doi, safe='')}?email=gerhard.fehr@fehradvice.com"
            req = ur.Request(url, headers={"User-Agent": "BEATRIX/1.0"})
            resp = json.loads(ur.urlopen(req, context=ctx2, timeout=8).read())
            if resp.get("is_oa"):
                best = resp.get("best_oa_location") or {}
                pdf_url = best.get("url_for_pdf") or best.get("url")
                if pdf_url:
                    add_result("Unpaywall", pdf_url, "pdf" if best.get("url_for_pdf") else "landing")
                for loc in resp.get("oa_locations", [])[:3]:
                    loc_url = loc.get("url_for_pdf") or loc.get("url")
                    if loc_url:
                        add_result(f"Unpaywall ({loc.get('host_type','')})", loc_url, "pdf" if loc.get("url_for_pdf") else "landing")
        except Exception as e:
            logger.debug(f"Unpaywall error: {e}")

    # 4. Semantic Scholar (DOI only! Title search returns wrong papers)
    if doi:
        try:
            url = f"https://api.semanticscholar.org/graph/v1/paper/DOI:{urllib.parse.quote(doi, safe='')}?fields=openAccessPdf,title"
            req = ur.Request(url, headers={"User-Agent": "BEATRIX/1.0"})
            resp = json.loads(ur.urlopen(req, context=ctx2, timeout=8).read())
            oa = resp.get("openAccessPdf") or {}
            if oa.get("url"):
                add_result("Semantic Scholar", oa["url"])
        except Exception as e:
            logger.debug(f"S2 DOI error: {e}")

    # 5. CORE (DOI only!)
    if doi:
        try:
            url = f"https://api.core.ac.uk/v3/search/works?q=doi:{urllib.parse.quote(doi, safe='')}%limit=2"
            req = ur.Request(url, headers={"User-Agent": "BEATRIX/1.0"})
            resp = json.loads(ur.urlopen(req, context=ctx2, timeout=8).read())
            for hit in resp.get("results", [])[:2]:
                dl = hit.get("downloadUrl")
                if dl:
                    add_result("CORE", dl)
        except Exception as e:
            logger.debug(f"CORE error: {e}")

    # 6. Fallback links
    if doi:
        add_result("DOI.org", f"https://doi.org/{doi}", "landing")
        if "nber" in doi.lower() or "3386" in doi:
            nber_id = doi.split("/")[-1].replace("w","")
            add_result("NBER", f"https://www.nber.org/papers/w{nber_id}", "landing")

    scholar_q = urllib.parse.quote(f'{author} {year} {title}'[:150])
    add_result("Google Scholar", f"https://scholar.google.com/scholar?q={scholar_q}", "search")

    return {
        "doi": doi, "title": title, "author": author, "year": year,
        "found": len([r for r in results if r["type"] == "pdf"]),
        "results": results
    }

@app.post("/api/papers/fetch-and-upload")
async def fetch_and_upload_pdf(request: Request, user=Depends(require_auth)):
    """Fetch a PDF from URL → full pipeline: extract text, detect paper, canonical name, GitHub, DB, embed."""
    import urllib.request as ur, ssl
    body = await request.json()
    pdf_url = body.get("url", "")
    paper_id = body.get("paper_id", "")
    hint_title = body.get("title", "")
    hint_author = body.get("author", "")
    hint_year = body.get("year", "")
    doi = body.get("doi", "")

    if not pdf_url:
        raise HTTPException(400, "URL required")

    ctx2 = ssl.create_default_context(); ctx2.check_hostname = False; ctx2.verify_mode = ssl.CERT_NONE

    # 1. Fetch file from URL
    try:
        req = ur.Request(pdf_url, headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) BEATRIX/1.0"})
        resp = ur.urlopen(req, context=ctx2, timeout=30)
        content_bytes = resp.read()
        content_type = resp.headers.get("Content-Type", "")
        if len(content_bytes) < 500:
            raise HTTPException(400, f"File too small ({len(content_bytes)} bytes)")
        if len(content_bytes) > MAX_FILE_SIZE:
            raise HTTPException(400, "File too large (>50MB)")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Could not fetch: {e}")

    # === HTML-to-PDF RESOLUTION (Caltech InvenioRDM, etc.) ===
    # If we got an HTML page instead of a PDF, try to extract the real PDF URL
    # from the citation_pdf_url meta tag (standard for academic repositories).
    is_html = 'html' in content_type.lower() or content_bytes[:50].lower().startswith((b'<!doctype', b'<html'))
    if is_html:
        import re as _re
        html_text = content_bytes.decode('utf-8', errors='ignore')
        # Look for <meta name="citation_pdf_url" content="https://...pdf">
        pdf_meta = _re.search(r'citation_pdf_url.+?content="([^"]+)"', html_text)
        if not pdf_meta:
            pdf_meta = _re.search(r"citation_pdf_url.+?content='([^']+)'", html_text)
        if pdf_meta:
            resolved_url = pdf_meta.group(1)
            logger.info(f"HTML-to-PDF resolved: {pdf_url[:60]} -> {resolved_url[:80]}")
            try:
                req2 = ur.Request(resolved_url, headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) BEATRIX/1.0"})
                resp2 = ur.urlopen(req2, context=ctx2, timeout=30)
                content_bytes = resp2.read()
                content_type = resp2.headers.get("Content-Type", "")
                pdf_url = resolved_url
                if len(content_bytes) < 500:
                    raise HTTPException(400, f"Resolved PDF too small ({len(content_bytes)} bytes)")
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(400, f"Could not fetch resolved PDF: {e}")
        else:
            raise HTTPException(400, "URL returned HTML instead of PDF. No citation_pdf_url meta tag found. Please use a direct PDF link.")


    is_pdf = content_bytes[:5] == b'%PDF-' or 'pdf' in content_type.lower()
    ext = "pdf" if is_pdf else "txt"
    original_filename = pdf_url.split("/")[-1].split("?")[0] or f"fetched.{ext}"

    # 2. Duplicate check
    content_hash = hashlib.sha256(content_bytes).hexdigest()
    db = get_db()
    try:
        existing = db.query(Document).filter(Document.content_hash == content_hash).first()
        if existing:
            return JSONResponse(status_code=409, content={
                "duplicate": True, "doc_id": str(existing.id),
                "title": existing.title or "Unbekannt",
                "date": existing.created_at.strftime('%d.%m.%Y %H:%M') if existing.created_at else "",
                "message": f"Bereits in der Datenbank: {existing.title}"
            })

        # 3. Save file locally
        file_id = str(uuid.uuid4())
        file_path = UPLOAD_DIR / f"{file_id}.{ext}"
        with open(file_path, "wb") as f:
            f.write(content_bytes)

        # 4. Extract text
        text_content = extract_text(str(file_path), ext)

        # 4b. TITLE VERIFICATION – check the fetched PDF is actually the right paper
        if hint_title and text_content:
            # Extract significant words from expected title (>3 chars, not stopwords)
            stopwords = {"the","and","for","with","from","that","this","into","than","also","have","been","were","they","their","which","about","would","there","could","other","more","some","these","only","over","such","after","most","then","them","will","each","made","like","long","many","much","very","just","before","between","through","during","without","within","against","among","across","under","above","below","along","since","upon","around","toward","towards"}
            title_words = [w.lower() for w in hint_title.split() if len(w) > 3 and w.lower() not in stopwords]
            if title_words:
                text_lower = text_content[:5000].lower()  # Check first 5000 chars (title page)
                matched = sum(1 for w in title_words if w in text_lower)
                match_ratio = matched / len(title_words) if title_words else 0
                logger.info(f"Title verification: {matched}/{len(title_words)} words matched ({match_ratio:.0%}) for '{hint_title[:50]}'")
                if match_ratio < 0.4:
                    # This is probably the WRONG paper – reject it
                    logger.warning(f"Title verification FAILED: expected '{hint_title[:50]}' but PDF text doesn't match (ratio={match_ratio:.0%})")
                    # Clean up temp file
                    try: os.remove(file_path)
                    except: pass
                    return JSONResponse(status_code=422, content={
                        "verified": False,
                        "match_ratio": round(match_ratio, 2),
                        "expected_title": hint_title,
                        "message": f"PDF enthält nicht das erwartete Paper. Titel-Match: {match_ratio:.0%}. Möglicherweise ein anderes Paper."
                    })

        # 5. Detect paper
        detection = detect_scientific_paper(text_content, original_filename)
        logger.info(f"Auto-fetch paper detection for '{original_filename}': score={detection['score']}, is_paper={detection['is_paper']}")

        # 6. Extract metadata + canonical key
        paper_meta = extract_paper_metadata(text_content, original_filename)
        # Override with hints from priority list if extraction failed
        if paper_meta["author"] in ("Unknown", "") and hint_author:
            paper_meta["author"] = hint_author
        if paper_meta["year"] in ("0000", "", "0") and hint_year:
            paper_meta["year"] = str(hint_year)
        if not paper_meta["title"] and hint_title:
            paper_meta["title"] = hint_title

        canonical_key = generate_canonical_key(
            paper_meta["author"], paper_meta["year"], paper_meta["title"],
            paper_meta.get("doc_type", "PAP"), ext
        )
        logger.info(f"Auto-fetch canonical: '{original_filename}' → '{canonical_key}'")

        # 7. Push to GitHub
        gh_result = push_to_github(canonical_key, content_bytes) or {}
        github_url = gh_result.get("url", None) if isinstance(gh_result, dict) else None
        gh_status = "indexed+github" if github_url else "indexed"

        # 8. Create DB entry
        doc = Document(
            id=file_id, title=canonical_key, content=text_content,
            source_type="file", file_type=ext, file_path=str(file_path),
            file_size=len(content_bytes), database_target="research",
            status=gh_status, github_url=github_url,
            uploaded_by=user.get("sub"), content_hash=content_hash,
            category="paper" if detection["is_paper"] else "document",
            language="en", tags=["auto-fetch", "priority-upload"],
            doc_metadata={
                "original_filename": original_filename,
                "fetch_url": pdf_url,
                "canonical_key": canonical_key,
                "content_length": len(text_content),
                "github": gh_result or {},
                "paper_detected": detection["is_paper"],
                "paper_score": detection["score"],
                "paper_signals": detection["signals"],
                "extracted_author": paper_meta["author"],
                "extracted_year": paper_meta["year"],
                "extracted_title": paper_meta["title"],
                "doc_type": paper_meta.get("doc_type", "PAP"),
                "doi": doi,
                "paper_id": paper_id,
                "hint_author": hint_author,
                "hint_year": hint_year,
                "hint_title": hint_title,
            }
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        # 9. Embed for vector search
        try:
            embed_document(db, doc.id)
        except Exception as e:
            logger.warning(f"Embedding failed for {canonical_key}: {e}")

        return {
            "ok": True,
            "doc_id": doc.id,
            "filename": canonical_key,
            "size": len(content_bytes),
            "text_length": len(text_content),
            "is_pdf": is_pdf,
            "is_paper": detection["is_paper"],
            "paper_score": detection["score"],
            "status": gh_status,
            "github_url": github_url,
            "html_url": gh_result.get("url", ""),
            "extracted": {
                "author": paper_meta["author"],
                "year": paper_meta["year"],
                "title": paper_meta["title"][:100]
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Auto-fetch pipeline error: {e}")
        raise HTTPException(500, f"Pipeline-Fehler: {e}")
    finally:
        db.close()




@app.get("/api/pdf-view")
async def pdf_view_proxy(url: str = ""):
    """Proxy a PDF and serve it with Content-Disposition: inline so Safari/iPad displays it."""
    if not url:
        raise HTTPException(400, "url parameter required")
    from urllib.parse import urlparse, quote, urlunparse
    import urllib.request as ur, ssl
    parsed = urlparse(url)
    # Re-encode spaces that FastAPI/Starlette decoded
    clean_url = urlunparse(parsed._replace(path=quote(parsed.path, safe='/')))
    domain = parsed.hostname or ""
    allowed = ['caltech.edu', 'uzh.ch', 'nber.org', 'ssrn.com', 'arxiv.org', 'iza.org',
               'princeton.edu', 'harvard.edu', 'stanford.edu', 'mit.edu', 'uchicago.edu',
               'yale.edu', 'berkeley.edu', 'upenn.edu', 'duke.edu', 'cmu.edu', 'cornell.edu',
               'ucla.edu', 'ucsd.edu', 'lse.ac.uk', 'warwick.ac.uk', 'nottingham.ac.uk',
               'unibocconi.it', 'unifr.ch', 'ceu.edu', 'pitt.edu', 'amazonaws.com']
    if not any(domain.endswith(a) for a in allowed):
        raise HTTPException(403, f"Domain not allowed: {domain}")
    ctx2 = ssl.create_default_context(); ctx2.check_hostname = False; ctx2.verify_mode = ssl.CERT_NONE
    try:
        req = ur.Request(clean_url, headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 BEATRIX/1.0"})
        resp = ur.urlopen(req, context=ctx2, timeout=60)
        content_type = resp.headers.get("Content-Type", "application/pdf")
        # Stream the response in chunks
        from starlette.responses import StreamingResponse
        def stream():
            while True:
                chunk = resp.read(65536)
                if not chunk:
                    break
                yield chunk
        fname = clean_url.split("/")[-1].split("?")[0] or "paper.pdf"
        return StreamingResponse(
            stream(),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'inline; filename="{fname}"',
                "Cache-Control": "public, max-age=3600",
                "X-Content-Type-Options": "nosniff"
            }
        )
    except Exception as e:
        raise HTTPException(502, f"Could not fetch PDF: {e}")


@app.post("/api/text/analyze")
async def analyze_text(request: Request, user=Depends(require_auth)):
    """Analyze pasted text: detect paper, classify Content Level (L0-L3),
    estimate Integration Level (I1-I5), extract EBF metadata via Claude."""
    body = await request.json()
    text = (body.get("text", "") or "").strip()
    total_length = body.get("total_length", len(text))  # Frontend sends actual text length
    if not text or len(text) < 100:
        return {"ok": False, "error": "Text zu kurz für Analyse"}
    return await _run_text_analysis(text, total_length)

@app.post("/api/file/analyze")
async def analyze_file(file: UploadFile = File(...), user=Depends(require_auth)):
    """Analyze an uploaded file (PDF, DOCX, TXT etc.) without storing it.
    Extracts text, runs paper detection, Content Level, Integration Level, EBF classification."""
    ext = file.filename.split(".")[-1].lower() if file.filename else ""
    if ext not in {"pdf", "txt", "md", "docx", "csv", "json"}:
        return {"ok": False, "error": f"Dateityp .{ext} nicht unterstützt"}
    content_bytes = await file.read()
    if len(content_bytes) > MAX_FILE_SIZE:
        return {"ok": False, "error": "Datei zu groß (max 50 MB)"}
    # Save to temp file for extraction
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
        tmp.write(content_bytes)
        tmp_path = tmp.name
    try:
        text = extract_text(tmp_path, ext)
        if not text or len(text) < 50:
            return {"ok": False, "error": "Kein Text extrahierbar"}
        total_length = len(text)
        # Send first 10k for analysis (same as text preview), full length for Content Level
        preview_text = text[:10000]
        result = await _run_text_analysis(preview_text, total_length, filename=file.filename)
        result["filename"] = file.filename
        result["file_ext"] = ext
        result["file_size"] = len(content_bytes)
        return result
    except Exception as e:
        logger.error(f"File analyze error for {file.filename}: {e}")
        return {"ok": False, "error": f"Analyse fehlgeschlagen: {str(e)}"}
    finally:
        import os
        try: os.unlink(tmp_path)
        except: pass

async def _run_text_analysis(text: str, total_length: int, filename: str = "") -> dict:
    """Shared analysis logic for both text and file analyze endpoints."""

    # ═══════════════════════════════════════════════════════
    # STEP 0a: Book detection (runs BEFORE paper detection)
    # ═══════════════════════════════════════════════════════
    book_detection = detect_book(text)
    if book_detection["is_book"]:
        # ISBN lookup for enrichment
        isbn_data = lookup_isbn(book_detection["isbn"]) if book_detection.get("isbn") else {}

        # Use Claude for book metadata if available
        title = isbn_data.get("title", "")
        language = "en"
        tags = isbn_data.get("subjects", [])[:5]
        summary = ""
        authors = isbn_data.get("authors", [])
        publisher = isbn_data.get("publishers", [None])[0] or book_detection.get("publisher", "")
        publish_date = isbn_data.get("publish_date", "")
        pages = isbn_data.get("number_of_pages")

        if ANTHROPIC_API_KEY:
            try:
                import urllib.request, ssl
                ctx = ssl.create_default_context()
                ctx.check_hostname = False
                ctx.verify_mode = ssl.CERT_NONE
                excerpt = text[:4000]
                prompt = f"""Analysiere dieses Buch. Antworte NUR mit validem JSON:
{{
  "title": "exakter Buchtitel",
  "authors": ["Nachname, Vorname", ...],
  "language": "de" oder "en",
  "tags": ["max 5 Schlagwörter"],
  "summary": "1-2 Sätze",
  "year": 2024,
  "publisher": "Verlag",
  "domains": ["Domänen wie finance, health, behavior, management, economics, general"]
}}

TEXT:
{excerpt}"""
                payload = json.dumps({
                    "model": ANTHROPIC_MODEL_STANDARD,
                    "max_tokens": 400,
                    "messages": [{"role": "user", "content": prompt}]
                }).encode()
                req = urllib.request.Request("https://api.anthropic.com/v1/messages",
                    data=payload, method="POST",
                    headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                              "Content-Type": "application/json"})
                resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
                raw = resp.get("content", [{}])[0].get("text", "")
                import re
                json_match = re.search(r'\{[^{}]*\}', raw, re.DOTALL)
                if json_match:
                    meta = json.loads(json_match.group())
                    if not title: title = meta.get("title", "")
                    if not authors: authors = meta.get("authors", [])
                    language = meta.get("language", "en")
                    if not tags: tags = meta.get("tags", [])
                    summary = meta.get("summary", "")
                    if not publisher: publisher = meta.get("publisher", "")
                    if not publish_date and meta.get("year"):
                        publish_date = str(meta["year"])
            except Exception as e:
                logger.warning(f"Book analysis Claude error: {e}")

        if not title:
            lines = [l.strip() for l in text[:500].split('\n') if l.strip()]
            title = lines[0][:100] if lines else "Untitled"

        return {
            "ok": True,
            "title": title,
            "language": language,
            "category": "book",
            "tags": tags[:5],
            "summary": summary,
            "database": "knowledge_base",
            "chars": total_length,
            # Book detection
            "book_detected": True,
            "book_score": book_detection["score"],
            "book_signals": book_detection["signals"][:8],
            "isbn": book_detection.get("isbn"),
            "isbn_data": isbn_data if isbn_data else None,
            # Book metadata
            "book_meta": {
                "authors": authors,
                "publisher": publisher,
                "publish_date": publish_date,
                "pages": pages,
                "cover_url": isbn_data.get("cover", ""),
                "openlibrary_url": isbn_data.get("url", ""),
            },
            # Not a paper
            "paper_detected": False,
            "paper_score": 0,
            "paper_signals": [],
            "content_level": None,
            "integration_level": None,
            "structural": None,
            "ebf": None,
            "paper_id": None,
        }

    # ═══════════════════════════════════════════════════════
    # STEP 0b: Study/Report detection
    # ═══════════════════════════════════════════════════════
    study_detection = detect_study_report(text, filename)
    if study_detection["is_study"]:
        # Use Claude for study metadata
        title = ""
        language = "en"
        tags = []
        summary = ""
        authors = []
        org = study_detection.get("org", "")
        study_type = study_detection.get("study_type", "report")

        study_type_labels = {
            "consulting": "Consulting-Studie",
            "think_tank": "Think-Tank-Report",
            "international_org": "Internationale Studie",
            "government": "Regierungsstudie",
            "policy_brief": "Policy Brief",
            "industry_report": "Branchenreport",
            "survey_study": "Umfrage-Studie",
            "report": "Report/Studie",
        }
        study_type_label = study_type_labels.get(study_type, "Report/Studie")

        if ANTHROPIC_API_KEY:
            try:
                import urllib.request, ssl
                ctx = ssl.create_default_context()
                ctx.check_hostname = False
                ctx.verify_mode = ssl.CERT_NONE
                excerpt = text[:4000]
                prompt = f"""Analysiere diese Studie/diesen Report. Antworte NUR mit validem JSON:
{{
  "title": "exakter Titel der Studie/des Reports",
  "authors": ["Nachname, Vorname" oder "Organisation"],
  "language": "de" oder "en",
  "tags": ["max 5 Schlagwörter"],
  "summary": "1-2 Sätze Zusammenfassung",
  "year": 2024,
  "organization": "herausgebende Organisation",
  "study_type": "consulting|government|think_tank|industry|survey|policy|whitepaper"
}}

TEXT:
{excerpt}"""
                payload = json.dumps({
                    "model": ANTHROPIC_MODEL_STANDARD,
                    "max_tokens": 400,
                    "messages": [{"role": "user", "content": prompt}]
                }).encode()
                req = urllib.request.Request("https://api.anthropic.com/v1/messages",
                    data=payload, method="POST",
                    headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                              "Content-Type": "application/json"})
                resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
                raw = resp.get("content", [{}])[0].get("text", "")
                import re
                json_match = re.search(r'\{[^{}]*\}', raw, re.DOTALL)
                if json_match:
                    meta = json.loads(json_match.group())
                    title = meta.get("title", "")
                    authors = meta.get("authors", [])
                    language = meta.get("language", "en")
                    tags = meta.get("tags", [])
                    summary = meta.get("summary", "")
                    if not org: org = meta.get("organization", "")
            except Exception as e:
                logger.warning(f"Study analysis Claude error: {e}")

        if not title:
            lines = [l.strip() for l in text[:500].split('\n') if l.strip()]
            title = lines[0][:100] if lines else "Untitled"

        return {
            "ok": True,
            "title": title,
            "language": language,
            "category": "study",
            "tags": tags[:5],
            "summary": summary,
            "database": "knowledge_base",
            "chars": total_length,
            # Study detection
            "study_detected": True,
            "study_score": study_detection["score"],
            "study_signals": study_detection["signals"][:8],
            "study_type": study_type,
            "study_type_label": study_type_label,
            "study_meta": {
                "authors": authors,
                "organization": org,
                "study_type": study_type,
                "study_type_label": study_type_label,
            },
            # Not a paper or book
            "book_detected": False,
            "paper_detected": False,
            "paper_score": 0,
            "paper_signals": [],
            "content_level": None,
            "integration_level": None,
            "structural": None,
            "ebf": None,
            "paper_id": None,
        }

    # ═══════════════════════════════════════════════════════
    # STEP 1: Paper detection (instant, no API call)
    # ═══════════════════════════════════════════════════════
    detection = detect_scientific_paper(text)

    # ═══════════════════════════════════════════════════════
    # STEP 2: Content Level Classification (L0-L3)
    # Based on: Appendix BM (METHOD-PAPERINT)
    # L0 = BibTeX only (no abstract)
    # L1 = Basic template (~2k chars, has abstract)
    # L2 = Full template (~6k chars, has structured YAML-worthy content)
    # L3 = Full text (>50k chars)
    # ═══════════════════════════════════════════════════════
    text_len = total_length  # Use actual document length, not truncated preview length
    has_abstract = any(s in text[:5000].lower() for s in ["abstract", "zusammenfassung"])
    # Implicit abstract detection: substantial text block before "Introduction" without explicit label
    if not has_abstract:
        import re
        intro_match = re.search(r'\n\s*introduction\s*\n', text[:8000], re.IGNORECASE)
        if intro_match:
            pre_intro = text[:intro_match.start()]
            # Find text after email/affiliation block (look for last @ or last country/institution mention)
            email_end = max(pre_intro.rfind('@'), pre_intro.rfind('.edu'), pre_intro.rfind('.ac.'),
                          pre_intro.rfind('.com'), pre_intro.rfind('.uk'))
            if email_end > 0:
                # Find end of that line
                line_end = pre_intro.find('\n', email_end)
                if line_end > 0:
                    candidate = pre_intro[line_end:].strip()
                    # A real abstract is typically 400-3000 chars of continuous prose
                    if 400 <= len(candidate) <= 3000:
                        has_abstract = True
    has_methodology = any(s in text.lower() for s in ["methodology", "method", "experimental design",
        "research design", "empirical", "randomized", "regression", "survey"])
    has_findings = any(s in text.lower() for s in ["findings", "results", "we find", "our results",
        "ergebnisse", "the evidence", "effect size", "significant"])
    has_references = any(s in text[-5000:].lower() for s in ["references\n", "bibliography",
        "literaturverzeichnis", "works cited"])

    if text_len > 50000:
        content_level = "L3"
        content_level_desc = "Volltext (Full Text)"
    elif text_len > 6000 and has_abstract and has_methodology:
        content_level = "L2"
        content_level_desc = "Volles Template (Structured Content)"
    elif text_len > 1500 and has_abstract:
        content_level = "L1"
        content_level_desc = "Basis-Template (Abstract vorhanden)"
    else:
        content_level = "L0"
        content_level_desc = "Minimal (kein Abstract erkannt)"

    # Structural characteristics (S1-S6)
    structural = {
        "S1_research_question": any(s in text[:8000].lower() for s in
            ["research question", "forschungsfrage", "we ask", "this paper asks",
             "we examine", "we investigate", "this study examines"]),
        "S2_methodology": has_methodology,
        "S3_sample_data": any(s in text.lower() for s in
            ["sample", "n =", "n=", "participants", "subjects", "respondents",
             "observations", "stichprobe", "dataset", "data set"]),
        "S4_findings": has_findings,
        "S5_validity": any(s in text.lower() for s in
            ["robustness", "sensitivity analysis", "validity", "external validity",
             "internal validity", "placebo test", "falsification"]),
        "S6_reproducibility": any(s in text.lower() for s in
            ["replication", "replicate", "reproducib", "pre-registered", "preregistered",
             "open data", "code availab"]),
    }

    # ═══════════════════════════════════════════════════════
    # STEP 3: Claude EBF Classification (enhanced prompt)
    # ═══════════════════════════════════════════════════════
    title = ""
    language = "de"
    category = "general"
    tags = []
    summary = ""
    ebf_data = {}

    if ANTHROPIC_API_KEY and detection["is_paper"]:
        try:
            import urllib.request, ssl
            ctx = ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = ssl.CERT_NONE

            excerpt_start = text[:4000]
            excerpt_end = text[-2000:] if text_len > 6000 else ""

            prompt = f"""Du bist BEATRIX, die wissenschaftliche Klassifikations-KI von FehrAdvice & Partners.
Analysiere dieses wissenschaftliche Paper und extrahiere Metadaten nach dem EBF-Schema (Empirical Behavioral Framework).

TEXT (Anfang):
{excerpt_start}

{"TEXT (Ende):" + chr(10) + excerpt_end if excerpt_end else ""}

Antworte NUR mit validem JSON:
{{
  "title": "exakter Titel des Papers",
  "authors": ["Nachname, Vorname", ...],
  "year": 2024,
  "language": "de" oder "en",
  "tags": ["max 5 Schlagwörter"],
  "summary": "1-2 Sätze Zusammenfassung",
  "journal_or_source": "Journal-Name oder Working Paper Serie oder null",
  "source_type": "journal_article" oder "working_paper" oder "book_chapter" oder "book" oder "conference_paper" oder "dissertation" oder "report",
  "doi": "DOI falls im Text vorhanden oder null",
  "evidence_tier": 1 oder 2 oder 3,
  "evidence_tier_reason": "Begründung: 1=Gold (Top-5 Journal, RCT, large N), 2=Silver (gutes Journal, solide Methodik), 3=Bronze (Working Paper, descriptive)",
  "methodology": {{
    "design": "experimental" oder "observational" oder "theoretical" oder "meta_analysis" oder "survey" oder "qualitative",
    "identification": "RCT" oder "IV" oder "DiD" oder "RDD" oder "matching" oder "descriptive" oder "theoretical"
  }},
  "ebf_dimensions": [
    {{
      "dimension": "CORE-WHO" oder "CORE-WHAT" oder "CORE-HOW" oder "CORE-WHEN" oder "CORE-WHERE" oder "CORE-AWARE" oder "CORE-READY" oder "CORE-STAGE" oder "CORE-HIERARCHY" oder "CORE-EIT",
      "connection": "kurze Begründung"
    }}
  ],
  "psi_dimensions": [
    {{
      "psi": "Ψ_I" oder "Ψ_S" oder "Ψ_C" oder "Ψ_K" oder "Ψ_E" oder "Ψ_T" oder "Ψ_M" oder "Ψ_P",
      "relevance": "kurze Begründung"
    }}
  ],
  "domains": ["finance" und/oder "health" und/oder "sustainability" und/oder "hr" und/oder "social_policy" und/oder "education" und/oder "behavior" und/oder "general"],
  "key_findings": [
    {{
      "finding": "Kernaussage",
      "effect_size": 0.5 oder null
    }}
  ]
}}"""

            payload = json.dumps({
                "model": ANTHROPIC_MODEL_STANDARD,
                "max_tokens": 1200,
                "messages": [{"role": "user", "content": prompt}]
            }).encode()

            req = urllib.request.Request(
                "https://api.anthropic.com/v1/messages",
                data=payload, method="POST",
                headers={
                    "x-api-key": ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json"
                }
            )
            resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=25).read())
            raw = resp.get("content", [{}])[0].get("text", "")

            import re
            # Find the outermost JSON object
            brace_count = 0
            start_idx = None
            end_idx = None
            for i, c in enumerate(raw):
                if c == '{':
                    if start_idx is None:
                        start_idx = i
                    brace_count += 1
                elif c == '}':
                    brace_count -= 1
                    if brace_count == 0 and start_idx is not None:
                        end_idx = i + 1
                        break
            if start_idx is not None and end_idx is not None:
                meta = json.loads(raw[start_idx:end_idx])
            else:
                meta = {}

            title = meta.get("title", "")
            language = meta.get("language", "de")
            category = "paper"
            tags = meta.get("tags", [])
            summary = meta.get("summary", "")
            ebf_data = {
                "authors": meta.get("authors", []),
                "year": meta.get("year"),
                "journal": meta.get("journal_or_source"),
                "source_type": meta.get("source_type", "journal_article"),
                "doi": meta.get("doi"),
                "evidence_tier": meta.get("evidence_tier", 3),
                "evidence_tier_reason": meta.get("evidence_tier_reason", ""),
                "methodology": meta.get("methodology", {}),
                "ebf_dimensions": meta.get("ebf_dimensions", []),
                "psi_dimensions": meta.get("psi_dimensions", []),
                "domains": meta.get("domains", []),
                "key_findings": meta.get("key_findings", [])
            }
        except Exception as e:
            logger.warning(f"Text analysis Claude error: {e}")
    elif ANTHROPIC_API_KEY and not detection["is_paper"]:
        # Simple analysis for non-papers
        try:
            import urllib.request, ssl
            ctx = ssl.create_default_context()
            ctx.check_hostname = False
            ctx.verify_mode = ssl.CERT_NONE
            excerpt = text[:3000]
            prompt = f"""Analysiere diesen Text. Antworte NUR mit JSON:
{{
  "title": "Titel des Dokuments",
  "language": "de" oder "en",
  "category": "paper" oder "axiom" oder "note" oder "general",
  "tags": ["max 5 Schlagwörter"],
  "summary": "1 Satz"
}}

TEXT:
{excerpt}"""
            payload = json.dumps({
                "model": ANTHROPIC_MODEL_STANDARD,
                "max_tokens": 300,
                "messages": [{"role": "user", "content": prompt}]
            }).encode()
            req = urllib.request.Request("https://api.anthropic.com/v1/messages",
                data=payload, method="POST",
                headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                          "Content-Type": "application/json"})
            resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
            raw = resp.get("content", [{}])[0].get("text", "")
            import re
            json_match = re.search(r'\{[^{}]*\}', raw, re.DOTALL)
            if json_match:
                meta = json.loads(json_match.group())
                title = meta.get("title", "")
                language = meta.get("language", "de")
                category = meta.get("category", "general")
                tags = meta.get("tags", [])
                summary = meta.get("summary", "")
        except Exception as e:
            logger.warning(f"Text analysis (simple) error: {e}")

    # Fallback title
    if not title:
        lines = [l.strip() for l in text[:500].split('\n') if l.strip()]
        title = lines[0][:100] if lines else "Untitled"

    # ═══════════════════════════════════════════════════════
    # STEP 4: Estimate Integration Level (I1-I5)
    # I1=MINIMAL, I2=STANDARD, I3=CASE, I4=THEORY, I5=FULL
    # For NEW papers, we estimate based on content quality
    # ═══════════════════════════════════════════════════════
    if detection["is_paper"]:
        s_count = sum(1 for v in structural.values() if v)
        tier = ebf_data.get("evidence_tier", 3)
        has_effect = any(f.get("effect_size") for f in ebf_data.get("key_findings", []))
        has_dims = len(ebf_data.get("ebf_dimensions", [])) > 0

        if s_count >= 5 and tier <= 1 and has_effect and has_dims:
            integration_level = "I3"
            integration_desc = "CASE-Ready (starke Evidenz, Effect Sizes, EBF-Dimensionen)"
        elif s_count >= 3 and has_dims:
            integration_level = "I2"
            integration_desc = "STANDARD (EBF-Dimensionen zugeordnet)"
        else:
            integration_level = "I1"
            integration_desc = "MINIMAL (Basis-Aufnahme)"

        # Generate paper_id suggestion
        first_author = ""
        if ebf_data.get("authors"):
            first_author = ebf_data["authors"][0].split(",")[0].strip().lower()
        year = ebf_data.get("year", "")
        title_word = ""
        for w in (title or "").split():
            if len(w) > 4 and w.lower() not in ["about", "their", "these", "which", "under", "between"]:
                title_word = w.lower()[:10]
                break
        paper_id = f"PAP-{first_author}{year}{title_word}" if first_author and year else ""
    else:
        integration_level = None
        integration_desc = None
        paper_id = ""

    # ═══════════════════════════════════════════════════════
    # STEP 5: Database target
    # ═══════════════════════════════════════════════════════
    db_target = "knowledge_base"
    if detection["is_paper"]:
        db_target = "research"
        category = "paper"

    return {
        "ok": True,
        "title": title,
        "language": language,
        "category": category,
        "tags": tags[:5],
        "summary": summary,
        "database": db_target,
        "chars": text_len,
        # Paper detection
        "paper_detected": detection["is_paper"],
        "paper_score": detection["score"],
        "paper_signals": detection["signals"][:8],
        # Content Level (L0-L3)
        "content_level": content_level if detection["is_paper"] else None,
        "content_level_desc": content_level_desc if detection["is_paper"] else None,
        # Integration Level (I1-I5)
        "integration_level": integration_level if detection["is_paper"] else None,
        "integration_level_desc": integration_desc if detection["is_paper"] else None,
        # Structural characteristics
        "structural": structural if detection["is_paper"] else None,
        # EBF classification
        "ebf": ebf_data if detection["is_paper"] and ebf_data else None,
        # Paper ID suggestion
        "paper_id": paper_id if detection["is_paper"] else None,
    }

@app.get("/api/documents")
async def list_documents(database: Optional[str] = None, limit: int = 50, user=Depends(require_auth)):
    db = get_db()
    try:
        query = db.query(Document).order_by(Document.created_at.desc())
        if database: query = query.filter(Document.database_target == database)
        return [DocumentResponse(id=d.id, title=d.title, source_type=d.source_type, file_type=d.file_type, database_target=d.database_target, status=d.status, created_at=d.created_at.isoformat(), github_url=d.github_url, metadata=d.doc_metadata, uploaded_by=d.uploaded_by) for d in query.limit(limit).all()]
    finally: db.close()

@app.get("/api/documents/{doc_id}")
async def get_document(doc_id: str, user=Depends(require_auth)):
    """Get full document details including content."""
    db = get_db()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc: raise HTTPException(404, "Nicht gefunden")
        return {
            "id": doc.id,
            "title": doc.title,
            "content": doc.content or "",
            "source_type": doc.source_type,
            "file_type": doc.file_type,
            "database_target": doc.database_target,
            "status": doc.status,
            "github_url": doc.github_url,
            "uploaded_by": doc.uploaded_by,
            "created_at": doc.created_at.isoformat(),
            "file_size": doc.file_size,
            "category": doc.category,
            "language": doc.language,
            "tags": doc.tags,
            "metadata": doc.doc_metadata or {}
        }
    finally: db.close()

@app.post("/api/documents/{doc_id}/classify")
async def classify_document(doc_id: str, user=Depends(require_auth)):
    """Run EBF classification on an existing document and update its metadata."""
    db = get_db()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc: raise HTTPException(404, "Nicht gefunden")

        text = doc.content or ""
        if len(text) < 100:
            return {"ok": False, "error": "Zu wenig Text für Klassifikation"}

        # Run paper detection
        detection = detect_scientific_paper(text)

        # Content Level
        text_len = len(text)
        has_abstract = any(s in text[:5000].lower() for s in ["abstract", "zusammenfassung"])
        # Implicit abstract detection: text block before "Introduction" without explicit label
        if not has_abstract:
            import re
            intro_match = re.search(r'\n\s*introduction\s*\n', text[:8000], re.IGNORECASE)
            if intro_match:
                pre_intro = text[:intro_match.start()]
                email_end = max(pre_intro.rfind('@'), pre_intro.rfind('.edu'), pre_intro.rfind('.ac.'),
                              pre_intro.rfind('.com'), pre_intro.rfind('.uk'))
                if email_end > 0:
                    line_end = pre_intro.find('\n', email_end)
                    if line_end > 0:
                        candidate = pre_intro[line_end:].strip()
                        if 400 <= len(candidate) <= 3000:
                            has_abstract = True
        has_methodology = any(s in text.lower() for s in ["methodology", "method", "experimental design",
            "research design", "empirical", "randomized", "regression", "survey"])
        has_findings = any(s in text.lower() for s in ["findings", "results", "we find", "our results",
            "ergebnisse", "the evidence", "effect size", "significant"])

        if text_len > 50000: content_level = "L3"
        elif text_len > 6000 and has_abstract and has_methodology: content_level = "L2"
        elif text_len > 1500 and has_abstract: content_level = "L1"
        else: content_level = "L0"

        # Structural characteristics
        structural = {
            "S1_research_question": any(s in text[:8000].lower() for s in
                ["research question", "forschungsfrage", "we ask", "this paper asks",
                 "we examine", "we investigate", "this study examines"]),
            "S2_methodology": has_methodology,
            "S3_sample_data": any(s in text.lower() for s in
                ["sample", "n =", "n=", "participants", "subjects", "respondents",
                 "observations", "stichprobe", "dataset", "data set"]),
            "S4_findings": has_findings,
            "S5_validity": any(s in text.lower() for s in
                ["robustness", "sensitivity analysis", "validity", "external validity",
                 "internal validity", "placebo test", "falsification"]),
            "S6_reproducibility": any(s in text.lower() for s in
                ["replication", "replicate", "reproducib", "pre-registered", "preregistered",
                 "open data", "code availab"]),
        }

        # Claude EBF classification
        ebf_data = {}
        if ANTHROPIC_API_KEY and detection["is_paper"]:
            try:
                import urllib.request as ureq
                import ssl as _ssl
                _ctx = _ssl.create_default_context(); _ctx.check_hostname = False; _ctx.verify_mode = _ssl.CERT_NONE

                excerpt_start = text[:4000]
                excerpt_end = text[-2000:] if text_len > 6000 else ""

                prompt = f"""Du bist BEATRIX, die wissenschaftliche Klassifikations-KI von FehrAdvice & Partners.
Analysiere dieses wissenschaftliche Paper und extrahiere Metadaten nach dem EBF-Schema.

TEXT (Anfang):
{excerpt_start}

{"TEXT (Ende):" + chr(10) + excerpt_end if excerpt_end else ""}

Antworte NUR mit validem JSON:
{{
  "title": "exakter Titel",
  "authors": ["Nachname, Vorname"],
  "year": 2024,
  "journal_or_source": "Journal oder null",
  "source_type": "journal_article|working_paper|book_chapter|book|conference_paper|dissertation|report",
  "doi": "DOI oder null",
  "evidence_tier": 1 oder 2 oder 3,
  "evidence_tier_reason": "kurze Begründung",
  "methodology": {{"design": "...", "identification": "..."}},
  "ebf_dimensions": [{{"dimension": "CORE-...", "connection": "..."}}],
  "psi_dimensions": [{{"psi": "Ψ_X", "relevance": "..."}}],
  "domains": ["finance|health|sustainability|hr|social_policy|education|behavior|general"],
  "key_findings": [{{"finding": "...", "effect_size": null}}],
  "summary": "1-2 Sätze"
}}"""

                payload = json.dumps({
                    "model": ANTHROPIC_MODEL_STANDARD, "max_tokens": 1200,
                    "messages": [{"role": "user", "content": prompt}]
                }).encode()
                req = ureq.Request("https://api.anthropic.com/v1/messages",
                    data=payload, method="POST",
                    headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                              "Content-Type": "application/json"})
                resp = json.loads(ureq.urlopen(req, context=_ctx, timeout=25).read())
                raw = resp.get("content", [{}])[0].get("text", "")

                # Parse outermost JSON
                brace_count = 0; start_idx = None; end_idx = None
                for i, c in enumerate(raw):
                    if c == '{':
                        if start_idx is None: start_idx = i
                        brace_count += 1
                    elif c == '}':
                        brace_count -= 1
                        if brace_count == 0 and start_idx is not None:
                            end_idx = i + 1; break
                if start_idx is not None and end_idx is not None:
                    ebf_data = json.loads(raw[start_idx:end_idx])
                else:
                    ebf_data = {}
            except Exception as e:
                logger.warning(f"Classify Claude error for {doc_id}: {e}")

        # Integration level estimate
        s_count = sum(1 for v in structural.values() if v)
        tier = ebf_data.get("evidence_tier", 3)
        has_effect = any(f.get("effect_size") for f in ebf_data.get("key_findings", []))
        has_dims = len(ebf_data.get("ebf_dimensions", [])) > 0

        if detection["is_paper"]:
            if s_count >= 5 and tier <= 1 and has_effect and has_dims:
                integration_level = "I3"
            elif s_count >= 3 and has_dims:
                integration_level = "I2"
            else:
                integration_level = "I1"
        else:
            integration_level = None

        # Paper ID
        paper_id = ""
        if detection["is_paper"] and ebf_data.get("authors"):
            first_author = ebf_data["authors"][0].split(",")[0].strip().lower()
            year = ebf_data.get("year", "")
            title_word = ""
            for w in (ebf_data.get("title", "") or doc.title or "").split():
                if len(w) > 4 and w.lower() not in ["about", "their", "these", "which", "under", "between"]:
                    title_word = w.lower()[:10]; break
            paper_id = f"PAP-{first_author}{year}{title_word}"

        # Update doc metadata
        meta = doc.doc_metadata or {}
        meta.update({
            "paper_detected": detection["is_paper"],
            "paper_score": detection["score"],
            "paper_signals": detection["signals"],
            "content_level": content_level,
            "integration_level": integration_level,
            "structural": structural,
            "ebf": ebf_data,
            "paper_id": paper_id,
            "classified_at": __import__('datetime').datetime.utcnow().isoformat()
        })
        doc.doc_metadata = meta

        # Update title if Claude found a better one
        if ebf_data.get("title") and (doc.title.startswith("Chat:") or doc.title in ["Untitled", ""]):
            doc.title = ebf_data["title"]

        # Update category
        if detection["is_paper"] and doc.category != "paper":
            doc.category = "paper"

        from sqlalchemy.orm.attributes import flag_modified
        flag_modified(doc, 'doc_metadata')
        db.commit()

        return {
            "ok": True,
            "doc_id": doc_id,
            "title": doc.title,
            "paper_detected": detection["is_paper"],
            "content_level": content_level,
            "integration_level": integration_level,
            "structural": structural,
            "ebf_summary": {
                "evidence_tier": ebf_data.get("evidence_tier"),
                "dimensions": len(ebf_data.get("ebf_dimensions", [])),
                "psi": len(ebf_data.get("psi_dimensions", [])),
                "findings": len(ebf_data.get("key_findings", []))
            } if ebf_data else None,
            "paper_id": paper_id
        }
    except HTTPException: raise
    except Exception as e:
        db.rollback()
        logger.error(f"Classify error: {e}")
        raise HTTPException(500, str(e))
    finally: db.close()

@app.post("/api/documents/classify-all")
async def classify_all_documents(request: Request, user=Depends(require_permission("document.classify_bulk"))):
    """Classify all unclassified documents. Returns progress summary."""
    body = await request.json() if request.headers.get("content-type", "").startswith("application/json") else {}
    force = body.get("force", False)  # Re-classify even if already classified
    db = get_db()
    try:
        docs = db.query(Document).order_by(Document.created_at.desc()).all()
        results = {"total": len(docs), "classified": 0, "skipped": 0, "errors": 0, "details": []}

        for doc in docs:
            # Skip if already classified (unless force)
            meta = doc.doc_metadata or {}
            if meta.get("classified_at") and not force:
                # Still fix missing titles even if skipping classification
                if doc.title in ["Untitled", "", "NBER WORKING PAPER SERIES"] or not doc.title or doc.title.lower().startswith("nber working") or doc.title.lower().startswith("working paper"):
                    text_for_title = doc.content or ""
                    lines = [l.strip() for l in text_for_title[:2000].split('\n') if l.strip() and len(l.strip()) > 3]
                    skip_patterns = ["nber working paper", "working paper series", "discussion paper",
                        "technical report", "research paper", "policy brief", "staff report",
                        "issn", "isbn", "http", "www.", "doi:", "jel class", "keywords:", "abstract"]
                    for line in lines:
                        lower = line.lower().strip()
                        if any(skip in lower for skip in skip_patterns): continue
                        if len(line) < 5: continue
                        if line.replace('.','').replace('-','').replace('/','').strip().isdigit(): continue
                        doc.title = line[:150]
                        results["details"].append({"id": doc.id, "title": line[:150], "action": "title_fixed"})
                        break
                results["skipped"] += 1
                continue

            text = doc.content or ""
            if len(text) < 100:
                results["skipped"] += 1
                continue

            try:
                # Lightweight classification (no Claude API to save tokens)
                detection = detect_scientific_paper(text)
                text_len = len(text)
                has_abstract = any(s in text[:5000].lower() for s in ["abstract", "zusammenfassung"])
                has_methodology = any(s in text.lower() for s in ["methodology", "method", "experimental design",
                    "research design", "empirical", "randomized", "regression", "survey"])

                if text_len > 50000: content_level = "L3"
                elif text_len > 6000 and has_abstract and has_methodology: content_level = "L2"
                elif text_len > 1500 and has_abstract: content_level = "L1"
                else: content_level = "L0"

                structural = {
                    "S1_research_question": any(s in text[:8000].lower() for s in
                        ["research question", "forschungsfrage", "we examine", "we investigate"]),
                    "S2_methodology": has_methodology,
                    "S3_sample_data": any(s in text.lower() for s in
                        ["sample", "n =", "participants", "observations", "dataset"]),
                    "S4_findings": any(s in text.lower() for s in
                        ["findings", "results", "we find", "our results", "ergebnisse"]),
                    "S5_validity": any(s in text.lower() for s in
                        ["robustness", "validity", "placebo test"]),
                    "S6_reproducibility": any(s in text.lower() for s in
                        ["replication", "reproducib", "pre-registered"]),
                }

                s_count = sum(1 for v in structural.values() if v)
                if detection["is_paper"]:
                    integration_level = "I2" if s_count >= 3 else "I1"
                else:
                    integration_level = None

                meta.update({
                    "paper_detected": detection["is_paper"],
                    "paper_score": detection["score"],
                    "paper_signals": detection["signals"],
                    "content_level": content_level,
                    "integration_level": integration_level,
                    "structural": structural,
                    "classified_at": __import__('datetime').datetime.utcnow().isoformat(),
                    "classification_method": "bulk_heuristic"
                })
                doc.doc_metadata = meta
                if detection["is_paper"] and doc.category != "paper":
                    doc.category = "paper"

                # Fix missing titles
                if doc.title in ["Untitled", ""] or not doc.title:
                    lines = [l.strip() for l in text[:2000].split('\n') if l.strip() and len(l.strip()) > 3]
                    skip_patterns = ["nber working paper", "working paper series", "discussion paper",
                        "technical report", "research paper", "policy brief", "staff report",
                        "issn", "isbn", "http", "www.", "doi:", "jel class", "keywords:", "abstract"]
                    for line in lines:
                        lower = line.lower().strip()
                        if any(skip in lower for skip in skip_patterns): continue
                        if len(line) < 5: continue
                        if line.replace('.','').replace('-','').replace('/','').strip().isdigit(): continue
                        doc.title = line[:150]; break

                from sqlalchemy.orm.attributes import flag_modified
                flag_modified(doc, 'doc_metadata')
                results["classified"] += 1
                results["details"].append({
                    "id": doc.id, "title": doc.title[:60],
                    "paper": detection["is_paper"], "level": content_level,
                    "integration": integration_level
                })
            except Exception as e:
                results["errors"] += 1
                logger.warning(f"Bulk classify error for {doc.id}: {e}")

        db.commit()
        return results
    except Exception as e:
        db.rollback()
        raise HTTPException(500, str(e))
    finally: db.close()

@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: str, user=Depends(require_permission("document.delete"))):
    db = get_db()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc: raise HTTPException(404, "Nicht gefunden")
        if doc.file_path and os.path.exists(doc.file_path): os.remove(doc.file_path)
        db.delete(doc); db.commit()
        return {"message": f"Geloescht: {doc.title}"}
    finally: db.close()

# ========== STARTUP: Embed existing documents ==========
@app.on_event("startup")
async def startup_embed():
    """On startup, embed any documents that don't have embeddings yet."""
    # ── Ensure admin accounts are active ──
    try:
        db = get_db()
        admin_emails = [e.strip().lower() for e in os.getenv("ADMIN_EMAILS", "").split(",") if e.strip()]
        for email in admin_emails:
            user = db.query(User).filter(User.email == email).first()
            if user:
                user.email_verified = True
                user.is_active = True
                if not user.role or user.role != "admin":
                    user.role = "admin"
                    logger.info(f"Startup: ensured admin role for {email}")
        db.commit()
        db.close()
    except Exception as e:
        logger.warning(f"Startup admin check error: {e}")

    if not VOYAGE_API_KEY:
        logger.info("No VOYAGE_API_KEY set — vector search disabled, using keyword fallback")
        return
    logger.info("Starting background embedding of existing documents...")
    try:
        db = get_db()
        count = embed_all_documents(db)
        db.close()
        logger.info(f"Startup embedding complete: {count} documents embedded")
    except Exception as e:
        logger.warning(f"Startup embedding error (non-critical): {e}")

# ========== ADMIN: Embedding management ==========
@app.get("/api/admin/embedding-stats")
async def embedding_stats(user=Depends(require_auth)):
    from sqlalchemy import text as sql_text
    db = get_db()
    try:
        total = db.query(Document).filter(Document.content.isnot(None)).count()
        try:
            embedded = db.execute(text("SELECT COUNT(*) FROM documents WHERE embedding IS NOT NULL")).scalar()
        except:
            db.rollback()
            try:
                embedded = db.execute(text("SELECT COUNT(*) FROM documents WHERE embedding_json IS NOT NULL AND embedding_json != ''")).scalar()
            except:
                db.rollback()
                embedded = 0
        return {
            "total_documents": total,
            "embedded": embedded,
            "not_embedded": total - embedded,
            "vector_search_enabled": bool(VOYAGE_API_KEY),
            "voyage_model": VOYAGE_MODEL
        }
    finally: db.close()

@app.post("/api/admin/embed-all")
async def admin_embed_all(user=Depends(require_permission("document.manage_embeddings"))):
    """Admin endpoint: Trigger embedding of all un-embedded documents."""
    db = get_db()
    try:
        u = db.query(User).filter(User.email == user["sub"]).first()
        if not u or not u.is_admin:
            raise HTTPException(403, "Nur Admins")
        if not VOYAGE_API_KEY:
            raise HTTPException(400, "VOYAGE_API_KEY nicht konfiguriert")
        count = embed_all_documents(db)
        return {"message": f"{count} Dokumente embedded", "embedded": count}
    finally: db.close()

# ── Leads API ────────────────────────────────────
# ══════════════════════════════════════════════════════
# ── Zefix (Swiss Commercial Register) via LINDAS SPARQL ──
# ═════════════════════════════════════════════════════════

ZEFIX_SPARQL_URL = "https://lindas.admin.ch/query"

def _zefix_sparql(query: str, timeout: int = 15) -> list:
    """Execute SPARQL query against LINDAS Zefix endpoint."""
    import urllib.parse, urllib.request, ssl, json as _json
    params = urllib.parse.urlencode({"query": query})
    url = f"{ZEFIX_SPARQL_URL}?{params}"
    headers = {"Accept": "application/sparql-results+json", "User-Agent": "BEATRIX/3.7"}
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    req = urllib.request.Request(url, headers=headers)
    resp = urllib.request.urlopen(req, context=ctx, timeout=timeout)
    data = _json.loads(resp.read().decode())
    return data.get("results", {}).get("bindings", [])

def _zefix_search(name: str, limit: int = 10) -> list:
    """Search Zefix by company name, return structured results."""
    # Escape single quotes
    safe_name = name.replace("'", "\\'")
    query = f"""
    PREFIX schema: <http://schema.org/>
    PREFIX admin: <https://schema.ld.admin.ch/>
    SELECT DISTINCT ?company ?legalName ?name ?uid ?legalForm ?street ?locality ?postalCode ?municipality ?desc
    FROM <https://lindas.admin.ch/foj/zefix>
    WHERE {{
      ?company a admin:ZefixOrganisation ;
               schema:legalName ?legalName .
      FILTER(CONTAINS(LCASE(?legalName), LCASE('{safe_name}')))
      OPTIONAL {{ ?company schema:name ?name . FILTER(LANG(?name) = '' || LANG(?name) = 'de') }}
      OPTIONAL {{ ?company schema:description ?desc }}
      OPTIONAL {{ ?company schema:identifier ?uidNode . FILTER(CONTAINS(STR(?uidNode), 'UID')) . ?uidNode schema:value ?uid }}
      OPTIONAL {{ ?company schema:additionalType/schema:name ?legalForm . FILTER(LANG(?legalForm) = 'de') }}
      OPTIONAL {{ ?company schema:address ?addr . ?addr schema:streetAddress ?street }}
      OPTIONAL {{ ?company schema:address ?addr2 . ?addr2 schema:addressLocality ?locality }}
      OPTIONAL {{ ?company schema:address ?addr3 . ?addr3 schema:postalCode ?postalCode }}
      OPTIONAL {{ ?company admin:municipality/schema:name ?municipality }}
    }}
    LIMIT {limit}
    """
    raw = _zefix_sparql(query)

    # Deduplicate by company IRI (multiple language names create dupes)
    seen = {}
    for r in raw:
        cid = r.get("company", {}).get("value", "")
        if cid not in seen:
            seen[cid] = {
                "zefix_id": cid.split("/")[-1] if cid else "",
                "legal_name": r.get("legalName", {}).get("value", ""),
                "name": r.get("name", {}).get("value", ""),
                "uid": r.get("uid", {}).get("value", ""),
                "legal_form": r.get("legalForm", {}).get("value", ""),
                "street": r.get("street", {}).get("value", ""),
                "locality": r.get("locality", {}).get("value", ""),
                "postal_code": r.get("postalCode", {}).get("value", ""),
                "municipality": r.get("municipality", {}).get("value", ""),
                "description": (r.get("desc", {}).get("value", "") or "")[:300],
            }
        else:
            # Merge: prefer German name, fill blanks
            existing = seen[cid]
            for field in ["name", "uid", "legal_form", "street", "locality", "postal_code", "municipality", "description"]:
                if not existing.get(field) and r.get(field, {}).get("value"):
                    val = r.get(field, {}).get("value", "")
                    existing[field] = val[:300] if field == "description" else val

    return list(seen.values())

@app.get("/api/zefix/search")
async def zefix_search(q: str = "", limit: int = 10, user=Depends(require_permission("crm.read"))):
    """Search Swiss Commercial Register (Zefix) via LINDAS SPARQL. Free, no auth needed."""
    if not q or len(q) < 2:
        raise HTTPException(400, "Query must be at least 2 characters")
    try:
        results = _zefix_search(q, min(limit, 25))
        return {"query": q, "count": len(results), "source": "zefix_lindas_sparql", "results": results}
    except Exception as e:
        logger.error(f"Zefix search error: {e}")
        raise HTTPException(502, f"Zefix SPARQL error: {str(e)[:200]}")

@app.get("/api/zefix/lookup/{uid}")
async def zefix_lookup_uid(uid: str, user=Depends(require_permission("crm.read"))):
    """Lookup a company by UID (CHE number)."""
    safe_uid = uid.replace("'", "").replace('"', '').strip()
    query = f"""
    PREFIX schema: <http://schema.org/>
    PREFIX admin: <https://schema.ld.admin.ch/>
    SELECT ?company ?legalName ?name ?legalForm ?street ?locality ?postalCode ?municipality ?desc
    FROM <https://lindas.admin.ch/foj/zefix>
    WHERE {{
      ?company a admin:ZefixOrganisation ;
               schema:legalName ?legalName ;
               schema:identifier ?uidNode .
      ?uidNode schema:value '{safe_uid}' .
      FILTER(CONTAINS(STR(?uidNode), 'UID'))
      OPTIONAL {{ ?company schema:name ?name }}
      OPTIONAL {{ ?company schema:description ?desc }}
      OPTIONAL {{ ?company schema:additionalType/schema:name ?legalForm . FILTER(LANG(?legalForm) = 'de') }}
      OPTIONAL {{ ?company schema:address ?addr . ?addr schema:streetAddress ?street }}
      OPTIONAL {{ ?company schema:address ?addr2 . ?addr2 schema:addressLocality ?locality }}
      OPTIONAL {{ ?company schema:address ?addr3 . ?addr3 schema:postalCode ?postalCode }}
      OPTIONAL {{ ?company admin:municipality/schema:name ?municipality }}
    }}
    LIMIT 5
    """
    try:
        raw = _zefix_sparql(query)
        if not raw:
            return {"uid": safe_uid, "found": False}
        r = raw[0]
        return {
            "uid": safe_uid, "found": True,
            "legal_name": r.get("legalName", {}).get("value", ""),
            "name": r.get("name", {}).get("value", ""),
            "legal_form": r.get("legalForm", {}).get("value", ""),
            "street": r.get("street", {}).get("value", ""),
            "locality": r.get("locality", {}).get("value", ""),
            "postal_code": r.get("postalCode", {}).get("value", ""),
            "municipality": r.get("municipality", {}).get("value", ""),
            "description": (r.get("desc", {}).get("value", "") or "")[:300],
        }
    except Exception as e:
        logger.error(f"Zefix UID lookup error: {e}")
        raise HTTPException(502, f"Zefix error: {str(e)[:200]}")

# ── Handelsregister.ai (German Commercial Register) ──
# ══════════════════════════════════════════════════════
HANDELSREGISTER_API_KEY = os.getenv("HANDELSREGISTER_API_KEY", "")

def _hr_api_request(endpoint: str, params: dict, timeout: int = 15) -> dict:
    """Make a request to handelsregister.ai API."""
    import urllib.parse, urllib.request, ssl, json as _json
    if not HANDELSREGISTER_API_KEY:
        raise ValueError("HANDELSREGISTER_API_KEY not configured")
    query_string = urllib.parse.urlencode(params, doseq=True)
    url = f"https://handelsregister.ai/api/v1/{endpoint}?{query_string}"
    headers = {
        "x-api-key": HANDELSREGISTER_API_KEY,
        "Accept": "application/json",
        "User-Agent": "BEATRIX/3.7"
    }
    ctx = ssl.create_default_context()
    req = urllib.request.Request(url, headers=headers)
    resp = urllib.request.urlopen(req, context=ctx, timeout=timeout)
    return _json.loads(resp.read().decode())

def _hr_search(name: str, limit: int = 10) -> list:
    """Search German Commercial Register via handelsregister.ai."""
    data = _hr_api_request("search-organizations", {"q": name, "limit": min(limit, 25)})
    results = []
    for item in (data if isinstance(data, list) else data.get("results", data.get("data", []))):
        if isinstance(item, dict):
            addr = item.get("address", {}) or {}
            reg = item.get("registration", {}) or {}
            results.append({
                "name": item.get("name", ""),
                "legal_form": item.get("legal_form", ""),
                "register_type": reg.get("register_type", item.get("register_type", "")),
                "register_number": reg.get("register_number", item.get("register_number", "")),
                "register_court": reg.get("court", item.get("register_court", "")),
                "status": item.get("status", ""),
                "street": addr.get("street", ""),
                "house_number": addr.get("house_number", ""),
                "zip_code": addr.get("zip_code", addr.get("postal_code", "")),
                "city": addr.get("city", ""),
                "country": addr.get("country_code", "DE"),
                "description": (item.get("purpose", item.get("description", "")) or "")[:400],
                "company_id": item.get("id", item.get("company_id", "")),
            })
    return results

def _hr_fetch(query: str) -> dict | None:
    """Fetch detailed company info from handelsregister.ai."""
    try:
        data = _hr_api_request("fetch-organization", {"q": query, "ai_search": "on-default"})
        if not data or (isinstance(data, dict) and data.get("error")):
            return None
        addr = data.get("address", {}) or {}
        reg = data.get("registration", {}) or {}
        cap = data.get("capital", {}) or {}
        mgmt = data.get("management", {}) or {}
        return {
            "name": data.get("name", ""),
            "legal_form": data.get("legal_form", ""),
            "register_type": reg.get("register_type", ""),
            "register_number": reg.get("register_number", ""),
            "register_court": reg.get("court", ""),
            "status": data.get("status", ""),
            "street": f"{addr.get('street', '')} {addr.get('house_number', '')}".strip(),
            "zip_code": addr.get("zip_code", addr.get("postal_code", "")),
            "city": addr.get("city", ""),
            "country": addr.get("country_code", "DE"),
            "description": (data.get("purpose", data.get("description", "")) or "")[:400],
            "capital_amount": cap.get("amount"),
            "capital_currency": cap.get("currency", "EUR"),
            "managing_director": mgmt.get("managing_director", {}).get("name", "") if isinstance(mgmt.get("managing_director"), dict) else "",
            "company_id": data.get("id", ""),
        }
    except Exception as e:
        logger.warning(f"HR fetch error: {e}")
        return None

@app.get("/api/handelsregister/search")
async def handelsregister_search(q: str = "", limit: int = 10, user=Depends(require_permission("crm.read"))):
    """Search German Commercial Register via handelsregister.ai API."""
    if not HANDELSREGISTER_API_KEY:
        raise HTTPException(503, "Handelsregister API key not configured. Set HANDELSREGISTER_API_KEY on Railway.")
    if not q or len(q) < 2:
        raise HTTPException(400, "Query must be at least 2 characters")
    try:
        results = _hr_search(q, min(limit, 25))
        return {"query": q, "count": len(results), "source": "handelsregister_ai", "results": results}
    except Exception as e:
        logger.error(f"Handelsregister search error: {e}")
        raise HTTPException(502, f"Handelsregister API error: {str(e)[:200]}")

@app.get("/api/handelsregister/lookup")
async def handelsregister_lookup(q: str = "", user=Depends(require_permission("crm.read"))):
    """Fetch detailed company info from handelsregister.ai."""
    if not HANDELSREGISTER_API_KEY:
        raise HTTPException(503, "Handelsregister API key not configured")
    if not q or len(q) < 2:
        raise HTTPException(400, "Query must be at least 2 characters")
    try:
        result = _hr_fetch(q)
        if not result:
            return {"query": q, "found": False}
        return {"query": q, "found": True, **result}
    except Exception as e:
        logger.error(f"Handelsregister lookup error: {e}")
        raise HTTPException(502, f"Handelsregister error: {str(e)[:200]}")

@app.get("/api/registry/search")
async def unified_registry_search(q: str = "", limit: int = 10, country: str = "", user=Depends(require_permission("crm.read"))):
    """Unified search across all connected company registers (CH: Zefix, DE: Handelsregister.ai)."""
    if not q or len(q) < 2:
        raise HTTPException(400, "Query must be at least 2 characters")
    all_results = []
    errors = []
    # Search Zefix (CH) - always free
    if country in ("", "CH", "ch"):
        try:
            ch_results = _zefix_search(q, min(limit, 10))
            for r in ch_results:
                r["_source"] = "zefix"
                r["_country"] = "CH"
                r["_flag"] = "🇨🇭"
            all_results.extend(ch_results)
        except Exception as e:
            errors.append(f"Zefix: {str(e)[:100]}")
    # Search Handelsregister (DE) - needs API key
    if country in ("", "DE", "de") and HANDELSREGISTER_API_KEY:
        try:
            de_results = _hr_search(q, min(limit, 10))
            for r in de_results:
                r["_source"] = "handelsregister_ai"
                r["_country"] = "DE"
                r["_flag"] = "🇩🇪"
            all_results.extend(de_results)
        except Exception as e:
            errors.append(f"Handelsregister: {str(e)[:100]}")
    return {
        "query": q, "count": len(all_results),
        "sources": {"zefix": country in ("","CH","ch"), "handelsregister": bool(HANDELSREGISTER_API_KEY) and country in ("","DE","de")},
        "results": all_results,
        "errors": errors if errors else None
    }

# ── CRM API ──────────────────────────────────────────
# ══════════════════════════════════════════════════════

CRM_STAGES = ['prospect', 'qualified', 'proposal', 'negotiation', 'won', 'active', 'dormant', 'churned', 'lost']
CRM_STAGE_PROB = {'prospect': 10, 'qualified': 25, 'proposal': 50, 'negotiation': 75, 'won': 100, 'active': 100, 'dormant': 5, 'churned': 0, 'lost': 0, 'closed_won': 100}

# ── Companies ──
@app.get("/api/crm/companies")
async def crm_get_companies(user=Depends(require_permission("crm.read"))):
    db = get_db()
    try:
        rows = db.execute(text("SELECT * FROM crm_companies ORDER BY updated_at DESC")).fetchall()
        return [dict(r._mapping) for r in rows]
    except: return []
    finally: db.close()

# GitHub-enriched companies – merges customer-registry.yaml with CRM data
_github_companies_cache = {"data": None, "ts": 0}

@app.get("/api/crm/companies/enriched")
async def crm_get_companies_enriched(user=Depends(require_permission("crm.read"))):
    """Fetch companies from GitHub customer-registry.yaml + customer profiles, merge with CRM DB."""
    import yaml, time as _time, urllib.request, ssl
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    db = get_db()
    try:
        # Cache GitHub data for 5 minutes
        now = _time.time()
        if _github_companies_cache["data"] and now - _github_companies_cache["ts"] < 300:
            github_customers = _github_companies_cache["data"]
        else:
            github_customers = []
            contacts_by_customer = {}
            try:
                # Fetch customer-registry.yaml
                gh_headers = {"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3.raw"}
                url = f"https://api.github.com/repos/{GH_REPO}/contents/data/customer-registry.yaml"
                req = urllib.request.Request(url, headers=gh_headers)
                content = urllib.request.urlopen(req, context=ctx, timeout=10).read().decode()
                registry = yaml.safe_load(content)
                github_customers = registry.get("customers", [])

                # Fetch customer-contacts.yaml for contact data
                try:
                    url2 = f"https://api.github.com/repos/{GH_REPO}/contents/data/customer-contacts.yaml"
                    req2 = urllib.request.Request(url2, headers=gh_headers)
                    contacts_content = urllib.request.urlopen(req2, context=ctx, timeout=10).read().decode()
                    contacts_data = yaml.safe_load(contacts_content)
                    for c in contacts_data.get("customers", []):
                        cid = c.get("customer_id", "")
                        contacts_by_customer[cid] = c
                except Exception as ce:
                    logger.warning(f"Contacts fetch: {ce}")

                # Fetch lead-database.yaml for full company list
                try:
                    url3 = f"https://api.github.com/repos/{GH_REPO}/contents/data/sales/lead-database.yaml"
                    req3 = urllib.request.Request(url3, headers=gh_headers)
                    leads_content = urllib.request.urlopen(req3, context=ctx, timeout=30).read().decode()
                    leads_data = yaml.safe_load(leads_content)
                    leads_list = leads_data.get("leads", [])
                    
                    # Build set of existing customer codes
                    existing_codes = {c.get("code","").upper() for c in github_customers}
                    
                    # Add companies from leads that aren't in customer-registry
                    for lead in leads_list:
                        co = lead.get("company", {})
                        if isinstance(co, str):
                            co = {"name": co, "short_name": co}
                        short = co.get("short_name", "")
                        if not short or short.upper() in existing_codes:
                            continue
                        existing_codes.add(short.upper())
                        
                        # Extract contacts from lead
                        lead_contacts = lead.get("contacts", [])
                        contact_count = len(lead_contacts) if isinstance(lead_contacts, list) else 0
                        
                        # Extract EBF/context data
                        ebf = lead.get("ebf_integration", {})
                        
                        github_customers.append({
                            "id": lead.get("id", ""),
                            "code": short,
                            "name": co.get("name", short),
                            "short_name": short,
                            "type": "lead",
                            "industry": lead.get("industry", ""),
                            "country": (lead.get("headquarters", {}) or {}).get("country", "") if isinstance(lead.get("headquarters"), dict) else "",
                            "status": lead.get("stage", ""),
                            "profile_path": None,
                            "context_path": None,
                            "notes": lead.get("notes", ""),
                            "lead_id": lead.get("id", ""),
                            "segment": lead.get("segment", ""),
                            "fit_score": lead.get("fit_score", 0),
                            "engagement_score": lead.get("engagement_score", 0),
                            "source": lead.get("source", ""),
                            "hot_lead": lead.get("hot_lead", False),
                            "priority": lead.get("priority", ""),
                            "next_action": lead.get("next_action", ""),
                            "next_action_date": lead.get("next_action_date", ""),
                            "tags": lead.get("tags", []),
                            "website": co.get("website", ""),
                            "linkedin": co.get("linkedin", ""),
                            "employee_count": lead.get("employee_count"),
                            "revenue_eur": lead.get("revenue_eur"),
                            "ebf_context_vectors": ebf.get("context_vectors") if isinstance(ebf, dict) else None,
                            "ebf_applicable_models": ebf.get("applicable_models") if isinstance(ebf, dict) else None,
                            "_lead_contacts": lead_contacts,
                            "_contact_count": contact_count,
                        })
                    logger.info(f"Lead-database merged: {len(leads_list)} leads, now {len(github_customers)} total companies")
                except Exception as le:
                    logger.warning(f"Lead-database fetch: {le}")

                # Enrich registry customers with contacts
                for cust in github_customers:
                    code = cust.get("code", "")
                    if code in contacts_by_customer:
                        cust["_contacts"] = contacts_by_customer[code]

                _github_companies_cache["data"] = github_customers
                _github_companies_cache["ts"] = now
                logger.info(f"GitHub customer registry loaded: {len(github_customers)} customers")
            except Exception as e:
                logger.error(f"GitHub customer fetch error: {e}")

        # Get CRM DB companies + deals for merge
        db_companies = {}
        try:
            rows = db.execute(text("SELECT * FROM crm_companies")).fetchall()
            for r in rows:
                m = dict(r._mapping)
                db_companies[m.get("id", "")] = m
                db_companies[m.get("name", "")] = m
                db_companies[m.get("domain", "")] = m
        except: pass

        # Get deal stats per company
        deal_stats = {}
        try:
            deals = db.execute(text("""SELECT d.company_id, d.company_name, d.stage, d.value, c.name as cname
                FROM crm_deals d LEFT JOIN crm_companies c ON d.company_id = c.id""")).fetchall()
            for d in deals:
                key = d.cname or d.company_name or d.company_id
                if not key: continue
                if key not in deal_stats:
                    deal_stats[key] = {"leads": 0, "active": 0, "won": 0, "pipeline": 0}
                deal_stats[key]["leads"] += 1
                if d.stage in ('Prospect','Qualified','Proposal','Negotiation','Active'):
                    deal_stats[key]["active"] += 1
                if d.stage == 'Won':
                    deal_stats[key]["won"] += 1
                deal_stats[key]["pipeline"] += float(d.value or 0)
        except: pass

        # Get contact counts
        contact_counts = {}
        try:
            contacts = db.execute(text("SELECT company_id, COUNT(*) as cnt FROM crm_contacts GROUP BY company_id")).fetchall()
            for c in contacts:
                contact_counts[c.company_id] = c.cnt
        except: pass

        # Get customer folders from GitHub
        customer_folders = _get_customer_folders()

        # Merge: GitHub as primary, enriched with CRM stats
        result = []
        for cust in github_customers:
            code = cust.get("code", "")
            name = cust.get("name", "")
            short = cust.get("short_name", code)
            db_match = db_companies.get(cust.get("id","")) or db_companies.get(f"CUS-{code}") or db_companies.get(name) or db_companies.get(code) or {}
            ds = deal_stats.get(name) or deal_stats.get(short) or deal_stats.get(code) or {}
            cc = cust.get("_contacts", {})

            result.append({
                "id": cust.get("id", ""),
                "code": code,
                "name": name,
                "short_name": short,
                "type": cust.get("type", ""),
                "industry": cust.get("industry", ""),
                "country": cust.get("country", ""),
                "status": cust.get("status", ""),
                "profile_path": cust.get("profile_path"),
                "context_path": cust.get("context_path"),
                "notes": cust.get("notes", ""),
                "has_profile": bool(cust.get("profile_path")),
                "has_context": bool(cust.get("context_path")),
                "has_customer_folder": code.lower() in [d.lower() for d in customer_folders],
                # CRM stats from DB
                "leads": ds.get("leads", 0),
                "active_leads": ds.get("active", 0),
                "won": ds.get("won", 0),
                "pipeline": ds.get("pipeline", 0),
                "db_contacts": contact_counts.get(db_match.get("id"), 0),
                # Contact info from GitHub customer-contacts
                "fa_owner": cc.get("fa_owner", "") or cust.get("_contacts", {}).get("fa_owner", ""),
                "relationship_status": cc.get("relationship_status", "") or cust.get("_contacts", {}).get("relationship_status", ""),
                "contact_persons": len(cc.get("contacts", [])) if isinstance(cc.get("contacts"), list) else cust.get("_contact_count", 0),
                # Lead-enriched fields (from lead-database.yaml)
                "lead_id": cust.get("lead_id", ""),
                "segment": cust.get("segment", ""),
                "fit_score": cust.get("fit_score", 0),
                "engagement_score": cust.get("engagement_score", 0),
                "source": cust.get("source", ""),
                "hot_lead": cust.get("hot_lead", False),
                "priority": cust.get("priority", ""),
                "next_action": cust.get("next_action", ""),
                "next_action_date": cust.get("next_action_date", ""),
                "tags": cust.get("tags", []),
                "website": cust.get("website", ""),
                "linkedin": cust.get("linkedin", ""),
                "employee_count": cust.get("employee_count"),
                "revenue_eur": cust.get("revenue_eur"),
                "ebf_context_vectors": cust.get("ebf_context_vectors"),
                "ebf_applicable_models": cust.get("ebf_applicable_models"),
                # Zefix enrichment (from DB)
                "uid": db_match.get("uid", ""),
                "legal_form": db_match.get("legal_form", ""),
                "zefix_id": db_match.get("zefix_id", ""),
                "postal_code": db_match.get("postal_code", ""),
                "locality": db_match.get("locality", ""),
                "canton": db_match.get("canton", ""),
                "address": db_match.get("address", ""),
                "zefix_description": db_match.get("zefix_description", ""),
            })

        # ── Group by parent company ──
        # Step 1: Identify parent codes (strategic/context type, or unique root codes)
        all_codes = [r["code"] for r in result]
        parent_candidates = set()
        for r in result:
            if r["type"] in ("strategic", "context"):
                parent_candidates.add(r["code"])

        # Step 2: For each company, find parent by prefix matching
        for r in result:
            r["parent_code"] = None
            if r["code"] in parent_candidates:
                r["parent_code"] = r["code"]  # is itself a parent
                continue
            # Check if code starts with a known parent code
            for pc in sorted(parent_candidates, key=len, reverse=True):
                if r["code"].upper().startswith(pc.upper()) and r["code"] != pc:
                    r["parent_code"] = pc
                    break

        # Step 3: Detect additional groups by name prefix (e.g. "DS Studio" entries)
        ungrouped = [r for r in result if not r["parent_code"]]
        from collections import Counter
        # Find names that share a common prefix (first word)
        first_words = Counter()
        for r in ungrouped:
            fw = r["name"].split()[0] if r["name"] else ""
            if len(fw) > 2:
                first_words[fw] += 1
        # Group if 2+ entries share same first word
        multi_groups = {fw for fw, cnt in first_words.items() if cnt >= 2}
        for r in result:
            if not r["parent_code"] and r["name"]:
                fw = r["name"].split()[0]
                if fw in multi_groups:
                    # Find or create parent: prefer strategic, else first entry
                    group_members = [x for x in result if x["name"].split()[0] == fw]
                    parent = next((x for x in group_members if x["type"] in ("strategic","context")), group_members[0])
                    r["parent_code"] = parent["code"]

        # Step 4: Entries without a group are their own parent
        for r in result:
            if not r["parent_code"]:
                r["parent_code"] = r["code"]

        # Build grouped output
        groups = {}
        for r in result:
            pc = r["parent_code"]
            if pc not in groups:
                groups[pc] = {"parent": None, "children": []}
            if r["code"] == pc:
                groups[pc]["parent"] = r
            else:
                groups[pc]["children"].append(r)

        # For groups without an explicit parent, promote first child
        for pc, g in groups.items():
            if not g["parent"] and g["children"]:
                g["parent"] = g["children"].pop(0)

        # Build final sorted list
        grouped_result = []
        for pc in sorted(groups.keys(), key=lambda k: groups[k]["parent"]["name"] if groups[k]["parent"] else ""):
            g = groups[pc]
            if not g["parent"]:
                continue
            parent = g["parent"]
            children = sorted(g["children"], key=lambda x: x["name"])
            # Aggregate stats
            parent["child_count"] = len(children)
            parent["total_leads"] = parent.get("leads", 0) + sum(c.get("leads", 0) for c in children)
            parent["total_contacts"] = parent.get("contact_persons", 0) + sum(c.get("contact_persons", 0) for c in children)
            parent["children"] = children
            grouped_result.append(parent)

        return grouped_result
    except Exception as e:
        logger.error(f"Enriched companies error: {e}")
        return []
    finally: db.close()

# ── PROJECT ENDPOINTS ──

@app.get("/api/projects")
async def get_projects(user=Depends(require_auth)):
    """List all projects from GitHub data/projects/*/project.yaml"""
    import urllib.request, ssl, yaml
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    try:
        url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects"
        req = urllib.request.Request(url, headers={"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"})
        items = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
        projects = []
        for item in items:
            if item["type"] != "dir":
                continue
            # Fetch project.yaml
            try:
                pyaml_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{item['name']}/project.yaml"
                req2 = urllib.request.Request(pyaml_url, headers={"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3.raw", "User-Agent": "BEATRIXLab"})
                content = urllib.request.urlopen(req2, context=ctx, timeout=15).read().decode()
                data = yaml.safe_load(content) or {}
                meta = data.get("metadata", {})
                client = data.get("client", {})
                project = data.get("project", {})
                timeline = data.get("timeline", {})
                team = data.get("team", {})
                # Handle timeline as list (milestone-style) vs dict
                if isinstance(timeline, list):
                    tl_start = timeline[0].get("date", "") if timeline else ""
                    tl_end = timeline[-1].get("date", "") if timeline else ""
                    tl_budget = None
                    tl_billing = ""
                else:
                    tl_start = timeline.get("start_date", "")
                    tl_end = timeline.get("end_date", "")
                    tl_budget = timeline.get("budget_chf", timeline.get("budget_eur"))
                    tl_billing = timeline.get("billing_type", "")
                # Handle team as dict or missing
                if not isinstance(team, dict):
                    team = {}
                # Extract fa_owner from various locations
                fa_owner = team.get("fa_owner", "")
                if not fa_owner and isinstance(team.get("fehradvice_team"), list):
                    for member in team["fehradvice_team"]:
                        if isinstance(member, dict) and member.get("role", "").lower() in ("lead", "owner", "projektleiter"):
                            fa_owner = member.get("name", member.get("code", ""))
                            break
                # Extract name from multiple possible fields
                prj_name = project.get("name") or meta.get("name") or item["name"]
                # Extract customer name
                cust_name = client.get("name") or client.get("short_name") or meta.get("customer", "")
                projects.append({
                    "project_id": meta.get("project_id", ""),
                    "project_code": meta.get("project_code", meta.get("short_code", item["name"])),
                    "slug": item["name"],
                    "name": prj_name,
                    "type": project.get("type", ""),
                    "status": meta.get("status", "planning").lower(),
                    "project_category": meta.get("project_category", ""),
                    "customer_code": client.get("customer_code", client.get("short_name", "")),
                    "customer_name": cust_name,
                    "start_date": tl_start,
                    "end_date": tl_end,
                    "budget_chf": tl_budget,
                    "billing_type": tl_billing,
                    "fa_owner": fa_owner,
                    "fa_team": team.get("fa_team", []),
                    "created": meta.get("created", ""),
                    "last_updated": meta.get("last_updated", ""),
                    "github_path": f"data/projects/{item['name']}/project.yaml"
                })
            except Exception as e2:
                logger.warning(f"Could not read project {item['name']}: {e2}")
                continue
        projects.sort(key=lambda x: (0 if x["status"] in ("active","aktiv") else 1 if x["status"]=="planning" else 2, x.get("start_date","") or ""))
        return projects
    except Exception as e:
        logger.error(f"Get projects error: {e}")
        return []

@app.post("/api/projects")
async def create_project(request: Request, user=Depends(require_permission("project.create"))):
    """Create a new project with unique atomic Projektkürzel"""
    try:
        import urllib.request, ssl, yaml
        ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
        body = await request.json()

        customer_code = body.get("customer_code", "").strip()
        name = body.get("name", "").strip()
        project_category = body.get("project_category", "mandat")  # mandat / lead / probono
        if not customer_code or not name:
            return JSONResponse({"error": "customer_code and name required"}, status_code=400)
        if project_category not in ("mandat", "lead", "probono"):
            return JSONResponse({"error": "project_category must be mandat, lead, or probono"}, status_code=400)

        gh = {"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"}
        today = datetime.utcnow().strftime("%Y-%m-%d")

        # ── Step 1: Load sequence registry ──
        seq_registry = {}
        prefix_mapping = {}
        seq_sha = None
        try:
            seq_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/config/project-sequences.yaml"
            req = urllib.request.Request(seq_url, headers=gh)
            seq_file = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
            seq_sha = seq_file["sha"]
            req2 = urllib.request.Request(seq_url, headers={**gh, "Accept": "application/vnd.github.v3.raw"})
            seq_data = yaml.safe_load(urllib.request.urlopen(req2, context=ctx, timeout=15).read().decode()) or {}
            seq_registry = seq_data.get("sequences", {})
            prefix_mapping = seq_data.get("prefix_mapping", {})
        except Exception as seq_err:
            logger.warning(f"Could not load sequence registry: {seq_err}")

        # ── Step 2: Determine prefix from customer_code ──
        prefix = prefix_mapping.get(customer_code.lower(), "")
        if not prefix:
            # Auto-generate: take first 3-5 uppercase chars
            prefix = customer_code.upper().replace("-", "")[:5]
            if len(prefix) < 2:
                prefix = name.upper().replace(" ", "")[:4]

        # ── Step 3: Get next sequence number ──
        # Start from registry value (includes legacy offset)
        next_seq = seq_registry.get(prefix, 1)

        # Also scan existing project codes to find highest used number
        existing_slugs = set()
        try:
            list_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects"
            req = urllib.request.Request(list_url, headers=gh)
            items = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
            for item in items:
                if item.get("type") == "dir":
                    existing_slugs.add(item["name"].lower())
                    try:
                        purl = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{item['name']}/project.yaml"
                        req2 = urllib.request.Request(purl, headers={**gh, "Accept": "application/vnd.github.v3.raw"})
                        pdata = yaml.safe_load(urllib.request.urlopen(req2, context=ctx, timeout=10).read().decode()) or {}
                        existing_code = (pdata.get("metadata", {}).get("project_code", "") or "").upper()
                        # Parse: if code starts with our prefix, extract number
                        if existing_code.startswith(prefix):
                            num_part = existing_code[len(prefix):]
                            try:
                                num = int(num_part)
                                if num >= next_seq:
                                    next_seq = num + 1
                            except ValueError:
                                pass
                    except:
                        pass
        except Exception as scan_err:
            logger.warning(f"Could not scan existing projects: {scan_err}")

        # ── Step 4: Generate project code ──
        project_code = f"{prefix}{next_seq:03d}"

        # ── Step 5: Update sequence registry on GitHub ──
        try:
            seq_registry[prefix] = next_seq + 1
            if seq_sha and seq_data:
                seq_data["sequences"] = seq_registry
                seq_data["metadata"]["updated"] = today
                # Add prefix mapping if new
                if customer_code.lower() not in prefix_mapping:
                    seq_data.setdefault("prefix_mapping", {})[customer_code.lower()] = prefix
                seq_yaml = yaml.dump(seq_data, default_flow_style=False, allow_unicode=True, sort_keys=False)
                seq_put = json.dumps({
                    "message": f"seq: {prefix} → {next_seq + 1} (nach {project_code})",
                    "content": base64.b64encode(seq_yaml.encode('utf-8')).decode(),
                    "sha": seq_sha, "branch": "main"
                }).encode('utf-8')
                seq_req = urllib.request.Request(
                    f"https://api.github.com/repos/{GH_REPO}/contents/data/config/project-sequences.yaml",
                    data=seq_put, method="PUT", headers={**gh, "Content-Type": "application/json"})
                urllib.request.urlopen(seq_req, context=ctx, timeout=15)
                logger.info(f"Sequence updated: {prefix} → {next_seq + 1}")
        except Exception as seq_update_err:
            logger.warning(f"Could not update sequence registry: {seq_update_err}")

        # ── Step 6: Generate slug (folder name) ──
        # Normalize German umlauts and special chars for URL-safe slugs
        umlaut_map = {'ä':'ae','ö':'oe','ü':'ue','ß':'ss','é':'e','è':'e','ê':'e','à':'a','â':'a','ô':'o','î':'i','ç':'c','ñ':'n'}
        slug_base = f"{customer_code}-{name}".lower()
        slug_base = "".join(umlaut_map.get(c, c) for c in slug_base)
        slug = "".join(c if (c.isascii() and c.isalnum()) or c == '-' else '-' for c in slug_base)
        slug = "-".join(part for part in slug.split("-") if part)[:60]
        if slug in existing_slugs:
            for n in range(2, 100):
                candidate_slug = f"{slug}-{n}"
                if candidate_slug not in existing_slugs:
                    slug = candidate_slug
                    break

        # ── Step 5: Category labels ──
        category_labels = {
            "mandat": "Bezahltes Mandatsprojekt",
            "lead": "Lead-Projekt (Akquise)",
            "probono": "Pro-Bono-Projekt"
        }
        category_icons = {
            "mandat": "💼",
            "lead": "🎯",
            "probono": "🤝"
        }

        # ── Step 6: Build project.yaml ──
        project_yaml = {
            "metadata": {
                "project_code": project_code,
                "project_id": f"PRJ-{project_code}",
                "slug": slug,
                "status": "planning",
                "project_category": project_category,
                "project_category_label": category_labels.get(project_category, project_category),
                "created": today,
                "created_by": user.get("sub", user.get("email", "")),
                "last_updated": today,
                "version": "1.0"
            },
            "client": {
                "customer_code": customer_code,
                "name": body.get("customer_name", customer_code),
            },
            "project": {
                "name": name,
                "type": body.get("type", "beratung"),
                "description": body.get("description", ""),
                "objectives": [],
            },
            "timeline": {
                "start_date": body.get("start_date", today),
                "end_date": body.get("end_date", ""),
                "budget_chf": body.get("budget_chf"),
                "billing_type": body.get("billing_type", "fixed"),
            },
            "team": {
                "fa_owner": body.get("fa_owner", "GF"),
                "fa_team": body.get("fa_team", []),
                "client_contacts": [],
            },
            "ebf_integration": {
                "bcm_models": [],
                "psi_dimensions": [],
                "behavioral_objectives": [],
            },
            "deliverables": [],
            "resources": [],
            "changelog": [{
                "date": today,
                "author": user.get("sub", user.get("email", "")),
                "action": f"Projekt eröffnet via BEATRIX ({category_labels.get(project_category, '')})"
            }]
        }

        # Enrich client name from companies cache
        try:
            reg_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/sales/customer-registry.yaml"
            req = urllib.request.Request(reg_url, headers={**gh, "Accept": "application/vnd.github.v3.raw"})
            reg_data = yaml.safe_load(urllib.request.urlopen(req, context=ctx, timeout=15).read().decode()) or {}
            for cust in reg_data.get("customers", []):
                if cust.get("code", "").upper() == customer_code.upper():
                    project_yaml["client"]["name"] = cust.get("name", customer_code)
                    project_yaml["client"]["industry"] = cust.get("industry", "")
                    project_yaml["client"]["country"] = cust.get("country", "")
                    break
        except Exception as enrich_err:
            logger.warning(f"Could not enrich client: {enrich_err}")

        # Push to GitHub
        yaml_content = yaml.dump(project_yaml, default_flow_style=False, allow_unicode=True, sort_keys=False)
        file_path = f"data/projects/{slug}/project.yaml"
        put_url = f"https://api.github.com/repos/{GH_REPO}/contents/{file_path}"
        put_data = json.dumps({
            "message": f"Projekt eröffnet: {project_code} – {name} ({customer_code}) [{project_category}]",
            "content": base64.b64encode(yaml_content.encode('utf-8')).decode(),
            "branch": "main"
        }).encode('utf-8')
        req = urllib.request.Request(put_url, data=put_data, method="PUT",
            headers={**gh, "Content-Type": "application/json"})
        result = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
        logger.info(f"Project created: {project_code} ({slug}) by {user.get('email','?')} [{project_category}]")
        audit_log(user, "project.create", "project", project_code, f"{slug} ({customer_code}) [{project_category}]")

        # ── Auto-trigger tasks for new project ──
        if project_category == "mandat":
            try:
                create_task_from_trigger("project_created", {
                    "customer_code": customer_code,
                    "project_slug": slug,
                    "fa_owner": body.get("fa_owner", "GF"),
                    "base_date": body.get("start_date", datetime.utcnow().strftime("%Y-%m-%d"))
                }, created_by=user.get("sub", user.get("email", "")))
            except Exception as te:
                logger.warning(f"Project task trigger failed: {te}")

        return {
            "ok": True,
            "project_code": project_code,
            "project_id": f"PRJ-{project_code}",
            "slug": slug,
            "project_category": project_category,
            "github_path": file_path,
            "sha": result.get("content", {}).get("sha", "")
        }
    except Exception as e:
        logger.error(f"Create project error: {e}")
        import traceback; traceback.print_exc()
        return JSONResponse({"error": f"Fehler: {str(e)}"}, status_code=500)

@app.get("/api/projects/{slug}/landing")
async def get_project_landing(slug: str, user=Depends(require_auth)):
    """Get project details + available customer resources for landing page"""
    import urllib.request, ssl, yaml
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    gh = {"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"}
    gh_raw = {**gh, "Accept": "application/vnd.github.v3.raw"}

    result = {"project": None, "resources": {}, "other_projects": []}

    # 1. Load project.yaml
    try:
        url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{slug}/project.yaml"
        req = urllib.request.Request(url, headers=gh_raw)
        project_data = yaml.safe_load(urllib.request.urlopen(req, context=ctx, timeout=15).read().decode()) or {}
        result["project"] = project_data
    except Exception as e:
        logger.error(f"Landing: project load error: {e}")
        return JSONResponse({"error": f"Projekt '{slug}' nicht gefunden"}, status_code=404)

    # 2. Determine customer code/folder
    client = project_data.get("client", {})
    customer_code = (client.get("short_name") or client.get("customer_code") or "").lower()
    if not customer_code:
        # Try to extract from slug
        customer_code = slug.split("-")[0] if "-" in slug else ""

    # 3. Scan customer folder for available resources
    if customer_code:
        try:
            cust_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/customers/{customer_code}"
            req2 = urllib.request.Request(cust_url, headers=gh)
            files = json.loads(urllib.request.urlopen(req2, context=ctx, timeout=15).read())

            file_names = [f["name"] for f in files]
            file_types = {f["name"]: f["type"] for f in files}
            file_sizes = {f["name"]: f.get("size", 0) for f in files}

            # Context vector
            context_files = [f for f in file_names if "context" in f.lower() and f.endswith(".yaml")]
            result["resources"]["context_vector"] = {
                "available": len(context_files) > 0,
                "count": len(context_files),
                "files": context_files
            }

            # Models
            model_files = [f for f in file_names if "model" in f.lower() and f.endswith(".yaml")]
            result["resources"]["models"] = {
                "available": len(model_files) > 0,
                "count": len(model_files),
                "files": model_files
            }

            # Profile
            profile_files = [f for f in file_names if "profile" in f.lower() and f.endswith(".yaml")]
            result["resources"]["profile"] = {
                "available": len(profile_files) > 0,
                "files": profile_files
            }

            # KPIs
            kpi_files = [f for f in file_names if "kpi" in f.lower()]
            result["resources"]["kpis"] = {
                "available": len(kpi_files) > 0,
                "files": kpi_files
            }

            # Scenarios
            scenario_files = [f for f in file_names if "scenario" in f.lower()]
            result["resources"]["scenarios"] = {
                "available": len(scenario_files) > 0,
                "count": len(scenario_files),
                "files": scenario_files
            }

            # Personas
            persona_files = [f for f in file_names if "persona" in f.lower()]
            result["resources"]["personas"] = {
                "available": len(persona_files) > 0,
                "files": persona_files
            }

            # Documents (PDFs, MDs)
            pdfs = [f for f in file_names if f.endswith(".pdf")]
            mds = [f for f in file_names if f.endswith(".md")]
            result["resources"]["documents"] = {
                "available": len(pdfs) + len(mds) > 0,
                "pdfs": len(pdfs),
                "mds": len(mds),
                "pdf_files": pdfs,
                "md_files": mds
            }

            # Data files (CSVs)
            csvs = [f for f in file_names if f.endswith(".csv")]
            result["resources"]["data_files"] = {
                "available": len(csvs) > 0,
                "count": len(csvs),
                "files": csvs
            }

            # Subfolders (project folders, etc.)
            dirs = [f for f in file_names if file_types.get(f) == "dir"]
            result["resources"]["folders"] = dirs

            # Total files
            result["resources"]["total_files"] = len(files)
            result["resources"]["customer_code"] = customer_code

        except urllib.error.HTTPError as he:
            if he.code == 404:
                result["resources"]["customer_folder_exists"] = False
            else:
                logger.warning(f"Landing: customer folder error: {he}")
        except Exception as e:
            logger.warning(f"Landing: customer scan error: {e}")

    # 4. Find other projects for this customer (from lead-database)
    try:
        leads_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/sales/lead-database.yaml"
        req3 = urllib.request.Request(leads_url, headers=gh_raw)
        leads_data = yaml.safe_load(urllib.request.urlopen(req3, context=ctx, timeout=15).read().decode()) or {}
        customer_name = (client.get("name") or "").lower()
        for lead in leads_data.get("leads", []):
            lead_company = (lead.get("company", {}).get("name", "") or "").lower()
            lead_code = (lead.get("company", {}).get("short_name", "") or lead.get("code", "") or "").lower()
            if customer_code and (customer_code in lead_code or customer_code in lead_company):
                for p in lead.get("projects", []):
                    p_name = p.get("name", "")
                    result["other_projects"].append({
                        "name": p_name,
                        "status": p.get("status", ""),
                        "id": p.get("id", ""),
                        "period": p.get("period", ""),
                        "deliverables_count": len(p.get("deliverables", p.get("deliverables_planned", []))),
                    })
    except Exception as e:
        logger.warning(f"Landing: other projects error: {e}")

    # 5. Find other BEATRIX projects with same customer (from data/projects/)
    if customer_code:
        try:
            prj_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects"
            req_prj = urllib.request.Request(prj_url, headers=gh)
            prj_dirs = json.loads(urllib.request.urlopen(req_prj, context=ctx, timeout=15).read())
            for pd in prj_dirs:
                if pd["type"] != "dir" or pd["name"] == slug:
                    continue
                # Check if slug starts with customer code
                if pd["name"].startswith(customer_code + "-") or pd["name"].startswith(customer_code.replace("-", "") + "-"):
                    # Load minimal project info
                    try:
                        p_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{pd['name']}/project.yaml"
                        req_p = urllib.request.Request(p_url, headers=gh_raw)
                        p_data = yaml.safe_load(urllib.request.urlopen(req_p, context=ctx, timeout=10).read().decode()) or {}
                        p_meta = p_data.get("metadata", {})
                        p_proj = p_data.get("project", {})
                        result["other_projects"].append({
                            "name": p_proj.get("name") or p_meta.get("name") or pd["name"],
                            "slug": pd["name"],
                            "project_code": p_meta.get("project_code") or "",
                            "status": p_meta.get("status") or "planning",
                            "source": "github"
                        })
                    except:
                        pass
        except Exception as e:
            logger.warning(f"Landing: project scan error: {e}")

    # 6. Load manually linked projects
    linked = project_data.get("linked_projects", [])
    result["linked_projects"] = []
    for link in linked:
        link_slug = link.get("slug", "")
        if not link_slug:
            continue
        try:
            lp_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{link_slug}/project.yaml"
            req_lp = urllib.request.Request(lp_url, headers=gh_raw)
            lp_data = yaml.safe_load(urllib.request.urlopen(req_lp, context=ctx, timeout=10).read().decode()) or {}
            lp_meta = lp_data.get("metadata", {})
            lp_proj = lp_data.get("project", {})
            result["linked_projects"].append({
                "name": lp_proj.get("name") or lp_meta.get("name") or link_slug,
                "slug": link_slug,
                "project_code": lp_meta.get("project_code") or "",
                "status": lp_meta.get("status") or "planning",
                "type": link.get("type", "related"),
                "note": link.get("note", "")
            })
        except:
            result["linked_projects"].append({
                "name": link_slug, "slug": link_slug,
                "project_code": "", "status": "unknown",
                "type": link.get("type", "related"), "note": link.get("note", "")
            })

    # ── Merge DB edits (project_edits table) ──
    try:
        db = get_db()
        edits = db.execute(text("SELECT section, data FROM project_edits WHERE slug = :s"), {"s": slug}).fetchall()
        if edits:
            project_data = result.get("project", {})
            for row in edits:
                section_name = row[0]
                section_data = row[1] if isinstance(row[1], (dict, list)) else json.loads(row[1])
                # Merge into project_data
                if section_name in ("objective", "scope", "team", "ebf_integration", "fehradvice_scope", "metadata", "budget"):
                    project_data.setdefault(section_name, {}).update(section_data)
                elif section_name in ("deliverables", "risks", "kpis"):
                    project_data[section_name] = section_data if isinstance(section_data, list) else section_data.get("items", [])
                elif section_name == "timeline":
                    project_data["timeline"] = section_data
                elif section_name == "sessions":
                    if isinstance(section_data, dict) and section_data.get("action") == "add":
                        sessions = project_data.get("sessions", [])
                        sessions.append(section_data.get("session", {}))
                        project_data["sessions"] = sessions
                    elif isinstance(section_data, list):
                        project_data["sessions"] = section_data
                else:
                    project_data[section_name] = section_data
            result["project"] = project_data
            result["has_local_edits"] = True
    except Exception as db_err:
        logger.warning(f"Landing: DB merge error: {db_err}")

    return result

# Cache customer folders list
_customer_folders_cache = {"data": [], "ts": 0}
def _get_customer_folders():
    import time as _time, urllib.request, ssl
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    now = _time.time()
    if _customer_folders_cache["data"] and now - _customer_folders_cache["ts"] < 300:
        return _customer_folders_cache["data"]
    try:
        gh_headers = {"Authorization": f"token {GH_TOKEN}"}
        url = f"https://api.github.com/repos/{GH_REPO}/contents/data/customers"
        req = urllib.request.Request(url, headers=gh_headers)
        items = json.loads(urllib.request.urlopen(req, context=ctx, timeout=10).read())
        folders = [i["name"] for i in items if i["type"] == "dir"]
        _customer_folders_cache["data"] = folders
        _customer_folders_cache["ts"] = now
        return folders
    except:
        return _customer_folders_cache["data"]

@app.post("/api/projects/{slug}/link")
async def link_project(slug: str, request: Request, user=Depends(require_auth)):
    """Add a linked project to this project's YAML"""
    import urllib.request, ssl, yaml
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    gh = {"Authorization": f"token {GH_TOKEN}", "Content-Type": "application/json", "User-Agent": "BEATRIXLab"}
    gh_raw = {**gh, "Accept": "application/vnd.github.v3.raw"}

    body = await request.json()
    target_slug = body.get("target_slug", "")
    link_type = body.get("type", "related")
    note = body.get("note", "")

    if not target_slug or target_slug == slug:
        raise HTTPException(400, "Ungültiges Ziel-Projekt")

    # Load current project.yaml
    try:
        url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{slug}/project.yaml"
        req = urllib.request.Request(url, headers=gh_raw)
        project_data = yaml.safe_load(urllib.request.urlopen(req, context=ctx, timeout=15).read().decode()) or {}
    except:
        raise HTTPException(404, f"Projekt '{slug}' nicht gefunden")

    # Add link (avoid duplicates)
    linked = project_data.get("linked_projects", [])
    if any(l.get("slug") == target_slug for l in linked):
        return {"status": "already_linked"}

    linked.append({"slug": target_slug, "type": link_type, "note": note, "added_by": user.email, "added_at": _time.strftime("%Y-%m-%dT%H:%M:%SZ", _time.gmtime())})
    project_data["linked_projects"] = linked

    # Save to GitHub
    try:
        import base64
        yaml_content = yaml.dump(project_data, allow_unicode=True, default_flow_style=False, sort_keys=False)
        sha_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{slug}/project.yaml"
        req_sha = urllib.request.Request(sha_url, headers={"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"})
        sha = json.loads(urllib.request.urlopen(req_sha, context=ctx, timeout=15).read()).get("sha")
        payload = json.dumps({"message": f"link: {slug} ↔ {target_slug} ({link_type})", "content": base64.b64encode(yaml_content.encode("utf-8")).decode(), "sha": sha})
        req_put = urllib.request.Request(sha_url, data=payload.encode("utf-8"), method="PUT", headers=gh)
        urllib.request.urlopen(req_put, context=ctx, timeout=15)
        logger.info(f"Project linked: {slug} → {target_slug} ({link_type}) by {user.email}")
    except Exception as e:
        logger.error(f"Link save error: {e}")
        raise HTTPException(500, f"Fehler beim Speichern: {e}")

    return {"status": "linked", "target": target_slug, "type": link_type}


@app.delete("/api/projects/{slug}/link/{target_slug}")
async def unlink_project(slug: str, target_slug: str, user=Depends(require_auth)):
    """Remove a linked project"""
    import urllib.request, ssl, yaml
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    gh = {"Authorization": f"token {GH_TOKEN}", "Content-Type": "application/json", "User-Agent": "BEATRIXLab"}
    gh_raw = {**gh, "Accept": "application/vnd.github.v3.raw"}

    try:
        url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{slug}/project.yaml"
        req = urllib.request.Request(url, headers=gh_raw)
        project_data = yaml.safe_load(urllib.request.urlopen(req, context=ctx, timeout=15).read().decode()) or {}
    except:
        raise HTTPException(404, f"Projekt '{slug}' nicht gefunden")

    linked = project_data.get("linked_projects", [])
    new_linked = [l for l in linked if l.get("slug") != target_slug]
    if len(new_linked) == len(linked):
        return {"status": "not_found"}

    project_data["linked_projects"] = new_linked

    try:
        import base64
        yaml_content = yaml.dump(project_data, allow_unicode=True, default_flow_style=False, sort_keys=False)
        sha_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/projects/{slug}/project.yaml"
        req_sha = urllib.request.Request(sha_url, headers={"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"})
        sha = json.loads(urllib.request.urlopen(req_sha, context=ctx, timeout=15).read()).get("sha")
        payload = json.dumps({"message": f"unlink: {slug} ↔ {target_slug}", "content": base64.b64encode(yaml_content.encode("utf-8")).decode(), "sha": sha})
        req_put = urllib.request.Request(sha_url, data=payload.encode("utf-8"), method="PUT", headers=gh)
        urllib.request.urlopen(req_put, context=ctx, timeout=15)
        logger.info(f"Project unlinked: {slug} → {target_slug} by {user.email}")
    except Exception as e:
        logger.error(f"Unlink save error: {e}")
        raise HTTPException(500, f"Fehler beim Speichern: {e}")

    return {"status": "unlinked", "target": target_slug}


@app.put("/api/projects/{slug}")
async def update_project(slug: str, request: Request, user=Depends(require_auth)):
    """Update project: save to PostgreSQL immediately, sync to GitHub in background"""
    body = await request.json()
    section = body.get("section", "")
    data = body.get("data", {})
    if not section or not data:
        return JSONResponse({"error": "section and data required"}, status_code=400)

    db = get_db()
    try:
        # 1. Save to PostgreSQL immediately (always reliable)
        db.execute(text("""
            INSERT INTO project_edits (slug, section, data, edited_by, synced_to_github)
            VALUES (:slug, :section, :data, :user, FALSE)
            ON CONFLICT (slug, section) DO UPDATE SET
                data = :data, edited_by = :user, synced_to_github = FALSE,
                created_at = CURRENT_TIMESTAMP
        """), {"slug": slug, "section": section, "data": json.dumps(data), "user": user.get("sub","unknown")})
        db.commit()
        logger.info(f"Project edit saved to DB: {slug}/{section} by {user.get('sub','?')}")

        # 2. Try GitHub sync (non-blocking, may fail)
        github_synced = False
        try:
            github_synced = _sync_project_to_github(slug, section, data, user)
        except Exception as gh_err:
            logger.warning(f"GitHub sync failed (will retry later): {gh_err}")

        return {"ok": True, "section": section, "slug": slug, "github_synced": github_synced}
    except Exception as e:
        logger.error(f"Project update error: {e}")
        db.rollback()
        return JSONResponse({"error": str(e)}, status_code=500)

def _sync_project_to_github(slug, section, data, user):
    """Try to sync project edit to GitHub. Returns True if successful."""
    import urllib.request, ssl, yaml, base64
    from datetime import date
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    file_path = f"data/projects/{slug}/project.yaml"
    url = f"https://api.github.com/repos/{GH_REPO}/contents/{file_path}"
    gh = {"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"}

    # Fetch current
    req = urllib.request.Request(url, headers=gh)
    existing = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
    sha = existing["sha"]

    req2 = urllib.request.Request(url, headers={**gh, "Accept": "application/vnd.github.v3.raw"})
    project = yaml.safe_load(urllib.request.urlopen(req2, context=ctx, timeout=15).read().decode()) or {}

    # Merge
    if section == "objective":
        project.setdefault("objective", {}).update(data)
    elif section == "scope":
        project.setdefault("scope", {}).update(data)
    elif section == "team":
        project.setdefault("team", {}).update(data)
    elif section == "timeline":
        project["timeline"] = data
    elif section == "deliverables":
        project["deliverables"] = data if isinstance(data, list) else data.get("items", [])
    elif section == "risks":
        project["risks"] = data if isinstance(data, list) else data.get("items", [])
    elif section == "kpis":
        project["kpis"] = data if isinstance(data, list) else data.get("items", [])
    elif section == "budget":
        if isinstance(project.get("timeline"), dict):
            project["timeline"].update(data)
        else:
            project.setdefault("budget", {}).update(data)
    elif section == "sessions":
        sessions = project.get("sessions", [])
        if isinstance(data, dict) and data.get("action") == "add":
            sessions.append(data.get("session", {}))
        elif isinstance(data, list):
            sessions = data
        project["sessions"] = sessions
    elif section == "ebf_integration":
        project.setdefault("ebf_integration", {}).update(data)
    elif section == "fehradvice_scope":
        project.setdefault("fehradvice_scope", {}).update(data)
    elif section == "metadata":
        project.setdefault("metadata", {}).update(data)
    else:
        project[section] = data

    project.setdefault("metadata", {})["last_updated"] = date.today().isoformat()

    # Push
    yaml_content = yaml.dump(project, default_flow_style=False, allow_unicode=True, sort_keys=False)
    commit_msg = f"update: {slug} – {section} (via BEATRIX)"
    put_data = json.dumps({"message": commit_msg, "content": base64.b64encode(yaml_content.encode('utf-8')).decode(), "sha": sha, "branch": "main"}).encode('utf-8')
    req3 = urllib.request.Request(url, data=put_data, method="PUT", headers={**gh, "Content-Type": "application/json"})
    urllib.request.urlopen(req3, context=ctx, timeout=15)

    # Mark as synced in DB
    try:
        db = get_db()
        db.execute(text("UPDATE project_edits SET synced_to_github = TRUE WHERE slug = :s AND section = :sec"), {"s": slug, "sec": section})
        db.commit()
    except: pass

    return True

@app.post("/api/projects/{slug}/sync")
async def sync_project_to_github(slug: str, user=Depends(require_auth)):
    """Retry syncing unsynced edits to GitHub"""
    db = get_db()
    edits = db.execute(text("SELECT section, data FROM project_edits WHERE slug = :s AND synced_to_github = FALSE"), {"s": slug}).fetchall()
    if not edits:
        return {"ok": True, "message": "Alles synchronisiert"}
    
    synced = []
    failed = []
    for row in edits:
        section = row[0]
        data = row[1] if isinstance(row[1], (dict, list)) else json.loads(row[1])
        try:
            _sync_project_to_github(slug, section, data, user)
            synced.append(section)
        except Exception as e:
            failed.append({"section": section, "error": str(e)})
    
    return {"ok": True, "synced": synced, "failed": failed}


# ========== PROJECT DELIVERABLE FILES ==========

@app.get("/api/projects/{slug}/deliverables/files")
async def list_deliverable_files(slug: str, user=Depends(require_auth)):
    """List all files in a project's deliverables folder on GitHub."""
    if not GH_TOKEN:
        return {"files": [], "error": "GitHub not configured"}
    import urllib.request as ur, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    gh = {"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3+json", "User-Agent": "BEATRIXLab"}
    path = f"data/projects/{slug}/deliverables"
    url = f"https://api.github.com/repos/{GH_REPO}/contents/{path}"
    try:
        req = ur.Request(url, headers=gh)
        items = json.loads(ur.urlopen(req, context=ctx, timeout=15).read())
        files = []
        for item in items:
            if item["type"] == "file":
                files.append({
                    "name": item["name"],
                    "size": item["size"],
                    "sha": item["sha"],
                    "download_url": item.get("download_url", ""),
                    "html_url": item.get("html_url", ""),
                })
        return {"files": files, "path": path}
    except Exception as e:
        if "404" in str(e):
            return {"files": [], "path": path}
        return {"files": [], "error": str(e)}


@app.post("/api/projects/{slug}/deliverables/upload")
async def upload_deliverable_file(slug: str, file: UploadFile = File(...),
                                   deliverable_name: Optional[str] = Form(None),
                                   user=Depends(require_auth)):
    """Upload a file to a project's deliverables folder on GitHub."""
    if not GH_TOKEN:
        raise HTTPException(500, "GitHub not configured")
    import urllib.request as ur, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    gh = {"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3+json",
          "User-Agent": "BEATRIXLab", "Content-Type": "application/json"}

    content_bytes = await file.read()
    if len(content_bytes) > 50 * 1024 * 1024:
        raise HTTPException(400, "Datei zu gross (max 50 MB)")

    # Sanitize filename
    import re
    safe_name = re.sub(r'[^\w\s\-.]', '', file.filename or "unnamed").strip()
    safe_name = re.sub(r'\s+', '_', safe_name)
    if not safe_name:
        safe_name = f"deliverable_{uuid.uuid4().hex[:8]}"

    path = f"data/projects/{slug}/deliverables/{safe_name}"
    url = f"https://api.github.com/repos/{GH_REPO}/contents/{path}"

    # Check if file exists (for update)
    sha = None
    try:
        req = ur.Request(url, headers={k: v for k, v in gh.items() if k != "Content-Type"})
        resp = json.loads(ur.urlopen(req, context=ctx, timeout=15).read())
        sha = resp.get("sha")
    except:
        pass

    payload = {
        "message": f"deliverable: {slug} – {safe_name}" + (f" ({deliverable_name})" if deliverable_name else ""),
        "content": base64.b64encode(content_bytes).decode(),
        "branch": "main"
    }
    if sha:
        payload["sha"] = sha

    try:
        req = ur.Request(url, data=json.dumps(payload).encode("utf-8"), method="PUT", headers=gh)
        resp = json.loads(ur.urlopen(req, context=ctx, timeout=30).read())
        result = {
            "ok": True,
            "name": safe_name,
            "size": len(content_bytes),
            "html_url": resp.get("content", {}).get("html_url", ""),
            "download_url": resp.get("content", {}).get("download_url", ""),
            "sha": resp.get("content", {}).get("sha", ""),
            "deliverable_name": deliverable_name,
        }
        logger.info(f"Deliverable file uploaded: {slug}/{safe_name} ({len(content_bytes)} bytes)")
        return result
    except Exception as e:
        logger.error(f"Deliverable upload failed: {e}")
        raise HTTPException(500, f"Upload fehlgeschlagen: {e}")


@app.delete("/api/projects/{slug}/deliverables/files/{filename}")
async def delete_deliverable_file(slug: str, filename: str, user=Depends(require_auth)):
    """Delete a file from a project's deliverables folder on GitHub."""
    if not GH_TOKEN:
        raise HTTPException(500, "GitHub not configured")
    import urllib.request as ur, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    gh = {"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3+json",
          "User-Agent": "BEATRIXLab", "Content-Type": "application/json"}

    path = f"data/projects/{slug}/deliverables/{filename}"
    url = f"https://api.github.com/repos/{GH_REPO}/contents/{path}"

    # Get SHA
    try:
        req = ur.Request(url, headers={k: v for k, v in gh.items() if k != "Content-Type"})
        resp = json.loads(ur.urlopen(req, context=ctx, timeout=15).read())
        sha = resp["sha"]
    except:
        raise HTTPException(404, "Datei nicht gefunden")

    payload = {"message": f"delete deliverable: {slug} – {filename}", "sha": sha, "branch": "main"}
    try:
        req = ur.Request(url, data=json.dumps(payload).encode("utf-8"), method="DELETE", headers=gh)
        ur.urlopen(req, context=ctx, timeout=15)
        logger.info(f"Deliverable file deleted: {slug}/{filename}")
        return {"ok": True, "deleted": filename}
    except Exception as e:
        raise HTTPException(500, f"Löschen fehlgeschlagen: {e}")


@app.post("/api/crm/companies")
async def crm_create_company(request: Request, user=Depends(require_permission("crm.write"))):
    import yaml as _yaml
    db = get_db()
    try:
        data = await request.json()
        name = data.get("name","").strip()
        if not name:
            raise HTTPException(400, "Firmenname erforderlich")
        code = data.get("code","").upper().strip()
        if not code:
            # Auto-generate code from name: first 3-5 chars
            code = ''.join(c for c in name.upper() if c.isalpha())[:4]
        cid = f"CUS-{code}"
        industry = data.get("industry","")
        country = data.get("country","CH")
        is_new = data.get("is_new_customer", True)

        # 1. Save to DB
        db.execute(text("""INSERT INTO crm_companies (id, name, domain, industry, size, website, address, notes, created_by)
            VALUES (:id, :name, :domain, :industry, :size, :website, :address, :notes, :cb)
            ON CONFLICT (id) DO NOTHING"""),
            {"id": cid, "name": name, "domain": data.get("domain",""),
             "industry": industry, "size": data.get("size",""),
             "website": data.get("website",""), "address": data.get("address",""),
             "notes": data.get("notes",""), "cb": user["sub"]})
        db.commit()

        # 2. Add to GitHub customer-registry.yaml
        gh_status = "skipped"
        try:
            import urllib.request, ssl, base64
            _ctx = ssl.create_default_context(); _ctx.check_hostname = False; _ctx.verify_mode = ssl.CERT_NONE
            gh_h = {"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIX"}
            # Get current registry
            req = urllib.request.Request(f"https://api.github.com/repos/{GH_REPO}/contents/data/customer-registry.yaml", headers=gh_h)
            resp = json.loads(urllib.request.urlopen(req, context=_ctx, timeout=15).read())
            sha = resp["sha"]
            registry = _yaml.safe_load(base64.b64decode(resp["content"]).decode())
            customers = registry.get("customers", [])
            # Check if already exists
            if not any(c.get("id") == cid for c in customers):
                new_entry = {
                    "id": cid, "code": code, "name": name, "short_name": data.get("short_name", code),
                    "type": "lead", "industry": industry, "country": country,
                    "status": "prospect",
                    "profile_path": None, "context_path": None,
                    "projects": [],
                    "created": datetime.utcnow().strftime("%Y-%m-%d"),
                    "notes": data.get("notes", "")
                }
                # Add optional fields
                if data.get("website"): new_entry["website"] = data["website"]
                if data.get("fa_owner"): new_entry["fa_owner"] = data["fa_owner"]
                customers.append(new_entry)
                registry["customers"] = customers
                registry["metadata"]["last_updated"] = datetime.utcnow().strftime("%Y-%m-%d")
                new_yaml = _yaml.dump(registry, default_flow_style=False, allow_unicode=True, sort_keys=False)
                put_data = json.dumps({
                    "message": f"feat: neue Firma {code} ({name}) via BEATRIX",
                    "content": base64.b64encode(new_yaml.encode('utf-8')).decode(),
                    "sha": sha, "branch": "main"
                }).encode()
                req2 = urllib.request.Request(
                    f"https://api.github.com/repos/{GH_REPO}/contents/data/customer-registry.yaml",
                    data=put_data, method="PUT", headers={**gh_h, "Content-Type": "application/json"})
                urllib.request.urlopen(req2, context=_ctx, timeout=15)
                gh_status = "added"
                # Also add to project-sequences.yaml prefix_mapping
                try:
                    req3 = urllib.request.Request(f"https://api.github.com/repos/{GH_REPO}/contents/data/config/project-sequences.yaml", headers=gh_h)
                    resp3 = json.loads(urllib.request.urlopen(req3, context=_ctx, timeout=10).read())
                    seq_data = _yaml.safe_load(base64.b64decode(resp3["content"]).decode())
                    pm = seq_data.get("prefix_mapping", {})
                    name_key = name.lower().split()[0] if name else code.lower()
                    if name_key not in pm and code.lower() not in pm:
                        pm[name_key] = code
                        pm[code.lower()] = code
                        seq_data["prefix_mapping"] = pm
                        if code not in seq_data.get("sequences", {}):
                            seq_data["sequences"][code] = {"next": 1, "label": name}
                        new_seq_yaml = _yaml.dump(seq_data, default_flow_style=False, allow_unicode=True, sort_keys=False)
                        put4 = json.dumps({
                            "message": f"feat: prefix mapping {code} für {name}",
                            "content": base64.b64encode(new_seq_yaml.encode('utf-8')).decode(),
                            "sha": resp3["sha"], "branch": "main"
                        }).encode()
                        req4 = urllib.request.Request(
                            f"https://api.github.com/repos/{GH_REPO}/contents/data/config/project-sequences.yaml",
                            data=put4, method="PUT", headers={**gh_h, "Content-Type": "application/json"})
                        urllib.request.urlopen(req4, context=_ctx, timeout=10)
                except Exception as se:
                    logger.warning(f"Sequence mapping: {se}")
            else:
                gh_status = "exists"
            # Invalidate cache
            _github_companies_cache["ts"] = 0
        except Exception as ge:
            logger.warning(f"GitHub registry: {ge}")
            gh_status = f"error: {ge}"

        # 3. Auto-enrich from Zefix (Swiss Commercial Register)
        zefix_data = None
        try:
            zefix_results = _zefix_search(name, 3)
            if zefix_results:
                best = None
                for zr in zefix_results:
                    if zr.get("legal_name", "").lower() == name.lower():
                        best = zr; break
                if not best:
                    best = zefix_results[0]
                zefix_data = best
                # Store Zefix data in dedicated columns
                zf_updates, zf_params = [], {"id": cid}
                for col, key in [("uid","uid"),("legal_form","legal_form"),("zefix_id","zefix_id"),
                                 ("postal_code","postal_code"),("locality","locality"),("canton","municipality")]:
                    if best.get(key):
                        zf_updates.append(f"{col} = :{col}"); zf_params[col] = best[key]
                if best.get("description"):
                    zf_updates.append("zefix_description = :zd"); zf_params["zd"] = best["description"][:500]
                if best.get("locality") and not data.get("address"):
                    addr_parts = [best.get("street",""), f"{best.get('postal_code','')} {best.get('locality','')}".strip()]
                    addr = ", ".join(p for p in addr_parts if p)
                    if addr:
                        zf_updates.append("address = :addr"); zf_params["addr"] = addr
                if zf_updates:
                    zf_updates.append("updated_at = CURRENT_TIMESTAMP")
                    db.execute(text(f"UPDATE crm_companies SET {', '.join(zf_updates)} WHERE id = :id"), zf_params)
                    db.commit()
        except Exception as ze:
            logger.warning(f"Zefix enrichment: {ze}")

        # 4. Auto-enrich from Handelsregister.ai (DE) if no Zefix match
        hr_data = None
        if not zefix_data and HANDELSREGISTER_API_KEY:
            try:
                hr_results = _hr_search(name, 3)
                if hr_results:
                    best = None
                    for hr in hr_results:
                        if hr.get("name", "").lower() == name.lower():
                            best = hr; break
                    if not best:
                        best = hr_results[0]
                    hr_data = best
                    # Store HR data in DB
                    hr_updates, hr_params = [], {"id": cid}
                    if best.get("legal_form"):
                        hr_updates.append("legal_form = :lf"); hr_params["lf"] = best["legal_form"]
                    if best.get("zip_code"):
                        hr_updates.append("postal_code = :pc"); hr_params["pc"] = best["zip_code"]
                    if best.get("city"):
                        hr_updates.append("locality = :loc"); hr_params["loc"] = best["city"]
                    if best.get("description"):
                        hr_updates.append("zefix_description = :zd"); hr_params["zd"] = best["description"][:500]
                    if best.get("register_number"):
                        reg_id = f"{best.get('register_type','')} {best.get('register_number','')}".strip()
                        if best.get("register_court"):
                            reg_id += f" ({best['register_court']})"
                        hr_updates.append("uid = :uid"); hr_params["uid"] = reg_id
                    if best.get("city") and not data.get("address"):
                        addr_parts = [f"{best.get('street','')} {best.get('house_number','')}".strip(),
                                      f"{best.get('zip_code','')} {best.get('city','')}".strip()]
                        addr = ", ".join(p for p in addr_parts if p)
                        if addr:
                            hr_updates.append("address = :addr"); hr_params["addr"] = addr
                    if hr_updates:
                        hr_updates.append("updated_at = CURRENT_TIMESTAMP")
                        db.execute(text(f"UPDATE crm_companies SET {', '.join(hr_updates)} WHERE id = :id"), hr_params)
                        db.commit()
            except Exception as he:
                logger.warning(f"Handelsregister enrichment: {he}")

        return {"id": cid, "code": code, "status": "created", "github": gh_status, "zefix": zefix_data, "handelsregister": hr_data}
    except HTTPException: raise
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

@app.put("/api/crm/companies/{cid}")
async def crm_update_company(cid: str, request: Request, user=Depends(require_permission("crm.write"))):
    db = get_db()
    try:
        data = await request.json()
        sets, params = [], {"id": cid}
        for f in ["name","domain","industry","size","website","address","notes"]:
            if f in data: sets.append(f"{f} = :{f}"); params[f] = data[f]
        if not sets: return {"status": "no changes"}
        sets.append("updated_at = CURRENT_TIMESTAMP")
        db.execute(text(f"UPDATE crm_companies SET {', '.join(sets)} WHERE id = :id"), params)
        db.commit()
        return {"status": "updated"}
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

# ── Zefix Enrichment ──
@app.post("/api/crm/companies/{cid}/zefix-enrich")
async def crm_zefix_enrich_company(cid: str, user=Depends(require_permission("crm.write"))):
    """Enrich a single company with Zefix data."""
    db = get_db()
    try:
        row = db.execute(text("SELECT * FROM crm_companies WHERE id = :id"), {"id": cid}).fetchone()
        if not row:
            raise HTTPException(404, "Firma nicht gefunden")
        company = dict(row._mapping)
        name = company.get("name", "")
        if not name:
            return {"status": "skip", "reason": "no name"}

        results = _zefix_search(name, 5)
        if not results:
            return {"status": "not_found", "name": name}

        # Find best match
        best = None
        for zr in results:
            if zr.get("legal_name", "").lower() == name.lower():
                best = zr; break
        if not best:
            # Try partial match on short name
            name_lower = name.lower()
            for zr in results:
                if name_lower in zr.get("legal_name", "").lower():
                    best = zr; break
        if not best:
            best = results[0]

        # Update DB
        zf_updates, zf_params = [], {"id": cid}
        for col, key in [("uid","uid"),("legal_form","legal_form"),("zefix_id","zefix_id"),
                         ("postal_code","postal_code"),("locality","locality"),("canton","municipality")]:
            if best.get(key):
                zf_updates.append(f"{col} = :{col}"); zf_params[col] = best[key]
        if best.get("description"):
            zf_updates.append("zefix_description = :zd"); zf_params["zd"] = best["description"][:500]
        if best.get("locality") and not company.get("address"):
            addr_parts = [best.get("street",""), f"{best.get('postal_code','')} {best.get('locality','')}".strip()]
            addr = ", ".join(p for p in addr_parts if p)
            if addr:
                zf_updates.append("address = :addr"); zf_params["addr"] = addr
        if zf_updates:
            zf_updates.append("updated_at = CURRENT_TIMESTAMP")
            db.execute(text(f"UPDATE crm_companies SET {', '.join(zf_updates)} WHERE id = :id"), zf_params)
            db.commit()
        return {"status": "enriched", "company": cid, "zefix": best, "fields_updated": len(zf_updates)-1}
    except HTTPException: raise
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

@app.post("/api/crm/companies/zefix-enrich-all")
async def crm_zefix_enrich_all(user=Depends(require_permission("crm.write"))):
    """Bulk-enrich all companies with Zefix data. Only processes CH companies without UID."""
    if not user.get("admin"):
        raise HTTPException(403, "Admin only")
    db = get_db()
    try:
        rows = db.execute(text("SELECT id, name, uid, address FROM crm_companies ORDER BY name")).fetchall()
        results = {"enriched": 0, "skipped": 0, "not_found": 0, "errors": 0, "details": []}

        for row in rows:
            company = dict(row._mapping)
            cid = company["id"]
            name = company.get("name", "")
            if not name or len(name) < 2:
                results["skipped"] += 1; continue
            # Skip if already has UID
            if company.get("uid"):
                results["skipped"] += 1; continue

            try:
                import time; time.sleep(0.3)  # Rate limit: ~3 req/sec
                search_results = _zefix_search(name, 5)
                if not search_results:
                    results["not_found"] += 1
                    results["details"].append({"id": cid, "name": name, "status": "not_found"})
                    continue

                # Find best match
                best = None
                name_lower = name.lower()
                for zr in search_results:
                    if zr.get("legal_name", "").lower() == name_lower:
                        best = zr; break
                if not best:
                    for zr in search_results:
                        if name_lower in zr.get("legal_name", "").lower() or zr.get("legal_name", "").lower() in name_lower:
                            best = zr; break
                if not best:
                    best = search_results[0]

                # Update
                zf_updates, zf_params = [], {"id": cid}
                for col, key in [("uid","uid"),("legal_form","legal_form"),("zefix_id","zefix_id"),
                                 ("postal_code","postal_code"),("locality","locality"),("canton","municipality")]:
                    if best.get(key):
                        zf_updates.append(f"{col} = :{col}"); zf_params[col] = best[key]
                if best.get("description"):
                    zf_updates.append("zefix_description = :zd"); zf_params["zd"] = best["description"][:500]
                if best.get("locality") and not company.get("address"):
                    addr_parts = [best.get("street",""), f"{best.get('postal_code','')} {best.get('locality','')}".strip()]
                    addr = ", ".join(p for p in addr_parts if p)
                    if addr:
                        zf_updates.append("address = :addr"); zf_params["addr"] = addr
                if zf_updates:
                    zf_updates.append("updated_at = CURRENT_TIMESTAMP")
                    db.execute(text(f"UPDATE crm_companies SET {', '.join(zf_updates)} WHERE id = :id"), zf_params)
                    db.commit()
                    results["enriched"] += 1
                    results["details"].append({"id": cid, "name": name, "status": "enriched",
                        "zefix_name": best.get("legal_name",""), "uid": best.get("uid","")})
                else:
                    results["not_found"] += 1
            except Exception as e:
                results["errors"] += 1
                results["details"].append({"id": cid, "name": name, "status": f"error: {str(e)[:100]}"})

        # Invalidate cache
        _github_companies_cache["ts"] = 0
        return results
    except HTTPException: raise
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

# ── Handelsregister Enrichment ──
@app.post("/api/crm/companies/{cid}/hr-enrich")
async def crm_hr_enrich_company(cid: str, user=Depends(require_permission("crm.write"))):
    """Enrich a single company with Handelsregister.ai data (German companies)."""
    if not HANDELSREGISTER_API_KEY:
        raise HTTPException(503, "HANDELSREGISTER_API_KEY not configured")
    db = get_db()
    try:
        row = db.execute(text("SELECT * FROM crm_companies WHERE id = :id"), {"id": cid}).fetchone()
        if not row:
            raise HTTPException(404, "Firma nicht gefunden")
        company = dict(row._mapping)
        name = company.get("name", "")
        if not name:
            return {"status": "skip", "reason": "no name"}

        results = _hr_search(name, 5)
        if not results:
            return {"status": "not_found", "name": name}

        # Find best match
        best = None
        name_lower = name.lower()
        for hr in results:
            if hr.get("name", "").lower() == name_lower:
                best = hr; break
        if not best:
            for hr in results:
                if name_lower in hr.get("name", "").lower() or hr.get("name", "").lower() in name_lower:
                    best = hr; break
        if not best:
            best = results[0]

        # Update DB
        hr_updates, hr_params = [], {"id": cid}
        if best.get("legal_form"):
            hr_updates.append("legal_form = :lf"); hr_params["lf"] = best["legal_form"]
        if best.get("zip_code"):
            hr_updates.append("postal_code = :pc"); hr_params["pc"] = best["zip_code"]
        if best.get("city"):
            hr_updates.append("locality = :loc"); hr_params["loc"] = best["city"]
        if best.get("description"):
            hr_updates.append("zefix_description = :zd"); hr_params["zd"] = best["description"][:500]
        if best.get("register_number"):
            reg_id = f"{best.get('register_type','')} {best.get('register_number','')}".strip()
            if best.get("register_court"):
                reg_id += f" ({best['register_court']})"
            hr_updates.append("uid = :uid"); hr_params["uid"] = reg_id
        if best.get("city") and not company.get("address"):
            addr_parts = [f"{best.get('street','')} {best.get('house_number','')}".strip(),
                          f"{best.get('zip_code','')} {best.get('city','')}".strip()]
            addr = ", ".join(p for p in addr_parts if p)
            if addr:
                hr_updates.append("address = :addr"); hr_params["addr"] = addr
        if hr_updates:
            hr_updates.append("updated_at = CURRENT_TIMESTAMP")
            db.execute(text(f"UPDATE crm_companies SET {', '.join(hr_updates)} WHERE id = :id"), hr_params)
            db.commit()
        return {"status": "enriched", "company": cid, "handelsregister": best, "fields_updated": len(hr_updates)-1}
    except HTTPException: raise
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

# ── Contacts ──
@app.get("/api/crm/contacts")
async def crm_get_contacts(user=Depends(require_permission("crm.read")), company_id: str = None):
    db = get_db()
    try:
        if company_id:
            rows = db.execute(text("SELECT * FROM crm_contacts WHERE company_id = :cid ORDER BY updated_at DESC"), {"cid": company_id}).fetchall()
        else:
            rows = db.execute(text("SELECT * FROM crm_contacts ORDER BY updated_at DESC")).fetchall()
        return [dict(r._mapping) for r in rows]
    except: return []
    finally: db.close()

@app.post("/api/crm/contacts")
async def crm_create_contact(request: Request, user=Depends(require_permission("crm.write"))):
    db = get_db()
    try:
        data = await request.json()
        cid = str(uuid.uuid4())[:8]
        db.execute(text("""INSERT INTO crm_contacts (id, company_id, name, email, phone, position, role_type, notes, created_by)
            VALUES (:id, :company_id, :name, :email, :phone, :position, :role_type, :notes, :cb)"""),
            {"id": cid, "company_id": data.get("company_id"), "name": data.get("name",""),
             "email": data.get("email",""), "phone": data.get("phone",""),
             "position": data.get("position",""), "role_type": data.get("role_type","kontakt"),
             "notes": data.get("notes",""), "cb": user["sub"]})
        db.commit()
        return {"id": cid, "status": "created"}
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

@app.put("/api/crm/contacts/{cid}")
async def crm_update_contact(cid: str, request: Request, user=Depends(require_permission("crm.write"))):
    db = get_db()
    try:
        data = await request.json()
        sets, params = [], {"id": cid}
        for f in ["company_id","name","email","phone","position","role_type","psi_profile","notes"]:
            if f in data: sets.append(f"{f} = :{f}"); params[f] = data[f] if f != "psi_profile" else json.dumps(data[f])
        if not sets: return {"status": "no changes"}
        sets.append("updated_at = CURRENT_TIMESTAMP")
        db.execute(text(f"UPDATE crm_contacts SET {', '.join(sets)} WHERE id = :id"), params)
        db.commit()
        return {"status": "updated"}
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

# ── Deals ──
@app.get("/api/crm/deals")
async def crm_get_deals(user=Depends(require_permission("lead.read")), stage: str = None):
    # CRM access check: must be @fehradvice.com with crm_access
    email = user.get("sub", "")
    crm_ok = user.get("crm_access", False)
    crm_role = user.get("crm_role", "none")
    owner_code = user.get("crm_owner_code", "")
    if not crm_ok:
        raise HTTPException(403, "Kein CRM-Zugang. Bitte Admin kontaktieren.")
    db = get_db()
    try:
        q = """SELECT d.*, c.name as company_name, ct.name as contact_name, ct.email as contact_email
               FROM crm_deals d
               LEFT JOIN crm_companies c ON d.company_id = c.id
               LEFT JOIN crm_contacts ct ON d.contact_id = ct.id"""
        params = {}
        conditions = []
        if stage:
            conditions.append("d.stage = :stage")
            params["stage"] = stage
        # Owner filtering: viewer sees only own leads, manager/admin sees all
        if crm_role not in ("admin", "manager") and not user.get("admin"):
            if owner_code:
                conditions.append("d.owner = :owner")
                params["owner"] = owner_code
            else:
                # No owner code assigned → see nothing
                return []
        if conditions:
            q += " WHERE " + " AND ".join(conditions)
        q += " ORDER BY d.updated_at DESC"
        rows = db.execute(text(q), params).fetchall()
        return [dict(r._mapping) for r in rows]
    except HTTPException: raise
    except Exception as e:
        logger.error(f"CRM deals fetch: {e}")
        return []
    finally: db.close()

@app.post("/api/crm/deals")
async def crm_create_deal(request: Request, user=Depends(require_permission("lead.create"))):
    db = get_db()
    try:
        data = await request.json()
        did = str(uuid.uuid4())[:8]
        stage = data.get("stage", "erstkontakt")
        prob = data.get("probability", CRM_STAGE_PROB.get(stage, 10))

        # Auto-create company if name provided but no company_id
        company_id = data.get("company_id")
        if not company_id and data.get("company_name"):
            company_id = str(uuid.uuid4())[:8]
            db.execute(text("""INSERT INTO crm_companies (id, name, created_by)
                VALUES (:id, :name, :cb)"""),
                {"id": company_id, "name": data["company_name"], "cb": user["sub"]})

        # Auto-create contact if name provided but no contact_id
        contact_id = data.get("contact_id")
        if not contact_id and data.get("contact_name"):
            contact_id = str(uuid.uuid4())[:8]
            db.execute(text("""INSERT INTO crm_contacts (id, company_id, name, email, created_by)
                VALUES (:id, :cid, :name, :email, :cb)"""),
                {"id": contact_id, "cid": company_id, "name": data["contact_name"],
                 "email": data.get("contact_email",""), "cb": user["sub"]})

        db.execute(text("""INSERT INTO crm_deals (id, company_id, contact_id, title, stage, value, probability, source, next_action, next_action_date, owner, context_id, notes, created_by)
            VALUES (:id, :company_id, :contact_id, :title, :stage, :value, :prob, :source, :next_action, :nad, :owner, :ctx, :notes, :cb)"""),
            {"id": did, "company_id": company_id, "contact_id": contact_id,
             "title": data.get("title",""), "stage": stage, "value": data.get("value",0),
             "prob": prob, "source": data.get("source",""),
             "next_action": data.get("next_action",""), "nad": data.get("next_action_date"),
             "owner": data.get("owner", user["sub"]), "ctx": data.get("context_id"),
             "notes": data.get("notes",""), "cb": user["sub"]})
        db.commit()
        logger.info(f"CRM deal created: {did} by {user['sub']}")
        return {"id": did, "company_id": company_id, "contact_id": contact_id, "status": "created"}
    except Exception as e:
        db.rollback(); logger.error(f"CRM deal create: {e}"); raise HTTPException(500, str(e))
    finally: db.close()

@app.put("/api/crm/deals/{did}")
async def crm_update_deal(did: str, request: Request, user=Depends(require_permission("lead.update"))):
    db = get_db()
    try:
        data = await request.json()
        sets, params = [], {"id": did}
        for f in ["company_id","contact_id","title","stage","value","probability","source","next_action","next_action_date","owner","context_id","notes","lost_reason","lead_code"]:
            if f in data:
                sets.append(f"{f} = :{f}"); params[f] = data[f]
        if "stage" in data:
            if data["stage"] == "won" or data["stage"] == "lost":
                sets.append("closed_at = CURRENT_TIMESTAMP")
            if data["stage"] in CRM_STAGE_PROB and "probability" not in data:
                sets.append("probability = :auto_prob")
                params["auto_prob"] = CRM_STAGE_PROB[data["stage"]]
        if not sets: return {"status": "no changes"}
        sets.append("updated_at = CURRENT_TIMESTAMP")
        db.execute(text(f"UPDATE crm_deals SET {', '.join(sets)} WHERE id = :id"), params)
        db.commit()
        return {"status": "updated"}
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

@app.delete("/api/crm/deals/{did}")
async def crm_delete_deal(did: str, user=Depends(require_permission("lead.delete"))):
    db = get_db()
    try:
        db.execute(text("DELETE FROM crm_activities WHERE deal_id = :id"), {"id": did})
        db.execute(text("DELETE FROM crm_deals WHERE id = :id"), {"id": did})
        db.commit()
        return {"status": "deleted"}
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

@app.post("/api/crm/deals/generate-lead-codes")
async def crm_generate_lead_codes(user=Depends(require_permission("lead.update"))):
    """Generate lead_codes for all deals that don't have one. Schema: L-{PREFIX}-{B/N}-{YY}-{NNN}"""
    # Prefix mapping: company name fragments → code
    PREFIX_MAP = {
        "ubs": "UBS", "alpla": "ALP", "porr": "PORR", "lukb": "LUKB", "helsana": "HELS",
        "helvetia": "HELV", "geberit": "GEB", "sbb": "SBB", "swisscom": "SWC", "roche": "ROC",
        "bmw": "BMW", "css": "CSS", "kpt": "KPT", "denner": "DEN", "post ": "POST",
        "postfinance": "POSTF", "a1 ": "A1", "a1 telekom": "A1", "migros-gen": "MIG",
        "migrosbank": "MGB", "assura": "ASSUR", "raiffeisen": "RBI", "ringier": "RIN",
        "kaufland": "KAUF", "innosuisse": "INNO", "mövenpick": "MOV", "energie": "ENS",
        "lindt": "LINDT", "ds studio": "DSS", "david schärer": "DSS", "sob": "SOB",
        "südostbahn": "SOB", "erz ": "ERZ", "entsorgung": "ERZ", "basel": "BST",
        "klima zürich": "KLZH", "kanton zürich": "KTZH", "ypsomed": "YPS",
        "hofer": "HOT", "bfe": "BFE", "bundesamt": "BFE", "industriellen": "IVOO",
        "iv-oö": "IVOO", "verband schweizer": "VSM", "puc": "PUC",
        "auto-suisse": "AUTS", "österreichische post": "POSTA",
        "schweiz. post": "POST", "schweizerische post": "POST",
    }
    # Known FehrAdvice customers (Bestandskunden)
    BESTANDS = {"UBS","ALP","PORR","LUKB","HELS","HELV","GEB","SBB","SWC","ROC","BMW","CSS",
                "KPT","DEN","POST","A1","MIG","MGB","RBI","RIN","LINDT","SOB","BFE","MOV",
                "DSS","HOT","YPS","ERZ","BST","POSTA","ENS","ASSUR","PUC","IVOO"}

    db = get_db()
    try:
        rows = db.execute(text("""SELECT d.id, COALESCE(c.name, d.title) as company_name, d.title, d.created_at, d.lead_code
            FROM crm_deals d LEFT JOIN crm_companies c ON d.company_id = c.id ORDER BY d.created_at""")).fetchall()
        counters = {}  # PREFIX-B/N-YY → count
        updated = 0
        results = []
        for r in rows:
            if r.lead_code:
                results.append({"id": r.id, "lead_code": r.lead_code, "status": "exists"})
                # Track counter
                parts = r.lead_code.split("-")
                if len(parts) >= 5:
                    key = f"{parts[1]}-{parts[2]}-{parts[3]}"
                    try: counters[key] = max(counters.get(key, 0), int(parts[4]))
                    except: pass
                continue
            # Map company name to prefix
            cname = (r.company_name or r.title or "").lower()
            prefix = None
            for frag, code in PREFIX_MAP.items():
                if frag in cname:
                    prefix = code; break
            if not prefix:
                # Try first word
                first = cname.split()[0].upper()[:4] if cname.strip() else "UNK"
                prefix = first
            # B or N
            typ = "B" if prefix in BESTANDS else "N"
            # Year from created_at
            yr = str(r.created_at)[:4][-2:] if r.created_at else "26"
            # Counter
            key = f"{prefix}-{typ}-{yr}"
            counters[key] = counters.get(key, 0) + 1
            num = str(counters[key]).zfill(3)
            lead_code = f"L-{prefix}-{typ}-{yr}-{num}"
            db.execute(text("UPDATE crm_deals SET lead_code = :lc, updated_at = CURRENT_TIMESTAMP WHERE id = :id"),
                       {"lc": lead_code, "id": r.id})
            updated += 1
            results.append({"id": r.id, "company": r.company_name, "lead_code": lead_code})
        db.commit()
        return {"updated": updated, "total": len(rows), "results": results}
    except Exception as e:
        db.rollback()
        logger.error(f"Lead code generation: {e}")
        raise HTTPException(500, str(e))
    finally: db.close()

@app.post("/api/crm/migrate-company-ids")
async def crm_migrate_company_ids(request: Request, user=Depends(require_permission("crm.write"))):
    """Remap LEAD-* company IDs to CUS-* superkeys. One-time migration."""
    if not user.get("admin"):
        raise HTTPException(403, "Admin only")
    db = get_db()
    try:
        data = await request.json()
        id_remap = data.get("remap", {})  # {"LEAD-030": "CUS-A1", ...}
        results = {"companies_updated": 0, "deals_updated": 0, "contacts_updated": 0, "duplicates_merged": 0}

        for old_id, new_id in id_remap.items():
            # Check if target CUS-* already exists in DB
            existing = db.execute(text("SELECT id FROM crm_companies WHERE id = :id"), {"id": new_id}).fetchone()

            if existing:
                # Merge: update deals/contacts to point to existing, delete old company
                deals_upd = db.execute(text("UPDATE crm_deals SET company_id = :new, updated_at = CURRENT_TIMESTAMP WHERE company_id = :old"),
                    {"new": new_id, "old": old_id}).rowcount
                contacts_upd = db.execute(text("UPDATE crm_contacts SET company_id = :new WHERE company_id = :old"),
                    {"new": new_id, "old": old_id}).rowcount
                acts_upd = db.execute(text("UPDATE crm_activities SET company_id = :new WHERE company_id = :old"),
                    {"new": new_id, "old": old_id}).rowcount
                db.execute(text("DELETE FROM crm_companies WHERE id = :old"), {"old": old_id})
                results["duplicates_merged"] += 1
                results["deals_updated"] += deals_upd
                results["contacts_updated"] += contacts_upd
            else:
                # Rename: update company ID + all references
                # Create new company entry with new ID
                old_data = db.execute(text("SELECT * FROM crm_companies WHERE id = :id"), {"id": old_id}).fetchone()
                if old_data:
                    row = dict(old_data._mapping)
                    db.execute(text("""INSERT INTO crm_companies (id, name, domain, industry, size, website, address, notes, created_by, created_at, updated_at)
                        VALUES (:id, :name, :domain, :industry, :size, :website, :address, :notes, :cb, :ca, CURRENT_TIMESTAMP)
                        ON CONFLICT (id) DO NOTHING"""),
                        {"id": new_id, "name": row.get("name",""), "domain": row.get("domain",""),
                         "industry": row.get("industry",""), "size": row.get("size",""),
                         "website": row.get("website",""), "address": row.get("address",""),
                         "notes": row.get("notes",""), "cb": row.get("created_by",""),
                         "ca": row.get("created_at")})
                    results["companies_updated"] += 1

                # Update all references
                deals_upd = db.execute(text("UPDATE crm_deals SET company_id = :new, updated_at = CURRENT_TIMESTAMP WHERE company_id = :old"),
                    {"new": new_id, "old": old_id}).rowcount
                contacts_upd = db.execute(text("UPDATE crm_contacts SET company_id = :new WHERE company_id = :old"),
                    {"new": new_id, "old": old_id}).rowcount
                acts_upd = db.execute(text("UPDATE crm_activities SET company_id = :new WHERE company_id = :old"),
                    {"new": new_id, "old": old_id}).rowcount
                results["deals_updated"] += deals_upd
                results["contacts_updated"] += contacts_upd

                # Delete old company
                db.execute(text("DELETE FROM crm_companies WHERE id = :old"), {"old": old_id})

        db.commit()
        # Invalidate cache
        _github_companies_cache["ts"] = 0
        return {"status": "migrated", **results, "remaps": len(id_remap)}
    except Exception as e:
        db.rollback()
        logger.error(f"Company ID migration: {e}")
        raise HTTPException(500, str(e))
    finally: db.close()

# ── Activities ──
@app.get("/api/crm/activities")
async def crm_get_activities(user=Depends(require_permission("crm.read")), deal_id: str = None, company_id: str = None):
    db = get_db()
    try:
        q = "SELECT * FROM crm_activities WHERE 1=1"
        params = {}
        if deal_id: q += " AND deal_id = :did"; params["did"] = deal_id
        if company_id: q += " AND company_id = :cid"; params["cid"] = company_id
        q += " ORDER BY created_at DESC LIMIT 100"
        rows = db.execute(text(q), params).fetchall()
        return [dict(r._mapping) for r in rows]
    except: return []
    finally: db.close()

@app.post("/api/crm/activities")
async def crm_create_activity(request: Request, user=Depends(require_permission("crm.write"))):
    db = get_db()
    try:
        data = await request.json()
        aid = str(uuid.uuid4())[:8]
        db.execute(text("""INSERT INTO crm_activities (id, deal_id, company_id, contact_id, type, subject, description, due_date, created_by)
            VALUES (:id, :did, :cid, :ctid, :type, :subject, :desc, :due, :cb)"""),
            {"id": aid, "did": data.get("deal_id"), "cid": data.get("company_id"),
             "ctid": data.get("contact_id"), "type": data.get("type","note"),
             "subject": data.get("subject",""), "desc": data.get("description",""),
             "due": data.get("due_date"), "cb": user["sub"]})
        db.commit()
        return {"id": aid, "status": "created"}
    except Exception as e:
        db.rollback(); raise HTTPException(500, str(e))
    finally: db.close()

# ── CRM Stats ──
@app.get("/api/crm/stats")
async def crm_stats(user=Depends(require_permission("crm.read"))):
    crm_ok = user.get("crm_access", False)
    crm_role = user.get("crm_role", "none")
    owner_code = user.get("crm_owner_code", "")
    if not crm_ok:
        return {"total_deals":0,"active_deals":0,"won_deals":0,"new_this_month":0,"pipeline_value":0,"total_value":0,"conversion_rate":0,"stages":{},"companies":0,"contacts":0}
    db = get_db()
    try:
        q = "SELECT stage, value, probability, created_at, closed_at FROM crm_deals"
        params = {}
        # Owner filtering for non-admin/manager
        if crm_role not in ("admin", "manager") and not user.get("admin"):
            if owner_code:
                q += " WHERE owner = :owner"
                params["owner"] = owner_code
            else:
                return {"total_deals":0,"active_deals":0,"won_deals":0,"new_this_month":0,"pipeline_value":0,"total_value":0,"conversion_rate":0,"stages":{},"companies":0,"contacts":0}
        deals = db.execute(text(q), params).fetchall()
        deals = [dict(r._mapping) for r in deals]
        active = [d for d in deals if d['stage'] not in ('won','lost')]
        won = [d for d in deals if d['stage'] == 'won']
        now = datetime.utcnow()
        this_month = [d for d in deals if d['created_at'] and d['created_at'].month == now.month and d['created_at'].year == now.year]
        pipeline_value = sum(d.get('value',0) * d.get('probability',10) / 100 for d in active)
        total_value = sum(d.get('value',0) for d in active)
        conv_rate = round(len(won) / len(deals) * 100) if deals else 0
        stages = {}
        for d in deals:
            s = d['stage']
            if s not in stages: stages[s] = {"count": 0, "value": 0}
            stages[s]["count"] += 1
            stages[s]["value"] += d.get('value',0)
        companies = db.execute(text("SELECT COUNT(*) FROM crm_companies")).scalar()
        contacts = db.execute(text("SELECT COUNT(*) FROM crm_contacts")).scalar()
        return {
            "total_deals": len(deals), "active_deals": len(active), "won_deals": len(won),
            "new_this_month": len(this_month), "pipeline_value": pipeline_value,
            "total_value": total_value, "conversion_rate": conv_rate,
            "stages": stages, "companies": companies or 0, "contacts": contacts or 0
        }
    except Exception as e:
        logger.error(f"CRM stats: {e}")
        return {"total_deals":0,"active_deals":0,"won_deals":0,"new_this_month":0,"pipeline_value":0,"total_value":0,"conversion_rate":0,"stages":{},"companies":0,"contacts":0}
    finally: db.close()

# ── Migrate old leads → CRM ──
@app.post("/api/crm/migrate-leads")
async def crm_migrate_leads(request: Request, user=Depends(require_permission("platform.admin_dashboard"))):
    db = get_db()
    try:
        leads = db.execute(text("SELECT * FROM leads")).fetchall()
        migrated = 0
        stage_map = {'kontakt':'erstkontakt','qualifiziert':'discovery','proposal':'proposal','won':'won','lost':'lost'}
        for lead in leads:
            l = dict(lead._mapping)
            # Check if already migrated
            existing = db.execute(text("SELECT id FROM crm_deals WHERE notes LIKE :ref"), {"ref": f"%migrated:lead:{l['id']}%"}).fetchone()
            if existing: continue
            # Create company
            cid = str(uuid.uuid4())[:8]
            db.execute(text("INSERT INTO crm_companies (id, name, created_by) VALUES (:id, :name, :cb)"),
                {"id": cid, "name": l.get('company','Unbekannt'), "cb": l.get('created_by','system')})
            # Create contact if exists
            ctid = None
            if l.get('contact'):
                ctid = str(uuid.uuid4())[:8]
                db.execute(text("INSERT INTO crm_contacts (id, company_id, name, email, created_by) VALUES (:id, :cid, :name, :email, :cb)"),
                    {"id": ctid, "cid": cid, "name": l['contact'], "email": l.get('email',''), "cb": l.get('created_by','system')})
            # Create deal
            did = str(uuid.uuid4())[:8]
            new_stage = stage_map.get(l.get('stage','kontakt'), 'erstkontakt')
            notes = f"{l.get('notes','')}\n[migrated:lead:{l['id']}]".strip()
            db.execute(text("""INSERT INTO crm_deals (id, company_id, contact_id, title, stage, value, probability, source, notes, created_by, created_at)
                VALUES (:id, :cid, :ctid, :title, :stage, :val, :prob, :src, :notes, :cb, :ca)"""),
                {"id": did, "cid": cid, "ctid": ctid, "title": l.get('company','Deal'),
                 "stage": new_stage, "val": l.get('value',0),
                 "prob": CRM_STAGE_PROB.get(new_stage, 10), "src": l.get('source',''),
                 "notes": notes, "cb": l.get('created_by','system'),
                 "ca": l.get('created_at', datetime.utcnow())})
            migrated += 1
        db.commit()
        logger.info(f"CRM migration: {migrated} leads migrated")
        return {"migrated": migrated, "total_leads": len(leads)}
    except Exception as e:
        db.rollback(); logger.error(f"CRM migration: {e}"); raise HTTPException(500, str(e))
    finally: db.close()

# ── GitHub YAML → PostgreSQL Sync ──
@app.post("/api/crm/sync-github")
async def crm_sync_github(user=Depends(require_permission("platform.admin_dashboard"))):
    """Sync leads, customers, contacts from GitHub YAML files into PostgreSQL CRM tables"""
    import yaml, urllib.request as ureq
    db = get_db()
    try:
        gh_token = os.getenv("GH_TOKEN", os.getenv("GITHUB_TOKEN", ""))
        repo = "FehrAdvice-Partners-AG/complementarity-context-framework"
        headers_gh = {"Authorization": f"token {gh_token}", "Accept": "application/vnd.github.v3.raw", "User-Agent": "BEATRIXLab/3.18"}
        ssl_ctx = __import__('ssl').create_default_context()
        ssl_ctx.check_hostname = False
        ssl_ctx.verify_mode = __import__('ssl').CERT_NONE

        def gh_read(path):
            req = ureq.Request(f"https://api.github.com/repos/{repo}/contents/{path}", headers=headers_gh)
            return ureq.urlopen(req, context=ssl_ctx).read().decode()

        stats = {"companies": 0, "contacts": 0, "deals": 0, "activities": 0, "skipped": 0}

        # ── 1. Read lead-database.yaml ──
        logger.info("CRM Sync: Reading lead-database.yaml...")
        lead_yaml = yaml.safe_load(gh_read("data/sales/lead-database.yaml"))
        leads_list = lead_yaml.get("leads", [])
        if not leads_list:
            # Try alternate structure
            leads_list = lead_yaml.get("pipeline", {}).get("leads", [])
        logger.info(f"CRM Sync: Found {len(leads_list)} leads")

        # ── 2. Read customer-registry.yaml ──
        logger.info("CRM Sync: Reading customer-registry.yaml...")
        cust_yaml = yaml.safe_load(gh_read("data/customer-registry.yaml"))
        customers_list = cust_yaml.get("customers", [])
        logger.info(f"CRM Sync: Found {len(customers_list)} customers in registry")

        # ── 3. Read person-registry.yaml ──
        logger.info("CRM Sync: Reading person-registry.yaml...")
        person_yaml = yaml.safe_load(gh_read("data/person-registry.yaml"))

        # ── 4. Sync customers from registry ──
        for cust in customers_list:
            cid = cust.get("id", cust.get("code", ""))
            if not cid: continue
            existing = db.execute(text("SELECT id FROM crm_companies WHERE id = :id"), {"id": cid}).fetchone()
            if existing:
                stats["skipped"] += 1
                continue
            db.execute(text("""INSERT INTO crm_companies (id, name, domain, industry, size, website, notes, created_by)
                VALUES (:id, :name, :domain, :industry, :size, :website, :notes, :cb)"""),
                {"id": cid, "name": cust.get("name",""), "domain": cust.get("code",""),
                 "industry": cust.get("industry",""), "size": cust.get("type",""),
                 "website": "", "notes": cust.get("notes",""), "cb": "github-sync"})
            stats["companies"] += 1

        # ── 5. Sync leads as deals + companies + contacts ──
        stage_map = {
            'PROSPECT': 'prospect', 'QUALIFIED': 'qualified', 'PROPOSAL': 'proposal',
            'NEGOTIATION': 'negotiation', 'WON': 'won', 'ACTIVE': 'active',
            'DORMANT': 'dormant', 'CHURNED': 'churned', 'LOST': 'lost', 'CLOSED_WON': 'won'
        }

        for lead in leads_list:
            lead_id = lead.get("id", "")
            if not lead_id: continue

            # Check if already synced
            existing = db.execute(text("SELECT id FROM crm_deals WHERE id = :id"), {"id": lead_id}).fetchone()
            if existing:
                stats["skipped"] += 1
                continue

            # Savepoint for each lead (so one failure doesn't kill the whole transaction)
            db.execute(text(f"SAVEPOINT sp_{lead_id.replace('-','_')}"))
            try:
                # Company info
                company = lead.get("company") or {}
                if isinstance(company, str):
                    company = {"name": company, "short_name": company}
                company_name = company.get("name", "") or ""
                short_name = company.get("short_name", "") or company_name or lead_id

                # Create or find company
                company_id = None
                ebf = lead.get("ebf_integration") or {}
                cus_ref = ebf.get("customer_registry_ref", "") or ""
                if cus_ref:
                    company_id = cus_ref
                    existing_co = db.execute(text("SELECT id FROM crm_companies WHERE id = :id"), {"id": cus_ref}).fetchone()
                    if not existing_co:
                        hq = lead.get("headquarters") or {}
                        db.execute(text("""INSERT INTO crm_companies (id, name, domain, industry, size, website, address, notes, created_by)
                            VALUES (:id, :name, :domain, :ind, :size, :web, :addr, :notes, :cb)"""),
                            {"id": cus_ref, "name": company_name, "domain": short_name,
                             "ind": lead.get("industry","") or "", "size": lead.get("segment","") or "",
                             "web": company.get("website","") or "",
                             "addr": f"{hq.get('city','')}, {hq.get('country','')}",
                             "notes": f"Employees: {lead.get('employee_count','?')}, Revenue: {lead.get('revenue_eur','?')} EUR",
                             "cb": "github-sync"})
                        stats["companies"] += 1
                else:
                    company_id = f"CO-{lead_id}"
                    existing_co = db.execute(text("SELECT id FROM crm_companies WHERE id = :id"), {"id": company_id}).fetchone()
                    if not existing_co:
                        db.execute(text("""INSERT INTO crm_companies (id, name, domain, industry, notes, created_by)
                            VALUES (:id, :name, :domain, :ind, :notes, :cb)"""),
                            {"id": company_id, "name": company_name, "domain": short_name,
                             "ind": lead.get("industry","") or "", "notes": "", "cb": "github-sync"})
                        stats["companies"] += 1

                # Create contacts
                first_contact_id = None
                for ct in (lead.get("contacts") or []):
                    if not ct or not isinstance(ct, dict): continue
                    ct_name = ct.get("name", "") or ""
                    if not ct_name or ct_name.startswith("["):
                        continue
                    ct_id = f"CT-{lead_id}-{ct_name[:10].replace(' ','-')}"
                    existing_ct = db.execute(text("SELECT id FROM crm_contacts WHERE id = :id"), {"id": ct_id}).fetchone()
                    if not existing_ct:
                        db.execute(text("""INSERT INTO crm_contacts (id, company_id, name, email, phone, position, role_type, notes, created_by)
                            VALUES (:id, :cid, :name, :email, :phone, :pos, :role, :notes, :cb)"""),
                            {"id": ct_id, "cid": company_id, "name": ct_name,
                             "email": ct.get("email","") or "", "phone": ct.get("phone","") or "",
                             "pos": ct.get("role","") or "", "role": "champion" if ct.get("is_champion") else "kontakt",
                             "notes": f"Relationship: {ct.get('relationship_strength','?')}", "cb": "github-sync"})
                        stats["contacts"] += 1
                    if not first_contact_id:
                        first_contact_id = ct_id

                # Create deal
                rel = lead.get("relationship") or {}
                stage_raw = lead.get("stage", "PROSPECT") or "PROSPECT"
                stage = stage_map.get(stage_raw, stage_raw.lower())
                prob = CRM_STAGE_PROB.get(stage, 10)
                value = 0
                try: value = int(rel.get("contract_value_eur", 0) or 0)
                except: pass
                owner = rel.get("owner", "GF") or "GF"

                # Build notes
                notes_parts = []
                if lead.get("notes"): notes_parts.append(str(lead["notes"]).strip())
                if lead.get("fit_score"): notes_parts.append(f"Fit Score: {lead['fit_score']}")
                if lead.get("engagement_score"): notes_parts.append(f"Engagement: {lead['engagement_score']}")
                if lead.get("hot_lead"): notes_parts.append(f"🔥 HOT LEAD: {lead.get('hot_lead_reason','')}")
                if lead.get("priority"): notes_parts.append(f"Priority: {lead['priority']} - {lead.get('priority_reason','')}")
                if lead.get("tags") and isinstance(lead["tags"], list): notes_parts.append(f"Tags: {', '.join(str(t) for t in lead['tags'])}")
                notes_text = "\n".join(notes_parts)

                source_info = lead.get("source") or {}
                if isinstance(source_info, dict):
                    source_str = f"{source_info.get('channel','')} {source_info.get('subchannel','')} {source_info.get('referrer','')}".strip()
                else:
                    source_str = str(source_info)

                # Handle next_action (can be string or dict)
                na_raw = lead.get("next_action") or ""
                if isinstance(na_raw, dict):
                    na_str = na_raw.get("description", "") or na_raw.get("type", "")
                    na_date = na_raw.get("date") or lead.get("next_action_date")
                else:
                    na_str = str(na_raw)
                    na_date = lead.get("next_action_date")
                if isinstance(na_date, dict): na_date = None
                if na_date: na_date = str(na_date)[:10]  # ensure string date

                # Ensure created_at is a string
                ca_raw = lead.get("created") or lead.get("created_at") or datetime.utcnow().isoformat()
                if isinstance(ca_raw, dict): ca_raw = datetime.utcnow().isoformat()
                ca_str = str(ca_raw)

                db.execute(text("""INSERT INTO crm_deals (id, company_id, contact_id, title, stage, value, probability, source,
                    next_action, next_action_date, owner, notes, created_by, created_at)
                    VALUES (:id, :cid, :ctid, :title, :stage, :val, :prob, :src, :na, :nad, :owner, :notes, :cb, :ca)"""),
                    {"id": lead_id, "cid": company_id, "ctid": first_contact_id,
                     "title": f"{short_name} – {lead.get('industry','') or ''}",
                     "stage": stage, "val": value, "prob": prob, "src": source_str,
                     "na": na_str, "nad": na_date,
                     "owner": f"OWN-{owner}", "notes": notes_text, "cb": "github-sync",
                     "ca": ca_str})
                stats["deals"] += 1

                # Import contact_log as activities
                for log in (lead.get("contact_log") or []):
                    if not log or not isinstance(log, dict): continue
                    act_id = f"ACT-{lead_id}-{str(log.get('date',''))[:10]}"
                    existing_act = db.execute(text("SELECT id FROM crm_activities WHERE id = :id"), {"id": act_id}).fetchone()
                    if existing_act: continue
                    db.execute(text("""INSERT INTO crm_activities (id, deal_id, company_id, type, subject, description, created_by, created_at)
                        VALUES (:id, :did, :cid, :type, :subj, :desc, :cb, :ca)"""),
                        {"id": act_id, "did": lead_id, "cid": company_id,
                         "type": log.get("type", "note") or "note", "subj": str(log.get("notes",""))[:200],
                         "desc": f"Contact: {log.get('contact_person','')}\nOutcome: {log.get('outcome','')}\n{log.get('notes','')}",
                         "cb": log.get("logged_by", "github-sync") or "github-sync",
                         "ca": str(log.get("date", datetime.utcnow().isoformat()))})
                    stats["activities"] += 1

            except Exception as lead_err:
                db.execute(text(f"ROLLBACK TO SAVEPOINT sp_{lead_id.replace('-','_')}"))
                logger.warning(f"CRM Sync: skipping {lead_id}: {lead_err}")
                stats["skipped"] += 1
                if "errors" not in stats: stats["errors"] = []
                if len(stats.get("errors",[])) < 10:
                    stats["errors"].append(f"{lead_id}: {str(lead_err)[:120]}")
                continue

        db.commit()
        logger.info(f"CRM Sync complete: {stats}")
        return {"status": "synced", "stats": stats}
    except Exception as e:
        db.rollback()
        logger.error(f"CRM Sync error: {e}")
        import traceback; traceback.print_exc()
        raise HTTPException(500, f"Sync failed: {str(e)}")
    finally: db.close()

# ── CURRENCY API ──────────────────────────────────────
@app.get("/api/currency/rates")
async def get_fx_rates(user=Depends(require_auth)):
    """Get current exchange rates (base: CHF). Cached for 6h."""
    fx = fetch_exchange_rates()
    # Return subset of most relevant currencies
    relevant = ["CHF", "EUR", "USD", "GBP", "SEK", "NOK", "DKK", "PLN", "CZK", "HUF",
                 "CAD", "AUD", "SGD", "AED", "JPY", "CNY", "INR"]
    rates_subset = {c: fx.get("to_chf", {}).get(c) for c in relevant if fx.get("to_chf", {}).get(c)}
    return {
        "reporting_currency": REPORTING_CURRENCY,
        "rates_to_chf": rates_subset,
        "rates_from_chf": {c: fx.get("rates", {}).get(c) for c in relevant if fx.get("rates", {}).get(c)},
        "date": fx.get("date", ""),
        "symbols": {c: CURRENCY_SYMBOLS.get(c, c) for c in relevant}
    }


@app.post("/api/currency/convert")
async def convert_currency_endpoint(request: Request, user=Depends(require_auth)):
    """Convert amount between currencies. Body: {amount, from, to?}"""
    body = await request.json()
    amount = body.get("amount", 0)
    from_cur = body.get("from", "CHF").upper()
    to_cur = body.get("to", "CHF").upper()

    if not amount:
        return {"error": "amount required"}

    result = convert_amount(float(amount), from_cur, to_cur)
    chf_equiv = convert_to_chf(float(amount), from_cur)

    return {
        "original": format_currency(amount, from_cur),
        "converted": format_currency(result["amount"], to_cur),
        "chf_equivalent": format_currency(chf_equiv["amount_chf"], "CHF"),
        "details": result,
        "chf_details": chf_equiv
    }


@app.get("/api/currency/customer/{customer_code}")
async def get_customer_currency_endpoint(customer_code: str, user=Depends(require_auth)):
    """Get the default currency for a customer based on their country."""
    currency = get_customer_currency(customer_code)
    cache = load_customer_context_from_github()
    c = cache["customers"].get(customer_code.lower(), {})

    return {
        "customer_code": customer_code,
        "country": c.get("country", "?"),
        "currency": currency,
        "symbol": CURRENCY_SYMBOLS.get(currency, currency),
        "is_reporting_currency": currency == REPORTING_CURRENCY,
        "fx_to_chf": convert_to_chf(1.0, currency) if currency != "CHF" else None
    }


# ══════════════════════════════════════════════════════════════
# TASK MANAGEMENT SYSTEM
# 4 assignee types: consultant, beatrix, external, bdm
# Auto-triggers from leads, projects, chat
# ══════════════════════════════════════════════════════════════

TASK_STATUSES = ["open", "in_progress", "waiting", "done", "cancelled", "blocked"]
TASK_PRIORITIES = ["low", "normal", "high", "urgent"]
TASK_ASSIGNEE_TYPES = ["consultant", "beatrix", "external", "bdm"]
TASK_CATEGORIES_BDM = ["vertrag", "rechnung", "reise", "termin", "dokument", "zugang", "offerte", "archiv", "allgemein"]

# ── Auto-Trigger Templates ──
TRIGGER_TEMPLATES = {
    "lead_created": [
        {"title": "Erstgespräch vorbereiten", "assignee_type": "consultant", "category": "sales", "priority": "high", "due_offset_days": 3},
        {"title": "Kundenprofil prüfen & ergänzen", "assignee_type": "beatrix", "category": "data", "action_type": "enrich_customer_profile", "priority": "normal", "due_offset_days": 1},
    ],
    "lead_won": [
        {"title": "Offerte/Vertrag erstellen", "assignee_type": "bdm", "category": "vertrag", "priority": "urgent", "due_offset_days": 2},
        {"title": "Projekt eröffnen", "assignee_type": "consultant", "category": "project", "priority": "high", "due_offset_days": 3},
        {"title": "Zugänge einrichten", "assignee_type": "bdm", "category": "zugang", "priority": "normal", "due_offset_days": 5},
    ],
    "project_created": [
        {"title": "Kickoff-Termin koordinieren", "assignee_type": "bdm", "category": "termin", "priority": "high", "due_offset_days": 5},
        {"title": "Vertrag finalisieren & versenden", "assignee_type": "bdm", "category": "vertrag", "priority": "urgent", "due_offset_days": 3},
        {"title": "Projektdokumentation anlegen", "assignee_type": "beatrix", "category": "dokument", "action_type": "create_project_docs", "priority": "normal", "due_offset_days": 2},
        {"title": "Kickoff-Präsentation vorbereiten", "assignee_type": "consultant", "category": "delivery", "priority": "high", "due_offset_days": 7},
    ],
    "project_completed": [
        {"title": "Schlussrechnung erstellen", "assignee_type": "bdm", "category": "rechnung", "priority": "high", "due_offset_days": 5},
        {"title": "Projekt archivieren", "assignee_type": "bdm", "category": "archiv", "priority": "normal", "due_offset_days": 14},
        {"title": "Kundenfeedback einholen", "assignee_type": "bdm", "category": "allgemein", "priority": "normal", "due_offset_days": 7},
        {"title": "Lessons Learned dokumentieren", "assignee_type": "consultant", "category": "delivery", "priority": "normal", "due_offset_days": 10},
        {"title": "Projekt-Summary für Knowledge Base generieren", "assignee_type": "beatrix", "category": "data", "action_type": "generate_project_summary", "priority": "normal", "due_offset_days": 3},
    ],
    "lead_next_action": [
        {"title": "{next_action}", "assignee_type": "consultant", "category": "sales", "priority": "high", "due_offset_days": 0},
    ],
}


def generate_task_id():
    """Generate unique task ID like TSK-20260211-A3F2"""
    from datetime import datetime
    import secrets
    d = datetime.utcnow().strftime("%Y%m%d")
    r = secrets.token_hex(2).upper()
    return f"TSK-{d}-{r}"


def create_task_from_trigger(trigger_event: str, context: dict, created_by: str = "system"):
    """Auto-create tasks from trigger events. Context has customer_code, project_slug, lead_id, assignee, etc."""
    from datetime import datetime, timedelta
    templates = TRIGGER_TEMPLATES.get(trigger_event, [])
    if not templates:
        return []

    db = get_db()
    created_tasks = []
    try:
        for tmpl in templates:
            title = tmpl["title"].format(**context) if "{" in tmpl["title"] else tmpl["title"]
            due_date = None
            if tmpl.get("due_offset_days") is not None:
                base_date = context.get("base_date")
                if base_date:
                    if isinstance(base_date, str):
                        base_date = datetime.strptime(base_date[:10], "%Y-%m-%d")
                else:
                    base_date = datetime.utcnow()
                due_date = (base_date + timedelta(days=tmpl["due_offset_days"])).strftime("%Y-%m-%d")

            task_id = generate_task_id()
            assignee = context.get("fa_owner", "GF") if tmpl["assignee_type"] == "consultant" else (
                "system" if tmpl["assignee_type"] == "beatrix" else (
                "BDM" if tmpl["assignee_type"] == "bdm" else context.get("external_contact", "")))

            db.execute(text("""INSERT INTO tasks
                (id, title, assignee_type, assignee, customer_code, project_slug, lead_id,
                 priority, status, due_date, category, source, trigger_event, action_type, action_config, created_by)
                VALUES (:id, :title, :at, :a, :cc, :ps, :lid,
                 :pri, 'open', :due, :cat, 'auto', :te, :act, :acfg, :cb)"""), {
                "id": task_id, "title": title,
                "at": tmpl["assignee_type"], "a": assignee,
                "cc": context.get("customer_code"), "ps": context.get("project_slug"),
                "lid": context.get("lead_id"),
                "pri": tmpl.get("priority", "normal"),
                "due": due_date, "cat": tmpl.get("category"),
                "te": trigger_event, "act": tmpl.get("action_type"),
                "acfg": json.dumps(tmpl.get("action_config")) if tmpl.get("action_config") else None,
                "cb": created_by
            })
            created_tasks.append({"id": task_id, "title": title, "assignee_type": tmpl["assignee_type"],
                                  "assignee": assignee, "due_date": due_date})
        db.commit()
        logger.info(f"⚡ Trigger '{trigger_event}' → {len(created_tasks)} tasks created")
    except Exception as e:
        db.rollback()
        logger.error(f"Task trigger error: {e}")
    finally:
        db.close()
    return created_tasks


# ── Task API Endpoints ──

@app.get("/api/tasks")
async def list_tasks(
    request: Request,
    status: str = None, assignee_type: str = None, assignee: str = None,
    customer_code: str = None, project_slug: str = None, lead_id: str = None,
    priority: str = None, due_before: str = None, due_after: str = None,
    limit: int = 50, offset: int = 0,
    user=Depends(require_auth)
):
    """List tasks with filters."""
    db = get_db()
    try:
        where = ["1=1"]
        params = {"lim": limit, "off": offset}
        if status:
            if status == "active":
                where.append("status IN ('open','in_progress','waiting','blocked')")
            else:
                where.append("status = :status"); params["status"] = status
        if assignee_type:
            where.append("assignee_type = :at"); params["at"] = assignee_type
        if assignee:
            where.append("assignee = :a"); params["a"] = assignee
        if customer_code:
            where.append("customer_code = :cc"); params["cc"] = customer_code
        if project_slug:
            where.append("project_slug = :ps"); params["ps"] = project_slug
        if lead_id:
            where.append("lead_id = :lid"); params["lid"] = lead_id
        if priority:
            where.append("priority = :pri"); params["pri"] = priority
        if due_before:
            where.append("due_date <= :db"); params["db"] = due_before
        if due_after:
            where.append("due_date >= :da"); params["da"] = due_after

        w = " AND ".join(where)
        rows = db.execute(text(f"""SELECT * FROM tasks WHERE {w}
            ORDER BY CASE priority WHEN 'urgent' THEN 0 WHEN 'high' THEN 1 WHEN 'normal' THEN 2 ELSE 3 END,
            due_date ASC NULLS LAST, created_at DESC
            LIMIT :lim OFFSET :off"""), params).fetchall()

        count_row = db.execute(text(f"SELECT COUNT(*) FROM tasks WHERE {w}"), params).fetchone()
        total = count_row[0] if count_row else 0

        tasks = []
        for r in rows:
            t = dict(r._mapping)
            for k in ["created_at","updated_at","completed_at","waiting_since","due_date","due_time"]:
                if t.get(k): t[k] = str(t[k])
            if t.get("action_config") and isinstance(t["action_config"], str):
                try: t["action_config"] = json.loads(t["action_config"])
                except: pass
            tasks.append(t)

        return {"tasks": tasks, "total": total, "limit": limit, "offset": offset}
    finally: db.close()


@app.post("/api/tasks")
async def create_task(request: Request, user=Depends(require_auth)):
    """Create a new task."""
    body = await request.json()
    db = get_db()
    try:
        task_id = generate_task_id()
        db.execute(text("""INSERT INTO tasks
            (id, title, description, assignee_type, assignee, customer_code, project_slug, lead_id, deal_id,
             priority, status, due_date, due_time, category, source, action_type, action_config,
             escalation_after_days, parent_task_id, sort_order, created_by)
            VALUES (:id, :title, :desc, :at, :a, :cc, :ps, :lid, :did,
             :pri, :st, :due, :duet, :cat, :src, :act, :acfg,
             :esc, :ptid, :so, :cb)"""), {
            "id": task_id,
            "title": body.get("title", ""),
            "desc": body.get("description"),
            "at": body.get("assignee_type", "consultant"),
            "a": body.get("assignee", ""),
            "cc": body.get("customer_code"),
            "ps": body.get("project_slug"),
            "lid": body.get("lead_id"),
            "did": body.get("deal_id"),
            "pri": body.get("priority", "normal"),
            "st": body.get("status", "open"),
            "due": body.get("due_date"),
            "duet": body.get("due_time"),
            "cat": body.get("category"),
            "src": body.get("source", "manual"),
            "act": body.get("action_type"),
            "acfg": json.dumps(body.get("action_config")) if body.get("action_config") else None,
            "esc": body.get("escalation_after_days"),
            "ptid": body.get("parent_task_id"),
            "so": body.get("sort_order", 0),
            "cb": user.get("email", user.get("sub", ""))
        })
        db.commit()

        # If external task, set waiting_since
        if body.get("assignee_type") == "external":
            db.execute(text("UPDATE tasks SET waiting_since = CURRENT_TIMESTAMP WHERE id = :id"), {"id": task_id})
            db.commit()

        return {"id": task_id, "status": "created"}
    except Exception as e:
        db.rollback()
        raise HTTPException(500, str(e))
    finally: db.close()


@app.put("/api/tasks/{task_id}")
async def update_task(task_id: str, request: Request, user=Depends(require_auth)):
    """Update a task."""
    body = await request.json()
    db = get_db()
    try:
        existing = db.execute(text("SELECT * FROM tasks WHERE id = :id"), {"id": task_id}).fetchone()
        if not existing:
            raise HTTPException(404, "Task not found")

        updates = []
        params = {"id": task_id}
        for field in ["title","description","assignee_type","assignee","customer_code","project_slug",
                       "lead_id","deal_id","priority","status","due_date","due_time","category",
                       "action_type","escalation_after_days","parent_task_id","sort_order"]:
            if field in body:
                updates.append(f"{field} = :{field}")
                params[field] = body[field]

        # Handle status transitions
        old_status = existing._mapping.get("status")
        new_status = body.get("status")
        if new_status and new_status != old_status:
            if new_status == "done":
                updates.append("completed_at = CURRENT_TIMESTAMP")
                updates.append("completed_by = :completed_by")
                params["completed_by"] = user.get("email", user.get("sub", ""))
            elif new_status == "waiting" and old_status != "waiting":
                updates.append("waiting_since = CURRENT_TIMESTAMP")
            elif old_status == "done" and new_status != "done":
                updates.append("completed_at = NULL")
                updates.append("completed_by = NULL")

        updates.append("updated_at = CURRENT_TIMESTAMP")
        set_clause = ", ".join(updates)
        db.execute(text(f"UPDATE tasks SET {set_clause} WHERE id = :id"), params)
        db.commit()
        return {"id": task_id, "status": "updated"}
    except HTTPException: raise
    except Exception as e:
        db.rollback()
        raise HTTPException(500, str(e))
    finally: db.close()


@app.delete("/api/tasks/{task_id}")
async def delete_task(task_id: str, user=Depends(require_auth)):
    """Delete a task."""
    db = get_db()
    try:
        db.execute(text("DELETE FROM tasks WHERE id = :id"), {"id": task_id})
        db.commit()
        return {"id": task_id, "status": "deleted"}
    finally: db.close()


@app.post("/api/tasks/{task_id}/complete")
async def complete_task(task_id: str, user=Depends(require_auth)):
    """Quick-complete a task."""
    db = get_db()
    try:
        db.execute(text("""UPDATE tasks SET status='done', completed_at=CURRENT_TIMESTAMP,
            completed_by=:cb, updated_at=CURRENT_TIMESTAMP WHERE id=:id"""),
            {"id": task_id, "cb": user.get("email", user.get("sub", ""))})
        db.commit()
        return {"id": task_id, "status": "done"}
    finally: db.close()


@app.post("/api/tasks/{task_id}/reopen")
async def reopen_task(task_id: str, user=Depends(require_auth)):
    """Reopen a completed task."""
    db = get_db()
    try:
        db.execute(text("""UPDATE tasks SET status='open', completed_at=NULL,
            completed_by=NULL, updated_at=CURRENT_TIMESTAMP WHERE id=:id"""), {"id": task_id})
        db.commit()
        return {"id": task_id, "status": "open"}
    finally: db.close()


@app.get("/api/tasks/stats")
async def task_stats(user=Depends(require_auth)):
    """Task dashboard stats."""
    db = get_db()
    try:
        stats = {}
        # Counts by status
        rows = db.execute(text("SELECT status, COUNT(*) as c FROM tasks GROUP BY status")).fetchall()
        stats["by_status"] = {r[0]: r[1] for r in rows}
        # Counts by assignee_type
        rows = db.execute(text("SELECT assignee_type, COUNT(*) as c FROM tasks WHERE status NOT IN ('done','cancelled') GROUP BY assignee_type")).fetchall()
        stats["by_type"] = {r[0]: r[1] for r in rows}
        # Overdue
        row = db.execute(text("SELECT COUNT(*) FROM tasks WHERE due_date < CURRENT_DATE AND status NOT IN ('done','cancelled')")).fetchone()
        stats["overdue"] = row[0] if row else 0
        # Due today
        row = db.execute(text("SELECT COUNT(*) FROM tasks WHERE due_date = CURRENT_DATE AND status NOT IN ('done','cancelled')")).fetchone()
        stats["due_today"] = row[0] if row else 0
        # Due this week
        row = db.execute(text("SELECT COUNT(*) FROM tasks WHERE due_date BETWEEN CURRENT_DATE AND CURRENT_DATE + 7 AND status NOT IN ('done','cancelled')")).fetchone()
        stats["due_this_week"] = row[0] if row else 0
        # Total active
        row = db.execute(text("SELECT COUNT(*) FROM tasks WHERE status NOT IN ('done','cancelled')")).fetchone()
        stats["active_total"] = row[0] if row else 0
        # By assignee (top 10)
        rows = db.execute(text("""SELECT assignee, assignee_type, COUNT(*) as c FROM tasks
            WHERE status NOT IN ('done','cancelled') GROUP BY assignee, assignee_type ORDER BY c DESC LIMIT 10""")).fetchall()
        stats["by_assignee"] = [{"assignee": r[0], "type": r[1], "count": r[2]} for r in rows]
        # Escalation warnings (external tasks waiting too long)
        rows = db.execute(text("""SELECT id, title, assignee, customer_code, waiting_since, escalation_after_days
            FROM tasks WHERE assignee_type = 'external' AND status IN ('open','waiting')
            AND waiting_since IS NOT NULL AND escalation_after_days IS NOT NULL
            AND waiting_since + (escalation_after_days || ' days')::interval < CURRENT_TIMESTAMP""")).fetchall()
        stats["escalations"] = [dict(r._mapping) for r in rows]
        for e in stats["escalations"]:
            for k in ["waiting_since"]: e[k] = str(e[k]) if e.get(k) else None

        return stats
    finally: db.close()


@app.post("/api/tasks/trigger")
async def trigger_tasks(request: Request, user=Depends(require_auth)):
    """Manually fire a trigger event. Body: {event, context}"""
    body = await request.json()
    event = body.get("event")
    context = body.get("context", {})
    if not event:
        raise HTTPException(400, "event required")
    tasks = create_task_from_trigger(event, context, created_by=user.get("email", user.get("sub", "")))
    return {"event": event, "tasks_created": len(tasks), "tasks": tasks}


# ── BEATRIX Memory (GitHub-backed) ──────────────────
@app.get("/api/memory")
async def get_memory(user=Depends(require_auth)):
    """Read BEATRIX extended memory from GitHub"""
    import urllib.request, ssl, yaml
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    try:
        url = f"https://api.github.com/repos/{GH_REPO}/contents/data/beatrix/memory.yaml"
        req = urllib.request.Request(url, headers={"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3.raw", "User-Agent": "BEATRIXLab"})
        content = urllib.request.urlopen(req, context=ctx, timeout=15).read().decode()
        data = yaml.safe_load(content) or {}
        return data
    except Exception as e:
        logger.error(f"Memory read error: {e}")
        return {"error": str(e)}

@app.post("/api/memory/add")
async def add_memory(request: Request, user=Depends(require_auth)):
    """Add a learning or decision to BEATRIX memory on GitHub"""
    import urllib.request, ssl, yaml, base64
    from datetime import date
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    body = await request.json()
    entry_type = body.get("type", "learning")  # learning or decision
    topic = body.get("topic", "").strip()
    content = body.get("content", "").strip()
    if not topic or not content:
        return JSONResponse({"error": "topic and content required"}, status_code=400)
    try:
        # Fetch current memory
        url = f"https://api.github.com/repos/{GH_REPO}/contents/data/beatrix/memory.yaml"
        req = urllib.request.Request(url, headers={"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab"})
        existing = json.loads(urllib.request.urlopen(req, context=ctx, timeout=15).read())
        sha = existing["sha"]
        raw_url = f"https://api.github.com/repos/{GH_REPO}/contents/data/beatrix/memory.yaml"
        req2 = urllib.request.Request(raw_url, headers={"Authorization": f"token {GH_TOKEN}", "Accept": "application/vnd.github.v3.raw", "User-Agent": "BEATRIXLab"})
        data = yaml.safe_load(urllib.request.urlopen(req2, context=ctx, timeout=15).read().decode()) or {}
        today = date.today().isoformat()
        if entry_type == "decision":
            decisions = data.get("decisions", [])
            new_id = f"D{len(decisions)+1:03d}"
            decisions.append({"id": new_id, "topic": topic, "decision": content, "date": today})
            data["decisions"] = decisions
        else:
            learnings = data.get("learnings", [])
            new_id = f"L{len(learnings)+1:03d}"
            learnings.append({"id": new_id, "topic": topic, "lesson": content, "date": today})
            data["learnings"] = learnings
        data["metadata"]["last_updated"] = today
        # Push back
        yaml_content = yaml.dump(data, default_flow_style=False, allow_unicode=True, sort_keys=False)
        put_data = json.dumps({"message": f"memory: {entry_type} {new_id} – {topic}", "content": base64.b64encode(yaml_content.encode('utf-8')).decode(), "sha": sha, "branch": "main"}).encode('utf-8')
        req3 = urllib.request.Request(url, data=put_data, method="PUT", headers={"Authorization": f"token {GH_TOKEN}", "Content-Type": "application/json", "User-Agent": "BEATRIXLab"})
        urllib.request.urlopen(req3, context=ctx, timeout=15)
        logger.info(f"Memory added: {new_id} by {user.get('sub','?')}")
        return {"ok": True, "id": new_id, "type": entry_type}
    except Exception as e:
        logger.error(f"Memory add error: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)

# ── Legacy Leads API (kept for backward compat) ──────
@app.get("/api/leads")
async def get_leads(user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        rows = db.execute(text("SELECT * FROM leads ORDER BY updated_at DESC")).fetchall()
        return [dict(r._mapping) for r in rows]
    except Exception as e:
        logger.error(f"Leads fetch error: {e}")
        return []
    finally: db.close()

@app.post("/api/leads")
async def create_lead(request: Request, user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        import uuid
        data = await request.json()
        lead_id = str(uuid.uuid4())[:8]
        db.execute(text("""INSERT INTO leads (id, company, contact, email, stage, value, source, notes, created_by)
            VALUES (:id, :company, :contact, :email, :stage, :value, :source, :notes, :created_by)"""),
            {"id": lead_id, "company": data.get("company",""), "contact": data.get("contact",""),
             "email": data.get("email",""), "stage": data.get("stage","kontakt"),
             "value": data.get("value", 0), "source": data.get("source",""),
             "notes": data.get("notes",""), "created_by": user["sub"]})
        db.commit()

        # ── Auto-trigger tasks for new lead ──
        try:
            trigger_ctx = {
                "customer_code": data.get("customer_code", ""),
                "lead_id": lead_id,
                "fa_owner": data.get("fa_owner", "GF"),
                "base_date": datetime.utcnow().strftime("%Y-%m-%d")
            }
            create_task_from_trigger("lead_created", trigger_ctx, created_by=user["sub"])
            # If next_action provided, create specific task
            if data.get("next_action"):
                trigger_ctx["next_action"] = data.get("next_action", "Follow-up")
                if data.get("next_action_date"):
                    trigger_ctx["base_date"] = data["next_action_date"]
                create_task_from_trigger("lead_next_action", trigger_ctx, created_by=user["sub"])
        except Exception as te:
            logger.warning(f"Lead task trigger failed: {te}")

        return {"id": lead_id, "status": "created"}
    except Exception as e:
        db.rollback()
        logger.error(f"Lead create error: {e}")
        raise HTTPException(500, str(e))
    finally: db.close()

@app.put("/api/leads/{lead_id}")
async def update_lead(lead_id: str, request: Request, user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        data = await request.json()

        # Check if stage changed to "won" for trigger
        old_stage = None
        if "stage" in data:
            old = db.execute(text("SELECT stage, company FROM leads WHERE id = :id"), {"id": lead_id}).fetchone()
            if old: old_stage = old[0]

        sets = []
        params = {"id": lead_id}
        for field in ["company", "contact", "email", "stage", "value", "source", "notes"]:
            if field in data:
                sets.append(f"{field} = :{field}")
                params[field] = data[field]
        if not sets: return {"status": "no changes"}
        sets.append("updated_at = CURRENT_TIMESTAMP")
        db.execute(text(f"UPDATE leads SET {', '.join(sets)} WHERE id = :id"), params)
        db.commit()

        # ── Auto-trigger tasks on stage change ──
        if data.get("stage") == "won" and old_stage != "won":
            try:
                create_task_from_trigger("lead_won", {
                    "customer_code": data.get("customer_code", ""),
                    "lead_id": lead_id,
                    "fa_owner": data.get("fa_owner", "GF"),
                    "base_date": datetime.utcnow().strftime("%Y-%m-%d")
                }, created_by=user["sub"])
            except Exception as te:
                logger.warning(f"Lead-won trigger failed: {te}")

        return {"status": "updated"}
    except Exception as e:
        db.rollback()
        logger.error(f"Lead update error: {e}")
        raise HTTPException(500, str(e))
    finally: db.close()

@app.delete("/api/leads/{lead_id}")
async def delete_lead(lead_id: str, user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        db.execute(text("DELETE FROM leads WHERE id = :id"), {"id": lead_id})
        db.commit()
        return {"status": "deleted"}
    except Exception as e:
        db.rollback()
        raise HTTPException(500, str(e))
    finally: db.close()

# ── Contexts / Ausgangslage API ────────────────────
@app.get("/api/contexts")
async def get_contexts(user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        rows = db.execute(text("SELECT * FROM contexts ORDER BY updated_at DESC")).fetchall()
        return [dict(r._mapping) for r in rows]
    except Exception as e:
        logger.error(f"Contexts fetch error: {e}")
        return []
    finally: db.close()

@app.post("/api/contexts")
async def create_context(request: Request, user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        import uuid
        data = await request.json()
        ctx_id = str(uuid.uuid4())[:8]
        db.execute(text("""INSERT INTO contexts (id, client, project, domain, status, situation, goal, constraints, created_by)
            VALUES (:id, :client, :project, :domain, :status, :situation, :goal, :constraints, :created_by)"""),
            {"id": ctx_id, "client": data.get("client",""), "project": data.get("project",""),
             "domain": data.get("domain","OTH"), "status": data.get("status","aktiv"),
             "situation": data.get("situation",""), "goal": data.get("goal",""),
             "constraints": data.get("constraints",""), "created_by": user["sub"]})
        db.commit()
        return {"id": ctx_id, "status": "created"}
    except Exception as e:
        db.rollback()
        logger.error(f"Context create error: {e}")
        raise HTTPException(500, str(e))
    finally: db.close()

@app.put("/api/contexts/{ctx_id}")
async def update_context(ctx_id: str, request: Request, user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        data = await request.json()
        sets = []
        params = {"id": ctx_id}
        for field in ["client", "project", "domain", "status", "situation", "goal", "constraints"]:
            if field in data:
                sets.append(f"{field} = :{field}")
                params[field] = data[field]
        if not sets: return {"status": "no changes"}
        sets.append("updated_at = CURRENT_TIMESTAMP")
        db.execute(text(f"UPDATE contexts SET {', '.join(sets)} WHERE id = :id"), params)
        db.commit()
        return {"status": "updated"}
    except Exception as e:
        db.rollback()
        logger.error(f"Context update error: {e}")
        raise HTTPException(500, str(e))
    finally: db.close()

@app.delete("/api/contexts/{ctx_id}")
async def delete_context(ctx_id: str, user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        db.execute(text("DELETE FROM contexts WHERE id = :id"), {"id": ctx_id})
        db.commit()
        return {"status": "deleted"}
    except Exception as e:
        db.rollback()
        raise HTTPException(500, str(e))
    finally: db.close()

# ── Ψ-Analysis API ─────────────────────────────────
@app.get("/api/psi-analyses")
async def get_psi_analyses(user=Depends(require_auth)):
    """Get all Ψ-analyses for current user, grouped by lineage with latest version first."""
    db = get_db()
    try:
        from sqlalchemy import text
        rows = db.execute(text("""
            SELECT id, lineage_id, version, question, mode, psi_profile, parameters,
                   synthesis, confidence, customer_code, project_slug, created_by, created_at
            FROM psi_analyses
            WHERE created_by = :email
            ORDER BY created_at DESC
        """), {"email": user["sub"]}).fetchall()
        return [dict(r._mapping) for r in rows]
    except Exception as e:
        logger.error(f"Psi analyses fetch error: {e}")
        return []
    finally: db.close()

@app.get("/api/psi-analyses/all")
async def get_all_psi_analyses(user=Depends(require_permission("platform.view_analytics"))):
    """Admin: Get all Ψ-analyses across all users for calibration comparison."""
    db = get_db()
    try:
        from sqlalchemy import text
        rows = db.execute(text("""
            SELECT id, lineage_id, version, question, mode, psi_profile, parameters,
                   synthesis, confidence, customer_code, project_slug, created_by, created_at
            FROM psi_analyses ORDER BY created_at DESC LIMIT 200
        """)).fetchall()
        return [dict(r._mapping) for r in rows]
    except Exception as e:
        logger.error(f"Psi analyses admin fetch: {e}")
        return []
    finally: db.close()

@app.get("/api/psi-analyses/lineage/{lineage_id}")
async def get_psi_lineage(lineage_id: str, user=Depends(require_auth)):
    """Get all versions of a specific analysis lineage (all users for comparison)."""
    db = get_db()
    try:
        from sqlalchemy import text
        rows = db.execute(text("""
            SELECT pa.*, u.name as user_name
            FROM psi_analyses pa LEFT JOIN users u ON pa.created_by = u.email
            WHERE pa.lineage_id = :lid ORDER BY pa.version DESC
        """), {"lid": lineage_id}).fetchall()
        return [dict(r._mapping) for r in rows]
    except Exception as e:
        logger.error(f"Psi lineage fetch error: {e}")
        return []
    finally: db.close()

@app.post("/api/psi-analyses")
async def create_psi_analysis(request: Request, user=Depends(require_auth)):
    """Create a new Ψ-analysis. Auto-detects lineage from question similarity."""
    db = get_db()
    try:
        from sqlalchemy import text
        import uuid, hashlib
        data = await request.json()

        question = data.get("question", "").strip()
        if not question:
            raise HTTPException(400, "Question required")

        # Determine lineage: check if same user has existing analysis with similar question
        lineage_id = data.get("lineage_id")  # explicit re-version
        version = 1

        if lineage_id:
            # Explicit new version of existing lineage
            existing = db.execute(text("""
                SELECT MAX(version) as max_v, id FROM psi_analyses
                WHERE lineage_id = :lid AND created_by = :email
                GROUP BY id ORDER BY max_v DESC LIMIT 1
            """), {"lid": lineage_id, "email": user["sub"]}).fetchone()
            if existing:
                version = (existing[0] or 0) + 1
        else:
            # Check for similar question by same user → auto-detect lineage
            existing = db.execute(text("""
                SELECT lineage_id, MAX(version) as max_v FROM psi_analyses
                WHERE created_by = :email AND LOWER(question) = LOWER(:q)
                GROUP BY lineage_id ORDER BY max_v DESC LIMIT 1
            """), {"email": user["sub"], "q": question}).fetchone()
            if existing:
                lineage_id = existing[0]
                version = (existing[1] or 0) + 1
            else:
                # New lineage: create ID from date + sequence
                today = datetime.utcnow().strftime("%Y-%m-%d")
                seq = db.execute(text("""
                    SELECT COUNT(*) FROM psi_analyses WHERE lineage_id LIKE :pattern
                """), {"pattern": f"PSI-{today}%"}).fetchone()[0]
                lineage_id = f"PSI-{today}-{seq+1:03d}"

        # Generate analysis ID
        analysis_id = f"{lineage_id}-v{version}"

        db.execute(text("""
            INSERT INTO psi_analyses (id, lineage_id, version, question, mode,
                macro_context, meso_context, micro_context, psi_profile, parameters,
                synthesis, implications, confidence, customer_code, project_slug,
                parent_version_id, created_by)
            VALUES (:id, :lid, :ver, :q, :mode, CAST(:macro AS jsonb), CAST(:meso AS jsonb),
                CAST(:micro AS jsonb), CAST(:psi AS jsonb), CAST(:params AS jsonb), :synth,
                CAST(:impl AS jsonb), :conf, :cust, :proj, :parent, :email)
        """), {
            "id": analysis_id,
            "lid": lineage_id,
            "ver": version,
            "q": question,
            "mode": data.get("mode", "schnell"),
            "macro": json.dumps(data.get("macro_context")) if data.get("macro_context") else None,
            "meso": json.dumps(data.get("meso_context")) if data.get("meso_context") else None,
            "micro": json.dumps(data.get("micro_context")) if data.get("micro_context") else None,
            "psi": json.dumps(data.get("psi_profile")) if data.get("psi_profile") else None,
            "params": json.dumps(data.get("parameters")) if data.get("parameters") else None,
            "synth": data.get("synthesis", ""),
            "impl": json.dumps(data.get("implications")) if data.get("implications") else None,
            "conf": data.get("confidence", ""),
            "cust": data.get("customer_code"),
            "proj": data.get("project_slug"),
            "parent": data.get("parent_version_id"),
            "email": user["sub"]
        })
        db.commit()

        # Async GitHub sync
        try:
            _sync_psi_to_github(analysis_id, lineage_id, version, question, data, user["sub"])
        except Exception as gh_err:
            logger.warning(f"Psi GitHub sync failed: {gh_err}")

        logger.info(f"Ψ-Analysis saved: {analysis_id} (lineage={lineage_id}, v{version}) by {user['sub']}")
        return {"id": analysis_id, "lineage_id": lineage_id, "version": version, "status": "created"}
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        logger.error(f"Psi analysis create error: {e}")
        raise HTTPException(500, str(e))
    finally: db.close()

@app.delete("/api/psi-analyses/{analysis_id}")
async def delete_psi_analysis(analysis_id: str, user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        db.execute(text("DELETE FROM psi_analyses WHERE id = :id AND created_by = :email"),
                   {"id": analysis_id, "email": user["sub"]})
        db.commit()
        return {"status": "deleted"}
    except Exception as e:
        db.rollback()
        raise HTTPException(500, str(e))
    finally: db.close()

def _sync_psi_to_github(analysis_id, lineage_id, version, question, data, user_email):
    """Push Ψ-analysis to GitHub as YAML."""
    import yaml, re as _re
    if not GH_TOKEN: return

    user_slug = _re.sub(r'[^a-z0-9]+', '-', user_email.split('@')[0].lower()).strip('-')
    yaml_content = {
        "id": analysis_id,
        "lineage_id": lineage_id,
        "version": version,
        "question": question,
        "mode": data.get("mode", "schnell"),
        "created_by": user_email,
        "created_at": datetime.utcnow().isoformat() + "Z",
        "macro_context": data.get("macro_context"),
        "meso_context": data.get("meso_context"),
        "micro_context": data.get("micro_context"),
        "psi_profile": data.get("psi_profile"),
        "parameters": data.get("parameters"),
        "synthesis": data.get("synthesis"),
        "implications": data.get("implications"),
        "confidence": data.get("confidence"),
        "customer_code": data.get("customer_code"),
        "project_slug": data.get("project_slug"),
    }

    yaml_str = yaml.dump(yaml_content, allow_unicode=True, default_flow_style=False, sort_keys=False)
    file_path = f"data/psi-analyses/{user_slug}/{lineage_id}-v{version}.yaml"

    import urllib.request, ssl, base64
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    gh_repo = os.environ.get("GH_REPO", "FehrAdvice-Partners-AG/complementarity-context-framework")
    headers = {"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIXLab/3.7", "Content-Type": "application/json"}

    # Check if file exists (get SHA for update)
    sha = None
    try:
        req = urllib.request.Request(f"https://api.github.com/repos/{gh_repo}/contents/{file_path}", headers=headers)
        resp = json.loads(urllib.request.urlopen(req, context=ctx).read())
        sha = resp.get("sha")
    except: pass

    payload = {
        "message": f"Ψ-Analysis: {lineage_id} v{version} by {user_slug}",
        "content": base64.b64encode(yaml_str.encode('utf-8')).decode(),
        "branch": "main"
    }
    if sha: payload["sha"] = sha

    try:
        req = urllib.request.Request(f"https://api.github.com/repos/{gh_repo}/contents/{file_path}",
            data=json.dumps(payload).encode("utf-8"), method="PUT", headers=headers)
        urllib.request.urlopen(req, context=ctx)
        logger.info(f"Ψ-Analysis synced to GitHub: {file_path}")
    except Exception as e:
        logger.warning(f"Ψ GitHub push failed for {file_path}: {e}")

# ── Feedback API ────────────────────────────────────
@app.post("/api/feedback")
async def create_feedback(request: Request, user=Depends(require_auth)):
    db = get_db()
    try:
        from sqlalchemy import text
        import uuid
        data = await request.json()
        fb_id = str(uuid.uuid4())[:8]
        db.execute(text("""INSERT INTO feedback (id, type, comment, screenshot, page, url, viewport, user_agent, created_by)
            VALUES (:id, :type, :comment, :screenshot, :page, :url, :viewport, :ua, :created_by)"""),
            {"id": fb_id, "type": data.get("type","bug"), "comment": data.get("comment",""),
             "screenshot": data.get("screenshot",""), "page": data.get("page",""),
             "url": data.get("url",""), "viewport": data.get("viewport",""),
             "ua": data.get("userAgent",""), "created_by": user["sub"]})
        db.commit()
        logger.info(f"Feedback from {user['sub']}: [{data.get('type')}] {data.get('comment','')[:60]}")
        return {"id": fb_id, "status": "created"}
    except Exception as e:
        db.rollback()
        logger.error(f"Feedback error: {e}")
        raise HTTPException(500, str(e))
    finally: db.close()

@app.post("/api/feedback/analyze")
async def analyze_feedback(request: Request, user=Depends(require_auth)):
    """AI-powered screen analysis: BEATRIX suggests improvements"""
    try:
        data = await request.json()
        screenshot = data.get("screenshot", "")
        page = data.get("page", "unknown")
        viewport = data.get("viewport", "")

        if not screenshot or not ANTHROPIC_API_KEY:
            return {"suggestions": []}

        # Extract base64 data from data URL
        img_data = screenshot.split(",")[1] if "," in screenshot else screenshot

        import urllib.request, ssl
        ctx = ssl.create_default_context()
        ctx.check_hostname = False
        ctx.verify_mode = ssl.CERT_NONE

        payload = json.dumps({
            "model": ANTHROPIC_MODEL_STANDARD,
            "max_tokens": 600,
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": img_data}},
                    {"type": "text", "text": f"""Du bist BEATRIX, eine Strategic Intelligence Suite. Analysiere diesen Screenshot der Seite '{page}' (Viewport: {viewport}).

Gib genau 3 konkrete, umsetzbare Verbesserungsvorschläge als JSON-Array zurück. Jeder Vorschlag hat:
- "type": "bug" | "ux" | "wunsch" | "performance"
- "title": kurzer Titel (max 40 Zeichen)
- "description": konkrete Beschreibung (max 80 Zeichen)
- "x": horizontale Position in Prozent (0-100) wo das Problem auf dem Screen ist
- "y": vertikale Position in Prozent (0-100) wo das Problem auf dem Screen ist
- "w": Breite des betroffenen Bereichs in Prozent (5-40)
- "h": Höhe des betroffenen Bereichs in Prozent (5-30)

Die Koordinaten sollen möglichst genau auf den relevanten UI-Bereich zeigen.
Fokussiere auf: UX-Verbesserungen, fehlende Features, visuelle Probleme, Barrierefreiheit.
Antworte NUR mit dem JSON-Array, kein anderer Text."""}
                ]
            }]
        }).encode()

        req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST", headers={
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json"
        })

        resp = urllib.request.urlopen(req, context=ctx, timeout=30)
        result = json.loads(resp.read())
        text = result.get("content", [{}])[0].get("text", "[]").strip()
        if text.startswith("```"): text = text.split("\n", 1)[1].rsplit("```", 1)[0]
        suggestions = json.loads(text)
        logger.info(f"Feedback analyze for {user['sub']}: {len(suggestions)} suggestions")
        return {"suggestions": suggestions[:3]}

    except Exception as e:
        logger.error(f"Feedback analyze error: {e}")
        return {"suggestions": []}

@app.get("/api/feedback")
async def get_feedback(user=Depends(require_permission("platform.admin_dashboard"))):
    db = get_db()
    try:
        from sqlalchemy import text
        rows = db.execute(text("SELECT id, type, comment, page, viewport, user_agent, status, created_by, created_at FROM feedback ORDER BY created_at DESC")).fetchall()
        return [dict(r._mapping) for r in rows]
    except Exception as e:
        logger.error(f"Feedback list error: {e}")
        return []
    finally: db.close()

@app.get("/api/feedback/{fb_id}/screenshot")
async def get_feedback_screenshot(fb_id: str, user=Depends(require_permission("platform.admin_dashboard"))):
    db = get_db()
    try:
        from sqlalchemy import text
        row = db.execute(text("SELECT screenshot FROM feedback WHERE id = :id"), {"id": fb_id}).fetchone()
        if not row: raise HTTPException(404, "Not found")
        return {"screenshot": row[0]}
    finally: db.close()

@app.put("/api/feedback/{fb_id}")
async def update_feedback(fb_id: str, request: Request, user=Depends(require_permission("platform.admin_dashboard"))):
    db = get_db()
    try:
        from sqlalchemy import text
        data = await request.json()
        status = data.get("status", "neu")
        db.execute(text("UPDATE feedback SET status = :status WHERE id = :id"), {"id": fb_id, "status": status})
        db.commit()
        return {"status": "updated"}
    finally: db.close()


# ============================================================================
# BEATRIX Review System – Behavioral Economics Paper Reviewer
# The 6th Reviewer for APE papers (and any paper in the KB)
# ============================================================================

BEATRIX_REVIEW_SYSTEM_PROMPT = """Du bist BEATRIX, ein Behavioral Economics Reviewer.
Du bist der 6. Reviewer in einem System mit 5 statistisch/methodisch orientierten Reviewern.

Deine EINZIGARTIGE Perspektive:
- Du bewertest Papers durch die Linse der Verhaltensökonomie
- Du nutzt das BCM 2.0 Framework (Behavioral Complementarity Model)
- Du analysierst die 8 Ψ-Dimensionen
- Du identifizierst behavioral mechanisms, die statistische Reviewer übersehen
- Du gibst praktische Gestaltungsempfehlungen für Policy-Interventionen

Du prüfst NICHT (das machen die anderen Reviewer):
- Statistische Korrektheit (Format Check, Regression Sanity)
- Identifikationsstrategie-Details (Pre-Trends, Placebo Tests)
- Formale Journal-Anforderungen

Du prüfst NUR:

## 1. BEHAVIORAL MECHANISM ANALYSIS
Welche behavioral mechanisms sind im Spiel? Welche werden diskutiert, welche fehlen?
(Present bias, defaults, social norms, framing, anchoring, loss aversion, 
choice architecture, commitment devices, identity utility, etc.)

## 2. Ψ-DIMENSIONS MAPPING (BCM 2.0)
Welche der 8 Ψ-Dimensionen sind relevant?
- Ψ₁ Risk Perception, Ψ₂ Time Preferences, Ψ₃ Social Preferences
- Ψ₄ Reference Dependence, Ψ₅ Attention/Salience, Ψ₆ Self-Control
- Ψ₇ Beliefs & Mental Models, Ψ₈ Identity & Self-Image
Tabellenformat: Dimension | Relevance | Addressed? | Gap

## 3. POLICY DESIGN IMPLICATIONS
Wie könnte die Intervention behavioral verbessert werden?
Commitment devices, framing, choice architecture, nudges, scaling predictions.
Heterogeneous treatment effects by behavioral type.

## 4. LITERATURE GAPS
Welche behavioral economics Literatur fehlt? Max. 7 Referenzen.
Nur wirklich relevante – keine generischen Textbook-Zitate.

## 5. SCORING
Tabellenformat:
| Dimension | Score (0-25) | Reasoning |
BCM Relevance, Methodological Rigor, Data Novelty, Practical Applicability
Total Score: X/100 → PURSUE (>65) / CONSIDER (45-65) / SKIP (<45)

## 6. VERDICT
ACCEPT / CONDITIONAL ACCEPT / MAJOR REVISION / REJECT
Immer aus behavioral Perspektive, nicht statistisch.

## 7. KB INTEGRATION TAGS
JSON-Block mit: bcm_dimensions, psi_dimensions, behavioral_mechanisms,
evidence_type, method, policy_domain, application_areas, connections

Antworte IMMER im Markdown-Format mit allen 7 Sections.
Sei konkret und spezifisch – beziehe dich auf exakte Textstellen.
Sprache: Englisch für den Review-Text, unabhängig von der Paper-Sprache."""


def call_claude_review(paper_content: str, paper_id: str, title: str,
                       method: str = "unknown", model_name: str = "unknown") -> str:
    """Call Claude API to generate a BEATRIX Behavioral Economics Review."""
    import urllib.request, ssl
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE

    if not ANTHROPIC_API_KEY:
        return "ERROR: Claude API not configured"

    user_prompt = f"""Please create a BEATRIX Behavioral Economics Review for:

**Paper ID:** {paper_id}
**Title:** {title}
**Method:** {method}
**Authoring Model:** {model_name}

**Paper Content:**
{paper_content[:35000]}

---
Create the review with all 7 sections. Focus on what statistical reviewers miss."""

    payload = json.dumps({
        "model": ANTHROPIC_MODEL_HEAVY,
        "max_tokens": 4096,
        "system": BEATRIX_REVIEW_SYSTEM_PROMPT,
        "messages": [{"role": "user", "content": user_prompt}]
    }).encode()

    req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
        headers={
            "x-api-key": ANTHROPIC_API_KEY,
            "anthropic-version": "2023-06-01",
            "Content-Type": "application/json",
            "User-Agent": "BEATRIXLab/3.7-review"
        })
    resp = json.loads(urllib.request.urlopen(req, context=ctx, timeout=120).read())
    return resp["content"][0]["text"]


def extract_review_score(review_text: str) -> int:
    """Extract total score from BEATRIX review."""
    import re
    patterns = [
        r"Total\s*(?:Score)?:\s*(\d+)/100",
        r"\*\*Total\s*(?:Score)?:\s*(\d+)",
        r"(\d+)/100\s*→\s*(PURSUE|CONSIDER|SKIP)"
    ]
    for pattern in patterns:
        match = re.search(pattern, review_text)
        if match:
            return int(match.group(1))
    return 0


def extract_review_verdict(review_text: str) -> str:
    """Extract verdict from BEATRIX review."""
    import re
    patterns = [
        r"(?:BEATRIX\s+)?Verdict[:\s]*\*?\*?(ACCEPT|CONDITIONAL ACCEPT|MAJOR REVISION|REJECT)",
        r"\*\*(ACCEPT|CONDITIONAL ACCEPT|MAJOR REVISION|REJECT)\s*[–-]",
    ]
    for pattern in patterns:
        match = re.search(pattern, review_text, re.IGNORECASE)
        if match:
            return match.group(1).upper()
    return "UNKNOWN"


def extract_review_tags(review_text: str) -> dict:
    """Extract KB integration tags JSON from review."""
    import re
    try:
        json_match = re.search(r'```json\s*(\{.*?\})\s*```', review_text, re.DOTALL)
        if json_match:
            return json.loads(json_match.group(1))
    except:
        pass
    return {}


@app.post("/api/documents/{doc_id}/review")
async def create_document_review(doc_id: str, user=Depends(require_auth)):
    """Generate a BEATRIX Behavioral Economics Review for a KB document."""
    db = get_db()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            raise HTTPException(404, "Document not found")

        # Extract method from tags
        method = "unknown"
        if doc.tags:
            tag_list = doc.tags if isinstance(doc.tags, list) else doc.tags.split(",")
            for t in tag_list:
                if t.strip() in ["DiD", "RDD", "RCT", "SCM", "synthetic_control", "IV", "regression_discontinuity"]:
                    method = t.strip()
                    break

        # Generate review
        review_text = call_claude_review(
            paper_content=doc.content[:35000] if doc.content else "",
            paper_id=doc_id,
            title=doc.title or "Untitled",
            method=method
        )

        score = extract_review_score(review_text)
        verdict = extract_review_verdict(review_text)
        tags = extract_review_tags(review_text)

        # Store review as new document in KB
        content_hash = hashlib.sha256(review_text.encode("utf-8")).hexdigest()
        review_doc = Document(
            title=f"BEATRIX Review: {(doc.title or 'Untitled')[:80]}",
            content=review_text,
            source_type="text",
            database_target="knowledge_base",
            category="study",
            language="en",
            tags=["BEATRIX-review", "behavioral-economics"] + tags.get("bcm_dimensions", [])[:5],
            status="indexed",
            uploaded_by=user.get("sub"),
            content_hash=content_hash,
            doc_metadata={
                "reviewed_doc_id": doc_id,
                "reviewed_doc_title": doc.title,
                "score": score,
                "verdict": verdict,
                "psi_dimensions": tags.get("psi_dimensions", []),
                "behavioral_mechanisms": tags.get("behavioral_mechanisms", []),
                "review_type": "beatrix_behavioral_economics"
            }
        )
        db.add(review_doc)
        db.commit()
        db.refresh(review_doc)

        # Embed for vector search
        try:
            embed_document(db, review_doc.id)
        except: pass

        # Push to GitHub
        gh_path = f"papers/reviews/{doc_id}/review_beatrix.md"
        gh_result = gh_put_file(GH_REPO, gh_path, review_text,
                                f"review: BEATRIX BE Review – {(doc.title or 'Untitled')[:60]}")

        return {
            "review_id": str(review_doc.id),
            "score": score,
            "verdict": verdict,
            "tags": tags,
            "github": gh_result,
            "review_preview": review_text[:500]
        }
    except HTTPException: raise
    except Exception as e:
        db.rollback()
        logger.error(f"Review generation failed for {doc_id}: {e}")
        raise HTTPException(500, f"Review generation failed: {e}")
    finally:
        db.close()


@app.post("/api/review/external")
async def review_external_paper(request: Request, user=Depends(require_auth)):
    """Review an external paper (e.g., APE GitHub) by URL or paper_id."""
    import urllib.request as ureq, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    body = await request.json()
    paper_url = body.get("paper_url", "")
    paper_id = body.get("paper_id", "")

    if not paper_url and not paper_id:
        raise HTTPException(400, "paper_url or paper_id required")

    # If paper_id looks like an APE paper, construct URL
    if paper_id.startswith("apep_") and not paper_url:
        paper_url = f"https://raw.githubusercontent.com/SocialCatalystLab/ape-papers/main/{paper_id}/v1/paper.tex"

    # Fetch content
    try:
        req = ureq.Request(paper_url, headers={"User-Agent": "BEATRIX"})
        content = ureq.urlopen(req, context=ctx, timeout=30).read().decode(errors="replace")
    except Exception as e:
        raise HTTPException(400, f"Could not fetch paper: {e}")

    # Try to get metadata for APE papers
    title = body.get("title", paper_id or "External Paper")
    method = body.get("method", "unknown")
    authoring_model = "unknown"

    if paper_id.startswith("apep_"):
        try:
            meta_url = f"https://raw.githubusercontent.com/SocialCatalystLab/ape-papers/main/{paper_id}/v1/metadata.json"
            req = ureq.Request(meta_url, headers={"User-Agent": "BEATRIX"})
            meta = json.loads(ureq.urlopen(req, context=ctx, timeout=10).read())
            title = meta.get("title", title)
            method = meta.get("method", method)
            authoring_model = meta.get("authoring_model", "unknown")
        except: pass

    # Generate review
    review_text = call_claude_review(
        paper_content=content[:35000],
        paper_id=paper_id or "external",
        title=title,
        method=method,
        model_name=authoring_model
    )

    score = extract_review_score(review_text)
    verdict = extract_review_verdict(review_text)
    tags = extract_review_tags(review_text)

    # Optionally store in KB
    store = body.get("store", True)
    review_id = None
    if store:
        db = get_db()
        try:
            content_hash = hashlib.sha256(review_text.encode("utf-8")).hexdigest()
            review_doc = Document(
                title=f"BEATRIX Review: {title[:80]}",
                content=review_text,
                source_type="text",
                database_target="knowledge_base",
                category="study",
                language="en",
                tags=["BEATRIX-review", "behavioral-economics", "external-review"] + tags.get("bcm_dimensions", [])[:5],
                status="indexed",
                uploaded_by=user.get("sub"),
                content_hash=content_hash,
                doc_metadata={
                    "reviewed_paper_id": paper_id,
                    "reviewed_paper_url": paper_url,
                    "score": score, "verdict": verdict,
                    "review_type": "beatrix_behavioral_economics"
                }
            )
            db.add(review_doc); db.commit(); db.refresh(review_doc)
            review_id = str(review_doc.id)
            try: embed_document(db, review_doc.id)
            except: pass
        except: db.rollback()
        finally: db.close()

    # Push to GitHub if APE paper
    gh_result = {}
    if paper_id.startswith("apep_"):
        gh_path = f"papers/external/ape/analysis/{paper_id}/review_beatrix.md"
        gh_result = gh_put_file(GH_REPO, gh_path, review_text,
                                f"review: BEATRIX BE Review – {paper_id}")

    return {
        "review_id": review_id,
        "paper_id": paper_id,
        "title": title,
        "score": score,
        "verdict": verdict,
        "tags": tags,
        "github": gh_result,
        "review": review_text
    }


@app.post("/api/review/batch")
async def batch_review_papers(request: Request, user=Depends(require_auth)):
    """Batch review multiple papers. Accepts list of paper_ids (APE) or doc_ids (KB)."""
    import urllib.request as ureq, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    body = await request.json()
    paper_ids = body.get("paper_ids", [])
    source = body.get("source", "ape")  # "ape" or "kb"
    store = body.get("store", True)

    if not paper_ids:
        raise HTTPException(400, "paper_ids required")
    if len(paper_ids) > 20:
        raise HTTPException(400, "Max 20 papers per batch")

    results = []
    for pid in paper_ids:
        try:
            if source == "ape" and pid.startswith("apep_"):
                # Fetch from APE GitHub
                tex_url = f"https://raw.githubusercontent.com/SocialCatalystLab/ape-papers/main/{pid}/v1/paper.tex"
                meta_url = f"https://raw.githubusercontent.com/SocialCatalystLab/ape-papers/main/{pid}/v1/metadata.json"

                req = ureq.Request(tex_url, headers={"User-Agent": "BEATRIX"})
                content = ureq.urlopen(req, context=ctx, timeout=30).read().decode(errors="replace")

                title = pid
                method = "unknown"
                try:
                    req = ureq.Request(meta_url, headers={"User-Agent": "BEATRIX"})
                    meta = json.loads(ureq.urlopen(req, context=ctx, timeout=10).read())
                    title = meta.get("title", pid)
                    method = meta.get("method", "unknown")
                except: pass

                review_text = call_claude_review(content[:30000], pid, title, method)

            elif source == "kb":
                # Fetch from BEATRIX DB
                db = get_db()
                try:
                    doc = db.query(Document).filter(Document.id == pid).first()
                    if not doc:
                        results.append({"paper_id": pid, "error": "not found"})
                        continue
                    review_text = call_claude_review(
                        (doc.content or "")[:30000], pid, doc.title or pid)
                    title = doc.title
                finally:
                    db.close()
            else:
                results.append({"paper_id": pid, "error": f"unknown source: {source}"})
                continue

            score = extract_review_score(review_text)
            verdict = extract_review_verdict(review_text)

            # Store if requested
            review_id = None
            if store:
                db = get_db()
                try:
                    content_hash = hashlib.sha256(review_text.encode("utf-8")).hexdigest()
                    review_doc = Document(
                        title=f"BEATRIX Review: {title[:80]}",
                        content=review_text,
                        source_type="text",
                        database_target="knowledge_base",
                        category="study",
                        language="en",
                        tags=["BEATRIX-review", "behavioral-economics", "batch-review"],
                        status="indexed",
                        uploaded_by=user.get("sub"),
                        content_hash=content_hash,
                        doc_metadata={"reviewed_paper_id": pid, "score": score, "verdict": verdict,
                                      "review_type": "beatrix_behavioral_economics"}
                    )
                    db.add(review_doc); db.commit(); db.refresh(review_doc)
                    review_id = str(review_doc.id)
                    try: embed_document(db, review_doc.id)
                    except: pass
                except: db.rollback()
                finally: db.close()

            results.append({
                "paper_id": pid,
                "title": title,
                "review_id": review_id,
                "score": score,
                "verdict": verdict,
                "review_preview": review_text[:300]
            })

        except Exception as e:
            results.append({"paper_id": pid, "error": str(e)})

    return {
        "reviews": results,
        "total": len(results),
        "successful": len([r for r in results if "score" in r]),
        "failed": len([r for r in results if "error" in r])
    }


# ============================================================================
# LERNFELD 1: Systematisches Ideen-Töten (Paper Triage)
# POST /api/documents/{doc_id}/triage
# ============================================================================

TRIAGE_SYSTEM_PROMPT = """Du bist ein Research-Evaluator für Behavioral Economics.
Bewerte dieses Paper/Dokument auf einer 100-Punkte-Skala.

Bewertungsdimensionen (je 0-25):
1. BCM-Relevanz: Wie relevant für das Behavioral Complementarity Model?
2. Methodische Rigorosität: DiD, RDD, RCT, Quasi-Experiment, Meta-Analyse?
3. Daten-Neuheit: Neue Datenquellen, Populationen, Kontexte?
4. Praktische Anwendbarkeit: Direkt umsetzbar für FehrAdvice-Projekte?

Basierend auf dem Total Score:
- PURSUE (>65): Sofort in KB integrieren, detaillierte Analyse
- CONSIDER (45-65): In Warteschlange, Review bei Bedarf
- SKIP (<45): Archivieren, nicht in aktive KB

WICHTIG: Vergib auch CONDITIONS – Bedingungen, unter denen das Paper relevant wird.
Beispiel: "PURSUE, WENN die Ergebnisse auf europäische Kontexte übertragbar sind"

Antworte NUR mit diesem JSON (kein anderer Text):
```json
{
  "score": 78,
  "recommendation": "PURSUE",
  "dimensions": {
    "bcm_relevance": {"score": 22, "reasoning": "..."},
    "methodological_rigor": {"score": 18, "reasoning": "..."},
    "data_novelty": {"score": 20, "reasoning": "..."},
    "practical_applicability": {"score": 18, "reasoning": "..."}
  },
  "conditions": ["Bedingung 1", "Bedingung 2"],
  "key_findings": ["Finding 1", "Finding 2"],
  "kill_reason": null,
  "one_line_summary": "Kurze Zusammenfassung in einem Satz"
}
```
Wenn SKIP, setze kill_reason auf den Hauptgrund."""


@app.post("/api/documents/{doc_id}/triage")
async def triage_document(doc_id: str, user=Depends(require_auth)):
    """Lernfeld 1: Systematisches Ideen-Töten – PURSUE/CONSIDER/SKIP mit Conditions."""
    import urllib.request as ureq, ssl, re
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    db = get_db()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            raise HTTPException(404, "Document not found")

        if not ANTHROPIC_API_KEY:
            raise HTTPException(501, "Claude API not configured")

        content = (doc.content or "")[:25000]
        user_msg = f"Titel: {doc.title}\n\nInhalt:\n{content}"

        payload = json.dumps({
            "model": ANTHROPIC_MODEL_LIGHT,
            "max_tokens": 1500,
            "system": TRIAGE_SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": user_msg}]
        }).encode()

        req = ureq.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
            headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                     "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-triage"})
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=60).read())
        answer = resp["content"][0]["text"]

        # Parse JSON
        triage_data = None
        json_match = re.search(r'```json\s*(\{.*?\})\s*```', answer, re.DOTALL)
        if json_match:
            try: triage_data = json.loads(json_match.group(1))
            except: pass
        if not triage_data:
            try: triage_data = json.loads(answer.strip())
            except: triage_data = {"score": 0, "recommendation": "ERROR", "error": answer[:500]}

        # Update document metadata
        meta = doc.doc_metadata or {}
        meta["triage"] = triage_data
        meta["triage_at"] = datetime.utcnow().isoformat()
        doc.doc_metadata = meta
        from sqlalchemy.orm.attributes import flag_modified
        flag_modified(doc, "doc_metadata")
        db.commit()

        # Push triage result to GitHub
        triage_md = f"""# Triage: {doc.title}

**Score:** {triage_data.get('score', '?')}/100 → **{triage_data.get('recommendation', '?')}**
**Date:** {datetime.utcnow().strftime('%Y-%m-%d')}

## Dimensions
| Dimension | Score | Reasoning |
|-----------|-------|-----------|
| BCM Relevance | {triage_data.get('dimensions',{}).get('bcm_relevance',{}).get('score','?')}/25 | {triage_data.get('dimensions',{}).get('bcm_relevance',{}).get('reasoning','?')} |
| Methodological Rigor | {triage_data.get('dimensions',{}).get('methodological_rigor',{}).get('score','?')}/25 | {triage_data.get('dimensions',{}).get('methodological_rigor',{}).get('reasoning','?')} |
| Data Novelty | {triage_data.get('dimensions',{}).get('data_novelty',{}).get('score','?')}/25 | {triage_data.get('dimensions',{}).get('data_novelty',{}).get('reasoning','?')} |
| Practical Applicability | {triage_data.get('dimensions',{}).get('practical_applicability',{}).get('score','?')}/25 | {triage_data.get('dimensions',{}).get('practical_applicability',{}).get('reasoning','?')} |

## Conditions
{chr(10).join('- ' + c for c in triage_data.get('conditions', [])) or '- None'}

## Key Findings
{chr(10).join('- ' + f for f in triage_data.get('key_findings', [])) or '- None'}

{f"## Kill Reason{chr(10)}{triage_data.get('kill_reason')}" if triage_data.get('kill_reason') else ''}

**Summary:** {triage_data.get('one_line_summary', '')}
"""
        gh_put_file(GH_REPO, f"papers/triage/{doc_id}/triage.md", triage_md,
                    f"triage: {triage_data.get('recommendation','?')} ({triage_data.get('score','?')}/100) – {(doc.title or '')[:50]}")

        return triage_data

    except HTTPException: raise
    except Exception as e:
        db.rollback()
        raise HTTPException(500, f"Triage failed: {e}")
    finally:
        db.close()


# ============================================================================
# LERNFELD 2: Pre-Commitment Plan Locking
# POST /api/projects/{project_id}/pre-commitment
# POST /api/pre-commitment
# ============================================================================

PRECOMMIT_SYSTEM_PROMPT = """Du bist ein Research-Design-Berater für Behavioral Economics.
Erstelle ein Pre-Commitment-Template für ein Forschungs- oder Beratungsprojekt.

Das Template wird auf GitHub mit Timestamp gepusht und ist danach NICHT mehr änderbar.
Es dient als Pre-Registration: Hypothesen und Erfolgskriterien werden VOR der Analyse festgelegt.

Erstelle das Template im folgenden Markdown-Format:

# Pre-Commitment: [Projekttitel]

**Status: 🔒 LOCKED – Do not modify after commit**
**Locked at:** [Timestamp]
**Author:** [Email]

## 1. Fragestellung
Was genau wollen wir wissen? (Eine klare, beantwortbare Frage)

## 2. Kontext-Hypothese (BCM/Ψ)
Was erwarten wir aufgrund des BCM und der Ψ-Dimensionen?
Welche Ψ-Dimensionen sind beteiligt?

## 3. Methode
Wie evaluieren wir die Intervention / das Projekt?
(Befragung, A/B-Test, Pre-Post, RCT, Beobachtung, etc.)

## 4. Vorab definierte Erfolgskriterien
Wann gilt die Intervention als wirksam?
(Konkrete Schwellenwerte, z.B. "Effektstärke d > 0.3")

## 5. Planned Checks
Welche Alternativerklärungen prüfen wir?
Welche Robustness Checks sind geplant?

## 6. Datenquellen
Woher kommen die Daten?

## 7. Timeline
Wann wird was gemacht?

## 8. Potential Concerns
Bekannte Einschränkungen und Risiken.

Fülle alle Sections basierend auf den Projekt-Informationen des Users aus.
Sei konkret und spezifisch – keine Platzhalter."""


@app.post("/api/pre-commitment")
async def create_pre_commitment(request: Request, user=Depends(require_auth)):
    """Lernfeld 2: Pre-Commitment Plan Locking – unmanipulierbares Analyse-Template."""
    import urllib.request as ureq, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    body = await request.json()
    project_title = body.get("title", "Untitled Project")
    description = body.get("description", "")
    context = body.get("context", "")

    if not ANTHROPIC_API_KEY:
        raise HTTPException(501, "Claude API not configured")

    user_msg = f"""Erstelle ein Pre-Commitment-Template für:

Projekt: {project_title}
Beschreibung: {description}
Kontext: {context}
Author: {user.get('sub', 'unknown')}
Timestamp: {datetime.utcnow().isoformat()}Z"""

    payload = json.dumps({
        "model": ANTHROPIC_MODEL_STANDARD,
        "max_tokens": 2500,
        "system": PRECOMMIT_SYSTEM_PROMPT,
        "messages": [{"role": "user", "content": user_msg}]
    }).encode()

    req = ureq.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
        headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                 "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-precommit"})
    resp = json.loads(ureq.urlopen(req, context=ctx, timeout=60).read())
    template = resp["content"][0]["text"]

    # Ensure LOCKED marker
    if "LOCKED" not in template:
        template = template.replace("# Pre-Commitment:",
            "# Pre-Commitment:").replace("\n",
            f"\n\n**Status: 🔒 LOCKED – Do not modify after commit**\n**Locked at:** {datetime.utcnow().isoformat()}Z\n**Author:** {user.get('sub', 'unknown')}\n",
            1)

    # Generate unique ID
    commit_id = hashlib.sha256(f"{project_title}{datetime.utcnow().isoformat()}".encode()).hexdigest()[:12]
    safe_title = project_title.replace(" ", "_").replace("/", "-")[:50]

    # Push to GitHub (immutable – no SHA override)
    gh_path = f"pre-commitments/{commit_id}_{safe_title}.md"
    gh_result = gh_put_file(GH_REPO, gh_path, template,
                            f"🔒 pre-commitment: {project_title[:60]} (LOCKED)")

    # Store in DB
    db = get_db()
    try:
        content_hash = hashlib.sha256(template.encode()).hexdigest()
        doc = Document(
            title=f"🔒 Pre-Commitment: {project_title}",
            content=template,
            source_type="text",
            database_target="knowledge_base",
            category="document",
            language="de",
            tags=["pre-commitment", "locked", "research-plan"],
            status="indexed+github",
            uploaded_by=user.get("sub"),
            content_hash=content_hash,
            github_url=gh_result.get("url", ""),
            doc_metadata={
                "type": "pre-commitment",
                "locked_at": datetime.utcnow().isoformat(),
                "commit_id": commit_id,
                "project_title": project_title
            }
        )
        db.add(doc); db.commit(); db.refresh(doc)
        try: embed_document(db, doc.id)
        except: pass
        doc_id = str(doc.id)
    except:
        db.rollback()
        doc_id = None
    finally:
        db.close()

    return {
        "commit_id": commit_id,
        "doc_id": doc_id,
        "github": gh_result,
        "locked_at": datetime.utcnow().isoformat(),
        "template_preview": template[:500]
    }


# ============================================================================
# LERNFELD 3: Adversarial Quality Control (Multi-Perspective Review)
# POST /api/documents/{doc_id}/multi-review
# ============================================================================

MULTI_REVIEW_PERSPECTIVES = {
    "analytical": {
        "name": "Analytischer Reviewer",
        "system": """Du bist ein strenger methodischer Reviewer. Du prüfst NUR:
- Sind empirische Claims durch Daten/Quellen gedeckt?
- Ist die Methodik angemessen?
- Stimmen die Schlussfolgerungen mit den Ergebnissen überein?
- Werden Limitationen ehrlich diskutiert?
Verdikt: ACCEPT / MINOR REVISION / MAJOR REVISION / REJECT
Antworte in 300-500 Wörtern."""
    },
    "practical": {
        "name": "Praxis-Reviewer",
        "system": """Du bist ein praxisorientierter Reviewer aus der Beratung. Du prüfst NUR:
- Kann ein Kunde/Unternehmen das umsetzen?
- Sind die Empfehlungen konkret genug?
- Fehlen wichtige Praxisaspekte (Budget, Timeline, Stakeholder)?
- Ist der Kosten-Nutzen klar?
Verdikt: ACCEPT / MINOR REVISION / MAJOR REVISION / REJECT
Antworte in 300-500 Wörtern."""
    },
    "devils_advocate": {
        "name": "Devil's Advocate",
        "system": """Du bist ein bewusst skeptischer Reviewer (Devil's Advocate). Du prüfst NUR:
- Was könnte schiefgehen?
- Welche Gegenargumente gibt es?
- Welche alternativen Erklärungen wurden ignoriert?
- Was sind die versteckten Annahmen?
Verdikt: ACCEPT / MINOR REVISION / MAJOR REVISION / REJECT
Antworte in 300-500 Wörtern."""
    }
}


@app.post("/api/documents/{doc_id}/multi-review")
async def multi_review_document(doc_id: str, user=Depends(require_auth)):
    """Lernfeld 3: Adversarial QC – 3 unabhängige Perspektiven + Disagreement-Analyse."""
    import urllib.request as ureq, ssl, re
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    db = get_db()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            raise HTTPException(404, "Document not found")

        if not ANTHROPIC_API_KEY:
            raise HTTPException(501, "Claude API not configured")

        content = (doc.content or "")[:20000]
        reviews = {}

        # Generate 3 independent reviews
        for perspective_key, perspective in MULTI_REVIEW_PERSPECTIVES.items():
            payload = json.dumps({
                "model": ANTHROPIC_MODEL_HEAVY,
                "max_tokens": 1200,
                "system": perspective["system"],
                "messages": [{"role": "user", "content": f"Review dieses Dokument:\n\nTitel: {doc.title}\n\n{content}"}]
            }).encode()

            req = ureq.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
                headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                         "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-multireview"})
            resp = json.loads(ureq.urlopen(req, context=ctx, timeout=60).read())
            review_text = resp["content"][0]["text"]

            # Extract verdict
            verdict = "UNKNOWN"
            for v in ["REJECT", "MAJOR REVISION", "MINOR REVISION", "ACCEPT"]:
                if v in review_text.upper():
                    verdict = v
                    break

            reviews[perspective_key] = {
                "name": perspective["name"],
                "verdict": verdict,
                "review": review_text
            }

        # Disagreement analysis
        verdicts = [r["verdict"] for r in reviews.values()]
        all_agree = len(set(verdicts)) == 1
        consensus_verdict = max(set(verdicts), key=verdicts.count) if verdicts else "UNKNOWN"

        # Generate synthesis
        synthesis_prompt = f"""Analysiere diese 3 unabhängigen Reviews des Dokuments "{doc.title}":

ANALYTISCHER REVIEWER ({reviews['analytical']['verdict']}):
{reviews['analytical']['review'][:800]}

PRAXIS-REVIEWER ({reviews['practical']['verdict']}):
{reviews['practical']['review'][:800]}

DEVIL'S ADVOCATE ({reviews['devils_advocate']['verdict']}):
{reviews['devils_advocate']['review'][:800]}

Erstelle eine Synthese:
1. Wo stimmen alle 3 überein? (HIGH PRIORITY Issues)
2. Wo widersprechen sich 2? (MEDIUM PRIORITY)
3. Was sagt nur einer? (LOW PRIORITY)
4. Gesamtverdikt basierend auf Konsens
Antworte in 200-400 Wörtern."""

        payload = json.dumps({
            "model": ANTHROPIC_MODEL_HEAVY,
            "max_tokens": 1000,
            "messages": [{"role": "user", "content": synthesis_prompt}]
        }).encode()
        req = ureq.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
            headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                     "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-synthesis"})
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=60).read())
        synthesis = resp["content"][0]["text"]

        # Build full report
        report = f"""# Multi-Perspective Review: {doc.title}

**Date:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC
**Document ID:** {doc_id}
**Consensus:** {'✅ All agree' if all_agree else '⚠️ Disagreement detected'} → **{consensus_verdict}**

## Verdicts
| Perspective | Verdict |
|-------------|---------|
| 🔬 {reviews['analytical']['name']} | {reviews['analytical']['verdict']} |
| 🏢 {reviews['practical']['name']} | {reviews['practical']['verdict']} |
| 😈 {reviews['devils_advocate']['name']} | {reviews['devils_advocate']['verdict']} |

## Analytical Review
{reviews['analytical']['review']}

## Practical Review
{reviews['practical']['review']}

## Devil's Advocate
{reviews['devils_advocate']['review']}

## Synthesis & Priority
{synthesis}
"""

        # Push to GitHub
        gh_result = gh_put_file(GH_REPO, f"papers/multi-review/{doc_id}/multi_review.md", report,
                    f"multi-review: {consensus_verdict} – {(doc.title or '')[:50]}")

        # Store in DB
        content_hash = hashlib.sha256(report.encode()).hexdigest()
        review_doc = Document(
            title=f"Multi-Review: {(doc.title or 'Untitled')[:80]}",
            content=report, source_type="text", database_target="knowledge_base",
            category="study", language="en",
            tags=["multi-review", "adversarial-qc", "analytical", "practical", "devils-advocate"],
            status="indexed", uploaded_by=user.get("sub"), content_hash=content_hash,
            doc_metadata={"reviewed_doc_id": doc_id, "consensus_verdict": consensus_verdict,
                          "all_agree": all_agree, "verdicts": {k: v["verdict"] for k, v in reviews.items()},
                          "review_type": "multi_perspective"}
        )
        db.add(review_doc); db.commit(); db.refresh(review_doc)
        try: embed_document(db, review_doc.id)
        except: pass

        return {
            "review_id": str(review_doc.id),
            "consensus": consensus_verdict,
            "all_agree": all_agree,
            "verdicts": {k: v["verdict"] for k, v in reviews.items()},
            "synthesis_preview": synthesis[:500],
            "github": gh_result
        }

    except HTTPException: raise
    except Exception as e:
        db.rollback()
        raise HTTPException(500, f"Multi-review failed: {e}")
    finally:
        db.close()


# ============================================================================
# LERNFELD 4: Integrity Scanning
# POST /api/documents/{doc_id}/integrity-check
# ============================================================================

INTEGRITY_SYSTEM_PROMPT = """Du bist ein Integrity Scanner für akademische und professionelle Dokumente.
Prüfe das Dokument auf folgende 6 Kategorien:

1. CLAIM_DATA_CONSISTENCY: Stimmen quantitative Claims im Text überein? (Z.B. Abstract sagt 23%, Discussion sagt 25%)
2. SOURCE_PROVENANCE: Sind Quellen für empirische Claims angegeben und nachvollziehbar?
3. EFFECT_SIZE_CONTEXT: Werden Effektgrössen im richtigen Kontext interpretiert? (Relative vs. absolute Werte)
4. METHODOLOGY_CLARITY: Ist die Methodik klar genug beschrieben, um repliziert zu werden?
5. SAMPLE_DESCRIPTION: Sind Stichprobe/Population ausreichend beschrieben?
6. LIMITATION_HONESTY: Werden Einschränkungen ehrlich diskutiert?

Für jede Kategorie:
- Flag-Level: CLEAN (kein Problem), LOW (minor), MEDIUM (should fix), HIGH (critical)
- Evidence: Konkrete Textstelle die das Problem zeigt
- Confidence: 0.0-1.0

Antworte NUR mit JSON:
```json
{
  "overall_verdict": "CLEAN oder SUSPICIOUS",
  "flags": [
    {"category": "CLAIM_DATA_CONSISTENCY", "severity": "MEDIUM", "evidence": "Abstract says 23% but Section 4 says 25%", "confidence": 0.85}
  ],
  "summary": "Kurze Zusammenfassung der Integrity-Analyse"
}
```"""


@app.post("/api/documents/{doc_id}/integrity-check")
async def integrity_check_document(doc_id: str, user=Depends(require_auth)):
    """Lernfeld 4: Integrity Scanning – automatische Konsistenz- und Quellenprüfung."""
    import urllib.request as ureq, ssl, re
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    db = get_db()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            raise HTTPException(404, "Document not found")

        if not ANTHROPIC_API_KEY:
            raise HTTPException(501, "Claude API not configured")

        content = (doc.content or "")[:30000]

        payload = json.dumps({
            "model": ANTHROPIC_MODEL_HEAVY,
            "max_tokens": 2000,
            "system": INTEGRITY_SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": f"Prüfe dieses Dokument:\n\nTitel: {doc.title}\n\n{content}"}]
        }).encode()

        req = ureq.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
            headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                     "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-integrity"})
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=60).read())
        answer = resp["content"][0]["text"]

        # Parse JSON
        integrity_data = None
        json_match = re.search(r'```json\s*(\{.*?\})\s*```', answer, re.DOTALL)
        if json_match:
            try: integrity_data = json.loads(json_match.group(1))
            except: pass
        if not integrity_data:
            try: integrity_data = json.loads(answer.strip())
            except: integrity_data = {"overall_verdict": "ERROR", "flags": [], "summary": answer[:500]}

        # Update document metadata
        meta = doc.doc_metadata or {}
        meta["integrity"] = integrity_data
        meta["integrity_at"] = datetime.utcnow().isoformat()
        doc.doc_metadata = meta
        from sqlalchemy.orm.attributes import flag_modified
        flag_modified(doc, "doc_metadata")
        db.commit()

        # Push to GitHub
        flags = integrity_data.get("flags", [])
        high_count = len([f for f in flags if f.get("severity") == "HIGH"])
        med_count = len([f for f in flags if f.get("severity") == "MEDIUM"])

        integrity_md = f"""# Integrity Check: {doc.title}

**Verdict:** {integrity_data.get('overall_verdict', '?')}
**Date:** {datetime.utcnow().strftime('%Y-%m-%d')}
**Flags:** {high_count} HIGH, {med_count} MEDIUM, {len(flags) - high_count - med_count} LOW

## Summary
{integrity_data.get('summary', '')}

## Flags
{''.join(f'''
### [{f.get("severity")}] {f.get("category")}
**Evidence:** {f.get("evidence", "?")}
**Confidence:** {f.get("confidence", "?")}
''' for f in flags) if flags else "No flags found."}
"""
        gh_put_file(GH_REPO, f"papers/integrity/{doc_id}/integrity_check.md", integrity_md,
                    f"integrity: {integrity_data.get('overall_verdict','?')} ({high_count}H/{med_count}M) – {(doc.title or '')[:50]}")

        return integrity_data

    except HTTPException: raise
    except Exception as e:
        db.rollback()
        raise HTTPException(500, f"Integrity check failed: {e}")
    finally:
        db.close()


# ============================================================================
# LERNFELD 5: Reply-to-Reviewers (Feedback Response + Knowledge Distillation)
# POST /api/documents/{doc_id}/reply-to-feedback
# ============================================================================

REPLY_SYSTEM_PROMPT = """Du bist ein wissenschaftlicher Autor, der auf Reviewer-Feedback antwortet.
Du erhältst ein Dokument und Feedback dazu.

Erstelle eine strukturierte Reply-to-Reviewers im folgenden Format:

# Reply to Feedback

## Summary of Changes
Kurze Übersicht der vorgenommenen Änderungen.

## Point-by-Point Response

### Feedback-Punkt 1: [Zusammenfassung]
**Feedback:** "[Zitat]"
**Response:** [Deine Antwort]
**Action:** [IMPLEMENTED / ACKNOWLEDGED / DECLINED with reason]
**Priority:** [HIGH / MEDIUM / LOW]

[Wiederhole für jeden Punkt]

## Changes NOT Implemented (with justification)
Punkte die bewusst nicht geändert wurden, mit Begründung.

## Knowledge Distilled
Was haben wir aus diesem Feedback gelernt? (2-3 Sätze)

Regeln:
- Sei ehrlich: Wenn der Reviewer Recht hat, sag es
- Sei konkret: Verweise auf exakte Stellen
- Sei konstruktiv: Auch bei Ablehnung, erkläre warum
- Destilliere Wissen: Jede Kritik ist eine Lerngelegenheit"""


@app.post("/api/documents/{doc_id}/reply-to-feedback")
async def reply_to_feedback(doc_id: str, request: Request, user=Depends(require_auth)):
    """Lernfeld 5: Reply-to-Reviewers – strukturierte Wissens-Destillation aus Feedback."""
    import urllib.request as ureq, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    body = await request.json()
    feedback = body.get("feedback", "")
    if not feedback:
        raise HTTPException(400, "feedback required")

    db = get_db()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            raise HTTPException(404, "Document not found")

        if not ANTHROPIC_API_KEY:
            raise HTTPException(501, "Claude API not configured")

        content = (doc.content or "")[:20000]

        user_msg = f"""Dokument: {doc.title}

Inhalt (Auszug):
{content}

---

Feedback das beantwortet werden muss:
{feedback}"""

        payload = json.dumps({
            "model": ANTHROPIC_MODEL_HEAVY,
            "max_tokens": 2500,
            "system": REPLY_SYSTEM_PROMPT,
            "messages": [{"role": "user", "content": user_msg}]
        }).encode()

        req = ureq.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
            headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                     "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-reply"})
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=60).read())
        reply = resp["content"][0]["text"]

        # Push to GitHub
        reply_count = 1
        # Check if there are existing replies
        try:
            import urllib.request
            check_url = f"https://api.github.com/repos/{GH_REPO}/contents/papers/replies/{doc_id}"
            check_req = urllib.request.Request(check_url,
                headers={"Authorization": f"token {GH_TOKEN}", "User-Agent": "BEATRIX"})
            existing = json.loads(urllib.request.urlopen(check_req, context=ctx, timeout=10).read())
            reply_count = len([f for f in existing if f["name"].startswith("reply_")]) + 1
        except: pass

        gh_result = gh_put_file(GH_REPO, f"papers/replies/{doc_id}/reply_{reply_count}.md", reply,
                    f"reply: Response to feedback #{reply_count} – {(doc.title or '')[:50]}")

        # Store as KB document (FAQ enrichment)
        content_hash = hashlib.sha256(reply.encode()).hexdigest()
        reply_doc = Document(
            title=f"Reply to Feedback #{reply_count}: {(doc.title or 'Untitled')[:70]}",
            content=reply, source_type="text", database_target="knowledge_base",
            category="document", language="de",
            tags=["reply-to-feedback", "knowledge-distillation", "quality-control"],
            status="indexed", uploaded_by=user.get("sub"), content_hash=content_hash,
            doc_metadata={"original_doc_id": doc_id, "reply_number": reply_count,
                          "feedback_preview": feedback[:200], "review_type": "reply_to_feedback"}
        )
        db.add(reply_doc); db.commit(); db.refresh(reply_doc)
        try: embed_document(db, reply_doc.id)
        except: pass

        return {
            "reply_id": str(reply_doc.id),
            "reply_number": reply_count,
            "github": gh_result,
            "reply_preview": reply[:500],
            "reply": reply
        }

    except HTTPException: raise
    except Exception as e:
        db.rollback()
        raise HTTPException(500, f"Reply generation failed: {e}")
    finally:
        db.close()


# ============================================================================
# EXTERNAL REVIEWERS – Multi-Model Review System (wie APE)
# POST /api/documents/{doc_id}/ape-review
# ============================================================================

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY", "")
XAI_API_KEY = os.getenv("XAI_API_KEY", "")

# Persistent key storage in DB (survives Railway restarts)
def _init_reviewer_keys_table():
    try:
        db = get_db()
        db.execute(text("""CREATE TABLE IF NOT EXISTS reviewer_keys (
            provider VARCHAR(50) PRIMARY KEY,
            api_key TEXT NOT NULL,
            updated_at TIMESTAMP DEFAULT NOW()
        )"""))
        db.commit()
        db.close()
    except: pass

def _load_keys_from_db():
    global OPENAI_API_KEY, GOOGLE_API_KEY, XAI_API_KEY
    try:
        db = get_db()
        rows = db.execute(text("SELECT provider, api_key FROM reviewer_keys")).fetchall()
        for row in rows:
            p, k = row[0], row[1]
            if p == "openai" and not OPENAI_API_KEY: OPENAI_API_KEY = k
            elif p == "google" and not GOOGLE_API_KEY: GOOGLE_API_KEY = k
            elif p == "xai" and not XAI_API_KEY: XAI_API_KEY = k
        db.close()
        loaded = [p for p, _ in rows]
        if loaded: print(f"[BOOT] Reviewer keys loaded from DB: {loaded}")
    except Exception as e:
        print(f"[BOOT] Key load from DB failed (first run?): {e}")

def _save_key_to_db(provider: str, api_key: str):
    try:
        db = get_db()
        db.execute(text("""INSERT INTO reviewer_keys (provider, api_key, updated_at)
            VALUES (:p, :k, NOW()) ON CONFLICT (provider) DO UPDATE
            SET api_key = :k, updated_at = NOW()"""), {"p": provider, "k": api_key})
        db.commit()
        db.close()
    except: pass

# Init on startup
try:
    _init_reviewer_keys_table()
    _load_keys_from_db()
except: pass

REVIEWER_CONFIGS = {
    "beatrix": {
        "name": "BEATRIX (Claude)",
        "icon": "🧠",
        "provider": "anthropic",
        "model": ANTHROPIC_MODEL_HEAVY,
        "focus": "Behavioral Economics, BCM, Ψ-Dimensions",
        "system": """You are BEATRIX, a Behavioral Economics reviewer. Focus EXCLUSIVELY on:
1. Missing behavioral mechanisms (present bias, loss aversion, social norms, default effects, etc.)
2. Ψ-Dimensions analysis (which of the 8 BCM psychological dimensions are relevant?)
3. Practical intervention design improvements
4. Policy implications from a behavioral science perspective
Score 0-100. Verdict: ACCEPT / CONDITIONAL ACCEPT / MAJOR REVISION / REJECT.
Be specific, cite mechanisms, suggest literature."""
    },
    "methodologist": {
        "name": "Methodologist (GPT)",
        "icon": "🔬",
        "provider": "openai",
        "model": "gpt-4o",
        "focus": "Statistical Methods, Identification Strategy, Causal Inference",
        "system": """You are a strict methodological reviewer for empirical research. Focus EXCLUSIVELY on:
1. Identification strategy: Is the causal claim credible? What are the threats?
2. Statistical methods: Are they appropriate? Standard errors correct? Power adequate?
3. Data quality: Sample size, measurement, selection bias, attrition
4. Robustness: What checks are missing? What would falsify the findings?
Score 0-100. Verdict: ACCEPT / MINOR REVISION / MAJOR REVISION / REJECT.
Be technically precise. Reference specific statistical tests and assumptions."""
    },
    "skeptic": {
        "name": "Skeptic (Grok)",
        "icon": "🦊",
        "provider": "xai",
        "model": "grok-3-mini",
        "focus": "Alternative Explanations, Hidden Assumptions, Replication Concerns",
        "system": """You are a deeply skeptical reviewer. Your job is to find what EVERYONE ELSE missed. Focus on:
1. Alternative explanations the authors didn't consider
2. Hidden assumptions that could invalidate the findings
3. Selection effects, confounders, reverse causality
4. Would this replicate? What's the file-drawer problem here?
5. Is the framing honest or does it oversell the findings?
Score 0-100. Verdict: ACCEPT / MINOR REVISION / MAJOR REVISION / REJECT.
Be direct and unflinching. No diplomatic hedging."""
    },
    "practitioner": {
        "name": "Practitioner (Gemini)",
        "icon": "💼",
        "provider": "google",
        "model": "gemini-2.0-flash",
        "focus": "Real-World Applicability, Implementation, Stakeholder Perspective",
        "system": """You are a practitioner reviewer from management consulting. Focus EXCLUSIVELY on:
1. Can these findings be implemented in practice? What's the cost-benefit?
2. What stakeholder perspectives are missing?
3. Is the effect size practically meaningful (not just statistically significant)?
4. What would a CEO/policy-maker need to know before acting on this?
5. External validity: Does this generalize beyond the specific context studied?
Score 0-100. Verdict: ACCEPT / MINOR REVISION / MAJOR REVISION / REJECT.
Be concrete about implementation barriers and opportunities."""
    },
    "contrarian": {
        "name": "Contrarian (Haiku)",
        "icon": "⚡",
        "provider": "anthropic",
        "model": ANTHROPIC_MODEL_LIGHT,
        "focus": "Quick Red-Team, Overlooked Angles, Fresh Perspective",
        "system": """You are a fast, sharp contrarian reviewer. Your job is to be the voice that says what nobody else dares. Focus on:
1. The ONE thing everyone else will miss
2. Flip the framing: What if the opposite conclusion is true?
3. Who benefits from these findings being believed? Who loses?
4. What's the simplest alternative explanation?
5. If this is wrong, what's the cost of acting on it?
Score 0-100. Verdict: ACCEPT / MINOR REVISION / MAJOR REVISION / REJECT.
Be brief, sharp, provocative. Max 400 words. No filler."""
    },
}


def call_external_reviewer(provider: str, model: str, system_prompt: str, content: str, title: str) -> dict:
    """Call an external AI model for review."""
    import urllib.request as ureq, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

    user_msg = f"Review this document:\n\nTitle: {title}\n\n{content[:20000]}"

    if provider == "anthropic":
        if not ANTHROPIC_API_KEY:
            return {"error": "Anthropic API key not configured", "available": False}
        payload = json.dumps({
            "model": model, "max_tokens": 1500, "system": system_prompt,
            "messages": [{"role": "user", "content": user_msg}]
        }).encode()
        req = ureq.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
            headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                     "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-ape"})
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=90).read())
        return {"text": resp["content"][0]["text"], "model": model, "available": True}

    elif provider == "openai":
        if not OPENAI_API_KEY:
            return {"error": "OpenAI API key not configured", "available": False}
        payload = json.dumps({
            "model": model, "max_tokens": 1500,
            "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_msg}]
        }).encode()
        req = ureq.Request("https://api.openai.com/v1/chat/completions", data=payload, method="POST",
            headers={"Authorization": f"Bearer {OPENAI_API_KEY}",
                     "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-ape"})
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=90).read())
        return {"text": resp["choices"][0]["message"]["content"], "model": model, "available": True}

    elif provider == "xai":
        if not XAI_API_KEY:
            return {"error": "xAI API key not configured", "available": False}
        payload = json.dumps({
            "model": model, "max_tokens": 1500, "stream": False,
            "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_msg}]
        }).encode()
        req = ureq.Request("https://api.x.ai/v1/chat/completions", data=payload, method="POST",
            headers={"Authorization": f"Bearer {XAI_API_KEY}",
                     "Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-ape"})
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=90).read())
        return {"text": resp["choices"][0]["message"]["content"], "model": model, "available": True}

    elif provider == "google":
        if not GOOGLE_API_KEY:
            return {"error": "Google API key not configured", "available": False}
        payload = json.dumps({
            "contents": [{"role": "user", "parts": [{"text": f"{system_prompt}\n\n{user_msg}"}]}],
            "generationConfig": {"maxOutputTokens": 1500}
        }).encode()
        req = ureq.Request(
            f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GOOGLE_API_KEY}",
            data=payload, method="POST",
            headers={"Content-Type": "application/json", "User-Agent": "BEATRIXLab/3.7-ape"})
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=90).read())
        text = resp.get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "")
        return {"text": text, "model": model, "available": True}

    return {"error": f"Unknown provider: {provider}", "available": False}


def extract_verdict_from_review(text: str) -> str:
    """Extract verdict from review text."""
    upper = text.upper()
    for v in ["REJECT", "MAJOR REVISION", "CONDITIONAL ACCEPT", "MINOR REVISION", "ACCEPT"]:
        if v in upper:
            return v
    return "UNKNOWN"


def extract_score_from_review(text: str) -> int:
    """Extract score from review text."""
    import re
    patterns = [
        r'(?:score|total|overall)[:\s]*(\d{1,3})\s*/?\s*100',
        r'(\d{1,3})\s*/\s*100',
        r'(?:score|total|overall)[:\s]*(\d{1,3})',
    ]
    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m:
            s = int(m.group(1))
            if 0 <= s <= 100:
                return s
    return 0


@app.get("/api/reviewers")
async def list_reviewers(user=Depends(require_auth)):
    """List available reviewers and their status."""
    result = []
    for key, config in REVIEWER_CONFIGS.items():
        available = False
        if config["provider"] == "anthropic" and ANTHROPIC_API_KEY: available = True
        elif config["provider"] == "openai" and OPENAI_API_KEY: available = True
        elif config["provider"] == "xai" and XAI_API_KEY: available = True
        elif config["provider"] == "google" and GOOGLE_API_KEY: available = True

        result.append({
            "key": key,
            "name": config["name"],
            "icon": config["icon"],
            "provider": config["provider"],
            "model": config["model"],
            "focus": config["focus"],
            "available": available
        })
    return result


@app.post("/api/documents/{doc_id}/ape-review")
async def ape_review_document(doc_id: str, request: Request, user=Depends(require_auth)):
    """APE-style multi-model review. Calls all available reviewers + generates synthesis."""
    body = await request.json() if request.headers.get("content-length", "0") != "0" else {}
    selected_reviewers = body.get("reviewers", None)  # None = all available

    db = get_db()
    try:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if not doc:
            raise HTTPException(404, "Document not found")

        content = (doc.content or "")[:20000]
        title = doc.title or "Untitled"
        reviews = {}
        errors = {}

        # Determine which reviewers to call
        to_call = selected_reviewers or list(REVIEWER_CONFIGS.keys())

        for reviewer_key in to_call:
            if reviewer_key not in REVIEWER_CONFIGS:
                continue
            config = REVIEWER_CONFIGS[reviewer_key]
            try:
                result = call_external_reviewer(
                    config["provider"], config["model"], config["system"], content, title
                )
                if result.get("available"):
                    review_text = result["text"]
                    reviews[reviewer_key] = {
                        "name": config["name"],
                        "icon": config["icon"],
                        "provider": config["provider"],
                        "model": result.get("model", config["model"]),
                        "focus": config["focus"],
                        "review": review_text,
                        "verdict": extract_verdict_from_review(review_text),
                        "score": extract_score_from_review(review_text)
                    }
                else:
                    errors[reviewer_key] = result.get("error", "Not available")
            except Exception as e:
                errors[reviewer_key] = str(e)[:200]

        if not reviews:
            raise HTTPException(500, f"No reviewers available. Errors: {errors}")

        # Disagreement analysis
        verdicts = {k: v["verdict"] for k, v in reviews.items()}
        scores = {k: v["score"] for k, v in reviews.items()}
        all_verdicts = list(verdicts.values())
        unique_verdicts = set(all_verdicts)
        consensus = max(set(all_verdicts), key=all_verdicts.count) if all_verdicts else "UNKNOWN"
        avg_score = sum(scores.values()) / len(scores) if scores else 0

        # Generate synthesis (using Claude as meta-reviewer)
        synthesis = ""
        if len(reviews) >= 2 and ANTHROPIC_API_KEY:
            import urllib.request as ureq, ssl
            ctx2 = ssl.create_default_context(); ctx2.check_hostname = False; ctx2.verify_mode = ssl.CERT_NONE
            review_summaries = "\n\n".join([
                f"### {v['icon']} {v['name']} ({v['verdict']}, {v['score']}/100):\n{v['review'][:600]}"
                for k, v in reviews.items()
            ])
            synth_prompt = f"""Analysiere diese {len(reviews)} unabhängigen Reviews von "{title}":

{review_summaries}

Erstelle eine Synthese (300 Wörter max):
1. HIGH PRIORITY: Wo stimmen alle/die meisten überein?
2. MEDIUM PRIORITY: Wo gibt es Disagreement?
3. LOW PRIORITY: Einzelne Reviewer-Meinungen
4. Gesamtempfehlung basierend auf Konsens
5. Top 3 Verbesserungsvorschläge"""

            payload = json.dumps({
                "model": ANTHROPIC_MODEL_HEAVY, "max_tokens": 1000,
                "messages": [{"role": "user", "content": synth_prompt}]
            }).encode()
            try:
                req = ureq.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
                    headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                             "Content-Type": "application/json"})
                resp = json.loads(ureq.urlopen(req, context=ctx2, timeout=60).read())
                synthesis = resp["content"][0]["text"]
            except: pass

        # Build full report
        report_lines = [f"# APE-Style Multi-Model Review: {title}\n"]
        report_lines.append(f"**Date:** {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC")
        report_lines.append(f"**Document:** {doc_id}")
        report_lines.append(f"**Reviewers:** {len(reviews)} active, {len(errors)} unavailable")
        report_lines.append(f"**Consensus:** {consensus} (avg score: {avg_score:.0f}/100)")
        report_lines.append(f"**Agreement:** {'✅ Unanimous' if len(unique_verdicts) == 1 else '⚠️ Disagreement'}\n")

        report_lines.append("## Verdict Summary\n| Reviewer | Verdict | Score |")
        report_lines.append("|----------|---------|-------|")
        for k, v in reviews.items():
            report_lines.append(f"| {v['icon']} {v['name']} | {v['verdict']} | {v['score']}/100 |")
        if errors:
            report_lines.append(f"\n*Unavailable:* {', '.join(f'{k} ({v})' for k, v in errors.items())}")

        for k, v in reviews.items():
            report_lines.append(f"\n## {v['icon']} {v['name']}\n**Focus:** {v['focus']}")
            report_lines.append(f"**Model:** {v['model']}\n**Verdict:** {v['verdict']} ({v['score']}/100)\n")
            report_lines.append(v['review'])

        if synthesis:
            report_lines.append(f"\n## 🔄 Synthesis & Priority Analysis\n{synthesis}")

        report = "\n".join(report_lines)

        # Push to GitHub
        gh_put_file(GH_REPO, f"papers/ape-reviews/{doc_id}/ape_review.md", report,
                    f"ape-review: {consensus} ({avg_score:.0f}/100) [{len(reviews)} reviewers] – {title[:50]}")

        # Store in DB
        content_hash = hashlib.sha256(report.encode()).hexdigest()
        review_doc = Document(
            title=f"APE Review: {title[:70]}",
            content=report, source_type="text", database_target="knowledge_base",
            category="study", language="en",
            tags=["ape-review", "multi-model"] + [k for k in reviews.keys()],
            status="indexed", uploaded_by=user.get("sub"), content_hash=content_hash,
            doc_metadata={
                "reviewed_doc_id": doc_id, "review_type": "ape_multi_model",
                "consensus": consensus, "avg_score": round(avg_score, 1),
                "reviewers": {k: {"verdict": v["verdict"], "score": v["score"]} for k, v in reviews.items()},
                "errors": errors, "unanimous": len(unique_verdicts) == 1
            }
        )
        db.add(review_doc); db.commit(); db.refresh(review_doc)
        try: embed_document(db, review_doc.id)
        except: pass

        return {
            "review_id": str(review_doc.id),
            "consensus": consensus,
            "avg_score": round(avg_score, 1),
            "unanimous": len(unique_verdicts) == 1,
            "reviewers": {k: {"name": v["name"], "icon": v["icon"], "verdict": v["verdict"],
                              "score": v["score"], "preview": v["review"][:300]} for k, v in reviews.items()},
            "errors": errors,
            "synthesis_preview": synthesis[:500] if synthesis else None,
            "github": f"papers/ape-reviews/{doc_id}/ape_review.md"
        }

    except HTTPException: raise
    except Exception as e:
        db.rollback()
        raise HTTPException(500, f"APE review failed: {e}")
    finally:
        db.close()


@app.post("/api/admin/reviewer-keys")
async def set_reviewer_keys(request: Request, user=Depends(require_auth)):
    """Admin: Set API keys for external reviewers (runtime only, not persisted)."""
    if not user.get("admin"):
        raise HTTPException(403, "Admin required")
    body = await request.json()
    global OPENAI_API_KEY, GOOGLE_API_KEY, XAI_API_KEY
    updated = []
    if "openai" in body and body["openai"]:
        OPENAI_API_KEY = body["openai"]
        _save_key_to_db("openai", body["openai"])
        updated.append("openai")
    if "google" in body and body["google"]:
        GOOGLE_API_KEY = body["google"]
        _save_key_to_db("google", body["google"])
        updated.append("google")
    if "xai" in body and body["xai"]:
        XAI_API_KEY = body["xai"]
        _save_key_to_db("xai", body["xai"])
        updated.append("xai")
    return {"updated": updated, "message": f"Keys set and persisted for: {', '.join(updated)}. Will survive restarts."}


# ============================================================================
# REAL MULTI-MODEL CHALLENGE – Echte 5 APIs, kein Fake
# POST /api/challenge
# ============================================================================

CHALLENGE_PROMPTS = {
    "beatrix": """Du bist BEATRIX, Behavioral Economics Reviewer. Challenge diese Aussage/Antwort:
- Welche Verhaltensökonomischen Mechanismen fehlen?
- Welche Biases werden ignoriert?
- Welche Ψ-Dimensionen sind relevant aber nicht erwähnt?
Max 150 Wörter. Sei direkt und spezifisch.""",

    "methodologist": """Du bist ein strenger Methodologe. Challenge diese Aussage/Antwort:
- Welche Evidenz fehlt?
- Welche Annahmen sind fragwürdig?
- Was müsste man empirisch prüfen bevor man das behaupten kann?
Max 150 Wörter. Sei technisch präzise.""",

    "skeptic": """Du bist ein radikaler Skeptiker. Challenge diese Aussage/Antwort:
- Was ist die einfachste alternative Erklärung?
- Was wenn das Gegenteil stimmt?
- Wer profitiert davon wenn man das glaubt?
Max 150 Wörter. Sei unbequem und direkt.""",

    "practitioner": """Du bist Management-Berater in der Praxis. Challenge diese Aussage/Antwort:
- Ist das umsetzbar? Was kostet es?
- Was würde ein CEO dazu sagen?
- Welche Stakeholder-Perspektive fehlt?
Max 150 Wörter. Sei konkret.""",

    "contrarian": """Du bist ein schneller Contrarian. Challenge diese Aussage/Antwort:
- Das EINE was alle anderen übersehen
- Dreh die Aussage um – was passiert dann?
- Wäre Nichtstun die bessere Option?
Max 100 Wörter. Kurz, scharf, provokant."""
}


@app.post("/api/challenge")
async def challenge_text(request: Request, user=Depends(require_auth)):
    """Real multi-model challenge: sends text to all 5 APIs independently."""
    body = await request.json()
    text = body.get("text", "")
    if not text or len(text) < 20:
        raise HTTPException(400, "Text too short to challenge")

    text = text[:4000]  # Limit
    challenges = {}
    errors = {}

    for key, prompt in CHALLENGE_PROMPTS.items():
        config = REVIEWER_CONFIGS.get(key)
        if not config:
            continue
        try:
            result = call_external_reviewer(
                config["provider"], config["model"],
                prompt,
                text, "Challenge"
            )
            if result.get("available"):
                challenges[key] = {
                    "name": config["name"],
                    "icon": config["icon"],
                    "model": result.get("model", config["model"]),
                    "challenge": result["text"]
                }
            else:
                errors[key] = result.get("error", "Not available")
        except Exception as e:
            errors[key] = str(e)[:100]

    if not challenges:
        raise HTTPException(500, f"No reviewers available. Errors: {errors}")

    # Synthesis (Claude meta-analysis)
    synthesis = ""
    if len(challenges) >= 2 and ANTHROPIC_API_KEY:
        import urllib.request as ureq, ssl
        ctx2 = ssl.create_default_context(); ctx2.check_hostname = False; ctx2.verify_mode = ssl.CERT_NONE
        challenge_text = "\n\n".join([
            f"{v['icon']} {v['name']}:\n{v['challenge'][:400]}"
            for k, v in challenges.items()
        ])
        synth_msg = f"""Analysiere diese {len(challenges)} unabhängigen Challenges und erstelle eine kurze Synthese (max 100 Wörter):
- Wo stimmen mehrere überein? (= WICHTIG)
- Was sagt nur einer? (= INTERESSANT)
- Was sollte der User als erstes beachten?

{challenge_text}"""
        try:
            payload = json.dumps({
                "model": ANTHROPIC_MODEL_HEAVY, "max_tokens": 400,
                "messages": [{"role": "user", "content": synth_msg}]
            }).encode()
            req = ureq.Request("https://api.anthropic.com/v1/messages", data=payload, method="POST",
                headers={"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
                         "Content-Type": "application/json"})
            resp = json.loads(ureq.urlopen(req, context=ctx2, timeout=30).read())
            synthesis = resp["content"][0]["text"]
        except: pass

    return {
        "challenges": challenges,
        "errors": errors,
        "synthesis": synthesis,
        "models_used": len(challenges)
    }

# ═══════════════════════════════════════════════════════════════════════════
# NOTEBOOKLM ENTERPRISE INTEGRATION v1.0
# Creates notebooks and adds paper text as source via Discovery Engine API
# ═══════════════════════════════════════════════════════════════════════════

def _gcp_get_access_token():
    """Get a fresh GCP access token using OAuth2 refresh token."""
    if not GCP_REFRESH_TOKEN or not GCP_OAUTH_CLIENT_ID or not GCP_OAUTH_CLIENT_SECRET:
        return None
    import urllib.request as ureq, urllib.parse, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    data = urllib.parse.urlencode({
        "client_id": GCP_OAUTH_CLIENT_ID,
        "client_secret": GCP_OAUTH_CLIENT_SECRET,
        "refresh_token": GCP_REFRESH_TOKEN,
        "grant_type": "refresh_token"
    }).encode()
    req = ureq.Request("https://oauth2.googleapis.com/token", data=data, method="POST",
        headers={"Content-Type": "application/x-www-form-urlencoded"})
    try:
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=15).read())
        return resp.get("access_token")
    except Exception as e:
        logger.error(f"GCP token refresh failed: {e}")
        return None

def _notebooklm_api(method, path, body=None, access_token=None):
    """Call NotebookLM Enterprise API via Discovery Engine."""
    token = access_token or _gcp_get_access_token()
    if not token:
        return {"error": "No GCP access token available. Configure GCP_REFRESH_TOKEN."}
    import urllib.request as ureq, ssl
    ctx = ssl.create_default_context(); ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE
    loc = GCP_NOTEBOOKLM_LOCATION
    base = f"https://{loc}-discoveryengine.googleapis.com/v1alpha/projects/{GCP_PROJECT_NUMBER}/locations/{loc}"
    url = f"{base}{path}"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    data = json.dumps(body).encode() if body else None
    req = ureq.Request(url, data=data, method=method, headers=headers)
    try:
        resp = json.loads(ureq.urlopen(req, context=ctx, timeout=30).read())
        return resp
    except Exception as e:
        error_body = ""
        if hasattr(e, 'read'):
            try: error_body = json.loads(e.read())
            except: pass
        return {"error": str(e), "details": error_body}

class NotebookLMRequest(BaseModel):
    title: str = "BEATRIX Paper"
    text: Optional[str] = None
    source_name: Optional[str] = None
    url: Optional[str] = None
    paper_id: Optional[str] = None

@app.get("/api/notebooklm/status")
async def notebooklm_status(user=Depends(require_auth)):
    """Check if NotebookLM Enterprise is configured."""
    configured = bool(GCP_REFRESH_TOKEN and GCP_OAUTH_CLIENT_ID and GCP_OAUTH_CLIENT_SECRET)
    return {
        "configured": configured,
        "project_number": GCP_PROJECT_NUMBER,
        "location": GCP_NOTEBOOKLM_LOCATION,
        "enterprise_url": f"https://notebooklm.cloud.google.com/{GCP_NOTEBOOKLM_LOCATION}/?project={GCP_PROJECT_NUMBER}",
        "missing": [k for k, v in {
            "GCP_OAUTH_CLIENT_ID": GCP_OAUTH_CLIENT_ID,
            "GCP_OAUTH_CLIENT_SECRET": GCP_OAUTH_CLIENT_SECRET,
            "GCP_REFRESH_TOKEN": GCP_REFRESH_TOKEN
        }.items() if not v]
    }

@app.post("/api/notebooklm/create")
async def notebooklm_create(body: NotebookLMRequest, user=Depends(require_auth)):
    """Create a NotebookLM notebook and add paper text/URL as source.
    
    Flow: Create notebook → Add source (text or URL) → Return notebook link.
    User clicks link → opens Enterprise UI → clicks Infographic/Slide Deck.
    """
    # If paper_id given, fetch text from database
    text_content = body.text
    source_name = body.source_name or body.title
    
    if body.paper_id and not text_content:
        session = SessionLocal()
        try:
            doc = session.execute(
                text("SELECT title, extracted_text FROM documents WHERE id = :id"),
                {"id": body.paper_id}
            ).fetchone()
            if doc:
                source_name = doc[0] or source_name
                text_content = doc[1]
                if not body.title or body.title == "BEATRIX Paper":
                    body.title = f"BEATRIX: {source_name[:80]}"
            else:
                raise HTTPException(404, "Paper not found")
        finally:
            session.close()
    
    if not text_content and not body.url:
        raise HTTPException(400, "Provide text, url, or paper_id")
    
    # Step 1: Create notebook
    logger.info(f"NotebookLM: Creating notebook '{body.title}'")
    nb_result = _notebooklm_api("POST", "/notebooks", {"title": body.title})
    
    if "error" in nb_result:
        raise HTTPException(502, f"Failed to create notebook: {nb_result}")
    
    notebook_id = nb_result.get("notebookId")
    if not notebook_id:
        raise HTTPException(502, f"No notebook ID returned: {nb_result}")
    
    logger.info(f"NotebookLM: Notebook created: {notebook_id}")
    
    # Step 2: Add source
    source_type_used = "text" if text_content else "url"
    src_result = {"error": "no source provided"}
    
    if text_content:
        source_body = {
            "userContents": [{
                "textContent": {
                    "sourceName": source_name[:200],
                    "content": text_content[:500000]  # API limit
                }
            }]
        }
        src_result = _notebooklm_api("POST", f"/notebooks/{notebook_id}/sources:batchCreate", source_body)
    elif body.url:
        # PDF URLs don't work with webContent - download and extract text
        if body.url.lower().endswith(".pdf"):
            logger.info(f"NotebookLM: PDF URL detected, downloading and extracting text...")
            source_type_used = "pdf_text"
            try:
                import requests as req_lib
                pdf_resp = req_lib.get(body.url, timeout=30)
                pdf_resp.raise_for_status()
                try:
                    import fitz  # PyMuPDF
                    doc = fitz.open(stream=pdf_resp.content, filetype="pdf")
                    pdf_text = "\n\n".join(page.get_text() for page in doc)
                    doc.close()
                except ImportError:
                    from io import BytesIO
                    try:
                        import pdfplumber
                        with pdfplumber.open(BytesIO(pdf_resp.content)) as pdf:
                            pdf_text = "\n\n".join(p.extract_text() or "" for p in pdf.pages)
                    except ImportError:
                        pdf_text = None
                
                if pdf_text and len(pdf_text.strip()) > 100:
                    source_body = {
                        "userContents": [{
                            "textContent": {
                                "sourceName": source_name[:200],
                                "content": pdf_text[:500000]
                            }
                        }]
                    }
                    src_result = _notebooklm_api("POST", f"/notebooks/{notebook_id}/sources:batchCreate", source_body)
                else:
                    logger.warning("NotebookLM: PDF text extraction yielded no usable text")
                    src_result = {"error": "PDF text extraction failed"}
            except Exception as e:
                logger.warning(f"NotebookLM: PDF download/extract failed: {e}")
                src_result = {"error": f"PDF processing failed: {str(e)}"}
        else:
            # Regular web URL
            source_type_used = "web_url"
            source_body = {
                "userContents": [{
                    "webContent": {
                        "url": body.url,
                        "sourceName": source_name[:200]
                    }
                }]
            }
            src_result = _notebooklm_api("POST", f"/notebooks/{notebook_id}/sources:batchCreate", source_body)
    
    if "error" in src_result:
        logger.warning(f"NotebookLM: Source add failed (notebook still created): {src_result}")
    
    # Build enterprise UI link
    loc = GCP_NOTEBOOKLM_LOCATION
    notebook_url = f"https://notebooklm.cloud.google.com/{loc}/?project={GCP_PROJECT_NUMBER}#notebook={notebook_id}"
    
    return {
        "success": True,
        "notebook_id": notebook_id,
        "notebook_url": notebook_url,
        "title": body.title,
        "source_added": "error" not in src_result,
        "source_type": source_type_used,
        "message": f"Notebook ready! Open the link and click 'Infographic' or 'Slide Deck' in the Studio panel."
    }

# ═══════════════════════════════════════════════════════════════════════════
# CONTENT TRANSFORM PIPELINE v1.0
# Generic content → NotebookLM transformation with auto-podcast generation
# Supports: papers, analyses, briefings, research summaries — any text content
# ═══════════════════════════════════════════════════════════════════════════

class TransformRequest(BaseModel):
    """Generic content transformation request."""
    title: str = "BEATRIX Content"
    text: Optional[str] = None
    url: Optional[str] = None
    paper_id: Optional[str] = None
    source_name: Optional[str] = None
    # Output options
    podcast: bool = True           # Generate audio overview
    podcast_focus: Optional[str] = None  # What to focus the podcast on
    podcast_language: Optional[str] = "de"  # Language code (de, en, etc.)
    # Content type hint
    content_type: Optional[str] = None  # paper, analysis, briefing, report, etc.

def _resolve_content(body, session=None):
    """Resolve content from text, URL, or paper_id. Returns (text, source_name, error)."""
    text_content = body.text
    source_name = body.source_name or body.title
    
    # Resolve from paper_id
    if body.paper_id and not text_content:
        if not session:
            session = SessionLocal()
            close_session = True
        else:
            close_session = False
        try:
            doc = session.execute(
                text("SELECT title, extracted_text FROM documents WHERE id = :id"),
                {"id": body.paper_id}
            ).fetchone()
            if doc:
                source_name = doc[0] or source_name
                text_content = doc[1]
                if not body.title or body.title in ("BEATRIX Content", "BEATRIX Paper"):
                    body.title = f"BEATRIX: {source_name[:80]}"
            else:
                return None, source_name, "Paper not found"
        finally:
            if close_session:
                session.close()
    
    # Resolve from PDF URL - download and extract
    if not text_content and body.url and body.url.lower().endswith(".pdf"):
        try:
            import requests as req_lib
            pdf_resp = req_lib.get(body.url, timeout=30)
            pdf_resp.raise_for_status()
            import fitz
            doc = fitz.open(stream=pdf_resp.content, filetype="pdf")
            text_content = "\n\n".join(page.get_text() for page in doc)
            doc.close()
            if not text_content or len(text_content.strip()) < 100:
                return None, source_name, "PDF text extraction yielded no usable text"
        except Exception as e:
            return None, source_name, f"PDF processing failed: {str(e)}"
    
    return text_content, source_name, None

def _create_notebook_with_source(title, text_content, source_name, url=None):
    """Create notebook and add source. Returns (notebook_id, source_added, error)."""
    # Create notebook
    nb_result = _notebooklm_api("POST", "/notebooks", {"title": title})
    if "error" in nb_result:
        return None, False, f"Failed to create notebook: {nb_result}"
    
    notebook_id = nb_result.get("notebookId")
    if not notebook_id:
        return None, False, f"No notebook ID returned: {nb_result}"
    
    # Add source
    if text_content:
        source_body = {
            "userContents": [{
                "textContent": {
                    "sourceName": source_name[:200],
                    "content": text_content[:500000]
                }
            }]
        }
    elif url:
        source_body = {
            "userContents": [{
                "webContent": {
                    "url": url,
                    "sourceName": source_name[:200]
                }
            }]
        }
    else:
        return notebook_id, False, "No content to add as source"
    
    src_result = _notebooklm_api("POST", f"/notebooks/{notebook_id}/sources:batchCreate", source_body)
    source_added = "error" not in src_result
    
    if not source_added:
        logger.warning(f"Transform: Source add failed: {src_result}")
    
    return notebook_id, source_added, None

def _trigger_audio_overview(notebook_id, focus=None, language_code=None):
    """Trigger audio overview generation. Returns (audio_id, error)."""
    # Note: episodeFocus and languageCode may not be available in all regions yet
    # Start with empty body, try with options if supported
    body = {}
    
    result = _notebooklm_api("POST", f"/notebooks/{notebook_id}/audioOverviews", body)
    
    # If empty body fails too, it's a real error
    if "error" in result and "Unknown name" not in str(result.get("details", "")):
        return None, f"Audio overview failed: {result}"
    
    audio_data = result.get("audioOverview", {})
    if audio_data.get("status") == "AUDIO_OVERVIEW_STATUS_IN_PROGRESS":
        return audio_data.get("audioOverviewId"), None
    elif "error" in result:
        return None, f"Audio overview failed: {result}"
    else:
        return audio_data.get("audioOverviewId"), None

@app.post("/api/transform")
async def transform_content(body: TransformRequest, user=Depends(require_auth)):
    """Generic content transformation pipeline.
    
    Takes any content (text, URL, paper_id) and transforms it into:
    - NotebookLM notebook with source
    - Auto-generated podcast (audio overview)
    - Links to Enterprise UI for Mind Map, Video Overview, Reports
    
    Works for: papers, BCM analyses, client briefings, research summaries,
    consulting reports, or any text content.
    """
    logger.info(f"Transform: Starting pipeline for '{body.title}' (type={body.content_type})")
    
    # Step 1: Resolve content
    text_content, source_name, error = _resolve_content(body)
    if error:
        raise HTTPException(400 if "not found" in error.lower() else 502, error)
    
    if not text_content and not body.url:
        raise HTTPException(400, "Provide text, url, or paper_id")
    
    # Step 2: Create notebook + add source
    notebook_id, source_added, error = _create_notebook_with_source(
        body.title, text_content, source_name, body.url
    )
    if error:
        raise HTTPException(502, error)
    
    logger.info(f"Transform: Notebook {notebook_id} created, source_added={source_added}")
    
    # Step 3: Generate podcast if requested and source was added
    audio_id = None
    audio_status = "skipped"
    if body.podcast and source_added:
        audio_id, audio_error = _trigger_audio_overview(
            notebook_id, 
            focus=body.podcast_focus,
            language_code=body.podcast_language
        )
        if audio_id:
            audio_status = "generating"
            logger.info(f"Transform: Audio overview {audio_id} generating")
        else:
            audio_status = f"failed: {audio_error}"
            logger.warning(f"Transform: Audio overview failed: {audio_error}")
    elif not source_added:
        audio_status = "skipped (no source)"
    
    # Build URLs
    loc = GCP_NOTEBOOKLM_LOCATION
    notebook_url = f"https://notebooklm.cloud.google.com/{loc}/?project={GCP_PROJECT_NUMBER}#notebook={notebook_id}"
    
    return {
        "success": True,
        "notebook_id": notebook_id,
        "notebook_url": notebook_url,
        "title": body.title,
        "content_type": body.content_type,
        "source": {
            "added": source_added,
            "name": source_name,
            "type": "text" if text_content else "url"
        },
        "podcast": {
            "status": audio_status,
            "audio_id": audio_id,
            "language": body.podcast_language
        },
        "available_outputs": {
            "podcast": audio_status == "generating",
            "mind_map": True,
            "video_overview": True,
            "reports": True,
            "infographic": False,  # Not yet in Enterprise API
            "slide_deck": False    # Not yet in Enterprise API
        },
        "message": "Content transformed! Podcast is generating. Open notebook for Mind Map, Video Overview, and Reports."
    }

@app.get("/api/transform/{notebook_id}/status")
async def transform_status(notebook_id: str, user=Depends(require_auth)):
    """Check the status of a transform job (especially podcast generation)."""
    loc = GCP_NOTEBOOKLM_LOCATION
    notebook_url = f"https://notebooklm.cloud.google.com/{loc}/?project={GCP_PROJECT_NUMBER}#notebook={notebook_id}"
    
    # Try to get audio overview status
    result = _notebooklm_api("GET", f"/notebooks/{notebook_id}/audioOverviews/default")
    
    audio_status = "unknown"
    if "error" in result:
        audio_status = "not_found_or_generating"
    else:
        audio_data = result.get("audioOverview", result)
        status = audio_data.get("status", "unknown")
        if "COMPLETE" in status.upper():
            audio_status = "ready"
        elif "PROGRESS" in status.upper():
            audio_status = "generating"
        elif "FAIL" in status.upper():
            audio_status = "failed"
        else:
            audio_status = status
    
    return {
        "notebook_id": notebook_id,
        "notebook_url": notebook_url,
        "podcast_status": audio_status,
        "message": "Open notebook link to access all outputs."
    }
